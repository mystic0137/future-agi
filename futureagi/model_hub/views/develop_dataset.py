import ast
import base64
import copy
import io
import json
import math
import os
import re
import shutil
import time
import traceback
import uuid
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from functools import partial
from queue import Queue

import json_repair
import numpy as np
import pandas as pd
import requests
import structlog
import weaviate
from django.conf import settings
from django.core.exceptions import ValidationError
from django.db import close_old_connections, connection, transaction
from django.db.models import (
    Case,
    Count,
    DateTimeField,
    F,
    FloatField,
    IntegerField,
    OuterRef,
    Prefetch,
    Q,
    Subquery,
    Sum,
    Value,
    When,
)
from django.db.models.functions import Cast, Coalesce
from django.forms import model_to_dict
from django.http import FileResponse, Http404
from django.shortcuts import get_object_or_404
from django.utils import timezone
from docx import Document
from pinecone import Pinecone
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from qdrant_client import QdrantClient
from rest_framework import serializers, viewsets
from rest_framework.decorators import action
from rest_framework.generics import CreateAPIView
from rest_framework.parsers import FormParser, JSONParser, MultiPartParser
from rest_framework.permissions import IsAuthenticated
from rest_framework.views import APIView
from weaviate import AuthApiKey

from accounts.models import OrgApiKey
from accounts.models.user import User
from accounts.serializers.org_api_key import OrgApiKeySerializer
from agentic_eval.core.embeddings.embedding_manager import (
    EmbeddingManager,
    model_manager,
)
from tfc.telemetry import wrap_for_thread

logger = structlog.get_logger(__name__)
from agentic_eval.core_evals.fi_evals import *  # noqa: F403
from agentic_eval.core_evals.fi_utils.token_count_helper import calculate_total_cost
from agentic_eval.core_evals.run_prompt.litellm_response import RunPrompt
from analytics.utils import (
    MixpanelEvents,
    get_mixpanel_properties,
    track_mixpanel_event,
)
from evaluations.constants import AGENT_EVALUATOR_TYPE_ID, FUTUREAGI_EVAL_TYPES
from model_hub.constants import (
    CREATE_KB_SDK_CODE,
    MAX_KB_SIZE,
    PYTHON_ADD_COLS,
    PYTHON_ADD_ROWS,
    UPDATE_KB_SDK_CODE,
    get_curl_ts_code,
)
from model_hub.models.api_key import ApiKey, SecretModel
from model_hub.models.choices import (
    BooleanChoices,
    CellStatus,
    DatasetSourceChoices,
    DatasetStatus,
    DataTypeChoices,
    DateTimeFormatChoices,
    EvalExplanationSummaryStatus,
    LiteLlmModelProvider,
    ModelChoices,
    ModelTypes,
    ProviderLogoUrls,
    SourceChoices,
    StatusType,
    determine_data_type,
)
from model_hub.models.develop_annotations import Annotations, AnnotationsLabels
from model_hub.models.develop_dataset import (
    Cell,
    Column,
    Dataset,
    Files,
    KnowledgeBaseFile,
    Row,
)
from model_hub.models.develop_optimisation import (
    OptimizationDataset,
)
from model_hub.models.evals_metric import EvalTemplate, Feedback, UserEvalMetric
from model_hub.models.experiments import ExperimentDatasetTable, ExperimentsTable
from model_hub.models.optimize_dataset import OptimizeDataset
from model_hub.models.run_prompt import PromptVersion, RunPrompter
from model_hub.serializers.develop_dataset import (
    ColumnSerializer,
    CompareDatasetSerializer,
    DatasetSerializer,
    FeedbackSerializer,
    FileSerializer,
    KnowledgeBaseFileSerializer,
    UploadFileForm,
)
from model_hub.serializers.develop_optimisation import EvalTemplateSerializer
from model_hub.serializers.eval_runner import UserEvalSerializer
from model_hub.serializers.experiments import DerivedDatasetSerializer
from model_hub.services.derived_variable_service import (
    cleanup_derived_variables_for_column,
    rename_derived_variables_for_column,
    rename_derived_variables_in_run_prompter,
)
from model_hub.tasks.develop_dataset import ingest_files_to_s3, remove_kb_files
from model_hub.types import ConversionResult
from model_hub.utils.eval_reasons import (
    MIN_ROWS_FOR_CRITICAL_ISSUES,
    get_explanation_summary,
)
from model_hub.utils.eval_result_columns import infer_eval_result_column_data_type
from model_hub.utils.evals import (
    FUNCTION_CONFIG_EVALS,
    NOT_UI_EVALS,
    USE_CASE_MAPPING,
)
from model_hub.utils.file_reader import FileProcessor
from model_hub.utils.function_eval_params import (
    has_function_params_schema,
    normalize_eval_runtime_config,
    params_with_defaults_for_response,
)
from model_hub.utils.kb_helpers import (
    cancel_kb_ingestion_workflow,
    schedule_kb_ingestion_on_commit,
)
from model_hub.utils.SQL_queries import SQLQueryHandler
from model_hub.utils.synthetic_task_manager import SyntheticTaskManager
from model_hub.utils.utils import contains_sql, get_diff
from model_hub.views.eval_runner import EvaluationRunner
from model_hub.views.run_prompt import PROVIDERS_WITH_JSON
from model_hub.views.utils.constants import EVAL_OUTPUT_TYPES
from model_hub.views.utils.evals import process_eval_for_single_row
from model_hub.views.utils.utils import (
    get_recommendations,
    update_column_id,
    validate_file_url,
)
from sdk.utils.helpers import _get_api_call_type
from tfc.settings.settings import BASE_URL, HUGGINGFACE_API_TOKEN

# Define a Temporal activity for running the evaluation
from tfc.ee_gates import strip_turing_from_config_options
from tfc.temporal import temporal_activity
from tfc.utils.error_codes import get_error_message
from tfc.utils.functions import (
    calculate_column_average,
    get_eval_stats,
    get_prompt_stats,
)
from tfc.utils.general_methods import GeneralMethods
from tfc.utils.parse_errors import parse_serialized_errors
from tfc.utils.storage import (
    delete_compare_folder,
    download_json_from_s3,
    upload_audio_to_s3,
    upload_audio_to_s3_duration,
    upload_compare_json_to_s3,
    upload_document_to_s3,
    upload_file_to_s3,
    upload_image_to_s3,
)
from tfc.constants.api_calls import APICallStatusChoices, APICallTypeChoices
try:
    from ee.usage.utils.usage_entries import ROW_LIMIT_REACHED_MESSAGE, log_and_deduct_cost_for_resource_request
except ImportError:
    ROW_LIMIT_REACHED_MESSAGE = None
    log_and_deduct_cost_for_resource_request = None

# =============================================================================
# Standalone helper functions for Temporal activities
# =============================================================================


def _process_other_datasets_impl(
    base_val,
    col_name,
    og_cell,
    columns_lookup,
    data_by_dataset,
    ds,
    dynamic_sources,
    i,
):
    """Standalone implementation of process_other_datasets for Temporal activities."""
    try:
        close_old_connections()
        ds_id = str(ds.id)
        if ds_id in data_by_dataset and base_val in data_by_dataset[ds_id]:
            cells_dict = data_by_dataset[ds_id][base_val]
            cell = cells_dict.get(col_name)
            if cell:
                col_obj = columns_lookup.get((ds.id, col_name)) or Column.objects.get(
                    dataset=ds, name=col_name
                )
                if col_obj.source in dynamic_sources:
                    return None
                value_infos = json.loads(cell.value_infos) if cell.value_infos else {}
                metadata = {}
                if isinstance(value_infos, dict):
                    metadata = value_infos.get("metadata") or {}
                    if isinstance(metadata, str):
                        metadata = json.loads(metadata)
                diff_value = None
                cell_value = cell.value
                included_sources = [
                    SourceChoices.EXTRACTED_JSON.value,
                    SourceChoices.CLASSIFICATION.value,
                    SourceChoices.EXTRACTED_ENTITIES.value,
                    SourceChoices.API_CALL.value,
                    SourceChoices.PYTHON_CODE.value,
                    SourceChoices.VECTOR_DB.value,
                    SourceChoices.CONDITIONAL.value,
                    SourceChoices.OTHERS.value,
                    SourceChoices.RUN_PROMPT.value,
                ]
                if (
                    not (i == 0 or "-reason" in col_name)
                    and (col_obj.source in included_sources)
                    and (
                        col_obj.data_type
                        not in [
                            DataTypeChoices.AUDIO.value,
                            DataTypeChoices.IMAGE.value,
                        ]
                    )
                ):
                    diff_value = get_diff(og_cell.value, cell.value)
                return {
                    "col_id": str(col_obj.id),
                    "cell_id": str(cell.id),
                    "cell_value": cell_value,
                    "cell_diff_value": diff_value,
                    "status": cell.status,
                    "value_infos": value_infos,
                    "metadata": metadata,
                    "cell_row_id": str(cell.row_id),
                }
        return None
    finally:
        close_old_connections()


def _prepare_compare_dataset_impl(
    dataset_id,
    common_base_values,
    base_column_name,
    data_by_dataset,
    comparison_datasets,
    columns_lookup,
    main_base_column,
    common_columns,
    compare_id,
    column_config,
    dataset_info,
    dynamic_sources,
):
    """Standalone implementation of prepare_compare_dataset for Temporal activities."""
    try:
        close_old_connections()
        # Build table rows using pre-fetched data
        table = []
        main_ds_id = str(dataset_id)
        with open(f"compare/{compare_id}/metadata.json", "w") as f:
            json.dump(
                {
                    "status": "processing",
                    "total_rows": len(common_base_values),
                    "total_pages": (len(common_base_values) + 9) // 10,
                    "total_processed": 0,
                    "dataset_info": dataset_info,
                    "base_column_name": base_column_name,
                    "base_dataset_id": main_ds_id,
                    "comparison_datasets": [
                        str(comparison_dataset.id)
                        for comparison_dataset in comparison_datasets
                    ],
                    "common_column_names": list(
                        common_columns.union({base_column_name})
                    ),
                },
                f,
                indent=4,
            )
            f.truncate()

        last_index_processed = None
        rowid_in_file = {}
        for index, base_val in enumerate(common_base_values):
            row_data = {"row_id": str(uuid.uuid4())}
            if (
                main_ds_id in data_by_dataset
                and base_val in data_by_dataset[main_ds_id]
            ):
                main_cells = data_by_dataset[main_ds_id][base_val]
                main_cell = main_cells.get(base_column_name)
                if main_cell:
                    value_infos = (
                        json.loads(main_cell.value_infos)
                        if main_cell.value_infos
                        else {}
                    )
                    metadata = value_infos.get("metadata") or {}
                    if isinstance(metadata, str):
                        metadata = json.loads(metadata)
                    row_data[str(main_base_column.id)] = {
                        "cell_value": main_cell.value,
                        "status": main_cell.status,
                        "value_infos": value_infos,
                        "metadata": metadata,
                        "cell_row_id": str(main_cell.row_id),
                    }

            for col_name in common_columns:
                main_cells = data_by_dataset[main_ds_id][base_val]
                og_cell = main_cells.get(col_name)

                with ThreadPoolExecutor(max_workers=10) as executor:
                    futures = []
                    for i, ds in enumerate(comparison_datasets):
                        future = executor.submit(
                            _process_other_datasets_impl,
                            base_val,
                            col_name,
                            og_cell,
                            columns_lookup,
                            data_by_dataset,
                            ds,
                            dynamic_sources,
                            i,
                        )
                        futures.append(future)

                    for future in as_completed(futures):
                        try:
                            result = future.result()
                            if result is not None:
                                row_data[result.pop("col_id")] = result
                        except Exception as e:
                            logger.exception(
                                f"Error in processing other datasets: {str(e)}"
                            )

            table.append(row_data)
            if index % 10 == 0 and index != 0:
                compare_json = {
                    "column_config": column_config,
                    "table": table[index - 10 : index],
                }

                logger.info(f"Writing page {index // 10} to file")
                rowid_in_file.update(
                    {
                        str(table["row_id"]): f"page_{index // 10}"
                        for table in compare_json["table"]
                    }
                )
                upload_compare_json_to_s3(
                    compare_id=compare_id,
                    compare_json=compare_json,
                    page_name=f"page_{index // 10}.json",
                )

                with open(f"compare/{compare_id}/metadata.json", "r+") as f:
                    metadata = json.load(f)
                    metadata["total_processed"] = index // 10
                    metadata["file_row_ids"] = rowid_in_file
                    f.seek(0)
                    json.dump(metadata, f, indent=4)
                    f.truncate()

                last_index_processed = index

        leftover_compare_json = None
        if last_index_processed and last_index_processed != index:
            leftover_compare_json = {
                "column_config": column_config,
                "table": table[last_index_processed:],
            }
            rowid_in_file.update(
                {
                    str(table["row_id"]): f"page_{index // 10 + 1}"
                    for table in leftover_compare_json["table"]
                }
            )
        else:
            leftover_compare_json = {"column_config": column_config, "table": table}
            rowid_in_file.update(
                {
                    str(table["row_id"]): f"page_{index // 10 + 1}"
                    for table in leftover_compare_json["table"]
                }
            )

        if leftover_compare_json:
            upload_compare_json_to_s3(
                compare_id=compare_id,
                compare_json=leftover_compare_json,
                page_name=f"page_{(index // 10) + 1}.json",
            )

        with open(f"compare/{compare_id}/metadata.json", "r+") as f:
            metadata = json.load(f)
            metadata["status"] = "completed"
            metadata["total_processed"] = len(common_base_values) // 10
            metadata["file_row_ids"] = rowid_in_file
            f.seek(0)
            json.dump(metadata, f, indent=4)
            f.truncate()

    except Exception as e:
        logger.exception(f"Error in preparing compare dataset: {str(e)}")
    finally:
        close_old_connections()


class AddRowsFromFile(CreateAPIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    parser_classes = (MultiPartParser, FormParser, JSONParser)

    def post(self, request, *args, **kwargs):
        try:
            form = UploadFileForm(request.POST, request.FILES)
            file = form.files.get("file")
            dataset_id = form.data.get("dataset_id")
            organization = (
                getattr(request, "organization", None) or request.user.organization
            )

            dataset = get_object_or_404(Dataset, id=dataset_id)

            if not file:
                return self._gm.bad_request(get_error_message("NO_FILE_UPLOADED"))

            # Check file size (10 MB limit, matching UI constraint)
            from model_hub.services.dataset_validators import MAX_FILE_SIZE_BYTES

            if file.size > MAX_FILE_SIZE_BYTES:
                return self._gm.bad_request("File size exceeds the 10 MB limit")

            # Process the file
            data, error = FileProcessor.process_file(file_obj=file)

            if error:
                return self._gm.bad_request(error)

            data.columns = data.columns.str.strip()
            data = data.map(lambda x: x.strip() if isinstance(x, str) else x)

            # --- Row Limit Check Start ---
            new_rows_count = data.shape[0]
            existing_rows_count = Row.objects.filter(
                dataset=dataset, deleted=False
            ).count()
            # total_rows_allowed = get_number_of_rows_allowed(organization)

            if log_and_deduct_cost_for_resource_request is not None:
                call_log_row = log_and_deduct_cost_for_resource_request(
                    organization,
                    api_call_type=APICallTypeChoices.ROW_ADD.value,
                    config={"total_rows": existing_rows_count + new_rows_count},
                    workspace=request.workspace,
                )
                if (
                    call_log_row is None
                    or call_log_row.status == APICallStatusChoices.RESOURCE_LIMIT.value
                ):
                    return self._gm.too_many_requests(ROW_LIMIT_REACHED_MESSAGE)
                call_log_row.status = APICallStatusChoices.SUCCESS.value
                call_log_row.save()
            # --- Row Limit Check End ---

            data = data.reset_index(drop=True)

            column_order = dataset.column_order
            column_config = dataset.column_config
            added = False
            # Process Columns, Rows, and Cells
            for column_name in data.columns:
                # Determine the data type dynamically
                data_type = determine_data_type(data[column_name])
                try:
                    column = Column.objects.get(
                        name=column_name,
                        dataset=dataset,
                    )

                except Column.DoesNotExist:
                    added = True
                    column = Column.objects.create(
                        id=uuid.uuid4(),
                        name=column_name,
                        data_type=data_type,
                        source=SourceChoices.OTHERS.value,
                        dataset=dataset,
                    )

                    column_order.append(str(column.id))

                    column_config[str(column.id)] = {
                        "is_visible": True,
                        "is_frozen": None,
                    }

                    rows = Row.objects.filter(dataset=dataset, deleted=False)

                    batch_size = 1000
                    batch = []

                    for row in rows:
                        batch.append(
                            Cell(
                                id=uuid.uuid4(),
                                dataset=dataset,
                                column=column,
                                row=row,
                                value="",
                            )
                        )

                        if len(batch) >= batch_size:
                            with transaction.atomic():
                                Cell.objects.bulk_create(batch)
                            batch = []

                    if batch:
                        with transaction.atomic():
                            Cell.objects.bulk_create(batch)

                if added:
                    dataset.column_order = column_order
                    dataset.column_config = column_config
                    dataset.save()

            # Get all columns for this dataset
            columns = Column.objects.filter(dataset=dataset, deleted=False).exclude(
                source__in=[
                    SourceChoices.EXPERIMENT.value,
                    SourceChoices.EXPERIMENT_EVALUATION.value,
                    SourceChoices.EXPERIMENT_EVALUATION_TAGS.value,
                ]
            )

            last_row = (
                Row.all_objects.filter(dataset=dataset).order_by("-created_at").first()
            )
            if last_row:
                max_order = last_row.order
            else:
                max_order = -1

            for index, row in data.iterrows():
                new_row = Row.objects.create(
                    id=str(uuid.uuid4()), dataset=dataset, order=max_order + 1 + index
                )

                for column in columns:
                    try:
                        value = row[column.name]
                    except KeyError:
                        value = None

                    if column.data_type == DataTypeChoices.IMAGE.value and value:
                        try:
                            # Generate a unique image key using dataset_id
                            image_key = f"images/{dataset_id}/{uuid.uuid4()}"
                            # Upload to S3 and get URL
                            image_url = upload_image_to_s3(
                                str(value), os.getenv("S3_FOR_DATA"), image_key
                            )
                            cell_value = image_url
                        except Exception as e:
                            logger.error(f"Error uploading image: {str(e)}")
                            cell_value = None

                    elif column.data_type == DataTypeChoices.AUDIO.value and value:
                        try:
                            audio_key = f"audio/{dataset_id}/{uuid.uuid4()}"
                            audio_url = upload_audio_to_s3(
                                str(value), os.getenv("S3_FOR_DATA"), audio_key
                            )
                            cell_value = audio_url
                        except Exception as e:
                            logger.exception(
                                f"Error in uploading audio to s3: {str(e)}"
                            )
                            cell_value = None

                    else:
                        cell_value = str(value)

                    Cell.objects.create(
                        id=uuid.uuid4(),
                        dataset=dataset,
                        column=column,
                        row=new_row,
                        value=cell_value,
                    )

            # insert_embeddings_task.delay(
            #     dataset_id=dataset_id,
            #     column_ids=[
            #         column.id
            #         for column in columns
            #         if column.data_type
            #         in [DataTypeChoices.IMAGE.value, DataTypeChoices.AUDIO.value]
            #     ],
            # )

            return self._gm.success_response(
                f"{data.shape[0]} Row(s) added successfully"
            )
        except Exception as e:
            logger.exception(f"Error in adding the rows from file: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_ADD_ROWS_FROM_FILE")
            )


class CloneDatasetView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    parser_classes = (MultiPartParser, FormParser, JSONParser)

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            if log_and_deduct_cost_for_resource_request is not None:
                call_log_row_entry = log_and_deduct_cost_for_resource_request(
                    organization=getattr(request, "organization", None)
                    or request.user.organization,
                    api_call_type=APICallTypeChoices.DATASET_ADD.value,
                    workspace=request.workspace,
                )
                if (
                    call_log_row_entry is None
                    or call_log_row_entry.status
                    == APICallStatusChoices.RESOURCE_LIMIT.value
                ):
                    return self._gm.too_many_requests(
                        get_error_message("DATASET_CREATE_LIMIT_REACHED")
                    )
                call_log_row_entry.status = APICallStatusChoices.SUCCESS.value
                call_log_row_entry.save()
            # Get the source dataset (org-scoped)
            source_dataset = get_object_or_404(
                Dataset,
                id=dataset_id,
                deleted=False,
                organization=getattr(request, "organization", None)
                or request.user.organization,
            )
            new_dataset_name = request.data.get(
                "new_dataset_name", f"Copy of {source_dataset.name}"
            )

            from model_hub.validators.dataset_validators import (
                validate_dataset_name_unique,
            )

            try:
                validate_dataset_name_unique(
                    new_dataset_name,
                    getattr(request, "organization", None) or request.user.organization,
                )
            except Exception as validation_err:
                return self._gm.bad_request(str(validation_err.detail[0]))

            # ------------------- Added Row Check -------------------
            row_count = Row.objects.filter(
                dataset=source_dataset, deleted=False
            ).count()
            if log_and_deduct_cost_for_resource_request is not None:
                call_log_row = log_and_deduct_cost_for_resource_request(
                    getattr(request, "organization", None) or request.user.organization,
                    api_call_type=APICallTypeChoices.ROW_ADD.value,
                    config={"total_rows": row_count},
                    workspace=request.workspace,
                )
                if (
                    call_log_row is None
                    or call_log_row.status == APICallStatusChoices.RESOURCE_LIMIT.value
                ):
                    return self._gm.too_many_requests(ROW_LIMIT_REACHED_MESSAGE)
                call_log_row.status = APICallStatusChoices.SUCCESS.value
                call_log_row.save()

            if source_dataset.dataset_config.get("eval_recommendations", None) is None:
                get_recommendations(source_dataset)

            # Create new dataset
            new_dataset_id = uuid.uuid4()
            new_dataset = Dataset.objects.create(
                id=new_dataset_id,
                name=new_dataset_name,
                organization=getattr(request, "organization", None)
                or request.user.organization,
                model_type=source_dataset.model_type,
                dataset_config=source_dataset.dataset_config,
                user=request.user,
            )

            # Create a mapping of old column IDs to new column IDs
            column_id_mapping = {}

            # Clone only columns with source "OTHERS"
            for old_column in Column.objects.filter(
                dataset=source_dataset, deleted=False, source=SourceChoices.OTHERS.value
            ):
                new_column_id = uuid.uuid4()
                column_id_mapping[str(old_column.id)] = str(new_column_id)

                Column.objects.create(
                    id=new_column_id,
                    dataset=new_dataset,
                    name=old_column.name,
                    data_type=old_column.data_type,
                    source=SourceChoices.OTHERS.value,
                )

            # Update column_order with only the cloned column IDs
            if source_dataset.column_order:
                new_column_order = [
                    column_id_mapping.get(col_id)
                    for col_id in source_dataset.column_order
                    if column_id_mapping.get(col_id)
                ]
                new_dataset.column_order = new_column_order

            # Update column_config with only the cloned column IDs
            if source_dataset.column_config:
                new_column_config = {}
                for old_col_id, config in source_dataset.column_config.items():
                    new_col_id = column_id_mapping.get(old_col_id)
                    if new_col_id:
                        new_column_config[new_col_id] = config
                new_dataset.column_config = new_column_config

            new_dataset.save()

            # Create a mapping of old row IDs to new row IDs
            row_id_mapping = {}

            # Clone rows and cells
            for old_row in Row.objects.filter(dataset=source_dataset, deleted=False):
                new_row_id = uuid.uuid4()
                row_id_mapping[str(old_row.id)] = str(new_row_id)

                # Create new row
                new_row = Row.objects.create(
                    id=new_row_id, dataset=new_dataset, order=old_row.order
                )

                # Clone cells for this row (only for cloned columns)
                cells_to_create = []
                for old_cell in Cell.objects.filter(
                    row=old_row,
                    deleted=False,
                    column_id__in=[
                        uuid.UUID(col_id) for col_id in column_id_mapping.keys()
                    ],
                ):
                    new_column_id = column_id_mapping.get(str(old_cell.column_id))
                    if new_column_id:
                        cells_to_create.append(
                            Cell(
                                id=uuid.uuid4(),
                                dataset=new_dataset,
                                column_id=new_column_id,
                                row=new_row,
                                value=old_cell.value,
                                value_infos=old_cell.value_infos,
                            )
                        )

                # Bulk create cells for better performance
                if cells_to_create:
                    Cell.objects.bulk_create(cells_to_create)

            return self._gm.success_response(
                {
                    "message": "Dataset cloned successfully",
                    "dataset_id": str(new_dataset.id),
                    "dataset_name": new_dataset.name,
                }
            )

        except Exception as e:
            logger.exception(f"Error in cloning the dataset: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_CLONE_DATASET")
            )


class AddAsNewDataset(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    parser_classes = (MultiPartParser, FormParser, JSONParser)

    def post(self, request, *args, **kwargs):
        try:
            dataset_id = request.data.get("dataset_id")
            if log_and_deduct_cost_for_resource_request is not None:
                call_log_row_entry = log_and_deduct_cost_for_resource_request(
                    organization=getattr(request, "organization", None)
                    or request.user.organization,
                    api_call_type=APICallTypeChoices.DATASET_ADD.value,
                    workspace=request.workspace,
                )
                if (
                    call_log_row_entry is None
                    or call_log_row_entry.status
                    == APICallStatusChoices.RESOURCE_LIMIT.value
                ):
                    return self._gm.too_many_requests(
                        get_error_message("DATASET_CREATE_LIMIT_REACHED")
                    )
                call_log_row_entry.status = APICallStatusChoices.SUCCESS.value
                call_log_row_entry.save()
            # Get the source dataset (org-scoped)
            _org = getattr(request, "organization", None) or request.user.organization
            source_dataset = Dataset.objects.filter(
                id=dataset_id, deleted=False, organization=_org
            ).first()
            exp_dataset = False
            if not source_dataset:
                source_dataset = get_object_or_404(
                    ExperimentDatasetTable,
                    id=dataset_id,
                    deleted=False,
                    experiments_datasets_created__dataset__organization=_org,
                )
                exp_dataset = True
            new_dataset_name = request.data.get(
                "name", f"Copy of {source_dataset.name}"
            )
            columns = request.data.get("columns", {})

            if len(set(columns.values())) != len(columns.values()):
                return self._gm.bad_request(get_error_message("DUPLICATE_COLUMN_NAME"))

            from model_hub.validators.dataset_validators import (
                validate_dataset_name_unique as _validate_name_unique,
            )

            try:
                _validate_name_unique(
                    new_dataset_name,
                    getattr(request, "organization", None) or request.user.organization,
                )
            except Exception as validation_err:
                return self._gm.bad_request(str(validation_err.detail[0]))

            if len(columns) < 1:
                return self._gm.bad_request(get_error_message("MISSING_COLUMN_MAPPING"))

            # ------------------- Added Row Check -------------------
            # total_rows_allowed = get_number_of_rows_allowed(getattr(request, "organization", None) or request.user.organization)
            if not exp_dataset:
                row_count = Row.objects.filter(
                    dataset=source_dataset, deleted=False
                ).count()
                if log_and_deduct_cost_for_resource_request is not None:
                    call_log_row = log_and_deduct_cost_for_resource_request(
                        getattr(request, "organization", None) or request.user.organization,
                        api_call_type=APICallTypeChoices.ROW_ADD.value,
                        config={"total_rows": row_count},
                        workspace=request.workspace,
                    )
                    if (
                        call_log_row is None
                        or call_log_row.status == APICallStatusChoices.RESOURCE_LIMIT.value
                    ):
                        return self._gm.too_many_requests(ROW_LIMIT_REACHED_MESSAGE)
                    call_log_row.status = APICallStatusChoices.SUCCESS.value
                    call_log_row.save()
            # ---------------------------------------------------------

            # Create new dataset
            new_dataset_id = uuid.uuid4()
            if exp_dataset:
                model_type = ModelTypes.GENERATIVE_LLM.value
            else:
                model_type = source_dataset.model_type
            new_dataset = Dataset.objects.create(
                id=new_dataset_id,
                name=new_dataset_name,
                organization=getattr(request, "organization", None)
                or request.user.organization,
                model_type=model_type,
                user=request.user,
            )
            if not exp_dataset:
                column_id_mapping = {}

                for col_id, new_name in columns.items():
                    old_column = Column.objects.get(id=col_id)
                    new_column_id = uuid.uuid4()
                    column_id_mapping[str(old_column.id)] = str(new_column_id)

                    new_column = Column.objects.create(
                        id=new_column_id,
                        dataset=new_dataset,
                        name=new_name,
                        data_type=old_column.data_type,
                        source=SourceChoices.OTHERS.value,
                    )

                # Update column_order with only the cloned column IDs
                if source_dataset.column_order:
                    new_column_order = [
                        column_id_mapping.get(col_id)
                        for col_id in source_dataset.column_order
                        if column_id_mapping.get(col_id)
                    ]
                    new_dataset.column_order = new_column_order

                # Update column_config with only the cloned column IDs
                if source_dataset.column_config:
                    new_column_config = {}
                    for old_col_id, config in source_dataset.column_config.items():
                        new_col_id = column_id_mapping.get(old_col_id)
                        if new_col_id:
                            new_column_config[new_col_id] = config
                    new_dataset.column_config = new_column_config

                new_dataset.save()

                # Create a mapping of old row IDs to new row IDs
                row_id_mapping = {}

                # Clone rows and cells
                for old_row in Row.objects.filter(
                    dataset=source_dataset, deleted=False
                ):
                    new_row_id = uuid.uuid4()
                    row_id_mapping[str(old_row.id)] = str(new_row_id)

                    # Create new row
                    new_row = Row.objects.create(
                        id=new_row_id, dataset=new_dataset, order=old_row.order
                    )

                    # Clone cells for this row (only for cloned columns)
                    cells_to_create = []
                    for old_cell in Cell.objects.filter(
                        row=old_row,
                        deleted=False,
                        column_id__in=[
                            uuid.UUID(col_id) for col_id in column_id_mapping.keys()
                        ],
                    ):
                        new_column_id = column_id_mapping.get(str(old_cell.column_id))
                        if new_column_id:
                            cells_to_create.append(
                                Cell(
                                    id=uuid.uuid4(),
                                    dataset=new_dataset,
                                    column_id=new_column_id,
                                    row=new_row,
                                    value=old_cell.value,
                                    value_infos=old_cell.value_infos,
                                )
                            )

                    # Bulk create cells for better performance
                    if cells_to_create:
                        Cell.objects.bulk_create(cells_to_create)

                get_recommendations(new_dataset)

            else:
                list(source_dataset.columns.all())
                experiment = source_dataset.experiments_datasets_created.first()
                user_eval_metric = list(experiment.user_eval_template_ids.all())
                total_columns = []
                for metric in user_eval_metric:
                    runner = EvaluationRunner(
                        user_eval_metric_id=metric.id,
                        is_only_eval=True,
                        format_output=True,
                    )
                    user_eval = UserEvalMetric.objects.get(id=metric.id)
                    cols_used = runner._get_all_column_ids_being_used(
                        user_eval_metric=user_eval
                    )
                    total_columns.extend(cols_used)

                total_columns.append(experiment.column.id)
                if experiment.column.source == SourceChoices.RUN_PROMPT.value:
                    run_prompt = RunPrompter.objects.get(id=experiment.column.source_id)
                    # Extract all column UUIDs from messages
                    message_column_ids = []
                    for message in run_prompt.messages:
                        content = message.get("content", "")
                        # Find all UUIDs between {{ and }}
                        column_ids = re.findall(r"\{\{([^}]*)\}\}", content)
                        # Clean up any whitespace and add to list
                        message_column_ids.extend(
                            [col_id.strip() for col_id in column_ids]
                        )

                    # Add unique column IDs to total_columns
                    total_columns.extend(
                        [uuid.UUID(col_id) for col_id in message_column_ids if col_id]
                    )

                experiment_columns = Column.objects.filter(id__in=columns.keys())
                column_id_mapping = {}
                row_id_mapping = {}
                row_order = 0
                column_order = []
                column_config = {}

                def handle_run_prompt(source_id, dataset):
                    source_run_prompter = RunPrompter.objects.get(id=source_id)

                    run_prompt_obj = RunPrompter.objects.create(
                        dataset=dataset,
                        **{
                            **model_to_dict(
                                source_run_prompter,
                                exclude=["id", "dataset", "organization", "tools"],
                            ),
                            "organization": getattr(request, "organization", None)
                            or request.user.organization,
                        },
                    )

                    run_prompt_obj.tools.set(source_run_prompter.tools.all())

                    run_prompt_obj.save()
                    return run_prompt_obj.id

                source_handlers = {
                    SourceChoices.RUN_PROMPT.value: handle_run_prompt,
                    SourceChoices.EVALUATION.value: lambda source_id, dataset: (
                        UserEvalMetric.objects.create(
                            dataset=dataset,
                            **{
                                **model_to_dict(
                                    UserEvalMetric.objects.get(id=source_id),
                                    exclude=[
                                        "id",
                                        "dataset",
                                        "organization",
                                        "config",
                                        "template",
                                    ],
                                ),
                                "organization": getattr(request, "organization", None)
                                or request.user.organization,
                                "template": UserEvalMetric.objects.get(
                                    id=source_id
                                ).template,
                                "config": UserEvalMetric.objects.get(
                                    id=source_id
                                ).config,
                                "user": request.user,
                            },
                        ).id
                    ),
                    SourceChoices.EXPERIMENT_EVALUATION.value: lambda source_id, dataset: (
                        UserEvalMetric.objects.create(
                            dataset=dataset,
                            **{
                                **model_to_dict(
                                    UserEvalMetric.objects.get(
                                        id=source_id.split("-sourceid-")[1]
                                    ),
                                    exclude=[
                                        "id",
                                        "dataset",
                                        "organization",
                                        "config",
                                        "template",
                                    ],
                                ),
                                "organization": getattr(request, "organization", None)
                                or request.user.organization,
                                "template": UserEvalMetric.objects.get(
                                    id=source_id.split("-sourceid-")[1]
                                ).template,
                                "config": UserEvalMetric.objects.get(
                                    id=source_id.split("-sourceid-")[1]
                                ).config,
                                "user": request.user,
                            },
                        ).id
                    ),
                }

                for column in experiment_columns:
                    new_column_id = uuid.uuid4()
                    column_id_mapping[str(column.id)] = str(new_column_id)
                    new_column = Column.objects.create(
                        id=new_column_id,
                        dataset=new_dataset,
                        source=(
                            column.source
                            if column.source
                            not in [SourceChoices.EXPERIMENT_EVALUATION.value]
                            else SourceChoices.EVALUATION.value
                        ),
                        source_id=(
                            source_handlers[column.source](
                                column.source_id, new_dataset
                            )
                            if column.source
                            in [
                                SourceChoices.RUN_PROMPT.value,
                                SourceChoices.EVALUATION.value,
                                SourceChoices.EXPERIMENT_EVALUATION.value,
                            ]
                            else None
                        ),
                        name=columns.get(str(column.id)),
                        **{
                            k: v
                            for k, v in model_to_dict(column).items()
                            if k not in ["id", "dataset", "source_id", "source", "name"]
                        },
                    )
                    column_order.append(str(new_column.id))
                    column_config[str(column.id)] = {
                        "is_visible": True,
                        "is_frozen": None,
                    }

                    logger.info(f"Created column: {column.name}")

                for column in experiment_columns:
                    if column.source == SourceChoices.RUN_PROMPT.value:
                        new_column = Column.objects.get(
                            id=column_id_mapping[str(column.id)]
                        )
                        run_prompt_col = RunPrompter.objects.get(
                            id=new_column.source_id
                        )
                        new_messages = []
                        for message in run_prompt_col.messages:
                            new_messages.append(
                                update_column_id(
                                    message=message, column_mapping=column_id_mapping
                                )
                            )
                        run_prompt_col.messages = new_messages
                        run_prompt_col.save()

                    elif column.source in [
                        SourceChoices.EVALUATION.value,
                        SourceChoices.EXPERIMENT_EVALUATION.value,
                    ]:
                        new_column = Column.objects.get(
                            id=column_id_mapping[str(column.id)]
                        )
                        user_eval_col = UserEvalMetric.objects.get(
                            id=new_column.source_id
                        )
                        config = user_eval_col.config
                        new_mapping_column_ids = {}
                        mapping_column_ids_initial = config.get("mapping", {})
                        for key in mapping_column_ids_initial.keys():
                            vals = mapping_column_ids_initial[key]
                            if isinstance(vals, list):
                                new_vals = []
                                for col_id in vals:
                                    new_col_id = column_id_mapping[str(col_id)]
                                    new_vals.append(new_col_id)
                            else:
                                new_vals = column_id_mapping[str(vals)]

                            new_mapping_column_ids.update({key: new_vals})

                        deterministic_column_ids = config.get("config", {}).get(
                            "input", []
                        )
                        deterministic_column_ids = [
                            column_id.strip() for column_id in deterministic_column_ids
                        ]
                        deterministic_column_ids = [
                            column_id.replace("{{", "").replace("}}", "")
                            for column_id in deterministic_column_ids
                        ]
                        new_deteministic_column_ids = []
                        for det_column_id in deterministic_column_ids:
                            new_deteministic_column_ids.append(
                                "{{" + column_id_mapping[str(det_column_id)] + "}}"
                            )

                        if new_mapping_column_ids:
                            config["mapping"] = new_mapping_column_ids

                        if deterministic_column_ids:
                            config["input"] = new_deteministic_column_ids

                        logger.info(
                            f"Updated column mapping and deterministic columns: {config}"
                        )
                        user_eval_col.config = config
                        user_eval_col.save()

                for column in experiment_columns:
                    for cell in column.cell_set.filter(
                        deleted=False, row__deleted=False, column__deleted=False
                    ):
                        if cell.row.deleted is False and cell.column.deleted is False:
                            if str(cell.row.id) not in row_id_mapping:
                                new_row_id = uuid.uuid4()
                                row_id_mapping[str(cell.row.id)] = str(new_row_id)
                                Row.objects.create(
                                    id=new_row_id, dataset=new_dataset, order=row_order
                                )
                                row_order += 1

                            Cell.objects.create(
                                id=uuid.uuid4(),
                                row_id=row_id_mapping[str(cell.row.id)],
                                column_id=column_id_mapping[str(cell.column.id)],
                                dataset=new_dataset,
                                **{
                                    k: v
                                    for k, v in model_to_dict(cell).items()
                                    if k not in ["id", "row", "column", "dataset"]
                                },
                            )

                get_recommendations(new_dataset)
                new_dataset.column_order = column_order
                new_dataset.column_config = column_config
                new_dataset.save(update_fields=["column_order", "column_config"])

            # insert_embeddings_task.delay(dataset_id=str(new_dataset.id))
            return self._gm.success_response(
                {
                    "message": "Created New Dataset Successsfully",
                    "dataset_id": str(new_dataset.id),
                    "dataset_name": new_dataset.name,
                }
            )

        except Exception as e:
            logger.exception(f"Error in Adding as New Dataset: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_CLONE_DATASET")
            )


class ColumnConfigView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def get(self, request, column_id):
        column = get_object_or_404(
            Column,
            id=column_id,
            dataset__organization=getattr(request, "organization", None)
            or request.user.organization,
        )

        if column.source == SourceChoices.EXPERIMENT.value:
            config = self.get_experiment_config(column)
        elif column.source == SourceChoices.OPTIMISATION.value:
            config = self.get_optimization_config(column)
        elif column.source == SourceChoices.RUN_PROMPT.value:
            config = self.get_run_prompt_config(column)
        elif column.source == SourceChoices.EVALUATION.value:
            config = self.get_evaluation_config(column)
        elif column.source == SourceChoices.EXPERIMENT_EVALUATION.value:
            config = self.get_experiment_evaluation_config(column)
        elif column.source == SourceChoices.OPTIMISATION_EVALUATION.value:
            config = self.get_optimisation_evaluation_config(column)
        else:
            return self._gm.bad_request("Invalid source")
        return self._gm.success_response(config)

    def get_optimisation_evaluation_config(self, column):
        evaluation_id = column.source_id.split("-sourceid-")[1]
        optimisation_id = column.source_id.split("-sourceid-")[0]
        optimisation = get_object_or_404(OptimizationDataset, id=optimisation_id)
        evaluation = get_object_or_404(UserEvalMetric, id=evaluation_id)
        return {
            "name": column.name,
            "template": evaluation.template.id,
            "template_config": evaluation.template.config,
            "description": evaluation.template.description,
            "config": evaluation.config,
            "status": evaluation.status,
            "optimisation_name": optimisation.name,
            "optimisation_config": optimisation.model_config,
        }

    def get_experiment_config(self, column):
        experiment = get_object_or_404(ExperimentsTable, id=column.source_id)
        user_eval_template_ids = list(
            experiment.user_eval_template_ids.values("id", "name")
        )

        return {
            "name": column.name,
            "prompt_config": experiment.prompt_config,
            "status": experiment.status,
            "user_eval_template_ids": user_eval_template_ids,
        }

    def get_optimization_config(self, column):
        optimization = get_object_or_404(OptimizationDataset, id=column.source_id)
        user_eval_template_ids = list(
            optimization.user_eval_template_ids.values("id", "name")
        )
        return {
            "name": column.name,
            "optimize_type": optimization.optimize_type,
            "optimized_k_prompts": optimization.optimized_k_prompts,
            "messages": optimization.messages,
            "status": optimization.status,
            "model_config": optimization.model_config,
            "user_eval_template_ids": user_eval_template_ids,
        }

    def get_run_prompt_config(self, column):
        run_prompt = get_object_or_404(RunPrompter, id=column.source_id)
        tool_names = list(run_prompt.tools.values_list("name", flat=True))
        return {
            "name": column.name,
            "model": run_prompt.model,
            "messages": run_prompt.messages,
            "output_format": run_prompt.output_format,
            "temperature": run_prompt.temperature,
            "frequency_penalty": run_prompt.frequency_penalty,
            "presence_penalty": run_prompt.presence_penalty,
            "max_tokens": run_prompt.max_tokens,
            "top_p": run_prompt.top_p,
            "response_format": run_prompt.response_format,
            "tool_choice": run_prompt.tool_choice,
            "tools": tool_names,
        }

    def get_evaluation_config(self, column):
        evaluation = get_object_or_404(UserEvalMetric, id=column.source_id)
        return {
            "name": column.name,
            "template": evaluation.template.id,
            "template_config": evaluation.template.config,
            "description": evaluation.template.description,
            "config": evaluation.config,
            "status": evaluation.status,
        }

    def get_experiment_evaluation_config(self, column):
        evaluation_id = column.source_id.split("-sourceid-")[1]
        experiment_id = column.source_id.split("-sourceid-")[0]
        evaluation = get_object_or_404(UserEvalMetric, id=evaluation_id)
        experiment_dataset = get_object_or_404(ExperimentDatasetTable, id=experiment_id)
        return {
            "name": column.name,
            "template": evaluation.template.id,
            "template_config": evaluation.template.config,
            "description": evaluation.template.description,
            "config": evaluation.config,
            "status": evaluation.status,
            "experiment_dataset": experiment_dataset.name,
            "experiment_dataset_config": experiment_dataset.prompt_config,
        }


class GetDatasetsView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):  # Changed from 'post' to 'get'
        try:
            # Get pagination and sorting parameters
            page_size = request.GET.get(
                "page_size", 10
            )  # Changed from request.data.get
            current_page = request.GET.get("page", 0)  # Changed from request.data.get
            search_text = request.GET.get(
                "search_text", ""
            )  # Changed from request.data.get
            sort_params = request.GET.get("sort", [])  # Changed from request.data.get

            # Convert string parameters to appropriate types
            try:
                page_size = int(page_size)
                current_page = int(current_page)
            except (ValueError, TypeError):
                page_size = 10
                current_page = 0

            # Clamp page_size to valid range (aligned with UI/MCP: 1-100)
            from model_hub.constants import MAX_PAGE_SIZE

            page_size = max(1, min(page_size, MAX_PAGE_SIZE))

            # Parse sort_params if it's a string
            if isinstance(sort_params, str):
                try:
                    sort_params = json.loads(sort_params)
                except (ValueError, TypeError):
                    sort_params = []

            # Base queryset with annotations for counts
            queryset = (
                Dataset.objects.filter(
                    organization=getattr(request, "organization", None)
                    or request.user.organization,
                    deleted=False,
                    source__in=[
                        DatasetSourceChoices.DEMO.value,
                        DatasetSourceChoices.BUILD.value,
                        DatasetSourceChoices.OBSERVE.value,
                    ],
                )
                .exclude(
                    # Exclude datasets that are linked to scenarios
                    scenarios__isnull=False
                )
                .annotate(
                    number_of_datapoints=Coalesce(
                        Subquery(
                            Row.objects.filter(dataset=OuterRef("pk"), deleted=False)
                            .values("dataset")
                            .annotate(count=Count("id"))
                            .values("count")[:1]
                        ),
                        Value(0),
                    ),
                    number_of_experiments=Coalesce(
                        Subquery(
                            ExperimentsTable.objects.filter(
                                dataset=OuterRef("pk"), deleted=False
                            )
                            .values("dataset")
                            .annotate(count=Count("id"))
                            .values("count")[:1]
                        ),
                        Value(0),
                    ),
                    number_of_optimisations=Coalesce(
                        Subquery(
                            OptimizeDataset.objects.filter(
                                column__dataset=OuterRef("pk"), deleted=False
                            )
                            .values("column__dataset")
                            .annotate(count=Count("id"))
                            .values("count")[:1]
                        ),
                        Value(0),
                    ),
                    derived_datasets_count=Coalesce(
                        Subquery(
                            ExperimentDatasetTable.objects.filter(
                                experiment__dataset=OuterRef("pk"),
                                experiment__deleted=False,
                                deleted=False,
                            )
                            .values("experiment__dataset")
                            .annotate(count=Count("id", distinct=True))
                            .values("count")[:1]
                        ),
                        Value(0),
                    ),
                )
            )
            # Apply search filter if search_text is provided
            if search_text:
                queryset = queryset.filter(name__icontains=search_text)
            # Apply sorting
            column_mapping = {
                "name": "name",
                "number_of_datapoints": "number_of_datapoints",
                "number_of_datapoints": "number_of_datapoints",
                "number_of_experiments": "number_of_experiments",
                "number_of_experiments": "number_of_experiments",
                "number_of_optimisations": "number_of_optimisations",
                "number_of_optimisations": "number_of_optimisations",
                "derived_datasets": "derived_datasets_count",
                "derived_datasets_count": "derived_datasets_count",
                "created_at": "created_at",
                "created_at": "created_at",
            }
            for sort_param in sort_params:
                column = sort_param.get("column_id")
                order = sort_param.get("type", "ascending")
                if column in column_mapping:
                    sort_field = column_mapping[column]
                    if order == "descending":
                        sort_field = f"-{sort_field}"
                    queryset = queryset.order_by(sort_field)
            # Get total count before pagination
            total_datasets = queryset.count()
            total_pages = (total_datasets + page_size - 1) // page_size
            # Apply pagination at database level
            start = current_page * page_size
            queryset = queryset[start : start + page_size]
            # Prepare response data
            datasets = []
            for dataset in queryset:
                datasets.append(
                    {
                        "id": str(dataset.id),
                        "name": dataset.name,
                        "number_of_datapoints": dataset.number_of_datapoints,
                        "number_of_experiments": dataset.number_of_experiments,
                        "number_of_optimisations": dataset.number_of_optimisations,
                        "derived_datasets": dataset.derived_datasets_count,
                        "created_at": dataset.created_at.strftime("%Y-%m-%d %H:%M"),
                        "dataset_type": dataset.model_type,
                    }
                )

            return self._gm.success_response(
                {
                    "datasets": datasets,
                    "total_pages": total_pages,
                    "total_count": total_datasets,
                }
            )
        except Exception as e:
            logger.exception(f"Error in fetching datasets: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_DATASETS")
            )


def calculate_column_avg_score(column_id):
    avg_score = None
    column = get_object_or_404(Column, id=column_id)
    if column.source == SourceChoices.EVALUATION.value:
        # Get all cells for this column
        cells = Cell.objects.filter(column=column, deleted=False)
        if cells.exists():
            values = []
            for cell in cells:
                try:
                    if cell.value is not None:
                        if cell.value.lower() in ["pass", "fail", "passed", "failed"]:
                            values.append(
                                cell.value.lower() == "pass"
                                or cell.value.lower() == "passed"
                            )
                        elif cell.value.isdigit():
                            values.append(float(cell.value))
                        else:
                            continue
                except (ValueError, TypeError):
                    continue

            if values:
                avg_score = sum(values) / len(values)
        else:
            logger.info("no values found")
    return avg_score


class GetDatasetTableView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    def _apply_filters(self, all_cells, rows, filters, error_messages, columns_map):
        # Apply filters if any
        for filter_item in filters:
            try:
                column_id = filter_item.get("column_id") or filter_item.get("columnId")
                filter_config = filter_item.get("filter_config", {}) or filter_item.get(
                    "filterConfig", {}
                )

                if not column_id or not filter_config:
                    continue

                filter_type = filter_config.get("filter_type") or filter_config.get(
                    "filterType"
                )
                filter_op = filter_config.get("filter_op") or filter_config.get(
                    "filterOp"
                )
                filter_value = filter_config.get("filter_value") or filter_config.get(
                    "filterValue"
                )
                cells = all_cells.filter(column_id=column_id)

                #

                if filter_type == "number":
                    operator_map = {
                        "greater_than": "gt",
                        "less_than": "lt",
                        "equals": "exact",
                        "not_equals": "exact",
                        "greater_than_or_equal": "gte",
                        "less_than_or_equal": "lte",
                    }

                    if filter_op in ["between", "not_in_between"]:
                        # Expect filter_value to be a list/array of [min, max]
                        if not isinstance(filter_value, list) or len(filter_value) != 2:
                            message = "Between operations \
                                    require an array of [min, max] values"
                            error_messages.append(message)
                            raise ValueError(message)

                        min_val, max_val = (
                            float(filter_value[0]),
                            float(filter_value[1]),
                        )

                        column_type = columns_map.get(column_id).data_type
                        if column_type == DataTypeChoices.AUDIO.value:
                            cells = cells.filter(
                                value__regex=r"^https?:\/\/[^\s]+$", deleted=False
                            ).annotate(
                                numeric_value=Cast(
                                    F("column_metadata__audio_duration_seconds"),
                                    output_field=FloatField(),
                                )
                            )
                            # cells = cells.filter(id__in=filtered_cell_ids, deleted=False)
                        else:
                            # Regular number filtering on value field
                            cells = cells.filter(
                                value__regex=r"^-?\d*\.?\d+$", deleted=False
                            ).annotate(numeric_value=Cast("value", FloatField()))

                        between_filter = Q(numeric_value__gte=min_val) & Q(
                            numeric_value__lte=max_val
                        )
                        if filter_op == "not_in_between":
                            cells = cells.filter(~between_filter, deleted=False)
                        else:
                            cells = cells.filter(between_filter, deleted=False)

                    elif filter_op in operator_map:
                        filter_value = float(filter_value)

                        column_type = columns_map.get(column_id).data_type
                        filter_kwargs = {
                            f"numeric_value__{operator_map[filter_op]}": filter_value
                        }
                        if column_type == DataTypeChoices.AUDIO.value:
                            # Fetching audio duration from column_metadata for filtering
                            cells = cells.filter(
                                value__regex=r"^https?:\/\/[^\s]+$"
                            ).annotate(
                                numeric_value=Cast(
                                    "column_metadata__audio_duration_seconds",
                                    output_field=FloatField(),
                                )
                            )
                        else:
                            # Regular number filtering on value field
                            cells = cells.filter(
                                value__regex=r"^-?\d*\.?\d+$", deleted=False
                            ).annotate(numeric_value=Cast("value", FloatField()))

                        if filter_op == "not_equals":
                            cells = cells.filter(~Q(**filter_kwargs), deleted=False)
                        else:
                            cells = cells.filter(**filter_kwargs, deleted=False)

                    else:
                        message = "Invalid filter operation. \
                        operations are: greater_than, less_than, equals, \
                        not_equals, greater_than_or_equal, less_than_or_equal, \
                        between, not_in_between"
                        error_messages.append(message)
                        raise ValueError(message)

                elif filter_type == "text" or filter_type == "array":
                    filter_value = filter_value.lower()
                    text_ops = {
                        "contains": {"value__icontains": filter_value},
                        "not_contains": {
                            "value__icontains": filter_value,
                            "negate": True,
                        },
                        "equals": {"value__iexact": filter_value},
                        "not_equals": {
                            "value__iexact": filter_value,
                            "negate": True,
                        },
                        "starts_with": {"value__istartswith": filter_value},
                        "ends_with": {"value__iendswith": filter_value},
                        "in": {"value__in": filter_value},
                        "not_in": {"value__in": filter_value, "negate": True},
                    }

                    if filter_op not in text_ops:
                        message = (
                            "Invalid filter operation. \
                            Allowed operations are: "
                            + ", ".join(text_ops.keys())
                        )
                        error_messages.append(message)
                        raise ValueError(message)

                    filter_kwargs = text_ops[filter_op]
                    if filter_kwargs.pop("negate", False):
                        cells = cells.filter(~Q(**filter_kwargs), deleted=False)
                    else:
                        cells = cells.filter(**filter_kwargs, deleted=False)

                elif filter_type == "boolean":
                    filter_value = filter_value.lower()
                    if filter_value not in ["true", "false"]:
                        raise ValueError(
                            "Invalid filter value. Allowed values are: true, false"
                        )
                    if filter_value == "true":
                        cells = cells.filter(
                            Q(value__icontains="true") | Q(value__iexact="Passed"),
                            deleted=False,
                        )
                    elif filter_value == "false":
                        cells = cells.filter(
                            Q(value__icontains="false") | Q(value__iexact="Failed"),
                            deleted=False,
                        )

                elif filter_type == "datetime":
                    valid_dates = []

                    if isinstance(filter_value, str):
                        try:
                            valid_dates = [
                                datetime.strptime(filter_value, "%Y-%m-%d %H:%M:%S")
                            ]
                        except Exception:
                            pass  # Silently skip invalid format

                    elif isinstance(filter_value, list):
                        for date_str in filter_value:
                            if date_str:
                                try:
                                    valid_dates.append(
                                        datetime.strptime(date_str, "%Y-%m-%d %H:%M:%S")
                                    )
                                except Exception:
                                    continue  # Silently skip invalid format or type

                    # No valid datetime to filter on
                    if not valid_dates and filter_op in [
                        "equals",
                        "not_equals",
                        "greater_than",
                        "less_than",
                        "greater_than_or_equal",
                        "less_than_or_equal",
                    ]:
                        return rows  # skip filtering

                    # Determine filter kwargs
                    if filter_op in ["between", "not_in_between"]:
                        start = valid_dates[0] if len(valid_dates) > 0 else None
                        end = valid_dates[1] if len(valid_dates) > 1 else None

                        if start is None and end is None:
                            return rows  # nothing to filter

                        elif start is None:
                            filter_kwargs = {"value__lte": end}
                        elif end is None:
                            filter_kwargs = {"value__gte": start}
                        else:
                            filter_kwargs = {"value__range": (start, end)}

                        if filter_op == "not_in_between":
                            cells = cells.filter(~Q(**filter_kwargs), deleted=False)
                        else:
                            cells = cells.filter(**filter_kwargs, deleted=False)

                    else:
                        datetime_ops = {
                            "equals": {"value": valid_dates[0]},
                            "not_equals": {"value": valid_dates[0], "negate": True},
                            "greater_than": {"value__gt": valid_dates[0]},
                            "less_than": {"value__lt": valid_dates[0]},
                            "greater_than_or_equal": {"value__gte": valid_dates[0]},
                            "less_than_or_equal": {"value__lte": valid_dates[0]},
                        }

                        if filter_op in datetime_ops:
                            filter_kwargs = datetime_ops[filter_op]
                            if filter_kwargs.pop("negate", False):
                                cells = cells.filter(~Q(**filter_kwargs), deleted=False)
                            else:
                                cells = cells.filter(**filter_kwargs, deleted=False)

                # Filter the rows based on matching cell row IDs
                rows = rows.filter(
                    id__in=cells.values_list("row_id", flat=True), deleted=False
                )
                # return

            except Exception as e:
                logger.exception(f"error in filter : {e}")
        return rows

    def _apply_sorting(
        self, all_cells, rows, sort_configs, error_messages, columns_map
    ):
        # Apply sorting
        for sort_item in sort_configs:
            # continue
            try:
                column_id = sort_item.get("column_id") or sort_item.get("columnId")
                sort_type = sort_item.get("type") or sort_item.get("type")

                if not column_id or not sort_type:
                    continue

                cells = all_cells.filter(column_id=column_id)
                column_type = columns_map.get(column_id).data_type
                sort_prefix = "-" if sort_type == "descending" else ""

                if (
                    column_type == DataTypeChoices.TEXT.value
                    or column_type == DataTypeChoices.BOOLEAN.value
                ):
                    cells = cells.order_by(f"{sort_prefix}value")

                elif (
                    column_type == DataTypeChoices.INTEGER.value
                    or column_type == DataTypeChoices.FLOAT.value
                ):
                    cells = cells.filter(
                        value__regex=r"^-?\d*\.?\d+$",
                    ).annotate(numeric_value=Cast("value", FloatField()))
                    cells = cells.order_by(f"{sort_prefix}numeric_value")

                elif column_type == DataTypeChoices.DATETIME.value:
                    cells = cells.annotate(
                        datetime_value=Cast("value", DateTimeField()),
                        sort_key=Case(
                            When(value__isnull=True, then=Value(1)),
                            default=Value(0),
                        ),
                    )
                    cells = cells.order_by("sort_key", f"{sort_prefix}datetime_value")

                elif column_type == DataTypeChoices.AUDIO.value:
                    cells = cells.annotate(
                        numeric_value=Cast(
                            "column_metadata__audio_duration_seconds",
                            output_field=FloatField(),
                        ),
                        sort_key=Case(
                            When(
                                column_metadata__audio_duration_seconds__isnull=True,
                                then=Value(1),
                            ),
                            default=Value(0),
                        ),
                    )
                    cells = cells.order_by("sort_key", f"{sort_prefix}numeric_value")

                else:
                    continue

                # Use sorted cell row_ids to order rows without massive CASE expression
                row_ids = list(cells.values_list("row_id", flat=True))
                if not row_ids:
                    rows = rows.none()
                else:
                    # Preserve cell-sorted order using a single subquery annotation
                    # instead of O(n) CASE/WHEN clauses
                    preserved = Case(
                        *[When(id=rid, then=Value(i)) for i, rid in enumerate(row_ids)],
                        default=Value(len(row_ids)),
                    )
                    rows = (
                        rows.filter(id__in=row_ids, deleted=False)
                        .annotate(_sort_order=preserved)
                        .order_by("_sort_order")
                    )

            except Exception as e:
                logger.error(f"error in sort : {e}")
        return rows

    def _apply_search(self, cells, rows, search, dataset_id, error_messages):
        search_key = search.get("key", "")
        search_value = search.get("type", ["text", "image", "audio"])

        if not search_key:
            return rows, {}

        if isinstance(search_key, str):
            search_key = search_key.lower()

        matched_cell_ids = set()
        search_results = {}

        # Text search
        if "text" in search_value:
            results = SQLQueryHandler.search_cells_by_text(search_key, dataset_id)

            for result in results:
                cell_id, key_exists, indices = result
                matched_cell_ids.add(cell_id)
                search_results[str(cell_id)] = {
                    "key_exists": key_exists,
                    "indices": indices,
                }

        # Image and audio search
        # for search_type in ["image", "audio"]:
        #     embedding_manager = EmbeddingManager()
        #     if search_type in search_value:
        #         if search_type == "image":
        #             has_image = Column.objects.filter(dataset_id=dataset_id, data_type=DataTypeChoices.IMAGE.value).exists()
        #             if not has_image:
        #                 logger.info(f"No image columns found for dataset_id: {dataset_id}")
        #                 continue
        #         elif search_type == "audio":
        #             has_audio = Column.objects.filter(dataset_id=dataset_id, data_type=DataTypeChoices.AUDIO.value).exists()
        #             if not has_audio:
        #                 logger.info(f"No audio columns found for dataset_id: {dataset_id}")
        #                 continue

        #         res = embedding_manager.retrieve_rag_cells(
        #             query=search_key,
        #             dataset_id=dataset_id,
        #             input_type="text",
        #             filter_by={"input_type": search_type},
        #         )

        #         if res is None:
        #             error_messages.append(
        #                 f"Failed to retrieve {search_type} embeddings for '{search_key}'"
        #             )
        #             continue

        #         matched_cell_ids.update(
        #             {
        #                 row["metadata"].get("cell_id")
        #                 for row in res
        #                 if row["metadata"].get("cell_id")
        #             }
        #         )

        #         for row in res:
        #             cell_id = row["metadata"].get("cell_id")
        #             if cell_id:
        #                 search_results[str(cell_id)] = {
        #                     "key_exists": True,
        #                     "indices": None,  # Not applicable for non-text
        #                 }

        # Find rows that contain the matched cells
        matched_cells = cells.filter(
            id__in=matched_cell_ids, deleted=False, dataset=dataset_id
        ).values("row_id")
        matched_row_ids = set(matched_cells.values_list("row_id", flat=True))
        filtered_rows = rows.filter(id__in=matched_row_ids, dataset=dataset_id)

        return filtered_rows, search_results

    def get(self, request, dataset_id, *args, **kwargs):
        try:
            # Get request parameters from query params instead of request.data
            filters = request.GET.get("filters", "[]") or request.GET.get(
                "filters", "[]"
            )
            sort_configs = request.GET.get("sort", "[]") or request.GET.get(
                "sort", "[]"
            )
            search = request.GET.get("search", "{}") or request.GET.get("search", "{}")
            from model_hub.services.dataset_validators import MAX_PAGE_SIZE

            page_size = min(
                int(request.GET.get("page_size", 10))
                or int(request.GET.get("pageSize", 10)),
                MAX_PAGE_SIZE,
            )
            current_page = int(request.GET.get("current_page_index", 0)) or int(
                request.GET.get("currentPageIndex", 0)
            )
            column_config_only = (
                request.GET.get("column_config_only", "false").lower() == "true"
                or request.GET.get("columnConfigOnly", "false").lower() == "true"
            )

            # Parse JSON parameters
            try:
                filters = json.loads(filters) if filters else []
                sort_configs = json.loads(sort_configs) if sort_configs else []
                search = json.loads(search) if search else {}
            except json.JSONDecodeError:
                filters = []
                sort_configs = []
                search = {}
            # Get base dataset and rows
            try:
                dataset = Dataset.objects.select_related("organization").get(
                    id=dataset_id,
                    organization=getattr(request, "organization", None)
                    or request.user.organization,
                    deleted=False,
                )
            except Dataset.DoesNotExist:
                return self._gm.bad_request(get_error_message("DATASET_NOT_FOUND"))
            rows = Row.objects.filter(dataset=dataset, deleted=False).order_by("order")
            existing_column_config = dataset.column_config
            error_messages = []
            # print("exiting sort")
            # Calculate pagination offsets
            start = current_page * page_size
            end = start + page_size

            # Get column configuration - fetch all non-deleted columns for
            # the dataset; use column_order only for sorting.
            column_order = dataset.column_order or []
            column_order_set = set(column_order)
            qs = Column.objects.filter(dataset=dataset, deleted=False)
            if dataset.source != DatasetSourceChoices.EXPERIMENT_SNAPSHOT.value:
                qs = qs.exclude(source__in=[
                    SourceChoices.EXPERIMENT.value,
                    SourceChoices.EXPERIMENT_EVALUATION.value,
                    SourceChoices.EXPERIMENT_EVALUATION_TAGS.value,
                ])
            all_columns = list(qs)
            columns_map = {str(col.id): col for col in all_columns}

            # Names of EVALUATION_REASON columns in this dataset. Reason columns
            # are always created as "{eval_column.name}-reason" (see
            # AddUserEvalView / EditAndRunUserEvalView / eval_group flows), so
            # pairing by name avoids coupling to the source_id format.
            reason_column_names = {
                _c.name
                for _c in all_columns
                if _c.source == SourceChoices.EVALUATION_REASON.value
            }

            # Sort: columns in column_order come first (in that order),
            # then any columns missing from column_order appended at the end.
            ordered_columns = []
            for col_id in column_order:
                if col_id in columns_map:
                    ordered_columns.append(columns_map[col_id])
            for col in all_columns:
                if str(col.id) not in column_order_set:
                    ordered_columns.append(col)
            columns = ordered_columns
            column_config = []

            cells = Cell.objects.filter(row__in=rows, column__in=columns, deleted=False)

            if filters:
                rows = self._apply_filters(
                    cells, rows, filters, error_messages, columns_map
                )
            search_results = {}
            if search:
                rows, search_results = self._apply_search(
                    cells, rows, search, dataset.id, error_messages
                )

            if sort_configs:
                rows = self._apply_sorting(
                    cells, rows, sort_configs, error_messages, columns_map
                )

            # Get total count using DB COUNT instead of loading all rows
            total_rows = rows.count()

            def process_single_column(
                column,
                existing_column_config,
                idx,
                result_queue,
                uem_map,
                run_prompter_map,
                optimisation_map,
                experiment_map,
                columns_map,
            ):
                try:
                    if column_config_only:
                        avg_score = calculate_column_average(column)
                        is_numeric_eval = avg_score.get("is_numeric_eval", False)
                        is_numeric_eval_percentage = avg_score.get(
                            "is_numeric_eval_percentage", False
                        )
                        avg_score = avg_score.get("average", None)
                    else:
                        avg_score = None
                        is_numeric_eval = False
                        is_numeric_eval_percentage = False

                    status = column.status
                    metadata = {"run_prompt": False, "run_prompt_id": []}
                    column_metadata = column.metadata
                    metadata.update(column_metadata)
                    eval_tag = []
                    choices_map = {}
                    reason_column_flag = False

                    # Handle EVALUATION source
                    if column.source == SourceChoices.EVALUATION.value:
                        # Source of truth: does a paired EVALUATION_REASON
                        # column actually exist for this eval column? Pair by
                        # name convention ("{eval_name}-reason") — independent
                        # of source_id format.
                        reason_column_flag = (
                            f"{column.name}-reason" in reason_column_names
                        )
                        user_eval_metric = uem_map.get(column.source_id)
                        if user_eval_metric:
                            status = user_eval_metric.status
                            eval_tag = user_eval_metric.template.eval_tags
                            tmpl_cfg = user_eval_metric.template.config or {}
                            choices_map = tmpl_cfg.get("choices_map", {})
                            # Fallbacks when no paired reason column exists yet:
                            # 1) Explicit template flag (legacy evals created
                            #    before physical reason columns were auto-
                            #    materialized, and composites which always
                            #    populate their aggregate reason cell via
                            #    CompositeEvaluationRunner).
                            # 2) AgentEvaluator / custom evals emit reason
                            #    inline into cell.value_infos rather than a
                            #    separate column.
                            if not reason_column_flag:
                                reason_column_flag = bool(
                                    tmpl_cfg.get("reason_column", False)
                                    or tmpl_cfg.get("eval_type_id")
                                    == AGENT_EVALUATOR_TYPE_ID
                                    or tmpl_cfg.get("custom_eval", False)
                                )
                            if column_config_only:
                                runner = EvaluationRunner(
                                    user_eval_metric_id=user_eval_metric.id,
                                    is_only_eval=True,
                                    format_output=True,
                                    source="dataset_evaluation",
                                    source_id=user_eval_metric.template.id,
                                )
                                cols_used = runner._get_all_column_ids_being_used(
                                    user_eval_metric=user_eval_metric
                                )
                                for eval_column_id in cols_used:
                                    eval_column = columns_map.get(eval_column_id)
                                    if (
                                        eval_column
                                        and eval_column.source
                                        == SourceChoices.RUN_PROMPT.value
                                    ):
                                        metadata["run_prompt"] = True
                                        metadata["run_prompt_id"].append(eval_column.id)

                    # Handle OPTIMISATION source
                    elif column.source == SourceChoices.OPTIMISATION.value:
                        if optimisation_map.get(column.source_id):
                            status = optimisation_map.get(column.source_id).status

                    # Handle OPTIMISATION_EVALUATION source
                    elif column.source == SourceChoices.OPTIMISATION_EVALUATION.value:
                        optimisation_id = column.source_id.split("-sourceid-")[0]
                        user_metric_id = column.source_id.split("-sourceid-")[1]
                        optimisation = optimisation_map.get(optimisation_id)
                        status = optimisation.status
                        if column_config_only:
                            user_eval = uem_map.get(user_metric_id)
                            if user_eval:
                                choices_map = user_eval.template.config.get(
                                    "choices_map", {}
                                )
                                reason_column_flag = bool(
                                    user_eval.template.config.get(
                                        "reason_column", False
                                    )
                                )
                                runner = EvaluationRunner(
                                    user_eval_metric_id=user_metric_id,
                                    is_only_eval=True,
                                    format_output=True,
                                    source="optimization",
                                    source_id=user_eval.template.id,
                                )
                                cols_used = runner._get_all_column_ids_being_used(
                                    user_eval_metric=user_eval
                                )
                                for eval_column_id in cols_used:
                                    eval_column = columns_map.get(eval_column_id)
                                    if (
                                        eval_column
                                        and eval_column.source
                                        == SourceChoices.RUN_PROMPT.value
                                    ):
                                        metadata["run_prompt"] = (
                                            True  # to show in feedback to improve prompt
                                        )
                                        metadata["run_prompt_id"].append(eval_column.id)

                    # Handle RUN_PROMPT source
                    elif column.source == SourceChoices.RUN_PROMPT.value:
                        model_name = "gpt-4o-mini"  # Default model name if run_prompter not present
                        run_prompter = run_prompter_map.get(column.source_id)
                        if run_prompter:
                            status = run_prompter.status
                            model_name = run_prompter.model
                        avg_latency = avg_tokens = avg_cost = 0

                        if status == StatusType.COMPLETED.value:
                            if column_config_only:
                                # Use DB aggregation instead of per-cell Python loop
                                agg = (
                                    Cell.objects.filter(
                                        column=column,
                                        deleted=False,
                                    )
                                    .exclude(
                                        status=CellStatus.ERROR.value,
                                    )
                                    .aggregate(
                                        cell_count=Count("id"),
                                        total_prompt_tokens=Sum("prompt_tokens"),
                                        total_completion_tokens=Sum(
                                            "completion_tokens"
                                        ),
                                        total_response_time=Sum("response_time"),
                                    )
                                )
                                cell_count = agg["cell_count"] or 0
                                if cell_count:
                                    total_prompt = agg["total_prompt_tokens"] or 0
                                    total_completion = (
                                        agg["total_completion_tokens"] or 0
                                    )
                                    total_response = agg["total_response_time"] or 0
                                    total_tokens = total_prompt + total_completion

                                    # Single cost calculation on totals (pricing is linear)
                                    cost_calculation = calculate_total_cost(
                                        model_name,
                                        {
                                            "prompt_tokens": total_prompt,
                                            "completion_tokens": total_completion,
                                        },
                                    )
                                    total_cost = (
                                        cost_calculation.get("total_cost", 0) or 0
                                    )

                                    metadata["average_latency"] = (
                                        total_response / cell_count
                                    )
                                    metadata["average_cost"] = total_cost / cell_count
                                    metadata["average_tokens"] = (
                                        total_tokens / cell_count
                                    )
                                else:
                                    metadata.update(
                                        {
                                            "average_latency": None,
                                            "average_cost": None,
                                            "average_tokens": None,
                                        }
                                    )

                    # Handle EXPERIMENT source
                    elif column.source == SourceChoices.EXPERIMENT.value:
                        if experiment_map.get(column.source_id):
                            status = experiment_map.get(column.source_id).status

                    result = {
                        "id": str(column.id),
                        "name": column.name,
                        "eval_tag": eval_tag,
                        "is_frozen": (existing_column_config or {})
                        .get(str(column.id), {})
                        .get("is_frozen", None),
                        "is_visible": (existing_column_config or {})
                        .get(str(column.id), {})
                        .get("is_visible", True),
                        "data_type": column.data_type,
                        "source_type": "text",
                        "origin_type": column.source,
                        "source_id": (
                            column.source_id.split("-sourceid-")[1]
                            if column.source_id and "-sourceid-" in column.source_id
                            else column.source_id
                        ),
                        "average_score": avg_score,
                        "order_index": idx,
                        "status": status,
                        "metadata": metadata,
                        "choices_map": choices_map,
                        "reason_column": reason_column_flag,
                        "is_numeric_eval": is_numeric_eval,
                        "is_numeric_eval_percentage": is_numeric_eval_percentage,
                    }
                    result_queue.put((idx, result))
                except Exception as e:
                    logger.exception(f"Error processing column {column.id}: {str(e)}")
                    result_queue.put((idx, None))

            # Combine queries for both UEM and RunPrompter objects
            source_ids = {
                "uem": set(),
                "run_prompter": set(),
                "optimization": set(),
                "experiment": set(),
            }

            for column in columns:
                if column.source == SourceChoices.EVALUATION.value:
                    source_ids["uem"].add(column.source_id)
                elif column.source == SourceChoices.OPTIMISATION_EVALUATION.value:
                    parts = column.source_id.split("-sourceid-")
                    if len(parts) == 2:
                        source_ids["optimization"].add(parts[0])
                        source_ids["uem"].add(parts[1])
                elif column.source == SourceChoices.RUN_PROMPT.value:
                    source_ids["run_prompter"].add(column.source_id)
                elif column.source == SourceChoices.OPTIMISATION.value:
                    source_ids["optimization"].add(column.source_id)
                elif column.source == SourceChoices.EXPERIMENT.value:
                    source_ids["experiment"].add(column.source_id)

            uem_map = {}
            run_prompter_map = {}
            optimisation_map = {}
            experiment_map = {}

            if source_ids["uem"]:
                uem_map = {
                    str(obj.id): obj
                    for obj in list(
                        UserEvalMetric.objects.filter(
                            id__in=source_ids["uem"],
                            deleted=False,
                            template__deleted=False,
                        ).select_related("template")
                    )
                }

            if source_ids["run_prompter"]:
                run_prompter_map = {
                    str(obj.id): obj
                    for obj in list(
                        RunPrompter.objects.filter(
                            id__in=source_ids["run_prompter"], deleted=False
                        )
                    )
                }

            if source_ids["optimization"]:
                optimisation_map = {
                    str(obj.id): obj
                    for obj in list(
                        OptimizationDataset.objects.filter(
                            id__in=source_ids["optimization"], deleted=False
                        )
                    )
                }

            if source_ids["experiment"]:
                experiment_map = {
                    str(obj.id): obj
                    for obj in list(
                        ExperimentsTable.objects.filter(
                            id__in=source_ids["experiment"], deleted=False
                        )
                    )
                }

            # Process columns sequentially (no thread overhead needed)
            result_queue = Queue()
            for idx, column in enumerate(columns):
                process_single_column(
                    column,
                    existing_column_config,
                    idx,
                    result_queue,
                    uem_map,
                    run_prompter_map,
                    optimisation_map,
                    experiment_map,
                    columns_map,
                )

            # Collect results in correct order
            results = []
            while not result_queue.empty():
                idx, result = result_queue.get()
                if result is not None:
                    results.append((idx, result))

            # Sort results by original index and append to column_config
            column_config.extend(
                [result for _, result in sorted(results, key=lambda x: x[0])]
            )

            if column_config_only:
                dataset_config = dataset.dataset_config
                file_processing_status = dataset_config.get(
                    "file_processing_status", ""
                )

                if (
                    len(column_config) == 0
                    and file_processing_status
                    and file_processing_status not in ["completed", "failed"]
                ):
                    return self._gm.success_response(
                        {"column_config": column_config, "is_processing_data": True}
                    )

                return self._gm.success_response({"column_config": column_config})

            # Prepare table data
            table_data = []
            processed_rows = set()
            all_cells = None
            # Paginate at DB level
            paginated_rows = list(rows[start:end])
            row_ids = [row.id for row in paginated_rows]

            # Fetch cells only for paginated rows
            column_ids = list(columns_map.keys())
            all_cells = Cell.objects.filter(
                row_id__in=row_ids, column_id__in=column_ids, deleted=False
            ).select_related("column")

            # Create a dictionary to group cells by row_id for faster lookup
            cells_by_row: dict[Any, Any] = {}
            for cell in list(all_cells):
                if cell.row_id not in cells_by_row:
                    cells_by_row[cell.row_id] = []
                cells_by_row[cell.row_id].append(cell)

            # Process rows
            row_processing_start = time.time()

            def process_row(row, cells_by_row, search, search_results):
                if str(row.id) in processed_rows:
                    return None

                row_data = {"row_id": str(row.id), "order": row.order}
                row_cells = cells_by_row.get(row.id, [])

                for cell in row_cells:
                    response_time_ms = None
                    token_count = None
                    value_infos = None

                    if cell.value_infos:
                        try:
                            metadata = {}
                            value_infos = json.loads(cell.value_infos)

                            if isinstance(value_infos.get("metadata", "{}"), str):
                                value_infos["metadata"] = json.loads(
                                    value_infos.get("metadata", "{}")
                                )
                            if isinstance(value_infos.get("metadata", {}), dict):
                                metadata = value_infos.get("metadata", {})
                                response_time_ms = metadata.get("response_time", None)
                                token_count = metadata.get("usage", {}).get(
                                    "total_tokens", None
                                )
                        except Exception:
                            response_time_ms = None
                            token_count = None
                            value_infos = cell.value_infos

                    # Handle special data types that need parsing
                    cell_value = cell.value
                    if cell.column.data_type == DataTypeChoices.PERSONA.value:
                        # Parse Python dict string format for persona data
                        if isinstance(
                            cell_value, str
                        ) and cell_value.strip().startswith("{"):
                            try:
                                parsed_dict = ast.literal_eval(cell_value)
                                # Convert to proper JSON string to ensure valid JSON format
                                cell_value = json.dumps(parsed_dict)
                            except (ValueError, SyntaxError):
                                # If parsing fails, keep as string
                                pass

                    row_data[str(cell.column_id)] = {
                        "cell_id": str(cell.id),
                        "cell_value": cell_value,
                        "metadata": {
                            "response_time_ms": response_time_ms,
                            "token_count": token_count,
                            "annotation": self._get_annotation_status(
                                cell.feedback_info, cell.value
                            ),
                        },
                        "status": cell.status,
                        "value_infos": value_infos if value_infos else cell.value_infos,
                        "feedback_info": cell.feedback_info,
                    }

                    if search and search_results and search_results.get(str(cell.id)):
                        row_data[str(cell.column_id)]["key_exists"] = search_results[
                            str(cell.id)
                        ].get("key_exists", False)
                        row_data[str(cell.column_id)]["indices"] = search_results[
                            str(cell.id)
                        ].get("indices", None)

                processed_rows.add(str(row.id))  # Mark this row as processed
                return row_data

            # Process rows sequentially (only ~page_size rows, no thread overhead needed)
            for row in paginated_rows:
                result = process_row(row, cells_by_row, search, search_results)
                if result:
                    table_data.append(result)

            logger.info(
                f"[TIMING] Row processing: {time.time() - row_processing_start:.2f}s"
            )
            task_manager = SyntheticTaskManager()

            dataset_config = dataset.dataset_config
            file_processing_status = dataset_config.get("file_processing_status", "")
            is_processing_data = False

            if (
                len(column_config) == 0
                and file_processing_status
                and file_processing_status not in ["completed", "failed"]
            ):
                is_processing_data = True

            response_data = {
                "metadata": {
                    "dataset_name": dataset.name,
                    "total_rows": total_rows,
                    "total_pages": (total_rows + page_size - 1) // page_size,
                    "error_messages": error_messages,
                    "status": self._get_dataset_status(dataset, all_cells),
                },
                "column_config": column_config,
                "table": table_data,
                "dataset_config": dataset.dataset_config,
                "synthetic_dataset": (
                    True if dataset.synthetic_dataset_config else False
                ),
                "synthetic_dataset_percentage": task_manager.get_progress(
                    dataset_id=dataset.id, request_uuid=None
                ),
                "synthetic_regenerate": task_manager.operation_regenerate_key(
                    op="get", dataset_id=dataset.id
                ),
                "is_processing_data": is_processing_data,
            }

            return self._gm.success_response(response_data)

        except Exception as e:
            logger.exception(f"Error in fetching dataset metadata: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_DATASET_METADATA")
            )

    def _get_dataset_status(self, dataset, all_cells=None):
        if all_cells is None:
            all_cells = Cell.objects.filter(
                row__dataset=dataset, row__deleted=False, deleted=False
            )
        if all_cells:
            if dataset.dataset_config.get("dataset_source_local"):
                columns_status = all_cells.filter(
                    column__data_type__in=[
                        DataTypeChoices.IMAGE.value,
                        DataTypeChoices.AUDIO.value,
                    ]
                ).values_list("column__status", flat=True)
                if columns_status and any(
                    status == StatusType.UPLOADING.value for status in columns_status
                ):
                    return {"dataset_status": DatasetStatus.PARTIAL_UPLOAD.value}

                failed_cells = all_cells.filter(
                    status=CellStatus.ERROR.value,
                    column__data_type__in=[
                        DataTypeChoices.IMAGE.value,
                        DataTypeChoices.AUDIO.value,
                    ],
                ).all()
                reasons = []
                for cell in failed_cells:
                    try:
                        value_infos = (
                            json.loads(cell.value_infos)
                            if cell.value_infos and isinstance(cell.value_infos, str)
                            else cell.value_infos
                        )
                        reasons.append(value_infos.get("reason"))
                    except (json.JSONDecodeError, AttributeError):
                        reasons.append(None)

                if failed_cells:
                    return {
                        "dataset_status": DatasetStatus.PARTIAL_EXTRACTED.value,
                        "failed_cells": failed_cells.count(),
                        "reason": reasons,
                    }
                else:
                    return {"dataset_status": DatasetStatus.COMPLETED.value}
            else:
                return {"dataset_status": DatasetStatus.COMPLETED.value}
        else:
            return {"dataset_status": DatasetStatus.RUNNING.value}

    def _get_annotation_status(self, feedback_info, cell_value):
        """Helper method to determine annotation status"""
        if not isinstance(feedback_info, dict) or "annotation" not in feedback_info:
            return None

        annotation = feedback_info["annotation"]
        if not isinstance(annotation, dict):
            return None

        if not annotation.get("auto_annotate"):
            return None

        if not cell_value:
            return None

        # Auto annotation case
        if "user_id" not in annotation:
            return "auto"

        # Verified case
        if "verified" in annotation and annotation["verified"]:
            return "verified"

        # Human annotation case
        if "user_id" in annotation and (
            "verified" not in annotation or not annotation["verified"]
        ):
            return "human"

        return None


class GetRowDataView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            # Get request parameters
            filters = request.data.get("filters", [])
            sort_configs = request.data.get("sort", [])
            row_id = request.data.get("row_id", None)

            # Get base dataset and rows
            dataset = get_object_or_404(Dataset, id=dataset_id, deleted=False)
            rows = Row.objects.filter(dataset=dataset, deleted=False).order_by("order")
            error_messages = []
            cells = Cell.objects.filter(row__in=rows, deleted=False)
            column_order = dataset.column_order or []
            qs = Column.objects.filter(dataset=dataset, deleted=False)
            if dataset.source != DatasetSourceChoices.EXPERIMENT_SNAPSHOT.value:
                qs = qs.exclude(source__in=[
                    SourceChoices.EXPERIMENT.value,
                    SourceChoices.EXPERIMENT_EVALUATION.value,
                    SourceChoices.EXPERIMENT_EVALUATION_TAGS.value,
                ])
            all_columns = list(qs)
            columns_map = {str(col.id): col for col in all_columns}

            # Apply filters if any
            if filters:
                rows = GetDatasetTableView()._apply_filters(
                    cells, rows, filters, error_messages, columns_map
                )

            # Apply sorting
            if sort_configs:
                rows = GetDatasetTableView()._apply_sorting(
                    rows, sort_configs, columns_map
                )

            # print("exiting sort")

            column_order = dataset.column_order or []

            current_row = get_object_or_404(
                rows, id=row_id, dataset=dataset, deleted=False
            )

            # Get the next 50 rows ordered after the current row
            next_row_ids = list(
                rows.filter(dataset=dataset, deleted=False, order__gt=current_row.order)
                .order_by("order")
                .values_list("id", flat=True)[:50]
            )

            # result = rows.annotate(
            #     next_row_id=StringAgg(
            #         Cast('id', output_field=CharField()),
            #         delimiter=',',
            #         ordering='order',
            #         filter=Q(order__gt=OuterRef('order'))
            #     )
            # ).filter(id=row_id).values('id', 'next_row_id').first()
            # previous_row_id = result['prev_row_id']

            # Prepare table data
            row_data_array = {}
            # row_data_array.update({"previous": { "row_id": str(previous_row_id) if previous_row_id else None}})
            row_data_array.update({"next": {"row_id": next_row_ids or []}})
            row_data = {"row_id": str(row_id)}

            current_cells = list(
                Cell.objects.filter(
                    row_id=row_id,
                    deleted=False,
                    dataset=dataset,
                    column__id__in=columns_map.keys(),
                )
            )
            for cell in current_cells:
                response_time_ms = None
                token_count = None
                value_infos = None
                if cell.value_infos:
                    try:
                        metadata = {}
                        value_infos = json.loads(cell.value_infos)
                        if isinstance(value_infos.get("metadata", "{}"), str):
                            value_infos["metadata"] = json.loads(
                                value_infos.get("metadata", "{}")
                            )
                        if isinstance(value_infos.get("metadata", {}), dict):
                            metadata = value_infos.get("metadata", {})
                            response_time_ms = metadata.get("response_time", None)
                            token_count = metadata.get("usage", {}).get(
                                "total_tokens", None
                            )
                    except Exception:
                        response_time_ms = None
                        token_count = None
                        value_infos = cell.value_infos

                row_data[str(cell.column_id)] = {
                    "cell_id": str(cell.id),
                    "cell_value": cell.value,
                    "metadata": {
                        "response_time_ms": response_time_ms,
                        "token_count": token_count,
                        "annotation": self._get_annotation_status(
                            cell.feedback_info, cell.value
                        ),
                    },
                    "status": cell.status,
                    "value_infos": value_infos,
                    "feedback_info": cell.feedback_info,
                }

            row_data_array.update({"current": row_data})
            return self._gm.success_response(row_data_array)

        except Exception as e:
            logger.exception(f"Error in fetching dataset metadata: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_DATASET_METADATA")
            )

    def _get_annotation_status(self, feedback_info, cell_value):
        """Helper method to determine annotation status"""
        if not isinstance(feedback_info, dict) or "annotation" not in feedback_info:
            return None

        annotation = feedback_info["annotation"]
        if not isinstance(annotation, dict):
            return None

        if not annotation.get("auto_annotate"):
            return None

        if not cell_value:
            return None

        # Auto annotation case
        if "user_id" not in annotation:
            return "auto"

        # Verified case
        if "verified" in annotation and annotation["verified"]:
            return "verified"

        # Human annotation case
        if "user_id" in annotation and (
            "verified" not in annotation or not annotation["verified"]
        ):
            return "human"

        return None


class GetExperimentDatasetTableView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    def is_valid_uuid4(self, s):
        try:
            val = uuid.UUID(s, version=4)
            # Ensure the string matches the canonical form
            return str(val) == s.lower()
        except (ValueError, AttributeError, TypeError):
            return False

    def get(self, request, experiment_dataset_id, *args, **kwargs):
        try:
            # Get request parameters
            page_size = int(request.GET.get("page_size", 10))
            current_page = int(request.GET.get("current_page_index", 0))

            # Get base dataset and rows
            dataset = get_object_or_404(
                ExperimentDatasetTable, id=experiment_dataset_id
            )
            rows = Row.objects.filter(dataset_id=experiment_dataset_id, deleted=False)
            # logger.exception(f"rows : {rows}")
            rows = rows.order_by("order")

            # Calculate pagination offsets
            start = current_page * page_size
            end = start + page_size

            # Get column configuration
            columns_in_experiment = list(dataset.columns.all())

            experiment = dataset.experiment
            user_eval_metric = list(experiment.user_eval_template_ids.all())
            total_columns = []
            for metric in user_eval_metric:
                runner = EvaluationRunner(
                    user_eval_metric_id=metric.id,
                    is_only_eval=True,
                    format_output=True,
                    source="experiment",
                    source_id=metric.template.id,
                )
                user_eval = UserEvalMetric.objects.get(id=metric.id)
                cols_used = runner._get_all_column_ids_being_used(
                    user_eval_metric=user_eval
                )
                logger.info(f"adding to total_columns cols_used: {cols_used}")
                total_columns.extend([col_id for col_id in cols_used if col_id])

            # V1 experiments have a single output column; V2 experiments do not.
            if experiment.column:
                logger.info(
                    f"adding to total_columns experiment.column.id: {experiment.column.id}"
                )
                total_columns.append(str(experiment.column.id))
                if experiment.column.source == SourceChoices.RUN_PROMPT.value:
                    run_prompt = RunPrompter.objects.get(id=experiment.column.source_id)
                    # Extract all column UUIDs from messages
                    message_column_ids = []
                    for message in run_prompt.messages:
                        content = message.get("content", "")
                        if isinstance(content, list):
                            for item in content:
                                if isinstance(item, dict) and item["type"] == "text":
                                    column_ids = re.findall(
                                        r"\{\{([^}]*)\}\}", item["text"]
                                    )
                                    message_column_ids.extend(
                                        [
                                            col_id.strip()
                                            for col_id in column_ids
                                            if col_id
                                        ]
                                    )
                        else:
                            # Find all UUIDs between {{ and }}
                            column_ids = re.findall(r"\{\{([^}]*)\}\}", content)
                            # Clean up any whitespace and add to list
                            message_column_ids.extend(
                                [col_id.strip() for col_id in column_ids if col_id]
                            )

                    # Add unique column IDs to total_columns
                    logger.info(
                        f"adding to total_columns message_column_ids: {message_column_ids}"
                    )
                    total_columns.extend(
                        [
                            col_id
                            for col_id in message_column_ids
                            if col_id not in [None, "None"]
                        ]
                    )
                    logger.info(f"TOTAL COLUMNS: {total_columns}")

            valid_total_columns = [s for s in total_columns if self.is_valid_uuid4(s)]
            working_dataset = experiment.snapshot_dataset or experiment.dataset
            columns_in_datasets = list(
                Column.objects.filter(
                    id__in=valid_total_columns,
                    deleted=False,
                    dataset=working_dataset,
                ).order_by("created_at")
            )
            column_config = []
            # Prepare table data - Modified to ensure each row appears only once
            table_data = []
            cells_by_row: dict[Any, Any] = {}

            columns_in_experiment = [
                col for col in columns_in_experiment if not col.name.endswith("reason")
            ]
            columns_in_datasets = [
                col for col in columns_in_datasets if not col.name.endswith("reason")
            ]

            for column in columns_in_datasets:
                column_config.append(
                    {
                        "id": str(column.id),
                        "name": column.name,
                        "origin_type": column.source,
                        "is_frozen": None,
                        "is_visible": True,
                        "data_type": column.data_type,
                        "status": "completed",
                    }
                )

                for cell in column.cell_set.filter(
                    deleted=False, row__deleted=False, column__deleted=False
                ):
                    if cell.row.deleted is False and cell.column.deleted is False:
                        if str(cell.row_id) not in cells_by_row:
                            cells_by_row[str(cell.row_id)] = {
                                "row_id": str(cell.row_id)
                            }

                        cells_by_row[str(cell.row_id)][str(column.id)] = {
                            "cell_value": cell.value or "",
                            "status": "completed",
                            "metadata": {},
                        }

            # Build a lookup of reason_column flags keyed by user_eval_metric id
            # so EXPERIMENT_EVALUATION columns can expose whether they have a
            # paired reason column.
            reason_column_by_metric = {}
            for metric in user_eval_metric:
                try:
                    reason_column_by_metric[str(metric.id)] = bool(
                        metric.template.config.get("reason_column", False)
                    )
                except Exception:
                    reason_column_by_metric[str(metric.id)] = False

            for column in columns_in_experiment:
                last_cell = column.cell_set.filter(
                    deleted=False, row__deleted=False, column__deleted=False
                ).last()
                origin_type = column.source
                avg_score = calculate_column_average(column.id)
                avg_score = avg_score.get("average", None)
                reason_column_flag = False
                if column.source == SourceChoices.EXPERIMENT_EVALUATION.value:
                    origin_type = "evaluation"
                    # source_id is typically "{edt_id}-sourceid-{user_eval_metric_id}"
                    sid = column.source_id or ""
                    metric_id = (
                        sid.split("-sourceid-")[1] if "-sourceid-" in sid else sid
                    )
                    reason_column_flag = reason_column_by_metric.get(
                        metric_id, False
                    )
                elif column.source == SourceChoices.EXPERIMENT_EVALUATION_TAGS.value:
                    origin_type = "evaluation_tags"

                column_config.append(
                    {
                        "id": str(column.id),
                        "name": column.name,
                        "is_frozen": None,
                        "is_visible": True,
                        "origin_type": origin_type,
                        "data_type": column.data_type,
                        "status": last_cell.status if last_cell else "NotStarted",
                        "average_score": avg_score,
                        "reason_column": reason_column_flag,
                    }
                )

                for cell in column.cell_set.filter(
                    deleted=False, row__deleted=False, column__deleted=False
                ):
                    if cell.row.deleted is False and cell.column.deleted is False:
                        try:
                            if str(cell.row_id) not in cells_by_row:
                                cells_by_row[str(cell.row_id)] = {
                                    "row_id": str(cell.row_id)
                                }

                            metadata = {}
                            value_infos = {}
                            value_infos = (
                                json.loads(cell.value_infos) if cell.value_infos else {}
                            )

                            if isinstance(value_infos.get("metadata", "{}"), str):
                                value_infos["metadata"] = json.loads(
                                    value_infos.get("metadata", "{}")
                                )
                                metadata = value_infos.get("metadata", {})
                            if isinstance(value_infos.get("metadata", {}), dict):
                                metadata = value_infos.get("metadata", {})

                            # Prepare cell data

                            cell_data = {
                                "cell_value": cell.value or "",
                                "status": (
                                    cell.status.lower()
                                    if hasattr(cell.status, "lower")
                                    else cell.status
                                ),
                                "metadata": {
                                    "response_time_ms": metadata.get(
                                        "response_time", 0
                                    ),
                                    "token_count": metadata.get("usage", {}).get(
                                        "total_tokens", 0
                                    ),
                                    "cost": metadata.get("cost", {}),
                                },
                            }

                            # Add value_infos if there's a reason
                            if value_infos.get("reason"):
                                cell_data["value_infos"] = {
                                    "reason": value_infos["reason"]
                                }

                            cells_by_row[str(cell.row_id)][str(column.id)] = cell_data

                        except Exception as e:
                            logger.error(e)
                            continue

            # logger.info(f"cells_BY_ROW: {cells_by_row}")
            table_data = list(cells_by_row.values())[start:end]

            response_data = {
                "metadata": {
                    "dataset_name": dataset.name,
                    "total_rows": len(table_data),
                    "total_pages": (len(table_data) + page_size - 1) // page_size,
                },
                "column_config": column_config,
                "table": table_data,
            }

            return self._gm.success_response(response_data)

        except Exception as e:
            logger.exception(f"Error in fetching experiment dataset metadata: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_EXP_DATASET_METADATA")
            )


class GetColumnDetailView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def get(self, request, dataset_id, *args, **kwargs):
        try:
            dataset = Dataset.objects.filter(
                id=dataset_id,
                deleted=False,
                organization=getattr(request, "organization", None)
                or request.user.organization,
            ).first()
            exp_dataset = ExperimentDatasetTable.objects.filter(
                id=dataset_id, deleted=False
            ).first()

            include_prompt = request.GET.get("include_prompt", False)
            source = request.GET.get("source", None)

            # Get column configuration
            if dataset:
                column_order = dataset.column_order or []
                if source and getattr(SourceChoices, source, None):
                    sources = [source]
                else:
                    if not include_prompt or include_prompt.lower() == "false":
                        sources = [SourceChoices.OTHERS.value]
                    else:
                        sources = [
                            SourceChoices.OTHERS.value,
                            SourceChoices.RUN_PROMPT.value,
                        ]
                columns = Column.objects.filter(
                    id__in=column_order, deleted=False, source__in=sources
                ).order_by(
                    Case(
                        *[
                            When(id=uuid.UUID(col_id), then=Value(idx))
                            for idx, col_id in enumerate(column_order)
                        ],
                        default=Value(len(column_order)),
                        output_field=IntegerField(),
                    )
                )
                column_config = []

                for column in columns:
                    column_config.append(
                        {
                            "id": str(column.id),
                            "name": column.name,
                            "data_type": column.data_type,
                        }
                    )
            elif exp_dataset:
                columns_in_experiment = list(exp_dataset.columns.all())
                experiment = exp_dataset.experiments_datasets_created.first()
                user_eval_metric = list(experiment.user_eval_template_ids.all())
                total_columns = []
                for metric in user_eval_metric:
                    runner = EvaluationRunner(
                        user_eval_metric_id=metric.id,
                        is_only_eval=True,
                        format_output=True,
                    )
                    user_eval = UserEvalMetric.objects.get(id=metric.id)
                    cols_used = runner._get_all_column_ids_being_used(
                        user_eval_metric=user_eval
                    )
                    total_columns.extend(cols_used)

                total_columns.append(experiment.column.id)
                if experiment.column.source == SourceChoices.RUN_PROMPT.value:
                    run_prompt = RunPrompter.objects.get(id=experiment.column.source_id)
                    # Extract all column UUIDs from messages
                    message_column_ids = []
                    for message in run_prompt.messages:
                        content = message.get("content", "")
                        if isinstance(content, list):
                            # Process all text items in the list
                            for item in content:
                                if isinstance(item, dict):
                                    text_content = item.get("text", "")
                                    # Find all UUIDs between {{ and }}
                                    column_ids = re.findall(
                                        r"\{\{([^}]*)\}\}", text_content
                                    )
                                    # Clean up any whitespace and add to list
                                    message_column_ids.extend(
                                        [col_id.strip() for col_id in column_ids]
                                    )
                        else:
                            # Handle case where content is not a list
                            if isinstance(content, dict):
                                content = content.get("text", "")
                            # Find all UUIDs between {{ and }}
                            column_ids = re.findall(r"\{\{([^}]*)\}\}", content)
                            # Clean up any whitespace and add to list
                            message_column_ids.extend(
                                [col_id.strip() for col_id in column_ids]
                            )

                    # Add unique column IDs to total_columns
                    total_columns.extend(
                        [uuid.UUID(col_id) for col_id in message_column_ids if col_id]
                    )

                columns_in_datasets = list(
                    Column.objects.filter(
                        id__in=total_columns,  # Convert UUIDs to strings
                        deleted=False,
                        dataset=experiment.dataset,
                    ).order_by("created_at")
                )
                experiment_columns = columns_in_experiment + columns_in_datasets
                experiment_columns = [
                    experiment_column
                    for experiment_column in experiment_columns
                    if not experiment_column.name.endswith("reason")
                ]
                column_config = []

                for column in experiment_columns:
                    column_config.append({"id": str(column.id), "name": column.name})

            else:
                return self._gm.bad_request("Dataset Not Found")

            return self._gm.success_response({"columns": column_config})

        except Exception as e:
            logger.exception(f"Error in fetching column: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_COLUMN")
            )


class GetDatasetsNamesView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    # def post(self, request, *args, **kwargs):
    #     try:
    #         excluded_datasets = request.data.get("excluded_dataset", [])
    #         search_text = request.data.get("search_text", "")

    #     # Query datasets for the current user's organization
    #     queryset = Dataset.objects.filter(
    #         organization=getattr(request, "organization", None) or request.user.organization, deleted=False
    #     )

    #     # Exclude specified datasets
    #     if excluded_datasets:
    #         queryset = queryset.exclude(id__in=excluded_datasets)

    #     # Apply search filter if provided
    #     if search_text:
    #         queryset = queryset.filter(name__icontains=search_text)

    #     # Format response
    #     datasets = [
    #         {"dataset_id": str(dataset.id), "name": dataset.name, "model_type": dataset.model_type}
    #         for dataset in queryset
    #     ]

    #     return self._gm.success_response({"datasets": datasets})

    # except Exception as e:
    #     import traceback
    #     traceback.print_exc()
    #     return self._gm.internal_server_error_response(str(e))
    def get(self, request, *args, **kwargs):
        try:
            excluded_datasets = request.GET.getlist("excluded_dataset", [])
            search_text = request.GET.get("search_text", "")
            include_experiments = (
                request.GET.get("include_experiments", "false").lower() == "true"
            )
            base_column_name = request.GET.get("base_column_name", None)
            datasets = []

            if base_column_name:
                base_columns = Column.objects.filter(
                    name=base_column_name, deleted=False
                ).select_related("dataset")
                dataset_ids = base_columns.values_list("dataset_id", flat=True)
                queryset = Dataset.objects.filter(
                    organization=getattr(request, "organization", None)
                    or request.user.organization,
                    id__in=dataset_ids,
                    deleted=False,
                    source__in=[
                        DatasetSourceChoices.DEMO.value,
                        DatasetSourceChoices.BUILD.value,
                        DatasetSourceChoices.OBSERVE.value,
                    ],
                )
            else:
                # Query datasets for the current user's organization
                queryset = Dataset.objects.filter(
                    organization=getattr(request, "organization", None)
                    or request.user.organization,
                    deleted=False,
                    source__in=[
                        DatasetSourceChoices.DEMO.value,
                        DatasetSourceChoices.BUILD.value,
                        DatasetSourceChoices.OBSERVE.value,
                    ],
                )
            # Exclude specified datasets
            if excluded_datasets:
                queryset = queryset.exclude(id__in=excluded_datasets)

            # Apply search filter if provided
            if search_text:
                queryset = queryset.filter(name__icontains=search_text)

            # Format response
            datasets.extend(
                {
                    "dataset_id": str(dataset.id),
                    "name": dataset.name,
                    "model_type": dataset.model_type,
                }
                for dataset in queryset.select_related("organization")
            )

            if include_experiments:
                experiments = ExperimentsTable.objects.filter(
                    dataset__in=queryset, deleted=False
                ).prefetch_related(
                    Prefetch(
                        "experiments_datasets",
                        queryset=ExperimentDatasetTable.objects.filter(deleted=False),
                        to_attr="filtered_datasets",
                    )
                )

                experiments_datasets = []
                for exp in experiments:
                    experiments_datasets.extend(exp.filtered_datasets)

                # Remove duplicates using a set of IDs to maintain uniqueness
                seen_ids = set()
                unique_datasets = []
                for dataset in experiments_datasets:
                    if dataset.id not in seen_ids:
                        seen_ids.add(dataset.id)
                        unique_datasets.append(dataset)

                datasets.extend(
                    {"dataset_id": str(dataset.id), "name": dataset.name}
                    for dataset in unique_datasets
                )

            return self._gm.success_response({"datasets": datasets})

        except Exception as e:
            logger.exception(f"Error in fetching dataset names: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_DATASETS_NAMES")
            )


class AddColumnsView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            columns_data = request.data.get("new_columns_data")

            if not columns_data or not isinstance(columns_data, list):
                return self._gm.bad_request(
                    get_error_message("MISSING_COLUMN_DATA_AS_DICT")
                )

            valid_data_types = {choice.value for choice in DataTypeChoices}

            for column in columns_data:
                col_name = column.get("name", "")
                if len(col_name) > 255:
                    return self._gm.bad_request(
                        get_error_message("COLUMN_NAME_TOO_LONG")
                    )
                col_type = column.get("data_type", "")
                if col_type and col_type not in valid_data_types:
                    return self._gm.bad_request(get_error_message("INVALID_DATA_TYPE"))

            col_names = [col.get("name", "") for col in columns_data]
            if len(col_names) != len(set(col_names)):
                return self._gm.bad_request(
                    get_error_message("DUPLICATE_COLUMN_NAMES_IN_REQUEST")
                )

            dataset = get_object_or_404(
                Dataset,
                id=dataset_id,
                organization=getattr(request, "organization", None)
                or request.user.organization,
                deleted=False,
            )

            if any(
                Column.objects.filter(
                    dataset=dataset, name=column["name"], deleted=False
                ).exists()
                for column in columns_data
            ):
                return self._gm.bad_request(get_error_message("COLUMN_NAME_EXISTS"))

            added_columns = []

            for column in columns_data:
                # Create new column
                new_column, created = Column.objects.get_or_create(
                    # id=column['id'],
                    name=column["name"],
                    data_type=column["data_type"],
                    source=(
                        column.get("source")
                        if column.get("source")
                        else SourceChoices.OTHERS.value
                    ),
                    dataset=dataset,
                )

                if created:
                    # Update column order
                    column_order = dataset.column_order or []
                    column_order.append(str(new_column.id))

                    dataset.column_order = column_order
                    dataset.save()

                    added_columns.append(new_column)

                    rows = Row.objects.filter(dataset=dataset, deleted=False)

                    if new_column.data_type == DataTypeChoices.FLOAT.value:
                        value = float(0)
                    elif new_column.data_type == DataTypeChoices.INTEGER.value:
                        value = 0
                    elif new_column.data_type == DataTypeChoices.BOOLEAN.value:
                        value = "false"
                    elif new_column.data_type == DataTypeChoices.ARRAY.value:
                        value = []
                    elif new_column.data_type == DataTypeChoices.JSON.value:
                        value = {}
                    else:
                        value = ""

                    for row in rows:
                        Cell.objects.create(
                            id=uuid.uuid4(),
                            dataset=dataset,
                            column=new_column,
                            row=row,
                            value=value,
                        )

            if len(added_columns) > 0:
                columns_data = ColumnSerializer(added_columns, many=True)
                return self._gm.success_response(
                    {
                        "message": f"{len(added_columns)} Columns added successfully",
                        "data": columns_data.data,
                    }
                )

            else:
                return self._gm.success_response("Column Already Exists")
        except Exception as e:
            logger.exception(f"Error in adding column: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_ADD_COLUMN")
            )


class AddEmptyColumnsView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            num_cols = request.data.get("num_cols", 0)

            dataset = get_object_or_404(Dataset, id=dataset_id)

            added_columns = []

            for i in range(num_cols):
                # Create new column
                new_column = Column.objects.create(
                    id=str(uuid.uuid4()),
                    name=f"Column{i + 1}",
                    data_type=DataTypeChoices.TEXT.value,
                    source=SourceChoices.OTHERS.value,
                    dataset=dataset,
                )

                # Update column order
                column_order = dataset.column_order or []
                column_order.append(str(new_column.id))

                dataset.column_order = column_order
                dataset.save()

                added_columns.append(new_column)

                rows = Row.objects.filter(dataset=dataset, deleted=False)

                for row in rows:
                    Cell.objects.create(
                        id=uuid.uuid4(),
                        dataset=dataset,
                        column=new_column,
                        row=row,
                        value="",
                    )

            if len(added_columns) > 0:
                columns_data = ColumnSerializer(added_columns, many=True)
                return self._gm.success_response(
                    {
                        "message": f"{len(added_columns)} Columns added successfully",
                        "data": columns_data.data,
                    }
                )

            else:
                return self._gm.success_response("Column Already Exists")
        except Exception as e:
            logger.exception(f"Error in adding empty column: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_ADD_EMPTY_COLUMN")
            )


class GetCellDataView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        try:
            # Get request parameters
            row_ids = request.data.get("row_ids", [])
            column_ids = request.data.get("column_ids", [])

            if not row_ids or not column_ids:
                return self._gm.bad_request_response(
                    "Both row_ids and column_ids are required"
                )

            # Get dataset
            # dataset = get_object_or_404(Dataset, id=dataset_id, deleted=False)

            # Get cells for specified rows and columns
            cells = Cell.objects.filter(
                # dataset=dataset,
                row_id__in=row_ids,
                column_id__in=column_ids,
                deleted=False,
                column__dataset__organization=getattr(request, "organization", None)
                or request.user.organization,
            ).select_related("column")

            # Organize data by row_id and column_id
            result = {}
            for cell in cells:
                row_id = str(cell.row_id)
                column_id = str(cell.column_id)

                if row_id not in result:
                    result[row_id] = {}

                value_infos = {}

                if cell.value_infos:
                    try:
                        value_infos = json.loads(cell.value_infos)

                    except Exception:
                        value_infos = cell.value_infos

                result[row_id][column_id] = {
                    "cell_value": cell.value,
                    "status": cell.status,
                    "value_infos": value_infos,
                    "feedback_info": cell.feedback_info,
                }

            return self._gm.success_response(result)

        except Exception as e:
            logger.exception(f"Error in fetching specific cells: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_CELLS")
            )


class AddStaticColumnView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            with transaction.atomic():
                new_column_name = request.data.get("new_column_name")
                column_type = request.data.get("column_type")
                source = request.data.get("source")

                if not new_column_name or not column_type:
                    return self._gm.bad_request(
                        get_error_message("MISSING_COLUMN_NAME_AND_TYPE")
                    )

                if len(new_column_name) > 255:
                    return self._gm.bad_request(
                        get_error_message("COLUMN_NAME_TOO_LONG")
                    )

                valid_data_types = {choice.value for choice in DataTypeChoices}
                if column_type not in valid_data_types:
                    return self._gm.bad_request(get_error_message("INVALID_DATA_TYPE"))

                dataset = get_object_or_404(
                    Dataset,
                    id=dataset_id,
                    organization=getattr(request, "organization", None)
                    or request.user.organization,
                    deleted=False,
                )

                if Column.objects.filter(
                    name=new_column_name,
                    dataset=dataset,
                    deleted=False,
                    dataset__organization=getattr(request, "organization", None)
                    or request.user.organization,
                ).exists():
                    return self._gm.bad_request(get_error_message("COLUMN_NAME_EXISTS"))

                # Create new column
                new_column = Column.objects.create(
                    id=uuid.uuid4(),
                    name=new_column_name,
                    data_type=column_type,
                    source=(
                        source
                        if source and getattr(SourceChoices, source, None)
                        else SourceChoices.OTHERS.value
                    ),
                    dataset=dataset,
                )

                # Update column order
                column_order = dataset.column_order or []
                column_order.append(str(new_column.id))
                dataset.column_order = column_order
                dataset.column_config[str(new_column.id)] = {
                    "is_visible": True,
                    "is_frozen": None,
                }
                dataset.save()

                # # for each row, create a cell with the new column and set the value to None
                rows = list(Row.objects.filter(dataset=dataset, deleted=False))

                # Prepare list of Cell instances (not saved yet)
                cells_to_create = [
                    Cell(
                        id=uuid.uuid4(),
                        dataset=dataset,
                        column=new_column,
                        row=row,
                        value=None,
                    )
                    for row in rows
                ]

                # Bulk create all cells in one query
                Cell.objects.bulk_create(cells_to_create)

            return self._gm.success_response("Column added successfully")

        except Exception as e:
            logger.exception(f"Error in adding static column: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_ADD_STATIC_COLUMN")
            )


class AddMultipleStaticColumnsView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, dataset_id, *args, **kwargs):
        """
        Add multiple static columns to a dataset at once.

        Expected request data:
        {
            "columns": [
                {
                    "new_column_name": "column1",
                    "column_type": "string",
                    "source": "OTHERS"  # optional
                },
                {
                    "new_column_name": "column2",
                    "column_type": "number",
                    "source": "OTHERS"  # optional
                }
            ]
        }
        """
        try:
            with transaction.atomic():
                columns_data = request.data.get("columns", [])

                if not columns_data or not isinstance(columns_data, list):
                    return self._gm.bad_request(
                        get_error_message("MISSING_COLUMNS_DATA")
                    )

                if len(columns_data) == 0:
                    return self._gm.bad_request("At least one column must be provided")

                # Validate all column data first
                new_column_names = []
                column_types = []
                sources = []

                for idx, column_data in enumerate(columns_data):
                    new_column_name = column_data.get("new_column_name")
                    column_type = column_data.get("column_type")
                    source = column_data.get("source")

                    if not new_column_name or not column_type:
                        return self._gm.bad_request(
                            f"Column at index {idx}: {get_error_message('MISSING_COLUMN_NAME_AND_TYPE')}"
                        )

                    new_column_names.append(new_column_name)
                    column_types.append(column_type)
                    sources.append(
                        source
                        if source and getattr(SourceChoices, source, None)
                        else SourceChoices.OTHERS.value
                    )

                # Check for duplicate names in the request
                if len(new_column_names) != len(set(new_column_names)):
                    return self._gm.bad_request(
                        "Duplicate column names found in request"
                    )

                dataset = get_object_or_404(
                    Dataset,
                    id=dataset_id,
                    organization=getattr(request, "organization", None)
                    or request.user.organization,
                    deleted=False,
                )

                # Check if any column names already exist
                existing_columns = Column.objects.filter(
                    name__in=new_column_names,
                    dataset=dataset,
                    deleted=False,
                    dataset__organization=getattr(request, "organization", None)
                    or request.user.organization,
                ).values_list("name", flat=True)

                if existing_columns:
                    return self._gm.bad_request(
                        f"Column name(s) already exist: {', '.join(existing_columns)}"
                    )

                # Create all new columns in bulk
                new_columns = []
                for name, col_type, source in zip(
                    new_column_names, column_types, sources
                ):
                    new_columns.append(
                        Column(
                            id=uuid.uuid4(),
                            name=name,
                            data_type=col_type,
                            source=source,
                            dataset=dataset,
                        )
                    )

                # Bulk create all columns
                created_columns = Column.objects.bulk_create(new_columns)

                # Update column order and column config for all new columns
                column_order = dataset.column_order or []
                column_config = dataset.column_config or {}

                for column in created_columns:
                    column_order.append(str(column.id))
                    column_config[str(column.id)] = {
                        "is_visible": True,
                        "is_frozen": None,
                    }

                dataset.column_order = column_order
                dataset.column_config = column_config
                dataset.save()

                # Get all rows for the dataset
                rows = list(Row.objects.filter(dataset=dataset, deleted=False))

                # Prepare all cells for bulk creation
                # For each row, create a cell for each new column
                cells_to_create = []
                for row in rows:
                    for column in created_columns:
                        cells_to_create.append(
                            Cell(
                                id=uuid.uuid4(),
                                dataset=dataset,
                                column=column,
                                row=row,
                                value=None,
                            )
                        )

                # Bulk create all cells in one query
                Cell.objects.bulk_create(cells_to_create)

                return self._gm.success_response("Columns added successfully")

        except Exception as e:
            logger.exception(f"Error in adding multiple static columns: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_ADD_STATIC_COLUMNS")
            )


class DeleteColumnView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    def delete(self, request, dataset_id, column_id, *args, **kwargs):
        try:
            dataset = get_object_or_404(Dataset, id=dataset_id)
            column = get_object_or_404(Column, id=column_id, dataset=dataset)

            # Delete associated source model based on source type
            if column.source_id:
                if column.source == SourceChoices.RUN_PROMPT.value:
                    RunPrompter.objects.filter(id=column.source_id).update(deleted=True)
                    # Clean up derived variables from associated prompt versions
                    try:
                        run_prompter = RunPrompter.objects.filter(
                            id=column.source_id
                        ).first()
                        if run_prompter and run_prompter.prompt_id:
                            # Get all versions for this prompt and clean up derived variables
                            prompt_versions = PromptVersion.objects.filter(
                                original_template_id=run_prompter.prompt_id,
                                deleted=False,
                            )
                            for version in prompt_versions:
                                if cleanup_derived_variables_for_column(
                                    version, column.name
                                ):
                                    version.save(update_fields=["metadata"])
                    except Exception as cleanup_error:
                        logger.warning(
                            f"Failed to cleanup derived variables for column {column.name}: {cleanup_error}"
                        )
                if column.source == SourceChoices.EVALUATION.value:
                    eval_metric = UserEvalMetric.objects.filter(
                        id=column.source_id
                    ).first()
                    if eval_metric and eval_metric.status in (
                        StatusType.RUNNING.value,
                        StatusType.NOT_STARTED.value,
                        StatusType.EXPERIMENT_EVALUATION.value,
                    ):
                        try:
                            from tfc.utils.distributed_state import (
                                evaluation_tracker,
                            )

                            evaluation_tracker.request_cancel(
                                eval_metric.id, reason="eval_column_deleted"
                            )
                        except Exception:
                            pass
                        from model_hub.utils.eval_cell_status import (
                            mark_eval_cells_stopped,
                        )

                        mark_eval_cells_stopped(
                            eval_metric,
                            reason="Evaluation column deleted by user",
                        )
                    if eval_metric:
                        eval_metric.deleted = True
                        eval_metric.save(update_fields=["deleted"])
                if column.source == SourceChoices.ANNOTATION_LABEL.value:
                    source_parts = column.source_id.split("-sourceid-")

                    annotation_id = source_parts[0]
                    label_id = source_parts[1]

                    annotation = Annotations.objects.get(id=annotation_id)
                    label = AnnotationsLabels.objects.get(id=label_id)

                    columns_to_delete = Column.objects.filter(
                        dataset=dataset,
                        source_id=f"{annotation.id}-sourceid-{label.id}",
                        deleted=False,
                    )

                    Cell.objects.filter(column__in=columns_to_delete).update(
                        deleted=True, deleted_at=timezone.now()
                    )

                    for col in columns_to_delete:
                        if str(col.id) in dataset.column_order:
                            dataset.column_order.remove(str(col.id))
                        if str(col.id) in dataset.column_config:
                            del dataset.column_config[str(col.id)]
                        annotation.columns.remove(col)

                    annotation.labels.remove(label)
                    columns_to_delete.update(deleted=True)

                    annotation.save()
                    dataset.save()

            # delete all cells associated with the column
            Cell.objects.filter(column=column).update(deleted=True)
            # Delete cells where source_id starts with column.id
            Cell.objects.filter(column__source_id__startswith=f"{column.id}").update(
                deleted=True
            )

            # Remove column from column_order
            if dataset.column_order:
                # Get columns to delete (including those with source_id starting with column.id)
                columns_to_delete = Column.objects.filter(
                    Q(id=column.id) | Q(source_id__startswith=f"{column.id}")
                ).values_list("id", flat=True)

                col_ids_to_remove = {str(c) for c in columns_to_delete}
                dataset.column_order = [
                    col_id
                    for col_id in dataset.column_order
                    if col_id not in col_ids_to_remove
                ]
                dataset.save(update_fields=["column_order"])

            # Update metrics BEFORE deleting columns — get_metrics_using_column
            # scopes by dataset via the Column row, which must still be
            # visible (deleted=False) for BaseModelManager to find it.
            metrics = UserEvalMetric.get_metrics_using_column(
                getattr(request, "organization", None) or request.user.organization.id,
                column_id,
            )
            if metrics:
                UserEvalMetric.objects.filter(
                    id__in=[m.id for m in metrics]
                ).update(column_deleted=True)

            # Now safe to delete columns
            Column.objects.filter(
                Q(id=column.id) | Q(source_id__startswith=f"{column.id}")
            ).update(deleted=True)

            return self._gm.success_response("Column deleted successfully")

        except Exception as e:
            logger.exception(f"Error in deleting the column: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_DELETE_COLUMN")
            )


class DeleteRowView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    def delete(self, request, dataset_id, *args, **kwargs):
        try:
            dataset = get_object_or_404(Dataset, id=dataset_id, deleted=False)
            row_ids = request.data.get("row_ids", [])
            selected_all_rows = request.data.get("selected_all_rows", False)

            if not row_ids and not selected_all_rows:
                return self._gm.bad_request(get_error_message("MISSING_ROW_IDS"))

            # Validate that provided row IDs exist in this dataset
            if not selected_all_rows and row_ids:
                existing_count = Row.objects.filter(
                    id__in=row_ids, dataset=dataset, deleted=False
                ).count()
                if existing_count != len(row_ids):
                    return self._gm.bad_request(
                        "Some row IDs were not found in this dataset"
                    )

            if selected_all_rows:
                if row_ids and len(row_ids) > 0:
                    Row.objects.filter(dataset=dataset, deleted=False).exclude(
                        id__in=row_ids
                    ).update(deleted=True)
                    Cell.objects.filter(dataset=dataset).exclude(
                        row_id__in=row_ids
                    ).update(deleted=True)
                else:
                    Row.objects.filter(dataset=dataset, deleted=False).update(
                        deleted=True
                    )
                    Cell.objects.filter(dataset=dataset).update(deleted=True)
            else:
                Row.objects.filter(id__in=row_ids, dataset=dataset).update(deleted=True)
                Cell.objects.filter(row_id__in=row_ids, dataset=dataset).update(
                    deleted=True
                )

            annotations = Annotations.objects.filter(dataset=dataset, deleted=False)

            for annot in annotations:
                for label in annot.labels.all():
                    if label.type == "text":
                        existing_metadata = label.metadata or {}
                        existing_metadata[str(dataset.id)] = {}
                        label.metadata = existing_metadata
                        label.save(update_fields=["metadata"])

            return self._gm.success_response("Row deleted successfully")

        except Exception as e:
            logger.exception(f"Error in deleting the row: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_DELETE_ROW")
            )


class AddEmptyRowsView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            from model_hub.services.dataset_validators import validate_num_rows

            num_rows, num_rows_err = validate_num_rows(
                request.data.get("num_rows", 1), max_allowed=100
            )
            if num_rows_err:
                return self._gm.bad_request(num_rows_err)

            # Get dataset and verify it exists
            dataset = get_object_or_404(Dataset, id=dataset_id)
            organization = (
                getattr(request, "organization", None) or request.user.organization
            )
            # --- Row Limit Check Start ---
            existing_rows_count = Row.objects.filter(
                dataset=dataset, deleted=False
            ).count()
            prospective_total = existing_rows_count + num_rows
            if log_and_deduct_cost_for_resource_request is not None:
                call_log_row = log_and_deduct_cost_for_resource_request(
                    organization,
                    api_call_type=APICallTypeChoices.ROW_ADD.value,
                    config={"total_rows": prospective_total},
                    workspace=request.workspace,
                )
                if (
                    call_log_row is None
                    or call_log_row.status == APICallStatusChoices.RESOURCE_LIMIT.value
                ):
                    return self._gm.too_many_requests(ROW_LIMIT_REACHED_MESSAGE)
                call_log_row.status = APICallStatusChoices.SUCCESS.value
                call_log_row.save()

            # Get all columns for this dataset
            columns = Column.objects.filter(dataset=dataset, deleted=False).exclude(
                source__in=[
                    SourceChoices.EXPERIMENT.value,
                    SourceChoices.EXPERIMENT_EVALUATION.value,
                    SourceChoices.EXPERIMENT_EVALUATION_TAGS.value,
                ]
            )

            last_row = (
                Row.all_objects.filter(dataset=dataset).order_by("-created_at").first()
            )
            if last_row:
                max_order = last_row.order
            else:
                max_order = -1

            # Create new rows and cells
            for i in range(num_rows):
                # Create row
                new_row = Row.objects.create(
                    id=uuid.uuid4(), dataset=dataset, order=max_order + 1 + i
                )

                # Create empty cells for each column
                for column in columns:
                    Cell.objects.create(
                        id=uuid.uuid4(),
                        dataset=dataset,
                        column=column,
                        row=new_row,
                        value=None,
                    )

            return self._gm.success_response(
                f"Successfully added {num_rows} empty row(s)"
            )

        except Exception as e:
            logger.exception(f"Error in adding empty rows: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_ADD_EMPTY_ROWS")
            )


class AddSDKRowsView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        try:
            dataset_name = request.data.get("dataset_name")
            dataset_id = request.data.get("dataset_id")

            # Get dataset and verify it exists
            if not dataset_name:
                dataset = get_object_or_404(Dataset, id=dataset_id)
            else:
                dataset = get_object_or_404(
                    Dataset,
                    name=dataset_name,
                    deleted=False,
                    organization=getattr(request, "organization", None)
                    or request.user.organization,
                )

            serialized_datset = DatasetSerializer(dataset)

            user_organization = (
                getattr(request, "organization", None) or request.user.organization
            )

            # --- Row Limit Check Start ---
            existing_rows_count = Row.objects.filter(
                dataset=dataset, deleted=False
            ).count()
            if log_and_deduct_cost_for_resource_request is not None:
                call_log_row = log_and_deduct_cost_for_resource_request(
                    getattr(request, "organization", None) or request.user.organization,
                    api_call_type=APICallTypeChoices.ROW_ADD.value,
                    config={"total_rows": existing_rows_count},
                    workspace=request.workspace,
                )
                if (
                    call_log_row is None
                    or call_log_row.status == APICallStatusChoices.RESOURCE_LIMIT.value
                ):
                    return self._gm.too_many_requests(ROW_LIMIT_REACHED_MESSAGE)
                call_log_row.status = APICallStatusChoices.SUCCESS.value
                call_log_row.save()
            # --- Row Limit Check End ---

            apiKeys = OrgApiKey.objects.filter(
                organization=user_organization,
                type="user",
                enabled=True,
                user=request.user,
            )
            if len(apiKeys) == 0:
                org_api_key = OrgApiKey.objects.create(
                    organization=user_organization,
                    type="user",
                    enabled=True,
                    user=request.user,
                )
                serialized_keys = OrgApiKeySerializer(
                    org_api_key,
                )

            else:
                apiKeys = OrgApiKey.objects.filter(
                    organization=user_organization,
                    type="user",
                    enabled=True,
                    user=request.user,
                )
                serialized_keys = OrgApiKeySerializer(
                    apiKeys[0],
                )

            (
                CURL_ADD_COLUMN_REQUEST,
                CURL_ADD_ROWS_REQUEST,
                TYPESCRIPT_ADD_COLUMNS,
                TYPESCRIPT_ADD_ROWS,
            ) = get_curl_ts_code(
                dataset_id=dataset.id,
                api_key=apiKeys[0].api_key,
                secret_key=apiKeys[0].secret_key,
                dataset_name=dataset.name,
            )

            response = {
                "api_keys": serialized_keys.data,
                "dataset": serialized_datset.data,
                "code": {
                    "python_add_row": PYTHON_ADD_ROWS.format(
                        apiKeys[0].api_key,
                        apiKeys[0].secret_key,
                        BASE_URL,
                        dataset.name,
                        dataset.name,
                    ),
                    "python_add_col": PYTHON_ADD_COLS.format(
                        apiKeys[0].api_key,
                        apiKeys[0].secret_key,
                        BASE_URL,
                        dataset.name,
                        dataset.name,
                    ),
                    "typescript_add_col": TYPESCRIPT_ADD_COLUMNS,
                    "typescript_add_row": TYPESCRIPT_ADD_ROWS,
                    "curl_add_col": CURL_ADD_COLUMN_REQUEST,
                    "curl_add_row": CURL_ADD_ROWS_REQUEST,
                },
            }

            return self._gm.success_response(response)

        except Exception as e:
            logger.exception(f"Error in adding sdk rows: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_ADD_SDK_ROWS")
            )


class ManuallyCreateDatasetView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        try:
            # Get and validate input parameters
            dataset_name = request.data.get("dataset_name")
            number_of_rows = int(request.data.get("number_of_rows", 1))
            number_of_columns = int(request.data.get("number_of_columns", 1))

            if log_and_deduct_cost_for_resource_request is not None:
                call_log_row_entry = log_and_deduct_cost_for_resource_request(
                    getattr(request, "organization", None) or request.user.organization,
                    api_call_type=APICallTypeChoices.DATASET_ADD.value,
                    workspace=request.workspace,
                )
                if (
                    call_log_row_entry is None
                    or call_log_row_entry.status == APICallStatusChoices.RESOURCE_LIMIT.value
                ):
                    return self._gm.too_many_requests(
                        get_error_message("DATASET_CREATE_LIMIT_REACHED")
                    )
                call_log_row_entry.status = APICallStatusChoices.SUCCESS.value
                call_log_row_entry.save()

            if not dataset_name:
                return self._gm.bad_request(get_error_message("MISSING_DATASET_NAME"))

            from model_hub.validators.dataset_validators import (
                validate_dataset_name_unique,
            )

            try:
                validate_dataset_name_unique(
                    dataset_name,
                    getattr(request, "organization", None) or request.user.organization,
                )
            except Exception as validation_err:
                return self._gm.bad_request(str(validation_err.detail[0]))

            if number_of_rows <= 0 or number_of_columns <= 0:
                return self._gm.bad_request(
                    get_error_message("INVALID_ROW_OR_COLUMN_COUNT")
                )

            # Enforce upper bounds (aligned with UI: max 100 rows, 100 columns)
            from model_hub.validators.dataset_validators import (
                validate_row_column_bounds,
            )

            try:
                validate_row_column_bounds(
                    rows=number_of_rows, columns=number_of_columns
                )
            except Exception as validation_err:
                return self._gm.bad_request(str(validation_err.detail[0]))

            # Check row limit
            organization = (
                getattr(request, "organization", None) or request.user.organization
            )

            if log_and_deduct_cost_for_resource_request is not None:
                call_log_row = log_and_deduct_cost_for_resource_request(
                    organization,
                    api_call_type=APICallTypeChoices.ROW_ADD.value,
                    config={"total_rows": number_of_rows},
                    workspace=request.workspace,
                )
                if (
                    call_log_row is None
                    or call_log_row.status == APICallStatusChoices.RESOURCE_LIMIT.value
                ):
                    return self._gm.too_many_requests(ROW_LIMIT_REACHED_MESSAGE)
                call_log_row.status = APICallStatusChoices.SUCCESS.value
                call_log_row.save()

            # Create dataset
            dataset = Dataset.objects.create(
                name=dataset_name,
                organization=getattr(request, "organization", None)
                or request.user.organization,
                source=DatasetSourceChoices.BUILD.value,
                user=request.user,
            )

            columns = []
            column_order = []
            column_config = {}

            # Create columns
            for i in range(number_of_columns):
                column_id = uuid.uuid4()
                column = Column.objects.create(
                    id=column_id,
                    name=f"Column {i + 1}",
                    data_type=DataTypeChoices.TEXT.value,
                    source=SourceChoices.OTHERS.value,
                    dataset=dataset,
                )
                columns.append(column)
                column_order.append(str(column_id))
                column_config[str(column_id)] = {"is_visible": True, "is_frozen": None}

            # Update dataset with column configuration
            dataset.column_order = column_order
            dataset.column_config = column_config
            dataset.dataset_config.update(
                {"eval_recommendations": ["Deterministic Evals"]}
            )
            dataset.save()

            # Create rows and cells in batches
            batch_size = 1000
            cells_to_create = []

            for i in range(number_of_rows):
                row = Row.objects.create(id=uuid.uuid4(), dataset=dataset, order=i)

                for column in columns:
                    cells_to_create.append(
                        Cell(
                            id=uuid.uuid4(),
                            dataset=dataset,
                            column=column,
                            row=row,
                            value=None,
                        )
                    )

                # Bulk create cells when batch size is reached
                if len(cells_to_create) >= batch_size:
                    Cell.objects.bulk_create(cells_to_create)
                    cells_to_create = []

            # Create any remaining cells
            if cells_to_create:
                Cell.objects.bulk_create(cells_to_create)
            return self._gm.success_response(
                {
                    "message": "Dataset created successfully",
                    "dataset_id": str(dataset.id),
                    "rows_created": number_of_rows,
                    "columns_created": number_of_columns,
                }
            )

        except ValueError:
            return self._gm.bad_request(get_error_message("INVALID_NUMBER_FORMAT"))
        except Exception as e:
            logger.exception(f"Error in creating dataset: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_CREATE_DATASET")
            )


class AddDataRowsView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    def post(self, request, dataset_id: str, *args, **kwargs):
        try:
            rows = request.data.get("rows", [])
            if not rows:
                return self._gm.bad_request(get_error_message("DATA_MISSING"))

            dataset = get_object_or_404(Dataset, id=dataset_id)

            # Validate row limit
            organization = (
                getattr(request, "organization", None) or request.user.organization
            )
            existing_rows_count = Row.objects.filter(
                dataset=dataset, deleted=False
            ).count()
            new_rows_count = len(rows)
            prospective_total = existing_rows_count + new_rows_count

            if log_and_deduct_cost_for_resource_request is not None:
                call_log_row = log_and_deduct_cost_for_resource_request(
                    organization,
                    api_call_type=APICallTypeChoices.ROW_ADD.value,
                    config={"total_rows": prospective_total},
                    workspace=request.workspace,
                )
                if (
                    call_log_row is None
                    or call_log_row.status == APICallStatusChoices.RESOURCE_LIMIT.value
                ):
                    return self._gm.too_many_requests(ROW_LIMIT_REACHED_MESSAGE)
                call_log_row.status = APICallStatusChoices.SUCCESS.value
                call_log_row.save()

            # Get valid columns for this dataset
            columns = Column.objects.filter(dataset=dataset, deleted=False).exclude(
                source__in=[
                    SourceChoices.EXPERIMENT.value,
                    SourceChoices.EXPERIMENT_EVALUATION.value,
                    SourceChoices.EXPERIMENT_EVALUATION_TAGS.value,
                ]
            )

            last_row = (
                Row.all_objects.filter(dataset=dataset).order_by("-created_at").first()
            )
            if last_row:
                max_order = last_row.order
            else:
                max_order = -1

            # Create rows and cells
            for index, row_data in enumerate(rows):
                new_row = Row.objects.create(
                    id=row_data.get("id", str(uuid.uuid4())),
                    dataset=dataset,
                    order=max_order + 1 + index,
                )

                provided_columns = {
                    cell["column_name"] for cell in row_data.get("cells", [])
                }

                # Process columns in a consistent order to ensure proper row alignment
                cells_to_create = []

                for column in columns:
                    cell_value = ""
                    cell_value_infos = None
                    column_metadata = {}

                    if column.name in provided_columns:
                        cell_data = next(
                            cell
                            for cell in row_data["cells"]
                            if cell["column_name"] == column.name
                        )
                        cell_value = cell_data.get("value", "")

                        try:
                            if (
                                column.data_type == DataTypeChoices.IMAGE.value
                                and cell_value
                            ):
                                cell_value = upload_image_to_s3(
                                    cell_data.get("value", ""),
                                    bucket_name="fi-customer-data-dev",
                                    object_key=f"images/{dataset_id}/{uuid.uuid4()}",
                                )

                            if (
                                column.data_type == DataTypeChoices.AUDIO.value
                                and cell_value
                            ):
                                cell_value, duration = upload_audio_to_s3_duration(
                                    cell_data.get("value", ""),
                                    bucket_name="fi-customer-data-dev",
                                    object_key=f"audio/{dataset_id}/{uuid.uuid4()}",
                                )
                                value_infos: dict[str, Any] = {}
                                column_metadata = {"audio_duration_seconds": duration}
                                cell_value_infos = json.dumps(value_infos)
                            elif (
                                column.data_type == DataTypeChoices.DOCUMENT.value
                                and cell_value
                            ):
                                # Handle document processing similar to image/audio
                                doc_key = f"documents/{dataset_id}/{uuid.uuid4()}"
                                name = cell_value
                                cell_value = upload_document_to_s3(
                                    cell_value,
                                    bucket_name="fi-customer-data-dev",
                                    object_key=doc_key,
                                )
                                cell_value_infos = json.dumps({"document_name": name})
                        except Exception as e:
                            logger.error(
                                f"Error processing {column.data_type} for column {column.name}: {str(e)}"
                            )
                            cell_value = ""  # Set to empty if upload fails

                    cells_to_create.append(
                        Cell(
                            id=uuid.uuid4(),
                            dataset=dataset,
                            column=column,
                            row=new_row,
                            value=cell_value,
                            value_infos=cell_value_infos,
                            column_metadata=column_metadata,
                        )
                    )

                # Bulk create all cells for this row at once to maintain consistency
                Cell.objects.bulk_create(cells_to_create)

            return self._gm.success_response(f"Successfully added {len(rows)} row(s)")

        except Exception as e:
            logger.exception(f"Error in adding data to rows: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_ADD_DATA_TO_ROWS")
            )


class DatasetDeleteSerializer(serializers.Serializer):
    dataset_ids = serializers.ListField(
        child=serializers.UUIDField(),
        required=True,
        allow_empty=False,
        max_length=50,
    )


class DeleteDatasetView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def delete(self, request, dataset_id=None, *args, **kwargs):
        try:
            # Handle bulk deletion
            serializer = DatasetDeleteSerializer(data=request.data)
            if not serializer.is_valid():
                return self._gm.bad_request(parse_serialized_errors(serializer))

            # Get datasets and verify they exist and belong to user's organization
            datasets = Dataset.objects.filter(
                id__in=serializer.validated_data["dataset_ids"],
                deleted=False,
                organization=getattr(request, "organization", None)
                or request.user.organization,
            )

            # Check if all requested datasets were found
            found_ids = {str(d.id) for d in datasets}
            requested_ids = {str(d) for d in serializer.validated_data["dataset_ids"]}
            missing_ids = requested_ids - found_ids

            if missing_ids:
                return self._gm.bad_request(
                    f"{get_error_message('DATASETS_NOT_FOUND')}: {', '.join(missing_ids)}"
                )

            # Cascade soft-delete experiments for each dataset
            from model_hub.signals import _cascade_soft_delete_dataset_experiments

            for ds in datasets:
                _cascade_soft_delete_dataset_experiments(ds)

            # Bulk soft-delete datasets
            updated_count = datasets.update(deleted=True)

            return self._gm.success_response(
                f"{updated_count} datasets deleted successfully"
            )

        except Exception as e:
            logger.exception(f"Error in deleting the dataset: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_DELETE_DATASET")
            )


class UpdateColumnNameView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    def put(self, request, dataset_id, column_id, *args, **kwargs):
        try:
            new_column_name = request.data.get("new_column_name")

            if not new_column_name:
                return self._gm.bad_request(
                    get_error_message("NEW_COLUMN_NAME_MISSING")
                )

            if len(new_column_name) > 255:
                return self._gm.bad_request(get_error_message("COLUMN_NAME_TOO_LONG"))

            dataset = get_object_or_404(
                Dataset,
                id=dataset_id,
                deleted=False,
                organization=getattr(request, "organization", None)
                or request.user.organization,
            )
            column = get_object_or_404(
                Column,
                id=column_id,
                dataset=dataset,
                dataset__organization=getattr(request, "organization", None)
                or request.user.organization,
            )

            if new_column_name != column.name:
                if Column.objects.filter(
                    name=new_column_name,
                    dataset=dataset,
                    dataset__organization=getattr(request, "organization", None)
                    or request.user.organization,
                    deleted=False,
                ).exists():
                    return self._gm.bad_request(get_error_message("COLUMN_NAME_EXISTS"))

                old_column_name = column.name

                # Update column name
                column.name = new_column_name
                column.save()

                # Update derived variable references if this is a run prompt column
                if column.source == SourceChoices.RUN_PROMPT.value and column.source_id:
                    self._update_derived_variables_for_renamed_column(
                        dataset, column, old_column_name, new_column_name
                    )

            return self._gm.success_response("Column name updated successfully")

        except Exception as e:
            logger.exception(f"Error in updating column name: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_UPDATE_COLUMN_NAME")
            )

    def _update_derived_variables_for_renamed_column(
        self, dataset, column, old_column_name: str, new_column_name: str
    ):
        """
        Update derived variable references when a run prompt column is renamed.

        This updates:
        1. RunPrompter.run_prompt_config["derived_variables"]
        2. PromptVersion.metadata["derived_variables"] for all associated versions
        """
        try:
            # Update RunPrompter
            run_prompter = RunPrompter.objects.filter(
                id=column.source_id,
                deleted=False,
            ).first()

            if run_prompter:
                if rename_derived_variables_in_run_prompter(
                    run_prompter, old_column_name, new_column_name
                ):
                    run_prompter.save(update_fields=["run_prompt_config"])

                # Update all prompt versions associated with this run prompter
                prompt_template = run_prompter.prompt
                if prompt_template:
                    from model_hub.models.prompts import PromptVersion

                    prompt_versions = PromptVersion.objects.filter(
                        prompt=prompt_template,
                        deleted=False,
                    )

                    for version in prompt_versions:
                        if rename_derived_variables_for_column(
                            version, old_column_name, new_column_name
                        ):
                            version.save(update_fields=["metadata"])

            logger.info(
                "Updated derived variables for renamed column",
                old_name=old_column_name,
                new_name=new_column_name,
                column_id=str(column.id),
            )

        except Exception as e:
            # Log but don't fail the rename operation
            logger.warning(
                "Failed to update derived variables for renamed column",
                old_name=old_column_name,
                new_name=new_column_name,
                error=str(e),
            )


class EditDatasetBehaviorView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    def put(self, request, dataset_id, *args, **kwargs):
        try:
            # Get request parameters
            dataset_name = request.data.get("dataset_name")
            column_order = request.data.get("column_order", [])
            column_config = request.data.get("column_config", {})
            dataset_config = request.data.get("dataset_config", {})

            # Get dataset
            dataset = get_object_or_404(Dataset, id=dataset_id, deleted=False)

            if dataset_config:
                dataset.dataset_config = {
                    "dismiss_banner": dataset_config.get("dismiss_banner")
                }

            # Update dataset name if provided
            if dataset_name:
                if dataset_name != dataset.name:
                    if Dataset.objects.filter(
                        name=dataset_name,
                        deleted=False,
                        organization=getattr(request, "organization", None)
                        or request.user.organization,
                    ).exists():
                        return self._gm.bad_request(
                            get_error_message("DATASET_NAME_EXISTS")
                        )
                    dataset.name = dataset_name

            # Get all valid column IDs for this dataset once (for both column_order and column_config validation)
            valid_column_ids = set(
                Column.objects.filter(dataset=dataset, deleted=False).values_list(
                    "id", flat=True
                )
            )
            valid_column_ids_str = {
                str(column_id) if not isinstance(column_id, str) else column_id
                for column_id in valid_column_ids
            }

            # Validate column_order if provided
            if column_order and len(column_order) > 0:
                column_order = [
                    str(column_id) if not isinstance(column_id, str) else column_id
                    for column_id in column_order
                ]
                invalid_columns = set(column_order) - valid_column_ids_str
                if invalid_columns:
                    return self._gm.bad_request(
                        f"{get_error_message('INVALID_COLUMN_IDS')}: {', '.join(invalid_columns)}"
                    )

                dataset.column_order = column_order

            # Validate and update column_config if provided
            if column_config:
                # Get existing config or initialize empty dict
                existing_config = dataset.column_config or {}

                # Validate all column IDs in column_config at once
                column_config_ids = set(column_config.keys())
                invalid_config_columns = column_config_ids - valid_column_ids_str
                if invalid_config_columns:
                    return self._gm.bad_request(
                        f"{get_error_message('INVALID_COLUMN_IDS')}: {', '.join(map(str, invalid_config_columns))}"
                    )

                for column_id, config in column_config.items():
                    # Update or create column configuration
                    column_settings = existing_config.get(column_id, {})

                    # Update visibility and frozen status if provided
                    column_settings["is_visible"] = config.get(
                        "is_visible", column_settings.get("is_visible", True)
                    )
                    column_settings["is_frozen"] = config.get(
                        "is_frozen", column_settings.get("is_frozen", None)
                    )

                    existing_config[column_id] = column_settings

                dataset.column_config = existing_config

            # Save all changes
            dataset.save()

            return self._gm.success_response("Dataset behavior updated successfully")

        except Exception as e:
            logger.exception(f"Error in editing the dataset behavior: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_EDIT_DATASET_BEHAVIOR")
            )


class UpdateCellValueView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    async def _handle_upload_file(self):
        pass

    def _convert_file_to_base64(self, request):
        file = request.FILES.get("new_value") or request.FILES.get("newValue")

        if file:
            try:
                # Reset file pointer to beginning in case it was read before
                if hasattr(file, "seek"):
                    file.seek(0)

                # Add debugging information
                logger.info(
                    f"Processing file: {file.name}, size: {file.size}, content_type: {file.content_type}"
                )

                file_content = file.read()

                base64_content = base64.b64encode(file_content).decode("utf-8")

                mime_type = file.content_type
                base64_string_with_mime = f"data:{mime_type};base64,{base64_content}"

                # Validate the base64 string
                try:
                    # Test if we can decode it back
                    base64.b64decode(base64_content)
                except Exception as decode_error:
                    logger.error(f"Base64 validation failed: {str(decode_error)}")
                    raise ValueError(  # noqa: B904
                        f"Generated base64 string is invalid: {str(decode_error)}"
                    )

                return base64_string_with_mime
            except Exception as e:
                logger.error(f"Error converting file to base64: {str(e)}")
                raise ValueError(f"Failed to convert file to base64: {str(e)}") from e
        else:
            # If no file, check if new_value is already a base64 string in request.data
            new_value = request.data.get("new_value") or request.data.get("newValue")
            if (
                new_value
                and isinstance(new_value, str)
                and new_value.startswith("data:")
            ):
                # It's already a base64 string with mime type
                return new_value
            elif (
                new_value
                and isinstance(new_value, str)
                and new_value.startswith("http")
            ):
                return new_value
            elif (
                new_value
                and isinstance(new_value, str)
                and new_value.startswith("data:application/")
            ):
                # It's already a base64 string with mime type (for documents)
                return new_value
            else:
                return None

    def _convert_to_base64(self, request):
        file = request.FILES.get("new_value") or request.FILES.get("newValue")

        if file:
            audio_content = file.read()
            base64_audio = base64.b64encode(audio_content).decode("utf-8")

            mime_type = file.content_type

            base64_string_with_mime = f"data:{mime_type};base64,{base64_audio}"

            return base64_string_with_mime
        else:
            return None

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            row_id = request.data.get("row_id") or request.data.get("rowId")
            column_id = request.data.get("column_id") or request.data.get("columnId")
            new_value = request.data.get("new_value", "") or request.data.get(
                "newValue", ""
            )

            if not all([row_id, column_id]):
                return self._gm.bad_request(
                    get_error_message("MISSING_ROW_ID_COLUMN_ID_AND_NEW_VALUE")
                )
            # Get the cell and verify it exists
            try:
                cell = get_object_or_404(
                    Cell, dataset_id=dataset_id, row_id=row_id, column_id=column_id
                )
            except Exception:
                cell = get_object_or_404(Cell, row_id=row_id, column_id=column_id)

            # Update the cell value
            # cell.value = str(new_value)
            column = get_object_or_404(Column, id=column_id)
            column_data_type = column.data_type

            # Check if column is editable (non-editable source types cannot be edited)
            from model_hub.services.dataset_validators import (
                MAX_CELL_VALUE_LENGTH,
                validate_column_is_editable,
            )

            is_editable, edit_err = validate_column_is_editable(column)
            if not is_editable:
                return self._gm.bad_request(edit_err)

            # Check max value length
            if isinstance(new_value, str) and len(new_value) > MAX_CELL_VALUE_LENGTH:
                return self._gm.bad_request(
                    f"Value exceeds maximum length of {MAX_CELL_VALUE_LENGTH} characters"
                )

            # Handle empty values for all data types
            if not new_value or (
                isinstance(new_value, str) and new_value.strip() == ""
            ):
                cell.value = None
                cell.value_infos = json.dumps({})
                cell.status = CellStatus.PASS.value
            elif column_data_type == DataTypeChoices.TEXT.value:
                cell.value = str(new_value)
                cell.value_infos = json.dumps({})
                cell.status = CellStatus.PASS.value
            elif column_data_type == DataTypeChoices.BOOLEAN.value:
                if new_value in BooleanChoices.TRUE_OPTIONS.value:
                    cell.value = "true"
                    cell.value_infos = json.dumps({})
                    cell.status = CellStatus.PASS.value
                elif new_value in BooleanChoices.FALSE_OPTIONS.value:
                    cell.value = "false"
                    cell.value_infos = json.dumps({})
                    cell.status = CellStatus.PASS.value
                else:
                    # Return back an error
                    return self._gm.bad_request(
                        f"{get_error_message('INVALID_BOOLEAN_VALUE')}"
                        " \
                    {}".format(
                            BooleanChoices.TRUE_OPTIONS.value
                            + BooleanChoices.FALSE_OPTIONS.value
                        )
                    )
            elif column_data_type == DataTypeChoices.INTEGER.value:
                # check if the value can be converted to an integer
                try:
                    int(float(new_value))
                except (ValueError, TypeError):
                    return self._gm.bad_request(get_error_message("INVALID_INTEGER"))
                cell.value = int(float(new_value))
                cell.value_infos = json.dumps({})
                cell.status = CellStatus.PASS.value
            elif column_data_type == DataTypeChoices.FLOAT.value:
                # check if the value can be converted to a float
                try:
                    float(new_value)
                except (ValueError, TypeError):
                    return self._gm.bad_request(
                        get_error_message("INVALID_FLOAT_VALUE")
                    )
                cell.value = float(new_value)
                cell.value_infos = json.dumps({})
                cell.status = CellStatus.PASS.value
            elif column_data_type == DataTypeChoices.DATETIME.value:
                found_format = False
                for date_format in DateTimeFormatChoices.OPTIONS.value:
                    try:
                        dt = datetime.strptime(new_value, date_format)
                        # Convert to standard format
                        cell.value = dt.strftime("%Y-%m-%d %H:%M:%S")
                        cell.value_infos = json.dumps({})
                        cell.status = CellStatus.PASS.value
                        cell.save()
                        found_format = True
                        break
                    except ValueError:
                        continue
                if not found_format:
                    return self._gm.bad_request(
                        f"Invalid datetime value. \
                    The value must be in one of the following formats: \
                    {DateTimeFormatChoices.OPTIONS.value}"
                    )

            elif column_data_type == DataTypeChoices.ARRAY.value:
                # check if the value can be converted to a json
                try:
                    if new_value.strip().startswith("["):
                        json_value = json.loads(new_value)
                    else:
                        return self._gm.bad_request(get_error_message("INVALID_ARRAY"))
                except (ValueError, TypeError, json.JSONDecodeError):
                    return self._gm.bad_request(
                        get_error_message("INVALID_ARRAY_VALUE")
                    )
                cell.value = json.dumps(json_value)
                cell.value_infos = json.dumps({})
                cell.status = CellStatus.PASS.value

            elif column_data_type == DataTypeChoices.JSON.value:
                # check if the value can be converted to a json
                try:
                    if new_value.strip().startswith("{"):
                        json_value = json.loads(new_value)
                    else:
                        return self._gm.bad_request(
                            get_error_message("INVALID_JSON_FORMAT")
                        )
                except (ValueError, TypeError, json.JSONDecodeError):
                    return self._gm.bad_request(get_error_message("INVALID_JSON_VALUE"))
                cell.value = json.dumps(json_value)
                cell.value_infos = json.dumps({})
                cell.status = CellStatus.PASS.value

            elif column_data_type == DataTypeChoices.IMAGE.value:
                try:
                    if not isinstance(new_value, str):
                        new_value = self._convert_to_base64(request)

                    # Handle empty image value
                    if not new_value or (
                        isinstance(new_value, str) and new_value.strip() == ""
                    ):
                        cell.value = None
                        cell.value_infos = json.dumps({})
                        cell.status = CellStatus.PASS.value
                    else:
                        image_key = f"images/{dataset_id}/{uuid.uuid4()}"
                        image_url = upload_image_to_s3(
                            new_value,
                            bucket_name="fi-customer-data-dev",
                            object_key=image_key,
                        )
                        cell.value = image_url
                        cell.value_infos = json.dumps({})
                        cell.status = CellStatus.PASS.value

                except Exception as e:
                    logger.error(f"ERROR: {e}")
                    cell.value = None
                    cell.status = CellStatus.ERROR.value
                    cell.value_infos = json.dumps({"reason": str(e)})
                    cell.save()
                    return self._gm.bad_request("Invalid image value")

            elif column_data_type == DataTypeChoices.AUDIO.value:
                try:
                    if not isinstance(new_value, str):
                        logger.info("Got AUDIO FILE")
                        new_value = self._convert_to_base64(request)
                    logger.info("GOT AUDIO BASE64")

                    # Handle empty audio value
                    if not new_value or (
                        isinstance(new_value, str) and new_value.strip() == ""
                    ):
                        cell.value = None
                        cell.value_infos = json.dumps({})
                        cell.status = CellStatus.PASS.value
                    else:
                        audio_key = f"audio/{dataset_id}/{uuid.uuid4()}"
                        audio_url, duration = upload_audio_to_s3_duration(
                            new_value,
                            bucket_name="fi-customer-data-dev",
                            object_key=audio_key,
                            duration_seconds=cell.column_metadata.get(
                                "audio_duration_seconds"
                            ),
                        )
                        value_infos = (
                            json.loads(cell.value_infos) if cell.value_infos else {}
                        )
                        if not isinstance(value_infos, dict):
                            value_infos = {}
                        value_infos["audio_url"] = audio_url
                        cell.value_infos = json.dumps(value_infos)
                        cell.status = CellStatus.PASS.value
                        cell.value = audio_url

                except Exception as e:
                    logger.error(f"ERROR: {e}")
                    cell.value = None
                    cell.status = CellStatus.ERROR.value
                    cell.value_infos = json.dumps({"reason": str(e)})
                    cell.save()
                    return self._gm.bad_request(str(e))

            elif column_data_type == DataTypeChoices.DOCUMENT.value:
                try:
                    logger.info("Got DOCUMENT FILE")
                    name = new_value
                    new_value = self._convert_file_to_base64(request)
                    logger.info("GOT DOCUMENT BASE64")

                    # Handle empty document value
                    if not new_value or (
                        isinstance(new_value, str) and new_value.strip() == ""
                    ):
                        cell.value = None
                        cell.value_infos = json.dumps({})
                        cell.status = CellStatus.PASS.value
                    else:
                        doc_key = f"documents/{dataset_id}/{uuid.uuid4()}"
                        doc_url = upload_document_to_s3(
                            new_value,
                            bucket_name="fi-customer-data-dev",
                            object_key=doc_key,
                        )
                        value_infos = (
                            json.loads(cell.value_infos) if cell.value_infos else {}
                        )
                        if not isinstance(value_infos, dict):
                            value_infos = {}
                        value_infos["document_url"] = doc_url
                        value_infos["document_name"] = name[:400]
                        cell.value_infos = json.dumps(value_infos)
                        cell.status = CellStatus.PASS.value
                        cell.value = doc_url

                except Exception as e:
                    logger.error(f"ERROR: {e}")
                    cell.value = None
                    cell.status = CellStatus.ERROR.value
                    cell.value_infos = json.dumps({"reason": str(e)})
                    cell.save()
                    return self._gm.bad_request(str(e))

            elif column_data_type == DataTypeChoices.IMAGES.value:
                try:
                    # Handle empty images value
                    if not new_value or (
                        isinstance(new_value, str) and new_value.strip() == ""
                    ):
                        cell.value = None
                        cell.value_infos = json.dumps({})
                        cell.status = CellStatus.PASS.value
                    else:
                        # Parse the incoming value - could be JSON array of URLs/base64
                        if isinstance(new_value, str):
                            try:
                                images_list = json.loads(new_value)
                            except json.JSONDecodeError:
                                # Single image as string
                                images_list = [new_value]
                        elif isinstance(new_value, list):
                            images_list = new_value
                        else:
                            images_list = [str(new_value)]

                        # Upload each image to S3
                        uploaded_urls = []
                        for img_value in images_list:
                            if img_value:
                                image_key = f"images/{dataset_id}/{uuid.uuid4()}"
                                image_url = upload_image_to_s3(
                                    img_value,
                                    bucket_name="fi-customer-data-dev",
                                    object_key=image_key,
                                )
                                uploaded_urls.append(image_url)

                        # Store as JSON array
                        cell.value = json.dumps(uploaded_urls)
                        cell.value_infos = json.dumps({})
                        cell.status = CellStatus.PASS.value

                except Exception as e:
                    logger.error(f"ERROR uploading images: {e}")
                    cell.value = None
                    cell.status = CellStatus.ERROR.value
                    cell.value_infos = json.dumps({"reason": str(e)})
                    cell.save()
                    return self._gm.bad_request("Invalid images value")

            else:
                cell.value = str(new_value)
                cell.value_infos = json.dumps({})
                cell.status = CellStatus.PASS.value

            cell.save()
            # insert_embeddings_task.delay(cell_id=cell.id)
            return self._gm.success_response("Cell value updated successfully")

        except Exception as e:
            logger.exception(f"Error in updating the cell value: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_UPDATE_CELL_VALUE")
            )


class UpdateRowView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def put(self, request, dataset_id, *args, **kwargs):
        try:
            rows = request.data.get("rows", [])
            if not rows:
                return self._gm.bad_request(get_error_message("DATA_MISSING"))

            dataset = get_object_or_404(Dataset, id=dataset_id)

            # Get valid columns for this dataset
            columns = Column.objects.filter(dataset=dataset, deleted=False).exclude(
                source__in=[
                    SourceChoices.EXPERIMENT.value,
                    SourceChoices.EXPERIMENT_EVALUATION.value,
                    SourceChoices.EXPERIMENT_EVALUATION_TAGS.value,
                ]
            )

            # Update rows and cells
            for row_data in rows:
                row_id = row_data.get("id")
                if not row_id:
                    return self._gm.bad_request(get_error_message("ROW_ID_MISSING"))

                # Get existing row
                try:
                    row = Row.objects.get(id=row_id, dataset=dataset, deleted=False)
                except Row.DoesNotExist:
                    return self._gm.bad_request(get_error_message("ROW_NOT_FOUND"))

                provided_columns = {
                    cell["column_name"] for cell in row_data.get("cells", [])
                }

                # Update cells for this row
                for column in columns:
                    cell_value = ""
                    if column.name in provided_columns:
                        cell_data = next(
                            cell
                            for cell in row_data["cells"]
                            if cell["column_name"] == column.name
                        )
                        cell_value = cell_data.get("value", "")

                    # Update or create cell
                    Cell.objects.update_or_create(
                        dataset=dataset,
                        column=column,
                        row=row,
                        defaults={"value": cell_value},
                    )

            return self._gm.success_response(f"Successfully updated {len(rows)} row(s)")

        except Exception as e:
            logger.exception(f"Error in updating rows: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_UPDATE_ROW")
            )


class UpdateColumnTypeView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def put(self, request, dataset_id, column_id, *args, **kwargs):
        try:
            new_data_type = request.data.get("new_column_type")
            preview = request.data.get("preview", True)
            force_update = request.data.get("force_update", False)

            if not new_data_type:
                return self._gm.bad_request(
                    get_error_message("MISSING_NEW_COLUMN_TYPE")
                )

            # Get dataset and column
            dataset = get_object_or_404(Dataset, id=dataset_id)
            column = get_object_or_404(Column, id=column_id, dataset=dataset)

            # Validate data type
            if column.source_id:
                return self._gm.bad_request(
                    get_error_message("INVALID_COLUMN_TYPE_CHANGE")
                )
            if new_data_type not in [choice.value for choice in DataTypeChoices]:
                return self._gm.bad_request(get_error_message("INVALID_DATA_TYPE"))

            # If preview mode, run validation synchronously
            if preview:
                return self._validate_conversion(column, new_data_type, force_update)

            # Update column status to running
            column.data_type = new_data_type
            column.status = StatusType.RUNNING.value
            Cell.objects.filter(
                column=column, deleted=False, row__deleted=False
            ).update(status=CellStatus.RUNNING.value, value_infos=json.dumps({}))
            column.save()

            force_update = True

            # Start async conversion in thread
            perform_conversion.apply_async(args=(column.id, new_data_type))

            return self._gm.success_response(
                {
                    "message": "Column type conversion started",
                    "column_id": str(column_id),
                    "new_data_type": new_data_type,
                    "status": StatusType.RUNNING.value,
                }
            )

        except Exception as e:
            logger.exception(f"Error in updating column type: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_UPDATE_COLUMN_TYPE")
            )

    def _validate_conversion(self, column, new_data_type, force_update):
        """Validate the conversion without making changes"""
        try:
            cells = Cell.objects.filter(column=column, deleted=False)
            validation_result = self._get_validation_result(cells, new_data_type)

            if validation_result.get("invalid_count", 0) > 0 and not force_update:
                return self._gm.bad_request(validation_result)

            return self._gm.success_response(validation_result)

        except Exception as e:
            logger.exception(f"Invalid conversion: {str(e)}")
            return self._gm.internal_server_error_response("Invalid_conversion")

    def _get_validation_result(self, cells, new_data_type):
        """Validate the conversion and return preview results"""
        try:
            validation_result = {
                "invalid_count": 0,
                "invalid_values": [],
                "valid_conversion_samples": {},
                "new_data_type": new_data_type,
            }

            if new_data_type == DataTypeChoices.TEXT.value:
                # All values are valid as text
                validation_result["valid_conversion_samples"] = {
                    cell.value: cell.value for cell in cells[:5]
                }

            elif new_data_type == DataTypeChoices.BOOLEAN.value:
                regex_pattern = r"^(true|True|TRUE|1|yes|Yes|YES|Passed|Passed|PASSED|false|False|FALSE|0|no|No|NO|Failed|Failed|FAILED)$"
                valid_cells = cells.filter(value__iregex=regex_pattern).all()[:100]
                invalid_cells = cells.exclude(value__iregex=regex_pattern).all()

                validation_result["valid_conversion_samples"] = {
                    cell.value: (
                        BooleanChoices.TRUE.value
                        if cell.value.lower() in BooleanChoices.TRUE_OPTIONS.value
                        else BooleanChoices.FALSE.value
                    )
                    for cell in valid_cells
                }
                validation_result["invalid_count"] = cells.exclude(
                    value__iregex=regex_pattern
                ).count()
                validation_result["invalid_values"] = [
                    cell.value for cell in invalid_cells[:100]
                ]

            elif new_data_type == DataTypeChoices.INTEGER.value:
                regex_pattern = r"^\-?\d*\.?\d+$"
                valid_cells = cells.filter(value__regex=regex_pattern).all()[:100]
                invalid_cells = cells.exclude(value__regex=regex_pattern).all()

                validation_result["valid_conversion_samples"] = {
                    cell.value: int(float(cell.value)) for cell in valid_cells
                }
                validation_result["invalid_count"] = cells.exclude(
                    value__regex=regex_pattern
                ).count()
                validation_result["invalid_values"] = [
                    cell.value for cell in invalid_cells[:100]
                ]

            elif new_data_type == DataTypeChoices.FLOAT.value:
                regex_pattern = r"^\-?\d*\.?\d+$"
                valid_cells = cells.filter(value__regex=regex_pattern).all()[:100]
                invalid_cells = cells.exclude(value__regex=regex_pattern).all()

                validation_result["valid_conversion_samples"] = {
                    cell.value: float(cell.value) for cell in valid_cells
                }
                validation_result["invalid_count"] = cells.exclude(
                    value__regex=regex_pattern
                ).count()
                validation_result["invalid_values"] = [
                    cell.value for cell in invalid_cells[:100]
                ]

            elif new_data_type == DataTypeChoices.DATETIME.value:
                valid_conversion_samples = {}
                invalid_values = []
                invalid_count = 0

                for cell in cells[:100]:  # Check first 100 cells for preview
                    if not cell.value:
                        continue

                    found_format = False
                    for date_format in DateTimeFormatChoices.OPTIONS.value:
                        try:
                            dt = datetime.strptime(cell.value, date_format)
                            valid_conversion_samples[cell.value] = dt.strftime(
                                "%Y-%m-%d %H:%M:%S"
                            )
                            found_format = True
                            break
                        except ValueError:
                            continue

                    if not found_format:
                        invalid_count += 1
                        invalid_values.append(cell.value)

                validation_result["valid_conversion_samples"] = valid_conversion_samples
                validation_result["invalid_count"] = invalid_count
                validation_result["invalid_values"] = invalid_values

            elif new_data_type == DataTypeChoices.ARRAY.value:
                valid_conversion_samples = {}
                invalid_values = []
                invalid_count = 0

                for cell in cells[:100]:  # Check first 100 cells for preview
                    try:
                        if cell.value and cell.value.strip().startswith("["):
                            json_value = json.loads(cell.value)
                            valid_conversion_samples[cell.value] = json.dumps(
                                json_value
                            )
                        else:
                            invalid_count += 1
                            invalid_values.append(cell.value)
                    except json.JSONDecodeError:
                        invalid_count += 1
                        invalid_values.append(cell.value)

                validation_result["valid_conversion_samples"] = valid_conversion_samples
                validation_result["invalid_count"] = invalid_count
                validation_result["invalid_values"] = invalid_values

            elif new_data_type == DataTypeChoices.JSON.value:
                valid_conversion_samples = {}
                invalid_values = []
                invalid_count = 0

                for cell in cells[:100]:  # Check first 100 cells for preview
                    try:
                        if cell.value and cell.value.strip().startswith("{"):
                            json_value = json.loads(cell.value)
                            valid_conversion_samples[cell.value] = json.dumps(
                                json_value
                            )
                        else:
                            invalid_count += 1
                            invalid_values.append(cell.value)
                    except json.JSONDecodeError:
                        invalid_count += 1
                        invalid_values.append(cell.value)

                validation_result["valid_conversion_samples"] = valid_conversion_samples
                validation_result["invalid_count"] = invalid_count
                validation_result["invalid_values"] = invalid_values

            elif new_data_type == DataTypeChoices.IMAGE.value:
                # For images, we can only validate URL format or base64 encoding
                valid_conversion_samples = {}
                invalid_values = []
                invalid_count = 0

                url_pattern = r"^(http|https)://"
                base64_pattern = r"^data:image\/[a-zA-Z]+;base64,"

                for cell in cells[:100]:  # Check first 100 cells for preview
                    if cell.value and (
                        re.match(url_pattern, cell.value)
                        or re.match(base64_pattern, cell.value)
                    ):
                        valid_conversion_samples[cell.value] = "Valid image URL/base64"
                    else:
                        invalid_count += 1
                        invalid_values.append(cell.value)

                validation_result["valid_conversion_samples"] = valid_conversion_samples
                validation_result["invalid_count"] = invalid_count
                validation_result["invalid_values"] = invalid_values

            elif new_data_type == DataTypeChoices.IMAGES.value:
                # For multiple images, validate JSON array of URLs or base64 strings
                valid_conversion_samples = {}
                invalid_values = []
                invalid_count = 0

                url_pattern = r"^(http|https)://"
                base64_pattern = r"^data:image\/[a-zA-Z]+;base64,"

                for cell in cells[:100]:  # Check first 100 cells for preview
                    if not cell.value or cell.value.strip() == "":
                        continue

                    try:
                        # Try to parse as JSON array
                        images_list = json.loads(cell.value)
                        if isinstance(images_list, list):
                            # Validate each image in the array
                            all_valid = True
                            for img in images_list:
                                if img and not (
                                    re.match(url_pattern, str(img))
                                    or re.match(base64_pattern, str(img))
                                ):
                                    all_valid = False
                                    break
                            if all_valid:
                                valid_conversion_samples[cell.value] = (
                                    "Valid images array"
                                )
                            else:
                                invalid_count += 1
                                invalid_values.append(cell.value)
                        else:
                            # Single value - check if it's a valid image
                            if re.match(url_pattern, str(cell.value)) or re.match(
                                base64_pattern, str(cell.value)
                            ):
                                valid_conversion_samples[cell.value] = (
                                    "Valid single image"
                                )
                            else:
                                invalid_count += 1
                                invalid_values.append(cell.value)
                    except json.JSONDecodeError:
                        # Not JSON - check if it's a single valid image URL/base64
                        if re.match(url_pattern, str(cell.value)) or re.match(
                            base64_pattern, str(cell.value)
                        ):
                            valid_conversion_samples[cell.value] = "Valid single image"
                        else:
                            invalid_count += 1
                            invalid_values.append(cell.value)

                validation_result["valid_conversion_samples"] = valid_conversion_samples
                validation_result["invalid_count"] = invalid_count
                validation_result["invalid_values"] = invalid_values

            if validation_result["invalid_count"] > 0:
                validation_result["message"] = (
                    f"Found {validation_result['invalid_count']} invalid entries. "
                    "Please correct these entries or use force_update to proceed."
                )

            return validation_result

        except Exception as e:
            raise ValueError(f"Validation failed: {str(e)}")  # noqa: B904


@temporal_activity(time_limit=3600, queue="tasks_s")
def perform_conversion(column_id, new_data_type):
    """
    Perform datatype conversion with all-or-nothing guarantee.
    Either all cells convert successfully, or none are modified.
    """
    column = None
    try:
        column = get_object_or_404(Column, id=column_id)
        dataset = column.dataset
        dataset_config = dataset.dataset_config

        if dataset_config.get("dataset_source_local"):
            dataset_config.pop("dataset_source_local")
            dataset.dataset_config = dataset_config
            dataset.save()

        cells = Cell.objects.filter(column=column, deleted=False, row__deleted=False)

        # Convert cells - in lenient mode, failed cells keep their original values
        # In strict mode, any failure aborts the entire operation
        converter = DatatypeConverter(
            new_data_type,
            column.dataset_id,
            allow_partial_failure=True,  # Keep original values for failed cells
        )
        converter.convert(cells)

        # Only reached if conversion succeeded
        column.data_type = new_data_type
        column.status = StatusType.COMPLETED.value
        column.error_message = None
        column.save()

    except Exception as e:
        logger.exception(
            f"Error in perform_conversion for column {column_id}: {str(e)}"
        )
        # No data was modified - cells are still in original state
        if column:
            column.status = StatusType.FAILED.value
            column.error_message = str(e)
            column.save()
        raise  # Re-raise so Temporal knows the activity failed


class DatatypeConverter:
    """
    Handles datatype conversions with validation-first approach.
    Only updates DB if ALL conversions succeed.
    """

    def __init__(self, new_data_type, dataset_id=None, allow_partial_failure=False):
        self.new_data_type = new_data_type
        self.dataset_id = dataset_id
        self.allow_partial_failure = allow_partial_failure

    def convert(self, cells_queryset):
        """
        Convert cells with validation-first approach.

        In strict mode: Only updates DB if ALL conversions succeed.
        In lenient mode: Updates successful conversions, preserves original values for failures.
        """
        # Phase 1: Load all cells and attempt conversion in memory
        conversion_results = []
        cells_dict = {}  # Store cell objects by ID for reuse

        for cell in cells_queryset.iterator(chunk_size=1000):
            # Store cell object for later use
            cells_dict[str(cell.id)] = cell

            result = self._convert_single_cell(cell)
            conversion_results.append(result)

        # Phase 2: Check if any conversions failed
        failed_conversions = [r for r in conversion_results if not r.success]

        if failed_conversions:
            if not self.allow_partial_failure:
                # Strict mode - abort if ANY cell fails
                error_summary = self._generate_error_summary(failed_conversions)
                raise ValueError(
                    f"Conversion failed for {len(failed_conversions)} cells. "
                    f"No data was modified. {error_summary}"
                )
            else:
                # Lenient mode - restore original values for failed cells
                logger.warning(
                    f"Conversion failed for {len(failed_conversions)} cells. "
                    f"Keeping original values for failed cells."
                )
                for result in conversion_results:
                    if not result.success:
                        # Restore original value from cells_dict
                        original_cell = cells_dict[result.cell_id]
                        result.new_value = original_cell.value
                        result.status = CellStatus.ERROR.value
                        # Keep the error info but preserve the value

        # Phase 3: Apply conversions (pass cells_dict to avoid re-querying)
        self._apply_conversions(conversion_results, cells_dict)

    def _convert_single_cell(self, cell) -> ConversionResult:
        """Convert a single cell value without touching the database"""
        try:
            converter_map = {
                DataTypeChoices.TEXT.value: self._convert_cell_to_text,
                DataTypeChoices.BOOLEAN.value: self._convert_cell_to_boolean,
                DataTypeChoices.INTEGER.value: self._convert_cell_to_integer,
                DataTypeChoices.FLOAT.value: self._convert_cell_to_float,
                DataTypeChoices.DATETIME.value: self._convert_cell_to_datetime,
                DataTypeChoices.ARRAY.value: self._convert_cell_to_array,
                DataTypeChoices.JSON.value: self._convert_cell_to_json,
                DataTypeChoices.IMAGE.value: self._convert_cell_to_image,
                DataTypeChoices.IMAGES.value: self._convert_cell_to_images,
                DataTypeChoices.AUDIO.value: self._convert_cell_to_audio,
                DataTypeChoices.DOCUMENT.value: self._convert_cell_to_document,
            }

            converter = converter_map.get(self.new_data_type)
            if not converter:
                raise ValueError(f"Unsupported datatype: {self.new_data_type}")

            new_value, value_infos = converter(cell)

            return ConversionResult(
                cell_id=str(cell.id),
                success=True,
                new_value=new_value,
                status=CellStatus.PASS.value,
                value_infos=value_infos or {},
            )

        except Exception as e:
            return ConversionResult(
                cell_id=str(cell.id),
                success=False,
                new_value=None,
                status=CellStatus.ERROR.value,
                value_infos={"reason": str(e)},
                error_message=str(e),
            )

    def _apply_conversions(
        self, conversion_results: list[ConversionResult], cells_dict: dict
    ):
        """Apply all successful conversions using bulk update"""
        cells_to_update = []
        for result in conversion_results:
            cell = cells_dict.get(result.cell_id)
            if cell:
                cell.value = result.new_value
                cell.status = result.status
                cell.value_infos = json.dumps(result.value_infos)
                cells_to_update.append(cell)

        # Single bulk update - all cells updated together
        if cells_to_update:
            Cell.objects.bulk_update(
                cells_to_update, ["value", "status", "value_infos"], batch_size=1000
            )

    def _generate_error_summary(
        self, failed_conversions: list[ConversionResult]
    ) -> str:
        """Generate human-readable error summary"""
        if not failed_conversions:
            return ""

        # Show first 5 errors as examples
        sample_errors = failed_conversions[:5]
        error_details = "\n".join(
            [f"  - Cell {r.cell_id}: {r.error_message}" for r in sample_errors]
        )

        summary = f"\nFirst {len(sample_errors)} errors:\n{error_details}"
        if len(failed_conversions) > 5:
            summary += f"\n  ... and {len(failed_conversions) - 5} more"

        return summary

    # Individual converter methods - pure functions, no DB access

    def _is_default_empty_value(self, value, data_type=None):
        """
        Check if a value is a default empty value from ANY data type.

        Default empty values by type:
        - TEXT: "" or None
        - ARRAY: "[]"
        - JSON: "{}"
        - IMAGE/AUDIO/DOCUMENT: None or ""
        - INTEGER/FLOAT/DATETIME/BOOLEAN: None or ""

        Note: "false" for BOOLEAN is NOT a default empty value - it's actual data.

        This function checks if the value is a default empty value from ANY type,
        not just the target type. This allows conversions like {} -> [] or [] -> {}
        to work properly.
        """
        if value is None or value == "":
            return True

        # Convert to string and strip to handle both string and object representations
        str_value = str(value).strip()

        # Check if it's a default empty value from any type
        if str_value in ("{}", "[]"):
            return True

        return False

    def _convert_cell_to_text(self, cell):
        """Convert to text - always succeeds"""
        return cell.value, {}

    def _convert_cell_to_boolean(self, cell):
        """Convert to boolean"""
        # Empty cells will be false by default
        if self._is_default_empty_value(cell.value, self.new_data_type):
            return BooleanChoices.FALSE.value, {}

        try:
            true_pattern = r"^(true|1|yes|passed)$"
            false_pattern = r"^(false|0|no|failed)$"

            value_str = str(cell.value).strip()

            if re.match(true_pattern, value_str, re.IGNORECASE):
                return BooleanChoices.TRUE.value, {}
            elif re.match(false_pattern, value_str, re.IGNORECASE):
                return BooleanChoices.FALSE.value, {}
            else:
                # Default to false for non-matching values
                return BooleanChoices.FALSE.value, {"note": "Defaulted to false"}
        except Exception:
            # If regex fails, default to false
            return BooleanChoices.FALSE.value, {
                "note": "Defaulted to false due to conversion error"
            }

    def _convert_cell_to_integer(self, cell):
        """Convert to integer"""
        # Check for empty or default empty values (but not "false" - that's real data)
        if self._is_default_empty_value(cell.value, self.new_data_type):
            raise ValueError("Empty value cannot be converted to integer")

        try:
            int_value = int(float(str(cell.value).strip()))
            return str(int_value), {}
        except (ValueError, TypeError) as e:
            raise ValueError(f"Cannot convert '{cell.value}' to integer") from e

    def _convert_cell_to_float(self, cell):
        """Convert to float"""
        # Check for empty or default empty values (but not "false" - that's real data)
        if self._is_default_empty_value(cell.value, self.new_data_type):
            raise ValueError("Empty value cannot be converted to float")

        try:
            float_value = float(str(cell.value).strip())
            # Use Python's general format to avoid unnecessary decimals
            formatted = f"{float_value:g}"
            return formatted, {}
        except (ValueError, TypeError) as e:
            raise ValueError(f"Cannot convert '{cell.value}' to float") from e

    def _convert_cell_to_datetime(self, cell):
        """Convert to datetime"""
        # Check for empty or default empty values
        if self._is_default_empty_value(cell.value, self.new_data_type):
            raise ValueError("Empty value cannot be converted to datetime")

        cell_value = str(cell.value).strip()

        # Try standard datetime formats
        for date_format in DateTimeFormatChoices.OPTIONS.value:
            try:
                dt = datetime.strptime(cell_value, date_format)
                return dt.strftime("%Y-%m-%d %H:%M:%S"), {}
            except ValueError:
                continue

        # Try Unix timestamps
        if cell_value.isdigit():
            try:
                timestamp = float(cell_value)
                if len(cell_value) == 10:  # Seconds
                    dt = datetime.fromtimestamp(timestamp)
                elif len(cell_value) == 13:  # Milliseconds
                    dt = datetime.fromtimestamp(timestamp / 1000)
                else:
                    raise ValueError("Invalid timestamp length")
                return dt.strftime("%Y-%m-%d %H:%M:%S"), {}
            except (ValueError, OverflowError, OSError):
                pass

        raise ValueError(f"Cannot parse datetime from '{cell.value}'")

    def _convert_cell_to_array(self, cell):
        """Convert to array (JSON array)"""
        # if empty cell or default empty value (like "{}"), return default empty array
        if self._is_default_empty_value(cell.value):
            return "[]", {}

        try:
            cell_value = str(cell.value).strip()

            if not cell_value.startswith("["):
                raise ValueError(f"Value does not look like an array: '{cell.value}'")

            # Try parsing as JSON
            try:
                json_value = json.loads(cell_value)
                if not isinstance(json_value, list):
                    raise ValueError("Parsed value is not an array")
                return json.dumps(json_value), {}
            except json.JSONDecodeError as e:
                raise ValueError(f"Invalid array format: {str(e)}") from e

        except Exception as e:
            # Catch any other unexpected errors
            raise ValueError(f"Failed to convert to array: {str(e)}") from e

    def _convert_cell_to_json(self, cell):
        """Convert to JSON (JSON object)"""
        # if empty cell or default empty value (like "[]"), return default empty object
        if self._is_default_empty_value(cell.value):
            return "{}", {}

        try:
            cell_value = str(cell.value).strip()

            # Try direct JSON parsing
            try:
                json_value = json.loads(cell_value)

                # array is still fine to view as json.
                # if isinstance(json_value, list):
                #     raise ValueError("Value is an array, not a JSON object")
                return json.dumps(json_value), {}
            except json.JSONDecodeError:
                try:
                    json_value = json_repair.loads(cell_value)
                    return json.dumps(json_value), {}
                except json.JSONDecodeError:
                    # Try ast.literal_eval for Python dicts
                    try:
                        parsed = ast.literal_eval(cell_value)
                        if isinstance(parsed, dict):
                            return json.dumps(parsed), {}
                    except (ValueError, SyntaxError):
                        pass

            raise ValueError(f"Cannot parse as valid JSON: '{cell.value[:100]}...'")

        except ValueError:
            # Re-raise ValueError as-is (our custom error messages)
            raise
        except Exception as e:
            # Catch any other unexpected errors
            raise ValueError(f"Failed to convert to JSON: {str(e)}") from e

    def _convert_cell_to_image(self, cell):
        """Convert to image - uploads to S3"""
        if self._is_default_empty_value(cell.value, self.new_data_type):
            return None, {}

        try:
            image_value = cell.value
            is_s3_url = False

            # Handle JSON array with single element (e.g., from images -> image conversion)
            if isinstance(image_value, str) and image_value.strip().startswith("["):
                try:
                    parsed = json.loads(image_value)
                    if isinstance(parsed, list) and len(parsed) == 1:
                        image_value = parsed[0]
                        is_s3_url = (
                            isinstance(image_value, str)
                            and "fi-customer-data" in image_value
                        )
                except (json.JSONDecodeError, TypeError):
                    pass

            # Always validate the URL, even if extracted from a JSON array
            validate_file_url(str(image_value), "image")

            # Skip re-upload only if it's already in our S3 bucket
            if is_s3_url:
                return image_value, {}

            image_key = f"images/{self.dataset_id}/{uuid.uuid4()}"
            image_url = upload_image_to_s3(
                image_value, bucket_name="fi-customer-data-dev", object_key=image_key
            )
            return image_url, {}
        except Exception as e:
            raise ValueError(f"Failed to upload image: {str(e)}") from e

    def _convert_cell_to_images(self, cell):
        """Convert to images (multiple) - uploads each image to S3 and stores as JSON array"""
        if self._is_default_empty_value(cell.value, self.new_data_type):
            return None, {}

        try:
            # Parse the value - could be JSON array or single image
            if isinstance(cell.value, str):
                try:
                    images_list = json.loads(cell.value)
                    if not isinstance(images_list, list):
                        images_list = [images_list]
                except json.JSONDecodeError:
                    # Single image as string
                    images_list = [cell.value]
            elif isinstance(cell.value, list):
                images_list = cell.value
            else:
                images_list = [str(cell.value)]

            # Upload each image to S3
            uploaded_urls = []
            for img_value in images_list:
                if img_value:
                    image_key = f"images/{self.dataset_id}/{uuid.uuid4()}"
                    image_url = upload_image_to_s3(
                        str(img_value),
                        bucket_name="fi-customer-data-dev",
                        object_key=image_key,
                    )
                    uploaded_urls.append(image_url)

            return json.dumps(uploaded_urls), {}
        except Exception as e:
            raise ValueError(f"Failed to upload images: {str(e)}") from e

    def _convert_cell_to_document(self, cell):
        """Convert to document - uploads to S3"""
        if self._is_default_empty_value(cell.value, self.new_data_type):
            return None, {}

        try:
            # Validate the document URL before attempting upload
            validate_file_url(str(cell.value), "document")

            doc_key = f"documents/{self.dataset_id}/{uuid.uuid4()}"
            doc_url = upload_document_to_s3(
                str(cell.value),
                bucket_name="fi-customer-data-dev",
                object_key=doc_key,
            )
            return doc_url, {}
        except Exception as e:
            raise ValueError(f"Failed to upload document: {str(e)}") from e

    def _convert_cell_to_audio(self, cell):
        """Convert to audio - uploads to S3"""
        if not cell.value:
            return None, {}

        try:
            # Validate the audio URL before attempting upload
            validate_file_url(str(cell.value), "audio")

            audio_key = f"audio/{self.dataset_id}/{uuid.uuid4()}"
            audio_url, duration_seconds = upload_audio_to_s3_duration(
                str(cell.value),
                bucket_name="fi-customer-data-dev",
                object_key=audio_key,
                duration_seconds=(
                    cell.column_metadata.get("audio_duration_seconds")
                    if hasattr(cell, "column_metadata")
                    else None
                ),
            )
            return audio_url, {}
        except Exception as e:
            raise ValueError(f"Failed to upload audio: {str(e)}") from e


class DownloadDatasetView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    def get(self, request, dataset_id, *args, **kwargs):
        try:
            # Get dataset and verify it exists
            dataset = get_object_or_404(Dataset, id=dataset_id)

            # Get all columns in the correct order
            _exp_sources = [
                SourceChoices.EXPERIMENT.value,
                SourceChoices.EXPERIMENT_EVALUATION.value,
                SourceChoices.EXPERIMENT_EVALUATION_TAGS.value,
            ]
            column_order = dataset.column_order or []
            if column_order:
                from django.db.models import Case, IntegerField, Value, When

                order_cases = [
                    When(id=uuid.UUID(col_id), then=Value(idx))
                    for idx, col_id in enumerate(column_order)
                ]
                columns = (
                    Column.objects.filter(dataset_id=dataset_id, deleted=False)
                    .annotate(
                        custom_order=Case(
                            *order_cases,
                            default=Value(len(column_order)),
                            output_field=IntegerField(),
                        )
                    )
                    .order_by("custom_order")
                )
            else:
                columns = Column.objects.filter(
                    dataset_id=dataset_id, deleted=False
                ).order_by("id")
            if dataset.source != DatasetSourceChoices.EXPERIMENT_SNAPSHOT.value:
                columns = columns.exclude(source__in=_exp_sources)

            # Create DataFrame
            rows = Row.objects.filter(dataset=dataset, deleted=False).order_by("order")

            # Initialize data dictionary with empty lists of the correct length
            num_rows = rows.count()
            data = {col.name: ["" for _ in range(num_rows)] for col in columns}

            # Fetch all cells in bulk to improve performance
            cells = Cell.objects.filter(
                row__in=rows, column__in=columns, deleted=False
            ).select_related("row", "column")

            # Create a mapping of (row_id, column_id) to cell value
            cell_mapping = {
                (str(cell.row_id), str(cell.column_id)): cell.value for cell in cells
            }

            # Fill the data dictionary
            for idx, row in enumerate(rows):
                for col in columns:
                    value = cell_mapping.get((str(row.id), str(col.id)), "")
                    data[col.name][idx] = value if value is not None else ""

            df = pd.DataFrame(data)

            # Convert to CSV buffer
            buffer = io.BytesIO()
            df.to_csv(buffer, index=False, encoding="utf-8")
            buffer.seek(0)

            # Create the response with the file
            filename = f"{dataset.name or 'dataset'}.csv"
            response = FileResponse(
                buffer, as_attachment=True, filename=filename, content_type="text/csv"
            )

            return response

        except Exception as e:
            logger.exception(f"Error in downloading the dataset: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_DOWNLOAD_DATASET")
            )


from rest_framework import serializers  # noqa: E402

from model_hub.models.choices import OwnerChoices  # noqa: E402


class TemplateEvalSerializer(serializers.Serializer):
    name = serializers.CharField(max_length=50)
    owner = serializers.CharField(max_length=50, default=OwnerChoices.SYSTEM.value)
    config = serializers.JSONField()  # JSONField is used for dictionary-like objects
    eval_tags = serializers.ListField(
        child=serializers.CharField(
            max_length=100
        ),  # Each tag is a string with max length of 100
        allow_empty=True,  # Allows empty lists
        required=False,  # Makes eval_tags optional
    )


class GetFunctionList(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        try:
            eval_templates = EvalTemplate.no_workspace_objects.filter(
                name__in=FUNCTION_CONFIG_EVALS
            ).all()

            # Serialize the eval templates
            serializer = EvalTemplateSerializer(eval_templates, many=True)

            return self._gm.success_response({"functions": serializer.data})
        except Exception as e:
            logger.error(f"Error in GetFunctionList: {str(e)}")
            return self._gm.bad_request(f"Error in GetFunctionList: {str(e)}")


class GetEvalsListView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def get(
        self, request, dataset_id=None, *args, **kwargs
    ):  # Changed from 'post' to 'get'
        try:
            # Get parameters from query params instead of request body
            search_text = request.GET.get("search_text", "").strip()
            eval_categories = request.GET.get("eval_categories")
            eval_type = request.GET.get("eval_type")
            eval_tags = request.GET.getlist("eval_tags[]") or request.GET.getlist(
                "eval_tags"
            )
            use_cases = request.GET.getlist("use_cases[]") or request.GET.getlist(
                "use_cases"
            )
            experiment_id = request.GET.get("experiment_id")
            order = request.GET.get("order")

            # Create a mock validated_data structure for compatibility
            validated_data = {
                "search_text": search_text,
                "eval_categories": eval_categories,
                "eval_type": eval_type,
                "eval_tags": eval_tags,
                "use_cases": use_cases,
            }
            search_text = validated_data.get("search_text", "").strip()
            eval_categories = validated_data.get("eval_categories")
            eval_type = validated_data.get("eval_type")
            eval_tags = validated_data.get("eval_tags")
            use_cases = validated_data.get("use_cases")

            all_evals = []

            if experiment_id and dataset_id:
                try:
                    experiment = ExperimentsTable.objects.get(
                        id=experiment_id, dataset_id=dataset_id, deleted=False
                    )
                except ExperimentsTable.DoesNotExist:
                    return self._gm.bad_request(
                        get_error_message("EXPERIMENT_NOT_FOUND")
                    )

                user_evals = (
                    experiment.user_eval_template_ids.all()
                    .filter(deleted=False, template__deleted=False)
                    .select_related("template")
                )
                if eval_type == "user":
                    all_evals.extend(
                        self._get_user_evals(
                            validated_data,
                            getattr(request, "organization", None)
                            or request.user.organization,
                            dataset_id,
                            search_text=search_text,
                            user_evals=user_evals,
                        )
                    )
            else:
                if eval_type:
                    if eval_type == "previously_configured":
                        all_evals.extend(
                            self._get_previously_configured_evals(
                                validated_data,
                                getattr(request, "organization", None)
                                or request.user.organization,
                                search_text=search_text,
                            )
                        )
                    elif eval_type == "user":
                        if not dataset_id:
                            return self._gm.bad_request(
                                get_error_message("MISSING_DATASET_ID")
                            )
                        all_evals.extend(
                            self._get_user_evals(
                                validated_data,
                                getattr(request, "organization", None)
                                or request.user.organization,
                                dataset_id,
                                search_text=search_text,
                            )
                        )
                elif eval_tags:
                    all_evals.extend(
                        self._get_preset_evals(validated_data, search_text=search_text)
                    )
                else:
                    if eval_categories == "futureagi_built":
                        all_evals.extend(
                            self._get_preset_evals(
                                validated_data, search_text=search_text
                            )
                        )
                    elif eval_categories == "user_built":
                        all_evals.extend(
                            self._custom_build_evals(
                                validated_data,
                                getattr(request, "organization", None)
                                or request.user.organization,
                                search_text=search_text,
                            )
                        )

                    else:
                        # Default to fetching all evals
                        all_evals.extend(
                            self._get_preset_evals(
                                validated_data, search_text=search_text
                            )
                        )
                        all_evals.extend(
                            self._custom_build_evals(
                                validated_data,
                                getattr(request, "organization", None)
                                or request.user.organization,
                                search_text=search_text,
                            )
                        )
            eval_recommendations = ["Deterministic Evals"]
            if dataset_id:
                try:
                    dataset = Dataset.objects.get(
                        id=dataset_id,
                        organization=getattr(request, "organization", None)
                        or request.user.organization,
                        deleted=False,
                    )
                    eval_recommendations = dataset.dataset_config.get(
                        "eval_recommendations", ["Deterministic Evals"]
                    )
                except Dataset.DoesNotExist:
                    pass
            if use_cases:
                all_items = []
                for use_case in use_cases:
                    all_items.extend(USE_CASE_MAPPING.get(use_case, []))

                eval_filters = set(all_items)
                filtered_evals = []
                for eval_item in all_evals:
                    if eval_item.get("name") in eval_filters:
                        filtered_evals.append(eval_item)
                    elif any(
                        use_case in eval_item.get("eval_template_tags", [])
                        for use_case in use_cases
                    ):
                        filtered_evals.append(eval_item)
                all_evals = filtered_evals

            # Apply ordering if specified
            if order == "simulate":
                all_evals = self._apply_simulate_ordering(all_evals)

            response = {
                "evals": all_evals,
                "eval_recommendations": eval_recommendations,
            }
            return self._gm.success_response(response)

        except Exception as e:
            logger.exception(f"Error in fetching the eval lists: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_EVAL_LISTS")
            )

    def _apply_simulate_ordering(self, evals_list):
        """
        Orders the evals list based on eval_template_tags for simulate workflow.
        Priority order: AUDIO > CONVERSATION > IMAGE > TEXT > SAFETY > RAG > HALLUCINATION > FUNCTION > LLMS > FUTURE_EVALS > CUSTOM
        """
        from model_hub.utils.evals import (
            AUDIO,
            CONVERSATION,
            CUSTOM,
            FUNCTION,
            FUTURE_EVALS,
            HALLUCINATION,
            IMAGE,
            LLMS,
            RAG,
            SAFETY,
            TEXT,
        )

        # Define the priority order for simulate
        tag_priority = {
            AUDIO: 0,
            CONVERSATION: 1,
            TEXT: 2,
            CUSTOM: 3,
            SAFETY: 4,
            RAG: 5,
            HALLUCINATION: 6,
            FUNCTION: 7,
            LLMS: 8,
            FUTURE_EVALS: 9,
            IMAGE: 10,
        }

        def get_sort_key(eval_item):
            """
            Returns a tuple (priority, name) for sorting.
            Priority is based on tag pairs first, then individual tags.
            If no matching tag is found, it's given the lowest priority.
            """
            tags = eval_item.get("eval_template_tags", [])

            # Define special pairs that get highest priority
            priority_pairs = [
                (
                    AUDIO,
                    CONVERSATION,
                ),  # AUDIO + CONVERSATION pair gets highest priority
                (AUDIO, CUSTOM),
                (AUDIO, TEXT),
                # Add more pairs here as needed
            ]

            # Check for priority pairs first
            if tags and len(tags) >= 2:
                for pair in priority_pairs:
                    if pair[0] in tags and pair[1] in tags:
                        # Return negative priority to ensure pairs come first
                        return (-1, eval_item.get("name", ""))

            # Find the highest priority tag (lowest number)
            min_priority = 999  # Default for items without matching tags

            if tags:
                for tag in tags:
                    if tag in tag_priority:
                        min_priority = min(min_priority, tag_priority[tag])

            # Return tuple for sorting: (priority, name)
            # Items with same priority will be sorted alphabetically by name
            return (min_priority, eval_item.get("name", ""))

        # Sort the evals list
        sorted_evals = sorted(evals_list, key=get_sort_key)

        return sorted_evals

    def _get_preset_evals(self, validated_data, search_text):
        from model_hub.utils.eval_list import (
            derive_eval_type,
            derive_output_type,
        )

        eval_templates = EvalTemplate.no_workspace_objects.filter(
            organization__isnull=True, deleted=False, visible_ui=True
        )

        # IMPORTANT: do NOT call insert_evals_template() here.
        # That helper uses a legacy hardcoded Python list
        # (`model_hub/utils/evals.py:evals_template`) which has empty
        # `config.code` for code evals. It overwrites every system eval's
        # config via bulk_update, wiping the real code that the
        # YAML-based `seed_system_evals` command installs — so users see
        # system code evals with no code and no variables in the picker.
        #
        # The authoritative seeder is now `seed_system_evals` (YAML-based)
        # which runs at deploy time. Reading from the DB directly is the
        # correct behaviour at request time.

        if search_text:
            eval_templates = eval_templates.filter(name__icontains=search_text)

        if validated_data.get("eval_tags"):
            eval_templates = eval_templates.filter(
                eval_tags__overlap=validated_data.get("eval_tags")
            )

        run_evals = []
        # Fetch the fields we need for the picker's display:
        # eval_type / output_type_normalized drive the type and output
        # badges; updated_at drives the Last Updated column; model is
        # used to compute is_model_required; owner is used for the
        # "created by" label.
        eval_templates = list(
            eval_templates.values(
                "id",
                "name",
                "config",
                "eval_tags",
                "description",
                "eval_type",
                "output_type_normalized",
                "updated_at",
                "created_at",
                "owner",
                "model",
            )
        )

        for template in eval_templates:
            is_model_present = False
            # for template in eval_templates:
            if template.get("config", {}).get("config", {}).get("model"):
                is_model_present = True
            if template.get("name") not in NOT_UI_EVALS:
                # Resolve the true eval_type (llm/code/agent). Prefer the
                # dedicated field; fall back to tag/config derivation for
                # legacy rows that pre-date the field.
                _eval_type = template.get("eval_type")
                if not _eval_type:
                    # Build a proxy object with the fields derive_eval_type
                    # reads — it only needs eval_type, config, eval_tags.
                    class _Proxy:
                        pass

                    proxy = _Proxy()
                    proxy.eval_type = None
                    proxy.config = template.get("config") or {}
                    proxy.eval_tags = template.get("eval_tags") or []
                    _eval_type = derive_eval_type(proxy)

                _output_type = template.get("output_type_normalized")
                if not _output_type:

                    class _Proxy2:
                        pass

                    proxy2 = _Proxy2()
                    proxy2.config = template.get("config") or {}
                    _output_type = derive_output_type(proxy2)

                run_evals.append(
                    {
                        "id": str(template.get("id")),
                        "name": template.get("name"),
                        "eval_template_name": template.get("name"),
                        "eval_required_keys": template.get("config", {}).get(
                            "required_keys", []
                        ),
                        "eval_template_tags": template.get("eval_tags"),
                        "description": template.get("description"),
                        "is_model_required": is_model_present,
                        "type": "futureagi_built",
                        # Fields the picker frontend needs:
                        "eval_type": _eval_type,
                        "output_type": _output_type,
                        "created_by_name": "System",
                        "owner": "system",
                        "updated_at": template.get("updated_at"),
                        "created_at": template.get("created_at"),
                        "model": template.get("model") or "",
                        "template_type": "single",
                    }
                )

        return run_evals

    def _get_user_evals(
        self, validated_data, organization, dataset_id, search_text, user_evals=None
    ):
        from model_hub.utils.eval_list import build_user_eval_list_items

        # When `user_evals` is provided the caller is in experiment scope
        # (evals keyed via experiment.user_eval_template_ids). Runtime state
        # for experiment evals lives on the EXPERIMENT_EVALUATION column, not
        # on UserEvalMetric.status (which stores a marker value).
        is_experiment_scope = user_evals is not None

        if user_evals is None:
            user_evals = UserEvalMetric.objects.select_related(
                "template", "eval_group"
            ).filter(
                dataset_id=dataset_id,
                show_in_sidebar=True,
                organization=organization,
                deleted=False,
                template__deleted=False,
                template__visible_ui=True,
            )
        else:
            # Filter by show_in_sidebar even when user_evals is provided
            user_evals = user_evals.filter(
                show_in_sidebar=True, template__visible_ui=True
            )

        if search_text:
            user_evals = user_evals.filter(name__icontains=search_text)

        if validated_data.get("eval_tags"):
            user_evals = user_evals.filter(
                eval_tags__overlap=validated_data.get("eval_tags")
            )

        return build_user_eval_list_items(
            user_evals, is_experiment_scope=is_experiment_scope
        )

    def _custom_build_evals(self, validated_data, organization, search_text):
        from model_hub.utils.eval_list import (
            derive_eval_type,
            derive_output_type,
            get_created_by_name,
        )

        eval_templates = EvalTemplate.objects.filter(
            organization=organization,
            owner=OwnerChoices.USER.value,
            deleted=False,
            visible_ui=True,
        ).prefetch_related("evaluators__user", "versions__created_by")

        if search_text:
            eval_templates = eval_templates.filter(Q(name__icontains=search_text))

        if validated_data.get("eval_tags"):
            eval_templates = eval_templates.filter(
                eval_tags__overlap=validated_data.get("eval_tags")
            )

        run_evals = []
        # For user evals we need the full model instance (not .values())
        # because get_created_by_name walks the `evaluators` / `versions`
        # relations to resolve the creator. The eval set is org-scoped so
        # the row count stays small.
        for template in eval_templates:
            is_model_present = False
            if (template.config or {}).get("config", {}).get("model"):
                is_model_present = True
            if template.name in NOT_UI_EVALS:
                continue

            # Real eval_type / output_type — prefer the dedicated fields,
            # fall back to tag-based derivation for legacy rows.
            _eval_type = template.eval_type or derive_eval_type(template)
            _output_type = template.output_type_normalized or derive_output_type(
                template
            )

            run_evals.append(
                {
                    "id": str(template.id),
                    "name": template.name,
                    "eval_template_name": template.name,
                    "eval_required_keys": (template.config or {}).get(
                        "required_keys", []
                    ),
                    "eval_template_tags": template.eval_tags,
                    "description": template.description,
                    "is_model_required": is_model_present,
                    "type": "user_built",
                    # Fields the picker frontend needs:
                    "eval_type": _eval_type,
                    "output_type": _output_type,
                    "created_by_name": get_created_by_name(template),
                    "owner": "user",
                    "updated_at": template.updated_at,
                    "created_at": template.created_at,
                    "model": template.model or "",
                    "template_type": template.template_type or "single",
                }
            )

        return run_evals

    def _get_previously_configured_evals(
        self, validated_data, organization, search_text
    ):
        eval_templates = EvalTemplate.objects.filter(
            organization=organization, deleted=False, visible_ui=True
        )
        if search_text:
            eval_templates = eval_templates.filter(Q(name__icontains=search_text))

        if validated_data.get("eval_tags"):
            eval_templates = eval_templates.filter(
                eval_tags__overlap=validated_data.get("eval_tags")
            )

        run_evals = []
        eval_templates = list(
            eval_templates.values("id", "name", "config", "eval_tags", "description")
        )
        for template in eval_templates:
            is_model_present = False
            if template.get("config", {}).get("config", {}).get("model"):
                is_model_present = True
            run_evals.append(
                {
                    "id": str(template.get("id")),
                    "name": template.get("name"),
                    "eval_template_name": template.get("name"),
                    "eval_required_keys": template.get("config", {}).get(
                        "required_keys", []
                    ),
                    "eval_template_tags": template.get("eval_tags"),
                    "description": template.get("description"),
                    "is_model_required": is_model_present,
                }
            )

        return run_evals


class GetEvalConfigView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        try:
            eval_id = request.query_params.get(
                "eval_id", None
            ) or request.query_params.get("evalId", None)

            if not eval_id:
                return self._gm.bad_request(get_error_message("MISSSING_EVAL_IDS"))

            try:
                template = EvalTemplate.no_workspace_objects.get(id=eval_id)
            except EvalTemplate.DoesNotExist:
                return self._gm.bad_request(get_error_message("MISSING_EVAL_TEMPLATE"))
            if template.owner == "user":
                eval_data = self._get_user_structure(
                    eval_id,
                    getattr(request, "organization", None) or request.user.organization,
                )
                eval_data.update({"owner": "user", "type": "user_built"})
                return self._gm.success_response(eval_data)
            else:
                eval_data = self._get_preset_structure(
                    eval_id,
                    getattr(request, "organization", None) or request.user.organization,
                )
                eval_data.update({"owner": "system", "type": "futureagi_built"})
                return self._gm.success_response(eval_data)

        except Exception as e:
            logger.exception(f"Error in fetching eval structure: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_EVAL_STRUCTURE")
            )

    def _get_preset_structure(self, template_id, organization):
        try:
            template = EvalTemplate.no_workspace_objects.get(id=template_id)

            final_config = template.config.get("config", {})
            function_params_schema, params = params_with_defaults_for_response(
                template.config, {}
            )
            if has_function_params_schema(template.config):
                final_config = {
                    key: value
                    for key, value in final_config.items()
                    if key not in function_params_schema
                }

            final_mapping = {}
            for key in template.config.get("required_keys", []):
                final_mapping[key] = ""

            eval_data = {
                "id": str(template.id),
                "template_id": str(template.id),
                "name": template.name,
                "description": template.description,
                "criteria": template.criteria,
                "eval_tags": template.eval_tags,
                "template_name": template.name,
                "required_keys": template.config.get("required_keys", []),
                "optional_keys": template.config.get("optional_keys", []),
                "variable_keys": template.config.get("variable_keys", []),
                "eval_type_id": template.config.get("eval_type_id", False),
                "function_eval": template.config.get("function_eval", False),
                "run_prompt_column": template.config.get("run_prompt_column", False),
                "mapping": final_mapping,
                "config": final_config,
                "params": params,
                "function_params_schema": function_params_schema,
                "model": template.model,
                "output": template.config.get("output", ""),
                "config_params_desc": template.config.get("config_params_desc", {}),
                "config_params_option": strip_turing_from_config_options(
                    template.config.get("config_params_option", {})
                ),
                "param_modalities": template.config.get("param_modalities", {}),
                "kb_id": None,
                "error_localizer": template.error_localizer_enabled,
                "api_key_available": (
                    True
                    if ApiKey.objects.filter(
                        organization=organization, provider="openai"
                    ).exists()
                    else False
                ),
            }

            return {"eval": eval_data}
        except Exception as e:
            logger.exception(f"Error in fetching eval structure: {str(e)}")
            raise e

    def _get_user_structure(self, eval_id, organization):
        try:
            template = get_object_or_404(
                EvalTemplate, id=eval_id, organization=organization
            )

            final_config = template.config.get("config", {})
            function_params_schema, params = params_with_defaults_for_response(
                template.config, {"params": template.config.get("params", {})}
            )
            if has_function_params_schema(template.config):
                final_config = {
                    key: value
                    for key, value in final_config.items()
                    if key not in function_params_schema
                }

            final_mapping = {}
            for key in template.config.get("required_keys", []):
                final_mapping[key] = template.config.get("mapping", {}).get(key, "")

            choices_dict = template.config.get("choices_map", None)
            if choices_dict:
                choices = list(choices_dict.keys())
                final_config["choices_map"] = choices_dict
            else:
                choices = (
                    ["Passed", "Failed"]
                    if template.config.get("output") == "Pass/Fail"
                    else []
                )
            eval_data = {
                "id": str(template.id),
                "template_id": str(template.id),
                "name": template.name,
                "reason_column": template.config.get("reason_column", False),
                "eval_type_id": template.config.get("eval_type_id", False),
                "function_eval": template.config.get("function_eval", False),
                "eval_tags": template.eval_tags,
                "description": template.description,
                "criteria": template.criteria,
                "model": template.model,
                "required_keys": template.config.get("required_keys", []),
                "optional_keys": template.config.get("optional_keys", []),
                "variable_keys": template.config.get("variable_keys", []),
                "run_prompt_column": template.config.get("run_prompt_column", False),
                "template_name": template.name,
                "mapping": final_mapping,
                "config": final_config,
                "params": params,
                "function_params_schema": function_params_schema,
                "output": template.config.get("output", ""),
                "config_params_desc": template.config.get("config_params_desc", {}),
                "config_params_option": strip_turing_from_config_options(
                    template.config.get("config_params_option", {})
                ),
                "param_modalities": template.config.get("param_modalities", {}),
                "choices": choices,
                "check_internet": template.config.get("check_internet", False),
            }

            return {"eval": eval_data}
        except Exception as e:
            logger.exception(f"Error in fetching eval structure: {str(e)}")
            raise e


class GetEvalStructureView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def get(
        self, request, eval_id, dataset_id=None, *args, **kwargs
    ):  # Changed from 'post' to 'get'
        try:
            eval_type = request.query_params.get(
                "eval_type"
            ) or request.query_params.get(
                "evalType"
            )  # Changed from request.data.get
            if not eval_type or eval_type not in [
                "preset",
                "user",
                "previously_configured",
            ]:
                return self._gm.bad_request(
                    get_error_message("INVALID_OR_MISSING_EVAL_TYPE", index=1)
                )

            if eval_type == "preset" or eval_type == "previously_configured":
                return self._get_preset_structure(
                    eval_id,
                    getattr(request, "organization", None) or request.user.organization,
                )
            else:  # user
                if not dataset_id:
                    return self._gm.bad_request(get_error_message("DATASET_ID_MISSING"))
                return self._get_user_structure(
                    eval_id,
                    dataset_id,
                    getattr(request, "organization", None) or request.user.organization,
                )

        except Exception as e:
            logger.exception(f"Error in fetching eval structure: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_EVAL_STRUCTURE")
            )

    def _get_preset_structure(self, template_id, organization):
        template = EvalTemplate.no_workspace_objects.get(id=template_id)

        final_config = template.config.get("config", {})
        function_params_schema, params = params_with_defaults_for_response(
            template.config, {}
        )
        if has_function_params_schema(template.config):
            final_config = {
                key: value
                for key, value in final_config.items()
                if key not in function_params_schema
            }
        eval_type_id = template.config.get("eval_type_id")

        final_mapping = {}
        for key in template.config.get("required_keys", []):
            final_mapping[key] = ""

        eval_data = {
            "id": str(template.id),
            "template_id": str(template.id),
            "name": template.name,
            "description": template.description,
            "eval_tags": template.eval_tags,
            "template_name": template.name,
            "required_keys": template.config.get("required_keys", []),
            "optional_keys": template.config.get("optional_keys", []),
            "variable_keys": template.config.get("variable_keys", []),
            "run_prompt_column": template.config.get("run_prompt_column", False),
            "mapping": final_mapping,
            "config": final_config,
            "params": params,
            "function_params_schema": function_params_schema,
            "eval_type_id": eval_type_id,
            "models": template.config.get("models", ""),
            "output": template.config.get("output", ""),
            "config_params_desc": template.config.get("config_params_desc", {}),
            "config_params_option": strip_turing_from_config_options(
                    template.config.get("config_params_option", {})
                ),
            "kb_id": None,
            "error_localizer": template.error_localizer_enabled,
            "choices": template.choices,
            "api_key_available": (
                True
                if ApiKey.objects.filter(
                    organization=organization, provider="openai"
                ).exists()
                else False
            ),
        }

        return self._gm.success_response({"eval": eval_data})

    def _get_user_structure(self, eval_id, dataset_id, organization):
        try:
            eval = get_object_or_404(
                UserEvalMetric,
                id=eval_id,
                dataset_id=dataset_id,
                organization=organization,
            )
        except Exception:
            return self._gm.bad_request(get_error_message("EVAL_STACK_UPDATED"))

        template = EvalTemplate.no_workspace_objects.get(id=eval.template_id)
        corresponding_column_id = Column.objects.filter(
            source_id=str(eval.id), deleted=False
        ).values_list("id", flat=True)
        if corresponding_column_id:
            corresponding_column_id = corresponding_column_id[0]

        final_config = template.config.get("config", {})
        normalized_runtime_config = normalize_eval_runtime_config(
            template.config,
            eval.config,
        )
        function_params_schema, params = params_with_defaults_for_response(
            template.config, normalized_runtime_config
        )
        if has_function_params_schema(template.config):
            final_config = {
                key: value
                for key, value in final_config.items()
                if key not in function_params_schema
            }
        for key in final_config:
            if key in eval.config.get("config", {}):
                final_config[key]["default"] = eval.config.get("config", {}).get(
                    key, ""
                )

        eval_type_id = template.config.get("eval_type_id")
        final_mapping = {}
        for key in template.config.get("required_keys", []):
            final_mapping[key] = eval.config.get("mapping", {}).get(key, "")

        eval_data = {
            "id": str(eval.id),
            "template_id": str(template.id),
            "name": eval.name,
            "eval_type_id": eval_type_id,
            "eval_type": template.eval_type,
            "reason_column": eval.config.get("reason_column", False),
            "eval_tags": template.eval_tags,
            "description": template.description,
            "required_keys": template.config.get("required_keys", []),
            "optional_keys": template.config.get("optional_keys", []),
            "variable_keys": template.config.get("variable_keys", []),
            "run_prompt_column": template.config.get("run_prompt_column", False),
            "template_name": template.name,
            "mapping": final_mapping,
            "config": final_config,
            "params": params,
            "function_params_schema": function_params_schema,
            "models": template.config.get("models", ""),
            "selected_model": eval.model,
            "error_localizer": eval.error_localizer,
            "kb_id": eval.kb_id,
            "output": template.config.get("output", ""),
            "config_params_desc": template.config.get("config_params_desc", {}),
            "config_params_option": strip_turing_from_config_options(
                    template.config.get("config_params_option", {})
                ),
            "run_config": eval.config.get("run_config", {}),
        }

        return self._gm.success_response({"eval": eval_data})


class StartEvalsProcess(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            user_eval_ids = request.data.get("user_eval_ids", [])
            experiment_id = request.data.get("experiment_id")

            # Experiment evals are orchestrated by a Temporal workflow — delegate
            # to the experiment rerun-cells view so the workflow fires correctly
            # rather than duplicating its ~120 lines of state-reset + dispatch.
            if experiment_id:
                if not user_eval_ids:
                    return self._gm.bad_request(get_error_message("MISSSING_EVAL_IDS"))
                from model_hub.views.experiments import ExperimentRerunCellsV2View

                request._full_data = {
                    "user_eval_metric_ids": user_eval_ids,
                    "failed_only": request.data.get("failed_only", False),
                }
                return ExperimentRerunCellsV2View().post(
                    request, experiment_id=experiment_id
                )

            eval_metrics = list(
                UserEvalMetric.objects.filter(
                    id__in=user_eval_ids,
                    dataset_id=dataset_id,
                    organization=getattr(request, "organization", None)
                    or request.user.organization,
                    deleted=False,
                ).select_related("dataset")
            )

            if not user_eval_ids:
                return self._gm.bad_request(get_error_message("MISSSING_EVAL_IDS"))

            for metric in eval_metrics:
                if metric.column_deleted:
                    return self._gm.bad_request(
                        f"{get_error_message('COLUMN_DELETED')} {metric.name}"
                    )

            # Update status for all specified evals
            updated = UserEvalMetric.objects.filter(
                id__in=user_eval_ids,
                dataset_id=dataset_id,
                organization=getattr(request, "organization", None)
                or request.user.organization,
                deleted=False,
            ).update(status=StatusType.NOT_STARTED.value)

            if updated == 0:
                return self._gm.bad_request(get_error_message("EVALS_NOT_FOUND"))

            Cell.objects.filter(
                column__source_id__in=user_eval_ids, deleted=False
            ).update(status=CellStatus.RUNNING.value)
            dataset = eval_metrics[0].dataset
            existing_columns = {
                str(col.source_id): col
                for col in list(
                    Column.objects.filter(
                        dataset=dataset, source=SourceChoices.EVALUATION.value
                    )
                )
            }

            column_order = dataset.column_order or []
            column_config = dataset.column_config or {}
            column_order_changed = False

            for user_eval_metric in eval_metrics:
                data_type = infer_eval_result_column_data_type(user_eval_metric.template)

                source_id = str(user_eval_metric.id)
                column = existing_columns.get(str(source_id))
                if not column:
                    column = Column.objects.create(
                        source_id=source_id,
                        name=user_eval_metric.name,
                        data_type=data_type,
                        source=SourceChoices.EVALUATION.value,
                        dataset=user_eval_metric.dataset,
                    )

                    reason_source_id = f"{column.id}-sourceid-{user_eval_metric.id}"
                    reason_column = Column.objects.create(
                        name=f"{user_eval_metric.name}-reason",
                        data_type=DataTypeChoices.TEXT.value,
                        source=SourceChoices.EVALUATION_REASON.value,
                        dataset=user_eval_metric.dataset,
                        source_id=reason_source_id,
                    )
                    column_order.extend([str(column.id), str(reason_column.id)])
                    column_config[str(column.id)] = {
                        "is_visible": True,
                        "is_frozen": None,
                    }
                    column_config[str(reason_column.id)] = {
                        "is_visible": True,
                        "is_frozen": None,
                    }
                    column_order_changed = True
                else:
                    # Reconcile: ensure reason column is in column_order
                    # (may have been dropped by a concurrent column_order overwrite)
                    reason_col = Column.objects.filter(
                        dataset=dataset,
                        source=SourceChoices.EVALUATION_REASON.value,
                        source_id__startswith=f"{column.id}-sourceid-",
                        deleted=False,
                    ).first()
                    if reason_col and str(reason_col.id) not in column_order:
                        eval_idx = column_order.index(str(column.id))
                        column_order.insert(eval_idx + 1, str(reason_col.id))
                        column_order_changed = True

                Cell.objects.filter(column__source_id=source_id, deleted=False).update(
                    status=CellStatus.RUNNING.value
                )

            if column_order_changed:
                dataset.column_order = column_order
                dataset.column_config = column_config
                dataset.save(update_fields=["column_order", "column_config"])

            return self._gm.success_response(
                f"Successfully updated {updated} eval(s) status"
            )

        except Exception as e:
            logger.exception(f"Error in starting evaluation process: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_START_EVAL_PROCESS")
            )


class DeleteEvalsView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def delete(self, request, dataset_id, eval_id, *args, **kwargs):
        try:
            delete_column = request.data.get("delete_column", False)
            experiment_id = request.data.get("experiment_id")
            organization = (
                getattr(request, "organization", None) or request.user.organization
            )
            # Experiment-scoped evals live under source_id=experiment_id, not
            # dataset_id. Branch the lookup so experiment eval deletion doesn't
            # 404 against the dataset-scoped record.
            lookup_kwargs = {"id": eval_id, "organization": organization}
            if experiment_id:
                lookup_kwargs["source_id"] = str(experiment_id)
            else:
                lookup_kwargs["dataset_id"] = dataset_id
            eval_metric = get_object_or_404(UserEvalMetric, **lookup_kwargs)

            # Experiments must retain at least one eval — the creation flow
            # validates this via CreateExperimentSerializer; mirror it here
            # so the delete path can't leave an experiment in a state the
            # create flow wouldn't accept.
            if experiment_id:
                experiment = ExperimentsTable.objects.filter(
                    id=experiment_id
                ).first()
                if experiment and (
                    experiment.user_eval_template_ids.filter(deleted=False).count()
                    <= 1
                ):
                    return self._gm.bad_request(
                        "Cannot delete the last evaluation. "
                        "An experiment must have at least one evaluation."
                    )

            # Stop any in-flight eval runner before deleting columns
            if delete_column and eval_metric.status in (
                StatusType.RUNNING.value,
                StatusType.NOT_STARTED.value,
                StatusType.EXPERIMENT_EVALUATION.value,
            ):
                try:
                    from tfc.utils.distributed_state import evaluation_tracker

                    evaluation_tracker.request_cancel(
                        eval_metric.id, reason="eval_deleted"
                    )
                except Exception:
                    pass
                from model_hub.utils.eval_cell_status import mark_eval_cells_stopped

                mark_eval_cells_stopped(
                    eval_metric, reason="Evaluation deleted by user"
                )

            if delete_column:
                if experiment_id:
                    # Experiment evals: one EXPERIMENT_EVALUATION column per EDT
                    # with source_id "{edt_id}-{col_id}-sourceid-{metric_id}".
                    # The dataset-style source_id=eval_metric.id lookup finds
                    # nothing here, so match by suffix across per-EDT + tag +
                    # reason columns.
                    per_edt_cols = list(
                        Column.objects.filter(
                            Q(
                                source__in=[
                                    SourceChoices.EXPERIMENT_EVALUATION.value,
                                    SourceChoices.EXPERIMENT_EVALUATION_TAGS.value,
                                ],
                                source_id__endswith=f"-sourceid-{eval_metric.id}",
                            )
                            | Q(
                                source=SourceChoices.EVALUATION_REASON.value,
                                source_id__endswith=f"-sourceid-{eval_metric.id}",
                            ),
                            deleted=False,
                        )
                    )
                    if per_edt_cols:
                        col_ids = [c.id for c in per_edt_cols]
                        snapshot_dataset = per_edt_cols[0].dataset
                        Cell.objects.filter(
                            column_id__in=col_ids, deleted=False
                        ).update(deleted=True)
                        Column.objects.filter(id__in=col_ids).update(deleted=True)
                        if snapshot_dataset.column_order:
                            col_id_strs = {str(cid) for cid in col_ids}
                            snapshot_dataset.column_order = [
                                cid
                                for cid in snapshot_dataset.column_order
                                if cid not in col_id_strs
                            ]
                            snapshot_dataset.save(update_fields=["column_order"])
                else:
                    # Check if column exists before attempting deletion
                    column = Column.objects.filter(
                        source_id=eval_metric.id, deleted=False
                    ).first()
                    if column:
                        # Delete all cells associated with the column and its dependent columns
                        Cell.objects.filter(
                            Q(column=column)
                            | Q(
                                column__source_id__startswith=f"{column.id}-sourceid-"
                            ),
                            deleted=False,
                        ).update(deleted=True)

                        dataset = column.dataset

                        # Remove columns from column_order
                        if dataset.column_order:
                            # Get all columns to delete (including those with source_id starting with column.id)
                            columns_to_delete = Column.objects.filter(
                                Q(id=column.id)
                                | Q(source_id__startswith=f"{column.id}-sourceid-"),
                                deleted=False,
                            ).values_list("id", flat=True)

                            col_ids_to_remove = {str(c) for c in columns_to_delete}
                            new_column_order = [
                                col_id
                                for col_id in dataset.column_order
                                if col_id not in col_ids_to_remove
                            ]
                            Dataset.objects.filter(id=dataset.id).update(
                                column_order=new_column_order
                            )

                        # Update metrics BEFORE deleting columns — the
                        # lookup scopes by dataset via the Column row, which
                        # must still be visible (deleted=False) for
                        # BaseModelManager to find it.
                        metrics = UserEvalMetric.get_metrics_using_column(
                            getattr(request, "organization", None)
                            or request.user.organization.id,
                            column.id,
                        )
                        if metrics:
                            UserEvalMetric.objects.filter(
                                id__in=[m.id for m in metrics]
                            ).update(column_deleted=True)

                        # Delete all related columns
                        Column.objects.filter(
                            Q(id=column.id)
                            | Q(source_id__startswith=f"{column.id}-sourceid-"),
                            deleted=False,
                        ).update(deleted=True)

                # Delete the eval_metric itself when delete_column is True
                eval_metric.deleted = True
                eval_metric.save()
            else:
                # Only hide from sidebar if delete_column is False
                eval_metric.show_in_sidebar = False
                eval_metric.save()

            return self._gm.success_response("Eval deleted successfully")

        except Http404:
            return self._gm.not_found("Eval not found")
        except Exception as e:
            logger.exception(f"Error in deleting evaluation: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_DELETE_EVALUATION")
            )


class DeleteTemplateEvalsView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def delete(self, request, dataset_id, eval_id, *args, **kwargs):
        try:
            # Get the eval and verify ownership
            eval_metric = EvalTemplate.no_workspace_objects.get(
                id=eval_id,
                organization=getattr(request, "organization", None)
                or request.user.organization,
                deleted=False,
            )
            eval_metric.deleted = True
            eval_metric.save()
            return self._gm.success_response("Eval deleted successfully")

        except Exception as e:
            logger.exception(f"Error in deleting eval template: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_DELETE_EVAL_TEMP")
            )


class EditAndRunUserEvalView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    @staticmethod
    def _coerce_bool(value):
        if isinstance(value, str):
            return value.strip().lower() not in {"false", "0", "no", ""}
        return bool(value)

    def post(self, request, dataset_id, eval_id, *args, **kwargs):
        from tfc.ee_gates import turing_oss_gate_for_template

        gate = turing_oss_gate_for_template(
            request.data.get("model"),
            template_id=request.data.get("template_id"),
            eval_type=request.data.get("eval_type"),
        )
        if gate is not None:
            return gate

        try:
            run = request.data.get("run", False)
            save_as_template = request.data.get("save_as_template", False)
            experiment_id = request.data.get("experiment_id")
            organization = (
                getattr(request, "organization", None) or request.user.organization
            )
            # When editing an eval attached to an experiment, the UserEvalMetric
            # is keyed by (id, source_id=experiment_id) — not (id, dataset_id).
            # Fall back to the dataset-scoped lookup for the dataset edit flow.
            lookup_kwargs = {"id": eval_id, "organization": organization}
            if experiment_id:
                lookup_kwargs["source_id"] = str(experiment_id)
            else:
                lookup_kwargs["dataset_id"] = dataset_id
            eval_metric = get_object_or_404(UserEvalMetric, **lookup_kwargs)
            if eval_metric.column_deleted:
                return self._gm.bad_request(
                    f"{get_error_message('COLUMN_DELETED')} {eval_metric.name}"
                )

            if save_as_template:
                template = eval_metric.template

                # Validate name format
                from model_hub.utils.eval_validators import validate_eval_name

                try:
                    template_name = validate_eval_name(request.data.get("name", ""))
                except ValueError as e:
                    return self._gm.bad_request(str(e))

                if (
                    EvalTemplate.objects.filter(
                        name=template_name,
                        organization=getattr(request, "organization", None)
                        or request.user.organization,
                        deleted=False,
                    ).exists()
                    or EvalTemplate.no_workspace_objects.filter(
                        name=template_name,
                        owner=OwnerChoices.SYSTEM.value,
                        deleted=False,
                    ).exists()
                ):
                    return self._gm.bad_request(get_error_message("EVAL_NAME_EXISTS"))

                new_template = EvalTemplate(
                    name=template_name,
                    description=template.description,
                    config=template.config,
                    eval_tags=template.eval_tags,
                    organization=getattr(request, "organization", None)
                    or request.user.organization,
                    owner=OwnerChoices.USER.value,
                    criteria=template.criteria,
                    choices=template.choices,
                    multi_choice=template.multi_choice,
                    error_localizer_enabled=template.error_localizer_enabled,
                )
                new_config = template.config
                runtime_config = normalize_eval_runtime_config(
                    template.config, request.data.get("config", {})
                )
                input_config = runtime_config.get("config", {})
                input_params = runtime_config.get("params", {})
                for key in input_config:
                    if key in new_config.get("config", {}):
                        new_config["config"][key]["default"] = input_config[key]
                if has_function_params_schema(new_config):
                    for key, value in input_params.items():
                        if key in new_config.get("function_params_schema", {}):
                            new_config["function_params_schema"][key]["default"] = value
                new_template.config = new_config
                new_template.save()
                eval_metric.template_id = new_template.id
                eval_metric.save()

            # Update the config if provided in request
            new_config = request.data.get("config")
            if new_config:
                new_config = normalize_eval_runtime_config(
                    eval_metric.template.config, new_config
                )
                # Default reason_column to True if not specified by caller, so
                # editing an eval never silently strips the reason column.
                if "reason_column" not in new_config:
                    new_config["reason_column"] = True
                eval_metric.config = new_config
            eval_metric.kb_id = request.data.get("kb_id") or eval_metric.kb_id
            if "error_localizer" in request.data:
                eval_metric.error_localizer = self._coerce_bool(
                    request.data.get("error_localizer")
                )
            elif (
                isinstance(request.data.get("config"), dict)
                and request.data["config"]
                .get("run_config", {})
                .get("error_localizer_enabled")
                is not None
            ):
                # Fallback: some callers (e.g. EvalPickerDrawer non-workbench
                # path) nest the flag inside config.run_config instead of
                # surfacing it at the top level.
                eval_metric.error_localizer = bool(
                    request.data["config"]["run_config"]["error_localizer_enabled"]
                )
            eval_metric.model = request.data.get("model") or eval_metric.model

            # Reason-column reconciliation differs by scope:
            #  * dataset: exactly one EVALUATION column (source_id == eval_metric.id)
            #  * experiment: one EXPERIMENT_EVALUATION column per EDT, source_id
            #    "{edt_id}-{col_id}-sourceid-{metric_id}" — doing a single
            #    Column.objects.get(source_id=eval_metric.id) here crashes.
            if new_config.get("reason_column"):
                if experiment_id:
                    per_edt_cols = Column.objects.filter(
                        source__in=[
                            SourceChoices.EXPERIMENT_EVALUATION.value,
                            SourceChoices.EXPERIMENT_EVALUATION_TAGS.value,
                        ],
                        source_id__endswith=f"-sourceid-{eval_metric.id}",
                        deleted=False,
                    )
                    for col in per_edt_cols:
                        Column.objects.get_or_create(
                            name=f"{eval_metric.name}-{col.name}-reason",
                            data_type=DataTypeChoices.TEXT.value,
                            source=SourceChoices.EVALUATION_REASON.value,
                            dataset=col.dataset,
                            source_id=f"{col.id}-sourceid-{eval_metric.id}",
                        )
                    # Rebuild snapshot_dataset.column_order so new reason cols
                    # surface in the grid.
                    try:
                        from model_hub.views.experiments import (
                            _build_and_save_v2_column_order,
                        )

                        experiment = ExperimentsTable.objects.filter(
                            id=experiment_id
                        ).first()
                        if experiment and experiment.snapshot_dataset_id:
                            _build_and_save_v2_column_order(
                                experiment, experiment.snapshot_dataset
                            )
                    except Exception:
                        logger.exception(
                            "Failed to rebuild column_order after edit-eval"
                        )
                else:
                    column = Column.objects.filter(
                        source_id=eval_metric.id, deleted=False
                    ).first()
                    if not column:
                        # Column doesn't exist yet (eval was added with run=false
                        # and never run). Skip reason-column creation — it will
                        # be created when the eval actually runs for the first time.
                        pass
                    else:
                        reason_column, created = Column.objects.get_or_create(
                            name=f"{eval_metric.name}-reason",
                            data_type=DataTypeChoices.TEXT.value,
                            source=SourceChoices.EVALUATION_REASON.value,
                            dataset=eval_metric.dataset,
                            source_id=f"{column.id}-sourceid-{eval_metric.id}",
                        )
                        if created:
                            column_order = eval_metric.dataset.column_order
                            column_order.append(str(reason_column.id))
                            eval_metric.dataset.column_order = column_order
                            eval_metric.dataset.save()

            # Eval columns live under different source types depending on scope:
            # dataset evals → SourceChoices.EVALUATION
            # experiment evals → SourceChoices.EXPERIMENT_EVALUATION
            column_source = (
                SourceChoices.EXPERIMENT_EVALUATION.value
                if experiment_id
                else SourceChoices.EVALUATION.value
            )
            if experiment_id:
                corresponding_column_ids = list(
                    Column.objects.filter(
                        source=column_source,
                        source_id__endswith=f"-sourceid-{eval_id}",
                        deleted=False,
                    ).values_list("id", flat=True)
                )
            else:
                corresponding_column_ids = list(
                    Column.objects.filter(
                        source=column_source,
                        source_id=str(eval_id),
                        deleted=False,
                    ).values_list("id", flat=True)
                )
            eval_metric.replace_column_id = (
                corresponding_column_ids[0] if corresponding_column_ids else None
            )

            # Set status to NOT_STARTED
            if run:
                eval_metric.status = StatusType.NOT_STARTED.value
                # Reset eval cells + their reason cells across all matching
                # columns (one per EDT for experiments, one for datasets).
                if corresponding_column_ids:
                    Cell.objects.filter(
                        column_id__in=corresponding_column_ids, deleted=False
                    ).update(status=CellStatus.RUNNING.value)
                    reason_source_ids = [
                        f"{cid}-sourceid-{eval_metric.id}"
                        for cid in corresponding_column_ids
                    ]
                    Cell.objects.filter(
                        column__source_id__in=reason_source_ids, deleted=False
                    ).update(status=CellStatus.RUNNING.value)
                else:
                    # Dataset fallback: reset cells under the base source_id
                    Cell.objects.filter(
                        column__source_id=str(eval_id), deleted=False
                    ).update(status=CellStatus.RUNNING.value)

            eval_metric.save()
            return self._gm.success_response(
                "Column evaluation updated and queued for processing"
            )

        except Http404:
            return self._gm.not_found("Eval not found")
        except ValueError as e:
            return self._gm.bad_request(str(e))
        except Exception as e:
            logger.exception(
                f"Error in updating the evaluation and process it: {str(e)}"
            )
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_UPDATE_EVALUATION_AND_PROCESS")
            )


class AddUserEvalView(CreateAPIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    @staticmethod
    def _coerce_bool(value):
        if isinstance(value, str):
            return value.strip().lower() not in {"false", "0", "no", ""}
        return bool(value)

    def post(self, request, dataset_id, *args, **kwargs):
        from tfc.ee_gates import turing_oss_gate_for_template

        gate = turing_oss_gate_for_template(
            request.data.get("model"),
            template_id=request.data.get("template_id"),
            eval_type=request.data.get("eval_type"),
        )
        if gate is not None:
            return gate

        organization = (
            getattr(request, "organization", None) or request.user.organization
        )

        # Validate dataset exists and belongs to user's organization
        try:
            dataset = Dataset.objects.get(id=dataset_id)
            if dataset.organization_id != organization.id:
                return self._gm.bad_request(
                    "Dataset does not belong to your organization"
                )
        except Dataset.DoesNotExist:
            return self._gm.not_found("Dataset not found")

        serializer = UserEvalSerializer(data=request.data)
        run = request.data.get("run", False)
        save_as_template = request.data.get("save_as_template", False)
        if serializer.is_valid():
            validated_data = serializer.validated_data
            id1 = dataset_id

            template_id = validated_data.get("template_id")
            if save_as_template:
                try:
                    if (
                        EvalTemplate.objects.filter(
                            name=validated_data.get("name"),
                            organization=organization,
                            deleted=False,
                        ).exists()
                        or EvalTemplate.no_workspace_objects.filter(
                            name=validated_data.get("name"),
                            owner=OwnerChoices.SYSTEM.value,
                            deleted=False,
                        ).exists()
                    ):
                        return self._gm.bad_request(
                            get_error_message("EVAL_NAME_EXISTS")
                        )

                    from model_hub.utils.eval_validators import (
                        validate_eval_template_org_access,
                    )

                    template = validate_eval_template_org_access(
                        validated_data.get("template_id"), organization
                    )
                    new_template = EvalTemplate(
                        name=validated_data.get("name"),
                        description=template.description,
                        config=template.config,
                        eval_tags=template.eval_tags,
                        organization=getattr(request, "organization", None)
                        or request.user.organization,
                        owner=OwnerChoices.USER.value,
                        criteria=template.criteria,
                        choices=template.choices,
                        multi_choice=template.multi_choice,
                        error_localizer_enabled=template.error_localizer_enabled,
                    )
                    new_config = template.config
                    runtime_config = normalize_eval_runtime_config(
                        template.config, validated_data.get("config", {})
                    )
                    input_config = runtime_config.get("config", {})
                    input_params = runtime_config.get("params", {})
                    for key in input_config:
                        if key in new_config.get("config", {}):
                            new_config["config"][key]["default"] = input_config[key]
                    if has_function_params_schema(new_config):
                        for key, value in input_params.items():
                            if key in new_config.get("function_params_schema", {}):
                                new_config["function_params_schema"][key][
                                    "default"
                                ] = value
                    new_template.config = new_config
                    new_template.save()
                    template_id = new_template.id
                except ValidationError as e:
                    return self._gm.bad_request(", ".join(e.messages))
                except ValueError as e:
                    return self._gm.bad_request(str(e))
                except Exception:
                    logger.exception("Error creating template")
                    return self._gm.bad_request(
                        get_error_message("FAILED_TO_CREATE_TEMPLATE")
                    )

            kb_id = validated_data.get("kb_id", None)
            if UserEvalMetric.objects.filter(
                name=validated_data.get("name"),
                organization=organization,
                dataset_id=id1,
                deleted=False,
            ).exists():
                return self._gm.bad_request(get_error_message("EVAL_NAME_EXISTS"))

            from model_hub.utils.eval_validators import (
                validate_eval_template_org_access,
            )

            template = validate_eval_template_org_access(template_id, organization)
            # Inherit template-level enablement unless caller explicitly overrides.
            if "error_localizer" in request.data:
                error_localizer = self._coerce_bool(
                    request.data.get("error_localizer", False)
                )
            else:
                error_localizer = bool(
                    getattr(template, "error_localizer_enabled", False)
                )

            # Validate required mapping keys. System evals stay strict —
            # every required key must be mapped. Custom evals allow
            # partial mappings; the shared validator at run time decides
            # whether to fail (all empty) or run with a warning.
            from model_hub.utils.eval_validators import validate_required_key_mapping

            mapping = validated_data.get("config", {}).get("mapping", {})
            required_keys = (
                template.config.get("required_keys", [])
                if template.config and isinstance(template.config, dict)
                else []
            )
            is_user_custom_eval = bool(
                template.config and template.config.get("custom_eval", False)
            )
            if not is_user_custom_eval:
                missing_keys = validate_required_key_mapping(
                    mapping, required_keys
                )
                if missing_keys:
                    return self._gm.bad_request(
                        f"Missing required mapping keys: {', '.join(missing_keys)}"
                    )

            try:
                validated_data["config"] = normalize_eval_runtime_config(
                    template.config, validated_data.get("config", {})
                )
            except ValueError as e:
                return self._gm.bad_request(str(e))

            # Default reason_column to True for single-eval creation so the
            # paired EVALUATION_REASON column is always materialized (matches
            # the composite / eval-group flow which hardcodes True).
            if "reason_column" not in validated_data["config"]:
                validated_data["config"]["reason_column"] = True

            user_eval_metric = UserEvalMetric.objects.create(
                name=validated_data.get("name"),
                organization=organization,
                dataset_id=id1,
                template_id=template_id,
                config=validated_data.get("config"),
                status=(
                    StatusType.NOT_STARTED.value if run else StatusType.INACTIVE.value
                ),
                error_localizer=error_localizer,
                kb_id=kb_id,
                user=request.user,
                model=validated_data.get("model", ModelChoices.TURING_LARGE.value),
                composite_weight_overrides=validated_data.get(
                    "composite_weight_overrides"
                ),
            )
            if kb_id:
                try:
                    KnowledgeBaseFile.objects.get(
                        id=kb_id,
                        organization=getattr(request, "organization", None)
                        or request.user.organization,
                    )
                except KnowledgeBaseFile.DoesNotExist:
                    pass
            else:
                pass

            if run:
                with transaction.atomic():
                    # Use no_workspace_objects manager to avoid the outer join issue with select_for_update
                    dataset = Dataset.no_workspace_objects.select_for_update().get(
                        id=id1
                    )
                    data_type = infer_eval_result_column_data_type(template)

                    column = Column.objects.create(
                        name=validated_data.get("name"),
                        data_type=data_type,
                        source=SourceChoices.EVALUATION.value,
                        dataset=dataset,
                        source_id=user_eval_metric.id,
                    )
                    column_order = dataset.column_order
                    column_order.append(str(column.id))
                    dataset.column_order = column_order
                    dataset.save()
                    if validated_data.get("config").get("reason_column"):
                        reason_column, created = Column.objects.get_or_create(
                            name=f"{user_eval_metric.name}-reason",
                            data_type=DataTypeChoices.TEXT.value,
                            source=SourceChoices.EVALUATION_REASON.value,
                            dataset=dataset,
                            source_id=f"{column.id}-sourceid-{user_eval_metric.id}",
                        )
                        column_order.append(str(reason_column.id))
                        dataset.column_order = column_order
                        dataset.save()

            return self._gm.success_response("success")
        return self._gm.bad_request(parse_serialized_errors(serializer))


class StopUserEvalView(APIView):
    """POST /develops/<dataset_id>/stop_user_eval/<eval_id>/
    Stops a running evaluation by setting its status to Completed.

    Accepts optional experiment_id in the body. When present, the eval is
    looked up via source_id=experiment_id (experiment-scoped UserEvalMetric)
    and cells are updated across both base columns (source_id=eval_id) and
    per-EDT columns (source_id ending with `-sourceid-{eval_id}`).
    """

    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, dataset_id, eval_id, *args, **kwargs):
        try:
            experiment_id = request.data.get("experiment_id")
            organization = (
                getattr(request, "organization", None) or request.user.organization
            )
            lookup_kwargs = {"id": eval_id, "organization": organization}
            if experiment_id:
                lookup_kwargs["source_id"] = str(experiment_id)
            else:
                lookup_kwargs["dataset_id"] = dataset_id
            eval_metric = get_object_or_404(UserEvalMetric, **lookup_kwargs)

            if eval_metric.status in (
                StatusType.RUNNING.value,
                StatusType.NOT_STARTED.value,
                StatusType.EXPERIMENT_EVALUATION.value,
            ):
                # Signal any in-flight workers to stop at the next row.
                # The runner polls evaluation_tracker.should_cancel per row
                # (see model_hub/tasks/user_evaluation.py:144-147).
                try:
                    from tfc.utils.distributed_state import evaluation_tracker

                    evaluation_tracker.request_cancel(
                        eval_metric.id, reason="user_stopped"
                    )
                except Exception as cancel_err:
                    logger.warning(
                        f"Failed to request cancel for eval {eval_id}: {cancel_err}"
                    )

                # Mark the eval as Error so list/sidebar show the failure
                # state alongside the per-cell error reason.
                eval_metric.status = StatusType.ERROR.value
                eval_metric.save(update_fields=["status", "updated_at"])

                # Flip running cells (eval + paired reason) to ERROR so the
                # UI doesn't leave a loading skeleton next to a stopped eval.
                from model_hub.utils.eval_cell_status import (
                    mark_eval_cells_stopped,
                )

                mark_eval_cells_stopped(
                    eval_metric, reason="Evaluation stopped by user"
                )

                # Experiment-scoped stop: if this was the only eval keeping
                # the experiment in RUNNING, flip it to COMPLETED so the UI
                # reflects reality. See services/experiment_utils.py.
                if experiment_id:
                    from model_hub.services.experiment_utils import (
                        maybe_complete_experiment_after_eval_stop,
                    )

                    maybe_complete_experiment_after_eval_stop(experiment_id)

            return self._gm.success_response("User evaluation stopped")
        except Exception as e:
            logger.exception(f"Error stopping eval: {e}")
            return self._gm.bad_request(str(e))


class PreviewRunEvalView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            config = request.data.get("config")
            if not config:
                return self._gm.bad_request("config is required")
            if "mapping" not in config:
                return self._gm.bad_request("mapping is required in config")
            template_id = request.data.get("template_id")
            model = request.data.get("model", ModelChoices.TURING_LARGE.value)
            call_type = config["mapping"].get("call_type", None)
            sdk_uuid = request.data.get("sdk_uuid", None)

            protect = False
            # Get protect_flash parameter from request (defaults to False)
            protect_flash = request.data.get("protect_flash", False)
            is_only_eval = True
            if call_type:
                if call_type == "protect":
                    protect = True
                    is_only_eval = False
                elif call_type == "protect_flash":
                    protect_flash = True
                    is_only_eval = False
            # Get dataset and selected rows
            rows = Row.objects.filter(dataset__id=dataset_id, deleted=False).order_by(
                "order"
            )[:3]

            source = Dataset.objects.get(id=dataset_id).source

            # get eval template of this eval metric
            eval_template = EvalTemplate.no_workspace_objects.get(id=template_id)

            eval_class = globals().get(eval_template.config.get("eval_type_id"))

            responses = []
            # Run the evaluation and get the result

            mappings = config.get("mapping")

            self.futureagi_eval = (
                True
                if eval_template.config.get("eval_type_id") in FUTUREAGI_EVAL_TYPES
                else False
            )

            eval_id = eval_template.config.get("eval_type_id")

            runner = EvaluationRunner(
                eval_id,
                is_only_eval=is_only_eval,
                format_output=True,
                futureagi_eval=self.futureagi_eval,
                source=request.data.get("source", "dataset_evaluation"),
                source_id=template_id,
                protect=protect,
                protect_flash=protect_flash,
                sdk_uuid=sdk_uuid,
            )
            data_config = config.get("config")

            run_prompt_column = eval_template.config.get("run_prompt_column", False)

            if source != DatasetSourceChoices.SDK.value:
                rows = rows[:3]

            # Wrap function with OTel context propagation for thread safety
            wrapped_process_eval = wrap_for_thread(process_eval_for_single_row)

            futures = []
            with ThreadPoolExecutor(max_workers=10) as executor:
                for row in rows:
                    future = executor.submit(
                        wrapped_process_eval,
                        runner,
                        request.user,
                        row,
                        mappings,
                        data_config,
                        run_prompt_column,
                        eval_class,
                        eval_template,
                        self.futureagi_eval,
                        source,
                        dataset_id,
                        model,
                    )
                    futures.append(future)

                for future in futures:
                    response = future.result()
                    responses.append(response)

            return self._gm.success_response({"responses": responses})

        except Exception as e:
            logger.exception(f"Error in preview the evaluation: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_PREVIEW_EVAL")
            )


class GetProviderStatusView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    def get(self, request, *args, **kwargs):
        try:
            # Get all providers from choices
            providers = LiteLlmModelProvider.get_choices()

            # Get existing API keys for the organization
            existing_keys = ApiKey.objects.filter(
                organization=getattr(request, "organization", None)
                or request.user.organization
            )

            # Create a dictionary for quick lookup of actual keys
            existing_keys_dict = {
                key.provider: {"id": str(key.id), "masked_key": key.masked_actual_key}
                for key in existing_keys
            }

            # Format response with provider status
            provider_status = [
                {
                    "provider": provider[0],  # Get the provider value from choice tuple
                    "display_name": provider[
                        1
                    ],  # Get the display name from choice tuple
                    "has_key": provider[0] in existing_keys_dict,
                    "masked_key": (
                        existing_keys_dict.get(provider[0]).get("masked_key")
                        if existing_keys_dict.get(provider[0])
                        else None
                    ),
                    "logo_url": ProviderLogoUrls.get_url_by_provider(provider[0]),
                    "type": (
                        "json"
                        if any(
                            provider[0].startswith(json_provider)
                            for json_provider in PROVIDERS_WITH_JSON
                        )
                        else "text"
                    ),
                    "id": (
                        existing_keys_dict.get(provider[0]).get("id")
                        if existing_keys_dict.get(provider[0])
                        else None
                    ),
                }
                for provider in providers
            ]

            return self._gm.success_response({"providers": provider_status})

        except Exception as e:
            logger.exception(f"Error in fetching the provider's status: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_PROVIDER_STATUS")
            )


class GetHuggingFaceDatasetListView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    parser_classes = (MultiPartParser, FormParser, JSONParser)

    def post(self, request, *args, **kwargs):
        try:
            search_query = request.data.get("search_query", "")

            filter_params = request.data.get("filter_params", {})
            try:
                parsed_filters = parse_huggingface_filter_params(filter_params)
            except Exception:
                logger.exception("Error parsing huggingface filter params")
                parsed_filters = {}

            API_URL = "https://huggingface.co/datasets-json"

            try:
                if search_query:
                    # Search for datasets with user query
                    query_params = {"search": search_query, **parsed_filters}
                else:
                    query_params = {**parsed_filters}
                # headers = {"Authorization": f"Bearer {HUGGINGFACE_API_TOKEN}"}

                response = requests.get(API_URL, params=query_params, timeout=30)
                if response.status_code != 200:
                    logger.error(
                        f"huggingface: Error from huggingface api: {response.status_code}, {response.text}"
                    )
                    return self._gm.bad_request(
                        get_error_message("FAILED_TO_FETCH_DATASETS")
                    )

                datasets = response.json()["datasets"]

                formatted_datasets = [
                    {
                        "id": dataset["id"],
                        "name": dataset.get("name", dataset["id"]),
                        "downloads": dataset.get("downloads", 0),
                        "likes": dataset.get("likes", 0),
                        "author": dataset.get("author"),
                    }
                    for dataset in datasets
                ]

                return self._gm.success_response(
                    {
                        "message": "Datasets retrieved successfully",
                        "total_datasets": response.json()["numTotalItems"],
                        "datasets": formatted_datasets,
                    }
                )

            except Exception as e:
                logger.exception(
                    f"huggingface: Error in loading the dataset from huggingface: {str(e)}"
                )
                return self._gm.internal_server_error_response(
                    f"Failed to fetch datasets list: {get_error_message('FAILED_TO_LOAD_DATASET_FROM_HUGGINGFACE')}"
                )

        except Exception as e:
            logger.exception(
                f"huggingface: Error in loading the dataset from huggingface: {str(e)}"
            )
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_LOAD_DATASET_FROM_HUGGINGFACE")
            )


def parse_huggingface_filter_params(filter_params):
    parsed_filters = {}

    modalities_filters = filter_params.get("modalities", [])
    modalities_filters_str = ""
    if modalities_filters and len(modalities_filters) > 0:
        possible_filters = [
            "3d",
            "timeseries",
            "audio",
            "geospatial",
            "text",
            "image",
            "tabular",
            "video",
        ]
        query = None

        for filter in modalities_filters:
            if filter in possible_filters:
                if query is None:
                    query = f"modality:{filter}"
                else:
                    query = query + "," + f"modality:{filter}"

        if query is not None:
            if modalities_filters_str == "":
                modalities_filters_str = query
            else:
                modalities_filters_str = modalities_filters_str + "," + query

    format_filters = filter_params.get("format", None)
    parsed_format_str = ""
    if format_filters and len(format_filters) > 0:
        possible_filters = [
            "csv",
            "json",
            "parquet",
            "imagefolder",
            "audiofolder",
            "webdataset",
            "arrow",
            "text",
        ]
        query = None

        for filter in format_filters:
            if filter in possible_filters:
                if query is None:
                    query = f"format:{filter}"
                else:
                    query = query + "," + f"format:{filter}"

        if query is not None:
            if parsed_format_str == "":
                parsed_format_str = query
            else:
                parsed_format_str = parsed_format_str + "," + query

    language_filters = filter_params.get("language", None)
    parsed_language_str = ""
    if language_filters and len(language_filters) > 0:
        possible_filters = []
        query = None

        for filter in language_filters:
            if filter in possible_filters:
                if query is None:
                    query = f"language:{filter}"
                else:
                    query = query + "," + f"language:{filter}"

        if query is not None:
            if parsed_language_str == "":
                parsed_language_str = query
            else:
                parsed_language_str = parsed_language_str + "," + query

    libraries_filters = filter_params.get("libraries", None)
    parsed_libraries_str = ""
    if libraries_filters and len(libraries_filters) > 0:
        possible_filters = []
        query = None

        for filter in libraries_filters:
            if filter in possible_filters:
                if query is None:
                    query = f"library:{filter}"
                else:
                    query = query + "," + f"library:{filter}"

        if query is not None:
            if parsed_libraries_str == "":
                parsed_libraries_str = query
            else:
                parsed_libraries_str = parsed_libraries_str + "," + query

    license_filters = filter_params.get("license", None)
    task_filter = filter_params.get("task_categories", None)
    sort_filters = filter_params.get("sort", None)
    size_categories = filter_params.get("size_categories", None)

    if license_filters:
        parsed_filters["license"] = f"license:{license_filters}"
    if parsed_language_str:
        parsed_filters["language"] = parsed_language_str
    if size_categories:
        parsed_filters["size_categories"] = size_categories
    if task_filter:
        parsed_filters["task_categories"] = f"task_categories:{task_filter}"
    if parsed_format_str:
        parsed_filters["format"] = parsed_format_str
    if modalities_filters_str:
        parsed_filters["modality"] = modalities_filters_str

    parsed_filters["sort"] = sort_filters if sort_filters else "trending"

    page_number = filter_params.get("page_number", None)
    if page_number:
        parsed_filters["p"] = page_number
    parsed_filters["withCount"] = True
    return parsed_filters


class GetHuggingFaceDatasetDetailView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    parser_classes = (MultiPartParser, FormParser, JSONParser)

    def post(self, request, *args, **kwargs):
        try:
            dataset_id = request.data.get("dataset_id", "")

            try:
                if not dataset_id:
                    return self._gm.bad_request(get_error_message("DATASET_ID_MISSING"))
                # Search for datasets
                API_URL = f"https://huggingface.co/api/datasets?id={dataset_id}"

                headers = {"Authorization": f"Bearer {HUGGINGFACE_API_TOKEN}"}

                response = requests.get(API_URL, headers=headers, timeout=30)
                if response.status_code != 200:
                    return self._gm.bad_request(
                        f"{get_error_message('FAILED_TO_FETCH_DATASET')}: {response.status_code} {response.text}"
                    )

                dataset = response.json()[0]
                formatted_datasets = {
                    "id": dataset["id"],
                    "name": dataset.get("name", dataset["id"]),
                    "description": dataset.get("description", ""),
                    "downloads": dataset.get("downloads", 0),
                    "likes": dataset.get("likes", 0),
                    "tags": dataset.get("tags", []),
                    "author": dataset.get("author"),
                }

                return self._gm.success_response(
                    {
                        "message": "Dataset details retrieved successfully",
                        "dataset": formatted_datasets,
                    }
                )

            except IndexError:
                return self._gm.bad_request(get_error_message("DATASET_NOT_FOUND"))

            except Exception:
                return self._gm.bad_request(
                    f"{get_error_message('FAILED_TO_FETCH_DATASET')}"
                )

        except Exception as e:
            logger.exception(
                f"huggingface: Error in loading the dataset from huggingface: {str(e)}"
            )
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_LOAD_DATASET_FROM_HUGGINGFACE")
            )


class ExtractJsonColumnView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def _process_cell(self, cell, json_key):
        """Process a single cell and extract JSON value"""
        try:
            close_old_connections()
            if cell.value:
                # Parse the string as a Python literal
                python_obj = ast.literal_eval(cell.value)
                # Convert Python object to JSON
                json_data = json.dumps(python_obj)
                json_data = json.loads(json_data)

                # Parse back to Python object if needed
                # print(cell.value,"cell.value*****")
                # json_data = json.loads(cell.value.replace("None", "null").replace("'", '"'))
                # Handle nested keys using dot notation (e.g., "key1.key2")
                value = json_data
                for key in json_key.split("."):
                    value = value.get(key, None)
                return str(value) if value is not None else None
            return None
        except Exception as e:
            logger.error("Error in json extract _process_cell", str(e))
            return None
        finally:
            close_old_connections()

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            column_id = request.data.get("column_id")
            json_key = request.data.get("json_key")
            new_column_name = request.data.get("new_column_name")
            concurrency = request.data.get(
                "concurrency", 5
            )  # Default to 5 concurrent workers

            if not all([column_id, json_key]):
                return self._gm.bad_request(
                    get_error_message("MISSING_COLUMN_ID_AND_JSON_KEY")
                )

            # Validate max_workers
            try:
                concurrency = int(concurrency)
                if concurrency < 1:
                    return self._gm.bad_request(
                        get_error_message("CONCURRENCY_NOT_POSITIVE")
                    )
                if concurrency > 10:
                    return self._gm.bad_request(
                        get_error_message("CONCURRENCY_EXCEEDS_MAX")
                    )
            except ValueError:
                return self._gm.bad_request(get_error_message("CONCURRENCY_INVALID"))

            org = getattr(request, "organization", None) or request.user.organization
            dataset = get_object_or_404(
                Dataset,
                id=dataset_id,
                organization=org,
                deleted=False,
            )

            if Column.objects.filter(
                name=new_column_name, dataset=dataset, deleted=False
            ).exists():
                return self._gm.bad_request(get_error_message("COLUMN_NAME_EXISTS"))
            # Get source column and verify it exists
            source_column = get_object_or_404(
                Column, id=column_id, dataset=dataset, deleted=False
            )

            # Create new column
            new_column = Column.objects.create(
                name=new_column_name or f"{source_column.name}_{json_key}",
                data_type=DataTypeChoices.TEXT.value,
                source=SourceChoices.EXTRACTED_JSON.value,
                dataset=dataset,
                metadata={
                    "source_column_id": str(source_column.id),
                    "json_key": json_key,
                    "concurrency": concurrency,
                },
            )

            # Update dataset's column order and config
            column_order = dataset.column_order or []
            column_order.append(str(new_column.id))

            column_config = dataset.column_config or {}
            column_config[str(new_column.id)] = {"is_visible": True, "is_frozen": None}

            dataset.column_order = column_order
            dataset.column_config = column_config
            dataset.save()
            Column.objects.filter(id=new_column.id).update(
                status=StatusType.RUNNING.value
            )
            extract_json_async.delay(
                column_id, json_key, concurrency, dataset_id, new_column.id
            )

            return self._gm.success_response(
                {
                    "message": "New column created successfully",
                    "new_column_id": str(new_column.id),
                    "new_column_name": new_column.name,
                }
            )

        except Exception as e:
            # If an error occurs, clean up the partially created column
            try:
                if "new_column" in locals():
                    new_column.delete()
            except Exception:
                pass
            logger.exception(f"Error in creation of json column: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_CREATE_JSON_COLUMN")
            )


@temporal_activity(time_limit=3600, queue="tasks_l")
def extract_json_async(column_id, json_key, concurrency, dataset_id, new_column_id):
    view = ExtractJsonColumnView()
    # Process cells concurrently
    Column.objects.filter(id=new_column_id).update(status=StatusType.RUNNING.value)
    source_cells = Cell.objects.filter(column_id=column_id, deleted=False)
    new_cells = []
    total_processed = 0
    failed_cells = 0

    # Wrap function with OTel context propagation for thread safety
    wrapped_process_cell = wrap_for_thread(view._process_cell)

    with ThreadPoolExecutor(max_workers=concurrency) as executor:
        # Submit all cells for processing
        future_to_cell = {
            executor.submit(wrapped_process_cell, cell, json_key): cell
            for cell in source_cells
        }

        # Process results as they complete
        for future in as_completed(future_to_cell):
            cell = future_to_cell[future]
            try:
                value = future.result()
                new_cells.append(
                    Cell(
                        dataset_id=dataset_id,
                        column_id=new_column_id,
                        row=cell.row,
                        value=value,
                    )
                )
                total_processed += 1
            except Exception as e:
                failed_cells += 1
                logger.error(f"Error processing cell: {str(e)}")
                # Create a failed cell with error information
                new_cells.append(
                    Cell(
                        dataset_id=dataset_id,
                        column_id=new_column_id,
                        row=cell.row,
                        value=None,
                        value_infos=json.dumps({"reason": str(e)}),
                        status=CellStatus.ERROR.value,
                    )
                )

    # Bulk create all cells at once
    if new_cells:
        Cell.objects.bulk_create(new_cells)
    Column.objects.filter(id=new_column_id).update(status=StatusType.COMPLETED.value)
    # insert_embeddings_task.delay(column_ids=[str(new_column_id)])


class ClassifyColumnView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def _classify_cell(self, cell, labels, model):
        """Classify a single cell's content into one of the given labels"""
        try:
            close_old_connections()
            if not cell.value:
                return None

            prompt = (
                f"Classify the following text into exactly one of these labels: {', '.join(labels)}.\n\n"
                f"Text: {cell.value}\n\n"
                f"Strictly return only the label, nothing else."
            )

            run_prompt = RunPrompt(
                model=model,
                organization_id=cell.column.dataset.organization.id,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,  # Lower temperature for more consistent results
                frequency_penalty=0.0,  # Default frequency penalty
                presence_penalty=0.0,  # Default presence penalty
                max_tokens=4000,  # Default max tokens
                top_p=1.0,  # Default top_p
                response_format={},  # Default response format
                tool_choice=None,  # Default tool choice
                tools=[],  # Default tools
                output_format="string",  # Default output format
                workspace_id=(
                    cell.column.dataset.workspace.id
                    if cell.column.dataset.workspace
                    else None
                ),
            )

            classification, value_infos = run_prompt.litellm_response()
            # Clean up response and verify it's in labels
            classification = classification.strip().lower()
            return (
                (
                    classification
                    if classification in [label.lower() for label in labels]
                    else None
                ),
                value_infos,
            )

        except Exception as e:
            logger.exception(f"Error classifying cell: {str(e)}")
            return str(e), {"reason": str(e)}
        finally:
            close_old_connections()

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            column_id = request.data.get("column_id")
            labels = request.data.get("labels", [])
            model = request.data.get("language_model_id", "gpt-4o")
            concurrency = request.data.get("concurrency", 5)
            new_column_name = request.data.get("new_column_name")

            # Validation
            if not column_id or not labels:
                return self._gm.bad_request(
                    get_error_message("MISSING_COLUMN_ID_AND_LABELS")
                )

            if not isinstance(labels, list) or len(labels) < 2:
                return self._gm.bad_request(get_error_message("LABELS_LIST_NOT_VALID"))

            # Get source column
            source_column = get_object_or_404(
                Column, id=column_id, dataset_id=dataset_id, deleted=False
            )

            if Column.objects.filter(
                name=new_column_name, dataset=dataset_id, deleted=False
            ).exists():
                return self._gm.bad_request(get_error_message("COLUMN_NAME_EXISTS"))

            # Create new column
            new_column = Column.objects.create(
                name=new_column_name or f"{source_column.name}_classification",
                data_type=DataTypeChoices.TEXT.value,
                source=SourceChoices.CLASSIFICATION.value,
                dataset_id=dataset_id,
                metadata={
                    "classification_labels": labels,
                    "model_used": model,
                    "source_column_id": str(column_id),
                    "concurrency": concurrency,
                },
            )

            # Update dataset's column order and config
            dataset = source_column.dataset
            column_order = dataset.column_order or []
            column_order.append(str(new_column.id))

            column_config = dataset.column_config or {}
            column_config[str(new_column.id)] = {"is_visible": True, "is_frozen": None}

            dataset.column_order = column_order
            dataset.column_config = column_config
            dataset.save()
            Column.objects.filter(id=new_column.id).update(
                status=StatusType.RUNNING.value
            )
            classify_column_async.delay(
                column_id, labels, model, concurrency, dataset_id, new_column.id
            )

            return self._gm.success_response(
                {
                    "message": "Classification column created successfully",
                    "new_column_id": str(new_column.id),
                    "new_column_name": new_column.name,
                }
            )

        except Exception as e:
            # Cleanup on error
            try:
                if "new_column" in locals():
                    new_column.delete()
            except Exception:
                pass
            logger.exception(f"Error in creating the classification column: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_CREATE_CLASSIFY_COLUMN")
            )


@temporal_activity(time_limit=3600, queue="tasks_l")
def classify_column_async(
    column_id, labels, model, concurrency, dataset_id, new_column_id
):
    view = ClassifyColumnView()
    # Process cells concurrently
    source_cells = Cell.objects.filter(column_id=column_id, deleted=False)
    new_cells = []
    total_processed = 0
    failed_cells = 0
    Column.objects.filter(id=new_column_id).update(status=StatusType.RUNNING.value)

    # Wrap function with OTel context propagation for thread safety
    wrapped_classify_cell = wrap_for_thread(view._classify_cell)

    with ThreadPoolExecutor(max_workers=concurrency) as executor:
        future_to_cell = {
            executor.submit(wrapped_classify_cell, cell, labels, model): cell
            for cell in source_cells
        }

        for future in as_completed(future_to_cell):
            cell = future_to_cell[future]
            try:
                classification, value_infos = future.result()
                if value_infos and "reason" in value_infos:
                    failed_cells += 1
                    new_cells.append(
                        Cell(
                            dataset_id=dataset_id,
                            column_id=new_column_id,
                            row=cell.row,
                            value=None,
                            value_infos=json.dumps(value_infos),
                            status=CellStatus.ERROR.value,
                        )
                    )
                else:
                    new_cells.append(
                        Cell(
                            dataset_id=dataset_id,
                            column_id=new_column_id,
                            row=cell.row,
                            value=classification,
                            value_infos=json.dumps(value_infos if value_infos else {}),
                        )
                    )
                    total_processed += 1
            except Exception:
                failed_cells += 1

    # Bulk create all cells
    if new_cells:
        Cell.objects.bulk_create(new_cells)
    Column.objects.filter(id=new_column_id).update(status=StatusType.COMPLETED.value)
    # insert_embeddings_task.delay(column_ids=[str(new_column_id)])


class ExtractEntitiesView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def _extract_entities(self, cell, instruction, model):
        """Extract entities from a single cell's content based on instruction"""
        try:
            close_old_connections()
            if not cell.value:
                return None, None

            prompt = f"""
            You are an AI assistant tasked with extracting entities from a given text based on a specific instruction. Your goal is to accurately identify and list the entities that match the given criteria. Follow these steps carefully:

1. First, you will be given an instruction for entity extraction. This instruction will guide what kind of entities you need to identify. Pay close attention to it:

<instruction>
{instruction}
</instruction>

2. Next, you will be presented with a text to analyze. Read through this text carefully, keeping the extraction instruction in mind:

<text>
{cell.value}
</text>

3. Extract the entities from the text according to the given instruction. Be thorough and accurate in your extraction.

4. Format your response as a JSON array of strings. Each entity you extract should be a separate string within this array.

5. Your output should strictly adhere to this format:
['entity1', 'entity2', 'entity3', ...]

6. Do not include any explanations, comments, or additional information. Your response should consist solely of the JSON array.

7. If no entities are found that match the instruction, return an empty array: []

8. Before finalizing your response, double-check the following:
   - Have you extracted all relevant entities according to the instruction?
   - Is your output formatted correctly as a JSON array of strings?
   - Have you removed any explanations or additional text?


9. Output your final response, ensuring it meets all the above requirements.

Remember, accuracy and adherence to the specified format are crucial. Your task is complete once you have provided the correctly formatted JSON array of extracted entities.

            """

            run_prompt = RunPrompt(
                model=model,
                organization_id=cell.column.dataset.organization.id,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                frequency_penalty=0.0,
                presence_penalty=0.0,
                max_tokens=4000,
                top_p=1.0,
                response_format={},
                tool_choice=None,
                tools=[],
                output_format="array",
                workspace_id=(
                    cell.column.dataset.workspace.id
                    if cell.column.dataset.workspace
                    else None
                ),
            )

            entities, value_infos = run_prompt.litellm_response()

            # Ensure the response is a valid array
            if isinstance(entities, list):
                return json.dumps(entities), value_infos
            elif isinstance(entities, dict) and "entities" in entities:
                return json.dumps(entities["entities"]), value_infos
            else:
                return None, None

        except Exception as e:
            logger.exception(f"Error extracting entities: {str(e)}")
            return str(e), {"reason": str(e)}
        finally:
            close_old_connections()

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            column_id = request.data.get("column_id")
            instruction = request.data.get("instruction")
            model = request.data.get("language_model_id", "gpt-4")
            concurrency = request.data.get("concurrency", 5)
            new_column_name = request.data.get("new_column_name")

            # Validation
            if not all([column_id, instruction]):
                return self._gm.bad_request(
                    get_error_message("MISSING_COLUMN_ID_AND_INSTRUCTIONS")
                )

            # Get source column
            source_column = get_object_or_404(
                Column, id=column_id, dataset_id=dataset_id, deleted=False
            )

            if Column.objects.filter(
                name=new_column_name, dataset=dataset_id, deleted=False
            ).exists():
                return self._gm.bad_request(get_error_message("COLUMN_NAME_EXISTS"))

            # Create new column
            new_column = Column.objects.create(
                name=new_column_name or f"{source_column.name}_entities",
                data_type=DataTypeChoices.ARRAY.value,
                source=SourceChoices.EXTRACTED_ENTITIES.value,
                dataset_id=dataset_id,
                metadata={
                    "instruction": instruction,
                    "model_used": model,
                    "source_column_id": str(column_id),
                    "concurrency": concurrency,
                },
            )

            # Update dataset's column order and config
            dataset = source_column.dataset
            column_order = dataset.column_order or []
            column_order.append(str(new_column.id))

            column_config = dataset.column_config or {}
            column_config[str(new_column.id)] = {"is_visible": True, "is_frozen": None}

            dataset.column_order = column_order
            dataset.column_config = column_config
            dataset.save()
            Column.objects.filter(id=new_column.id).update(
                status=StatusType.RUNNING.value
            )

            extract_async.delay(
                column_id, instruction, model, concurrency, dataset_id, new_column.id
            )

            return self._gm.success_response(
                {
                    "message": "Entity extraction completed successfully",
                }
            )

        except Exception as e:
            # Cleanup on error
            try:
                if "new_column" in locals():
                    new_column.delete()
            except Exception:
                pass
            logger.exception(f"Error in extracting the entities: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_EXTRACT_ENTITY")
            )


@temporal_activity(time_limit=3600, queue="tasks_l")
def extract_async(
    source_column_id, instruction, model, concurrency, dataset_id, new_column_id
):
    # Process cells concurrently
    source_cells = Cell.objects.filter(column_id=source_column_id, deleted=False)
    new_cells = []
    Column.objects.filter(id=new_column_id).update(status=StatusType.RUNNING.value)
    total_processed = 0
    failed_cells = 0
    view = ExtractEntitiesView()

    # Wrap function with OTel context propagation for thread safety
    wrapped_extract_entities = wrap_for_thread(view._extract_entities)

    with ThreadPoolExecutor(max_workers=concurrency) as executor:
        future_to_cell = {
            executor.submit(wrapped_extract_entities, cell, instruction, model): cell
            for cell in source_cells
        }

        for future in as_completed(future_to_cell):
            cell = future_to_cell[future]
            try:
                entities, value_infos = future.result()
                if value_infos and "reason" in value_infos:
                    failed_cells += 1
                    new_cells.append(
                        Cell(
                            dataset_id=dataset_id,
                            column_id=new_column_id,
                            row=cell.row,
                            value=None,
                            value_infos=json.dumps(value_infos),
                            status=CellStatus.ERROR.value,
                        )
                    )
                else:
                    new_cells.append(
                        Cell(
                            dataset_id=dataset_id,
                            column_id=new_column_id,
                            row=cell.row,
                            value=entities,
                            value_infos=json.dumps(value_infos if value_infos else {}),
                        )
                    )
                    total_processed += 1
            except Exception as e:
                logger.error(f"Failed to process cell: {str(e)}")
                failed_cells += 1

    # Bulk create all cells
    if new_cells:
        Cell.objects.bulk_create(new_cells)

    Column.objects.filter(id=new_column_id).update(status=StatusType.COMPLETED.value)
    # insert_embeddings_task.delay(column_ids=[str(new_column_id)])


class AddApiColumnView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def _replace_variables(self, value, row):
        """Replace variables in a string with actual cell values"""
        if isinstance(value, str) and re.search(r"\{{.*?\}}", value):
            matches = re.findall(r"\{{(.*?)\}}", value)
            for match in matches:
                try:
                    cell = Cell.objects.get(column__id=match, row=row)
                    value = value.replace(
                        f"{{{{{match}}}}}",
                        str(cell.value) if cell.value is not None else "",
                    )
                except Exception as e:
                    logger.error(f"Error replacing variable: {str(e)}")
        return value

    def _make_api_call(self, cell, config):
        """Make API call for a single cell and return the result"""
        try:
            # Process parameters
            processed_params = {}
            for param_name, param_config in config.get("params", {}).items():
                if param_config["type"] == "PlainText":
                    processed_params[param_name] = param_config["value"]
                elif param_config["type"] == "Secret":
                    processed_params[param_name] = SecretModel.objects.get(
                        id=param_config["value"]
                    ).actual_key
                elif param_config["type"] == "Variable":
                    try:
                        param_cell = Cell.objects.get(
                            column__id=self._replace_variables(
                                param_config["value"], cell.row
                            ),
                            row=cell.row,
                        )
                        processed_params[param_name] = param_cell.value
                    except Exception as e:
                        logger.error(f"Error replacing variable: {str(e)}")

            # Process headers
            processed_headers = {}
            for header_name, header_config in config.get("headers", {}).items():
                if header_config["type"] == "PlainText":
                    processed_headers[header_name] = header_config["value"]
                elif header_config["type"] == "Secret":
                    # secret = get_object_or_404(Secret, id=header_config['value'])
                    processed_headers[header_name] = SecretModel.objects.get(
                        id=header_config["value"]
                    ).actual_key
                elif header_config["type"] == "Variable":
                    try:
                        header_cell = Cell.objects.get(
                            column__id=self._replace_variables(
                                header_config["value"], cell.row
                            ),
                            row=cell.row,
                        )
                        processed_headers[header_name] = header_cell.value
                    except Exception as e:
                        logger.error(f"Error replacing variable: {str(e)}")

            # Process body if it exists
            body = {}
            if "body" in config:
                for key, values in config["body"].items():
                    if not cell:
                        body[key] = values
                    else:
                        body[key] = self._replace_variables(values, cell.row)

            # Process URL
            url = config["url"]

            # Make the API call
            response = requests.request(
                method=config["method"],
                url=url,
                params=processed_params,
                headers=processed_headers,
                json=body if isinstance(body, dict) else None,
                data=body if isinstance(body, str) else None,
                timeout=30,
            )
            response.raise_for_status()

            # Process response based on output_type
            if config["output_type"] == "string":
                return str(response.text), {"response_status": response.status_code}
            elif config["output_type"] == "object":
                return response.json(), {"response_status": response.status_code}
            elif config["output_type"] == "array":
                return json.dumps(response.json()), {
                    "response_status": response.status_code
                }
            elif config["output_type"] == "number":
                return str(response.text), {"response_status": response.status_code}
            else:
                return response.text, {"response_status": response.status_code}

        except Exception as e:
            logger.exception(f"API call error: {str(e)}")
            return str(e), {"response_status": 400}

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            column_name = request.data.get("column_name")
            config = request.data.get("config")  # URL, method, params, headers, body
            concurrency = request.data.get("concurrency", 5)

            if not all([column_name, config]):
                return self._gm.bad_request(
                    get_error_message("MISSING_COLUMN_NAME_AND_CONFIG")
                )

            # Validate config
            required_fields = ["url", "method", "output_type"]
            if not all(field in config for field in required_fields):
                return self._gm.bad_request(
                    f"Config must include: {', '.join(required_fields)}"
                )

            if Column.objects.filter(
                name=column_name, dataset=dataset_id, deleted=False
            ).exists():
                return self._gm.bad_request(get_error_message("COLUMN_NAME_EXISTS"))

            # Create new column
            new_column = Column.objects.create(
                name=column_name,
                data_type=DataTypeChoices.TEXT.value,  # You might want to make this configurable
                source=SourceChoices.API_CALL.value,
                dataset_id=dataset_id,
                metadata={
                    "api_config": {
                        "url": config["url"],
                        "method": config["method"],
                        "output_type": config["output_type"],
                        "params": config.get("params", {}),
                        "headers": config.get("headers", {}),
                        "body": config.get("body", {}),
                    },
                    "concurrency": concurrency,
                },
            )

            # Update dataset's column order and config
            dataset = Dataset.objects.get(id=dataset_id)
            column_order = dataset.column_order or []
            column_order.append(str(new_column.id))

            column_config = dataset.column_config or {}
            column_config[str(new_column.id)] = {"is_visible": True, "is_frozen": None}

            dataset.column_order = column_order
            dataset.column_config = column_config
            dataset.save()
            Column.objects.filter(id=new_column.id).update(
                status=StatusType.RUNNING.value
            )
            add_api_column_async.delay(config, dataset_id, concurrency, new_column.id)

            return self._gm.success_response(
                {
                    "message": "API column created successfully",
                    "new_column_id": str(new_column.id),
                    "new_column_name": new_column.name,
                }
            )

        except Exception as e:
            # Clean up on error
            try:
                if "new_column" in locals():
                    new_column.delete()
            except Exception:
                pass
            logger.exception(f"Error in creating the api column: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_CREATE_API_COLUMN")
            )


@temporal_activity(time_limit=3600, queue="tasks_l")
def add_api_column_async(config, dataset_id, concurrency, new_column_id):
    view = AddApiColumnView()
    # Process all rows
    rows = Row.objects.filter(dataset_id=dataset_id, deleted=False)
    new_cells = []
    total_processed = 0
    failed_cells = 0
    Column.objects.filter(id=new_column_id).update(status=StatusType.RUNNING.value)

    # Wrap function with OTel context propagation for thread safety
    wrapped_make_api_call = wrap_for_thread(view._make_api_call)

    with ThreadPoolExecutor(max_workers=concurrency) as executor:
        # Create a cell for each row and submit API calls
        future_to_cell = {}
        for row in rows:
            cell = Cell(
                dataset_id=dataset_id, column_id=new_column_id, row=row, value=None
            )
            future_to_cell[executor.submit(wrapped_make_api_call, cell, config)] = cell

        # Process results as they complete
        for future in as_completed(future_to_cell):
            cell = future_to_cell[future]
            try:
                value, value_infos = future.result()
                cell.value = value
                cell.value_infos = json.dumps(value_infos) if value_infos else None
                new_cells.append(cell)
                total_processed += 1
            except Exception as e:
                failed_cells += 1
                logger.error(f"Error processing cell: {str(e)}")

    # Bulk create all cells
    if new_cells:
        Cell.objects.bulk_create(new_cells)
    Column.objects.filter(id=new_column_id).update(status=StatusType.COMPLETED.value)
    # insert_embeddings_task.delay(column_ids=[str(new_column_id)])


class ExecutePythonCodeView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def _execute_python_code(self, row, code):
        """Execute restricted Python code for a single row

        WARNING: This executes user-provided code. Use with caution and only
        for trusted users. Consider implementing proper sandboxing.
        """
        try:
            if contains_sql(code):
                return "Raw SQL queries are not allowed in Python code.", {
                    "reason": "Raw SQL queries are not allowed in Python code."
                }

            # Additional security checks
            dangerous_imports = [
                "os",
                "sys",
                "subprocess",
                "eval",
                "exec",
                "__import__",
            ]
            for dangerous in dangerous_imports:
                if dangerous in code:
                    return f"Import of '{dangerous}' is not allowed.", {
                        "reason": f"Import of '{dangerous}' is not allowed for security reasons."
                    }

            # Fetch cells for the row with column names
            cells = Cell.objects.filter(row=row, deleted=False).select_related("column")

            # Create kwargs from cell data
            kwargs = {cell.column.name: cell.value for cell in cells}

            # Restricted globals - only safe builtins
            safe_builtins = {
                "abs": abs,
                "all": all,
                "any": any,
                "bool": bool,
                "dict": dict,
                "enumerate": enumerate,
                "filter": filter,
                "float": float,
                "int": int,
                "len": len,
                "list": list,
                "map": map,
                "max": max,
                "min": min,
                "range": range,
                "round": round,
                "set": set,
                "sorted": sorted,
                "str": str,
                "sum": sum,
                "tuple": tuple,
                "zip": zip,
            }
            global_namespace = {"__builtins__": safe_builtins}
            local_namespace = {}

            # Execute the provided code with restricted globals
            # WARNING: This still has security implications and should be properly sandboxed
            # exec(code, global_namespace, local_namespace)  # nosec B102 - sandboxed execution

            # Validate presence of `main()` function
            if "main" not in local_namespace or not callable(local_namespace["main"]):
                raise ValueError("Code must define a callable 'main' function.")

            result = local_namespace["main"](**kwargs)

            return str(result), None

        except Exception as e:
            traceback.format_exc()
            return str(e), {"reason": str(e)}

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            code = request.data.get("code")
            new_column_name = request.data.get("new_column_name")
            concurrency = request.data.get("concurrency", 5)

            # Validation
            if not all([code]):
                return self._gm.bad_request(get_error_message("CODE_MISSING"))

            if Column.objects.filter(
                name=new_column_name, dataset=dataset_id, deleted=False
            ).exists():
                return self._gm.bad_request(get_error_message("COLUMN_NAME_EXISTS"))

            # Create new column
            new_column = Column.objects.create(
                name=new_column_name if new_column_name else "Python Code Output",
                data_type=DataTypeChoices.TEXT.value,
                source=SourceChoices.PYTHON_CODE.value,
                dataset_id=dataset_id,
                metadata={"python_code": code, "concurrency": concurrency},
            )

            # Update dataset's column order and config
            dataset = Dataset.objects.get(id=dataset_id)
            column_order = dataset.column_order or []
            column_order.append(str(new_column.id))

            column_config = dataset.column_config or {}
            column_config[str(new_column.id)] = {"is_visible": True, "is_frozen": None}

            dataset.column_order = column_order
            dataset.column_config = column_config
            dataset.save()
            Column.objects.filter(id=new_column.id).update(
                status=StatusType.RUNNING.value
            )
            execute_python_code_async.delay(
                code, dataset_id, concurrency, new_column.id
            )

            return self._gm.success_response(
                {
                    "message": "Python code execution completed successfully",
                    "new_column_id": str(new_column.id),
                    "new_column_name": new_column.name,
                }
            )

        except Exception as e:
            # Cleanup on error
            try:
                if "new_column" in locals():
                    new_column.delete()
            except Exception:
                pass
            logger.exception(f"Error in execution of code: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_EXECUTE_CODE")
            )


@temporal_activity(time_limit=3600, queue="tasks_s")
def execute_python_code_async(code, dataset_id, concurrency, new_column_id):
    view = ExecutePythonCodeView()
    # Process rows concurrently
    rows = Row.objects.filter(dataset_id=dataset_id, deleted=False)
    new_cells = []
    total_processed = 0
    failed_cells = 0
    Column.objects.filter(id=new_column_id).update(status=StatusType.RUNNING.value)

    # Wrap function with OTel context propagation for thread safety
    wrapped_execute_python_code = wrap_for_thread(view._execute_python_code)

    with ThreadPoolExecutor(max_workers=concurrency) as executor:
        future_to_row = {
            executor.submit(wrapped_execute_python_code, row, code): row for row in rows
        }

        for future in as_completed(future_to_row):
            row = future_to_row[future]
            try:
                value, value_infos = future.result()
                if value_infos and "reason" in value_infos:
                    # Handle case where function returned error
                    new_cells.append(
                        Cell(
                            dataset_id=dataset_id,
                            column_id=new_column_id,
                            row=row,
                            value=None,
                            value_infos=json.dumps(value_infos),
                            status=CellStatus.ERROR.value,
                        )
                    )
                    failed_cells += 1
                else:
                    new_cells.append(
                        Cell(
                            dataset_id=dataset_id,
                            column_id=new_column_id,
                            row=row,
                            value=value,
                            value_infos=json.dumps(value_infos if value_infos else {}),
                        )
                    )
                    total_processed += 1
            except Exception as e:
                logger.exception(f"Failed to process row: {str(e)}")
                failed_cells += 1
                # Create a failed cell with error information
                new_cells.append(
                    Cell(
                        dataset_id=dataset_id,
                        column_id=new_column_id,
                        row=row,
                        value=None,
                        value_infos=json.dumps({"reason": str(e)}),
                        status=CellStatus.ERROR.value,
                    )
                )

    # Bulk create all cells
    if new_cells:
        Cell.objects.bulk_create(new_cells)
    Column.objects.filter(id=new_column_id).update(status=StatusType.COMPLETED.value)
    # insert_embeddings_task.delay(column_ids=[str(new_column_id)])


class ConditionalColumnView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def _evaluate_condition(self, condition, row, org_id=None):
        """Evaluate a condition for a given row"""
        try:
            # Replace variables in condition with actual values
            if isinstance(condition, str) and re.search(r"\{{.*?\}}", condition):
                matches = re.findall(r"\{{(.*?)\}}", condition)
                for match in matches:
                    try:
                        cell = Cell.objects.get(column__id=match, row=row)
                        value = cell.value
                        if value:
                            try:
                                json_data = json.loads(value)
                                value = json.dumps(
                                    json_data
                                )  # Keep JSON data as string for evaluation
                            except json.JSONDecodeError:
                                value = str(value).strip().lower()
                        condition = condition.replace(
                            f"{{{{{match}}}}}",
                            str(cell.value) if cell.value is not None else "",
                        )
                    except Exception as e:
                        logger.error(f"Error replacing variable in condition: {str(e)}")
                        return False

            prompt = f"""
            You are an AI expert that has to evaluate if the given condition according to provided information is True or False . You will be given condition just strictly return True or False according to the condition.
            condition: {condition}
            """
            try:
                run_prompt = RunPrompt(
                    model="gpt-4o-mini",
                    organization_id=row.dataset.organization.id,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3,  # Lower temperature for more consistent results
                    frequency_penalty=0.0,  # Default frequency penalty
                    presence_penalty=0.0,  # Default presence penalty
                    max_tokens=4000,  # Default max tokens
                    top_p=1.0,  # Default top_p
                    response_format={},  # Default response format
                    tool_choice=None,  # Default tool choice
                    tools=[],  # Default tools
                    output_format="string",  # Default output format
                    workspace_id=(
                        row.dataset.workspace.id if row.dataset.workspace else None
                    ),
                )

            except Exception:
                run_prompt = RunPrompt(
                    model="gpt-4o-mini",
                    organization_id=org_id,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3,  # Lower temperature for more consistent results
                    frequency_penalty=0.0,  # Default frequency penalty
                    presence_penalty=0.0,  # Default presence penalty
                    max_tokens=4000,  # Default max tokens
                    top_p=1.0,  # Default top_p
                    response_format={},  # Default response format
                    tool_choice=None,  # Default tool choice
                    tools=[],  # Default tools
                    output_format="string",  # Default output format
                    # workspace_id not available in fallback case
                )

            condition_status, value_infos = run_prompt.litellm_response()

            condition_status = True if condition_status == "True" else False
            return condition_status

        except Exception as e:
            logger.error(f"Error evaluating condition: {str(e)}")
            logger.error("traceback : ", traceback.format_exc())
            return False

    def _process_branch(self, row, branch_config, org_id=None):
        """Process a single branch configuration"""
        try:
            node_config = branch_config.get("branch_node_config", {})
            output_type = node_config.get("type")
            config = node_config.get("config", {})

            logger.exception(f"CONFIG:  {config}")

            # Get source column if specified
            source_column = None
            if config.get("column_id"):
                source_column = Column.objects.get(id=config["column_id"])

            if output_type == "static_value":
                return config.get("value"), None

            elif output_type == "column_value":
                if not config.get("column_id"):
                    return None, None
                cell = Cell.objects.get(column_id=config["column_id"], row=row)
                return cell.value, None

            elif output_type == "classification":
                if not source_column:
                    return None, None
                classifier = ClassifyColumnView()
                return classifier._classify_cell(
                    cell=Cell.objects.get(column=source_column, row=row),
                    labels=config.get("labels", []),
                    model=config.get("language_model_id"),
                )

            elif output_type == "extract_entities":
                if not source_column:
                    return None, None
                extractor = ExtractEntitiesView()
                return extractor._extract_entities(
                    cell=Cell.objects.get(column=source_column, row=row),
                    instruction=config.get("instruction"),
                    model=config.get("language_model_id"),
                )

            elif output_type == "extract_json":
                if not source_column:
                    return None, None
                json_extractor = ExtractJsonColumnView()
                result = json_extractor._process_cell(
                    cell=Cell.objects.get(column=source_column, row=row),
                    json_key=config.get("json_key"),
                )
                return result, None

            elif output_type == "extract_code":
                executor = ExecutePythonCodeView()
                return executor._execute_python_code(row, config.get("code"))

            elif output_type == "api_call":
                executor = AddApiColumnView()
                if not source_column:
                    return executor._make_api_call(cell=None, config=config)
                return executor._make_api_call(
                    cell=Cell.objects.get(column=source_column, row=row), config=config
                )

            elif output_type == "run_prompt":
                executor = RunPrompt(
                    model=config.get("model"),
                    organization_id=org_id,
                    messages=config.get("messages"),
                    temperature=config.get(
                        "temperature"
                    ),  # Lower temperature for more consistent results
                    frequency_penalty=config.get(
                        "frequency_penalty"
                    ),  # Default frequency penalty
                    presence_penalty=config.get(
                        "presence_penalty"
                    ),  # Default presence penalty
                    max_tokens=config.get("max_tokens"),  # Default max tokens
                    top_p=config.get("top_p"),  # Default top_p
                    response_format=config.get(
                        "response_format"
                    ),  # Default response format
                    tool_choice=config.get("tool_choice"),  # Default tool choice
                    tools=config.get("tools"),  # Default tools
                    output_format=config.get("output_format"),  # Default output format
                    workspace_id=(
                        row.dataset.workspace.id if row.dataset.workspace else None
                    ),
                )
                return executor.litellm_response()

            elif output_type == "retrieval":
                executor = AddVectorDBColumnView()
                return executor._process_row(row, config, org_id)

            else:
                return None, None

        except Exception as e:
            logger.exception(f"Error processing branch: {str(e)}")
            return None, None

    def _process_row(self, row, config, org_id=None):
        """Process a single row through all conditions"""
        try:
            final_value = None
            final_value_infos = None

            # Track if any condition has been met
            condition_met = False

            logger.exception(f"CONFIG_FIRST: {config}")

            for branch in config:
                if not isinstance(branch, dict):
                    logger.error(f"Invalid branch format: {branch}")
                    continue

                branch_type = branch.get("branch_type", "").lower()
                condition = branch.get("condition", "")

                should_execute = False

                if branch_type == "if":
                    condition_met = self._evaluate_condition(condition, row, org_id)
                    should_execute = condition_met

                elif branch_type == "elif":
                    if (
                        not condition_met
                    ):  # Only check elif if no previous condition was true
                        condition_met = self._evaluate_condition(condition, row, org_id)
                        should_execute = condition_met

                elif branch_type == "else":
                    should_execute = not condition_met

                if should_execute:
                    value, value_infos = self._process_branch(
                        row, branch, org_id=org_id
                    )
                    if value is not None:
                        final_value = value
                        final_value_infos = value_infos
                        break  # Exit after first matching condition

            return final_value, final_value_infos

        except Exception as e:
            logger.error("traceback : ", traceback.format_exc())
            logger.error(f"Error processing row: {str(e)}")
            return str(e), {"reason": str(e)}

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            config = request.data.get("config", [])
            new_column_name = request.data.get("new_column_name")
            concurrency = request.data.get("concurrency", 5)
            self.organization_id = (
                getattr(request, "organization", None) or request.user.organization.id
            )

            if not config:
                return self._gm.bad_request(get_error_message("CONFIG_MISSING"))
            if not new_column_name:
                return self._gm.bad_request(
                    get_error_message("NEW_COLUMN_NAME_MISSING")
                )

            dataset = Dataset.objects.get(id=dataset_id)

            if Column.objects.filter(
                name=new_column_name, dataset=dataset_id, deleted=False
            ).exists():
                return self._gm.bad_request(get_error_message("COLUMN_NAME_EXISTS"))

            # Create new column
            new_column = Column.objects.create(
                name=new_column_name,
                data_type=DataTypeChoices.TEXT.value,
                source=SourceChoices.CONDITIONAL.value,
                dataset_id=dataset_id,
                metadata={"conditional_config": config, "concurrency": concurrency},
            )

            # Update dataset configuration
            column_order = dataset.column_order or []
            column_order.append(str(new_column.id))

            column_config = dataset.column_config or {}
            column_config[str(new_column.id)] = {"is_visible": True, "is_frozen": None}

            dataset.column_order = column_order
            dataset.column_config = column_config
            dataset.save()
            Column.objects.filter(id=new_column.id).update(
                status=StatusType.RUNNING.value
            )

            conditional_column_async.delay(
                config, dataset_id, concurrency, new_column.id
            )

            return self._gm.success_response(
                {
                    "message": "Conditional column created successfully",
                    "new_column_id": str(new_column.id),
                    "new_column_name": new_column.name,
                }
            )

        except Exception as e:
            try:
                if "new_column" in locals():
                    new_column.delete()
            except Exception:
                pass
            logger.exception(f"Error in creating the conditional column: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_CRETE_CONDITIONAL_COLUMN")
            )


@temporal_activity(time_limit=3600, queue="tasks_l")
def conditional_column_async(config, dataset_id, concurrency, new_column_id):
    view = ConditionalColumnView()
    # Process rows
    rows = Row.objects.filter(dataset_id=dataset_id, deleted=False)
    organization_id = Dataset.objects.get(id=dataset_id).organization.id
    new_cells = []
    total_processed = 0
    failed_cells = 0
    Column.objects.filter(id=new_column_id).update(status=StatusType.RUNNING.value)

    # Wrap function with OTel context propagation for thread safety
    wrapped_process_row = wrap_for_thread(view._process_row)

    with ThreadPoolExecutor(max_workers=concurrency) as executor:
        future_to_row = {
            executor.submit(wrapped_process_row, row, config, organization_id): row
            for row in rows
        }

        for future in as_completed(future_to_row):
            row = future_to_row[future]
            try:
                value, value_infos = future.result()
                if value_infos and "reason" in value_infos:
                    new_cells.append(
                        Cell(
                            dataset_id=dataset_id,
                            column_id=new_column_id,
                            row=row,
                            value=None,
                            value_infos=json.dumps(value_infos),
                            status=CellStatus.ERROR.value,
                        )
                    )
                else:
                    new_cells.append(
                        Cell(
                            dataset_id=dataset_id,
                            column_id=new_column_id,
                            row=row,
                            value=value,
                            value_infos=json.dumps(value_infos if value_infos else {}),
                        )
                    )
                    total_processed += 1
            except Exception as e:
                logger.error("traceback : ", traceback.format_exc())
                logger.error(f"Error processing row: {str(e)}")
                failed_cells += 1
                # Create a failed cell with error information
                new_cells.append(
                    Cell(
                        dataset_id=dataset_id,
                        column_id=new_column_id,
                        row=row,
                        value=None,
                        value_infos=json.dumps({"reason": str(e)}),
                        status=CellStatus.ERROR.value,
                    )
                )

    if new_cells:
        Cell.objects.bulk_create(new_cells)
    Column.objects.filter(id=new_column_id).update(status=StatusType.COMPLETED.value)
    # insert_embeddings_task.delay(column_ids=[str(new_column_id)])


class AddVectorDBColumnView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    organization_id = None

    def _get_api_key_for_provider(self, organization_id, workspace_id, provider):
        """Get API key for a provider, filtering by workspace"""
        api_key_entry = ApiKey.objects.filter(
            organization_id=organization_id,
            workspace_id=workspace_id,
            provider=provider,
        ).first()
        if not api_key_entry:
            raise ValueError(
                f"API key not configured for {provider}. Please add your API key in settings."
            )
        return api_key_entry.actual_key

    def _query_vector_db(self, text_input, config, organization_id, workspace_id=None):
        """Query vector database using text input"""
        try:
            # 1. Set up the embedding model based on config
            embedding_config = config.get("embedding_config", {})
            embedding_type = embedding_config.get("type", "")
            if embedding_type == "openai":
                embedding_vector = model_manager.get_embeddings(
                    text_input,
                    "openai",
                    embedding_config.get("model", "text-embedding-3-small"),
                    model_params={
                        "api_key": self._get_api_key_for_provider(
                            organization_id, workspace_id, "openai"
                        )
                    },
                )
            elif embedding_type == "huggingface":
                embedding_vector = model_manager.get_embeddings(
                    text_input,
                    "huggingface",
                    embedding_config.get("model", "all-mpnet-base-v2"),
                    model_params={
                        "api_key": self._get_api_key_for_provider(
                            organization_id, workspace_id, "huggingface"
                        )
                    },
                )
            elif embedding_type == "sentence_transformers":
                embedding_vector = model_manager.get_embeddings(
                    text_input,
                    "sentence_transformers",
                    embedding_config.get("model", "all-mpnet-base-v2"),
                    model_params={
                        "api_key": self._get_api_key_for_provider(
                            organization_id, workspace_id, "huggingface"
                        )
                    },
                )
            else:
                raise ValueError(f"Unsupported embedding type: {embedding_type}")

            if isinstance(embedding_vector, str):
                embedding_vector = [embedding_vector]

            # Ensure the vector has 512 dimensions
            if isinstance(embedding_vector, list) or hasattr(
                embedding_vector, "tolist"
            ):
                embedding_vector = np.array(
                    embedding_vector
                )  # Convert to numpy array for processing
                embedding_vector = embedding_vector[
                    : config.get("vector_length", 512)
                ]  # Truncate to 512 dimensions
                embedding_vector = (
                    embedding_vector.tolist()
                )  # Convert back to list for Weaviate

                # Debug logging
                logger.info(
                    f"Generated embedding vector with {len(embedding_vector)} dimensions"
                )
                logger.info(f"First 5 values: {embedding_vector[:5]}")

            # 3. Initialize vector store based on type
            sub_type = config.get("sub_type")

            if sub_type == "pinecone":
                return self._query_pinecone(embedding_vector, config)

            elif sub_type == "qdrant":
                return self._query_qdrant(embedding_vector, config)

            elif sub_type == "weaviate":
                return self._query_weaviate(
                    embedding_vector, text_input, config, organization_id, workspace_id
                )

            else:
                raise ValueError(f"Unsupported vector database type: {sub_type}")

        except Exception as e:
            logger.exception("vector_embedding_database_error", sub_type=sub_type)
            return str(e)

    def _query_pinecone(self, query, config):
        """
                {
            "indexName": "apicall2",
            "namespace": "garvit",
            "topK": 2,
            "queryKey": "vector",
            "embeddingConfig": {
                "model": "text-embedding-3-small",
                "type": "openai"
            },
            "key": "text",
            "concurrency": 2,
            "vectorLength": 512
        }

        """
        pc = Pinecone(api_key=SecretModel.objects.get(id=config["api_key"]).actual_key)
        index = pc.Index(config["index_name"])
        query_object = {}

        # Validate that query is not None or empty
        if not query or (isinstance(query, list) and len(query) == 0):
            raise ValueError("Query vector is empty or None")

        # Ensure query is a list of numbers
        if not isinstance(query, list):
            raise ValueError(f"Query must be a list, got {type(query)}")

        # Validate that all elements are numbers
        if not all(isinstance(x, int | float) for x in query):
            raise ValueError("All elements in query vector must be numbers")

        query_object["vector"] = query
        query_object["top_k"] = config["top_k"]
        query_object["namespace"] = config.get("namespace", "default")
        # Add include_metadata=True to get metadata
        query_object["include_metadata"] = True
        # Add include_values=True if you want vector values
        query_object["include_values"] = True

        results = index.query(**query_object)
        if not results["matches"]:
            return "No matches found"

        metadata_list = []
        for result in results["matches"]:
            metadata = result.get("metadata", {})
            if config.get("key"):
                metadata = metadata.get(config["key"])
            metadata_list.append(metadata)

        return metadata_list

        # Return both ID and full match info including metadata
        return metadata

    def _query_qdrant(self, query, config):
        """Query Qdrant vector database

                {
            "subType": "qdrant",
            "newColumnName": "valuss",
            "columnId": "dddb8ffb-1de1-4b98-a489-de6dd33eb884",
            "apiKey": "your-qdrant-api-key",
            "topK": 2,
            "embeddingConfig": {
                "model": "text-embedding-3-small",
                "type": "openai"
            },
            "concurrency": 2,
            "url": "https://0763e883-fc79-4636-a8e3-cdeb47f024be.us-east-1-0.aws.cloud.qdrant.io:6333",
            "collectionName": "mid",
            "key": "name",
            "vectorLength": 512
        }

                Args:
                    query: Vector to search for
                    config: Dictionary containing:
                        - api_key: Qdrant API key
                        - url: Qdrant instance URL (including port)
                        - collection_name: Name of collection to search
                        - top_k: Number of results to return (default: 1)
                        - query_key: Key for vector query (default: 'vector')

                Returns:
                    Dictionary containing search results and metadata
        """
        try:
            # Validate required config parameters
            required_params = ["api_key", "url", "collection_name"]
            missing_params = [p for p in required_params if not config.get(p)]
            if missing_params:
                raise ValueError(
                    f"Missing required parameters: {', '.join(missing_params)}"
                )

            # Initialize Qdrant client
            client = QdrantClient(
                url=config["url"],
                api_key=SecretModel.objects.get(id=config["api_key"]).actual_key,
            )
            # query=[0.2, 0.1, 0.9, 0.7]

            # Build search query
            query_object = {
                "collection_name": config["collection_name"],
                "query_vector": query,  # Qdrant expects 'query_vector' instead of 'vector'
                "limit": config.get("top_k", 5),  # Default to 1 if not specified
                "with_payload": True,  # Always get metadata
                "with_vectors": False,  # Don't return vectors by default for efficiency
            }

            # Execute search
            results = client.search(**query_object)

            if not results:
                return "No matches found"

            metadata_list = []
            for result in results:
                metadata = result.payload
                if config.get("key"):
                    metadata = metadata.get(config.get("key"))
                metadata_list.append(metadata)
            return metadata_list

        except Exception as e:
            logger.error(f"Error querying Qdrant: {str(e)}")
            raise ValueError(f"Failed to query Qdrant: {str(e)}")  # noqa: B904

    def get_client(self, config, organization_id, workspace_id=None, use_hybrid=False):
        embedding_config = config.get("embedding_config", {})
        embedding_type = embedding_config.get("type", "")
        key = None
        if embedding_type:
            key = self._get_api_key_for_provider(
                organization_id, workspace_id, embedding_type
            )

        auth = AuthApiKey(
            api_key=SecretModel.objects.get(id=config["api_key"]).actual_key
        )
        connect_kwargs = {"auth_client_secret": auth}
        if key and use_hybrid:
            connect_kwargs["additional_headers"] = {"X-OpenAI-Api-Key": key}

        return weaviate.connect_to_wcs(
            cluster_url=config["url"],
            auth_credentials=auth,
            # additional_headers=additional_headers or None
        )

    def _query_weaviate(
        self, query, text_input, config, organization_id, workspace_id=None
    ):
        try:
            client = self.get_client(
                config,
                organization_id,
                workspace_id,
                config.get("search_type") == "hybrid",
            )

            search_type = config.get("search_type", "semantic_search")
            limit = config.get("top_k", 5)
            class_name = config["collection_name"]
            return_field = config.get("key")

            if search_type == "hybrid" and text_input:
                resp = (
                    client.query.get(class_name, [return_field])
                    .with_hybrid(query=text_input, alpha=0.5)
                    .with_limit(limit)
                    .do()
                )
            else:
                if not query:
                    raise ValueError("Vector query is required for semantic search")
                resp = (
                    client.query.get(class_name, [return_field])
                    .with_near_vector(
                        {"vector": query, "certainty": config.get("certainty", 0.0)}
                    )
                    .with_limit(limit)
                    .do()
                )

            data = resp.get("data", {})
            hits = data.get("Get", {}).get(class_name, [])
            if not hits:
                err = resp.get("errors", [{}])[0].get("message", "No results")
                return err

            results = [
                (
                    m.get(return_field)
                    if return_field
                    else {k: v for k, v in m.items() if k not in ("id", "_additional")}
                )
                for m in hits
            ]
            return results

        except Exception as e:
            logger.error("Error in _query_weaviate", exc_info=True)
            raise ValueError(f"Failed to query Weaviate: {e}")  # noqa: B904

    def _process_row(self, row, column, config, organization_id, workspace_id=None):
        try:
            input_cell = Cell.objects.get(row=row, column=column)
            query = input_cell.value
            result_info = self._query_vector_db(
                query, config, organization_id, workspace_id
            )
            return result_info, {}
        except Exception as e:
            logger.error("traceback : ", traceback.format_exc())
            logger.error(f"Error processing row: {str(e)}")
            return str(e), {"reason": str(e)}

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            config = request.data
            self.organization_id = (
                getattr(request, "organization", None) or request.user.organization.id
            )
            column_id = config.get("column_id")
            new_column_name = config.get("new_column_name", "Vector DB Result")
            concurrency = config.get("concurrency", 5)

            if not all([column_id, config.get("sub_type"), config.get("api_key")]):
                return self._gm.bad_request(
                    get_error_message("MISSING_COLUMN_ID_SUB_TYPE_AND_API_KEY")
                )

            input_column = get_object_or_404(
                Column, id=column_id, dataset_id=dataset_id
            )

            if Column.objects.filter(
                name=new_column_name, dataset=dataset_id, deleted=False
            ).exists():
                return self._gm.bad_request(get_error_message("COLUMN_NAME_EXISTS"))

            new_column = Column.objects.create(
                name=new_column_name,
                data_type=DataTypeChoices.ARRAY.value,
                source=SourceChoices.VECTOR_DB.value,
                dataset_id=dataset_id,
                metadata={
                    "vector_db_config": {
                        "sub_type": config["sub_type"],
                        "collection_name": config.get("collection_name"),
                        "url": config.get("url"),
                        "search_type": config.get("search_type"),
                        "key": config.get("key"),
                        "limit": config.get("limit", 1),
                        "index_name": config.get("index_name"),
                        "top_k": config.get("top_k", 1),
                        "namespace": config.get("namespace"),
                        "api_key": config.get("api_key"),
                        "embedding_config": config.get("embedding_config"),
                    },
                    "source_column_id": str(column_id),
                    "concurrency": concurrency,
                },
            )

            dataset = Dataset.objects.get(id=dataset_id)
            column_order = dataset.column_order or []
            column_order.append(str(new_column.id))

            column_config = dataset.column_config or {}
            column_config[str(new_column.id)] = {"is_visible": True, "is_frozen": None}

            dataset.column_order = column_order
            dataset.column_config = column_config
            dataset.save()

            # Ensure config is JSON serializable
            serializable_config = json.loads(json.dumps(config))
            Column.objects.filter(id=new_column.id).update(
                status=StatusType.RUNNING.value
            )

            add_vector_db_column_async.delay(
                serializable_config,
                dataset_id,
                concurrency,
                str(input_column.id),
                getattr(request, "organization", None) or request.user.organization.id,
                new_column.id,
                str(request.workspace.id) if request.workspace else None,
            )

            return self._gm.success_response(
                {
                    "message": "Vector DB column created successfully",
                    "new_column_id": str(new_column.id),
                    "new_column_name": new_column.name,
                }
            )

        except Exception as e:
            try:
                if "new_column" in locals():
                    new_column.delete()
            except Exception:
                pass
            logger.exception(f"Error in creating vector db column: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_CREATE_VECTOR_DB_COLUMN")
            )


@temporal_activity(time_limit=3600, queue="tasks_l")
def add_vector_db_column_async(
    config,
    dataset_id,
    concurrency,
    input_column_id,
    org_id,
    new_column_id,
    workspace_id=None,
):
    # Ensure config is properly deserialized if needed
    if isinstance(config, str):
        config = json.loads(config)
    view = AddVectorDBColumnView()
    input_column = Column.objects.get(id=input_column_id)
    rows = Row.objects.filter(dataset_id=dataset_id, deleted=False)
    new_cells = []
    total_processed = 0
    failed_cells = 0
    Column.objects.filter(id=new_column_id).update(status=StatusType.RUNNING.value)

    # Wrap function with OTel context propagation for thread safety
    wrapped_process_row = wrap_for_thread(view._process_row)

    with ThreadPoolExecutor(max_workers=concurrency) as executor:
        future_to_row = {
            executor.submit(
                wrapped_process_row, row, input_column, config, org_id, workspace_id
            ): row
            for row in rows
        }

        for future in as_completed(future_to_row):
            row = future_to_row[future]
            try:
                value, value_infos = future.result()
                if value_infos and "reason" in value_infos:
                    # Handle case where function returned error
                    new_cells.append(
                        Cell(
                            dataset_id=dataset_id,
                            column_id=new_column_id,
                            row=row,
                            value=None,
                            value_infos=json.dumps(value_infos),
                            status=CellStatus.ERROR.value,
                        )
                    )
                    failed_cells += 1
                else:
                    new_cells.append(
                        Cell(
                            dataset_id=dataset_id,
                            column_id=new_column_id,
                            row=row,
                            value=value,
                            value_infos=json.dumps(value_infos if value_infos else {}),
                        )
                    )
                    total_processed += 1
            except Exception as e:
                logger.exception(f"Failed to process row: {str(e)}")
                failed_cells += 1
                # Create a failed cell with error information
                new_cells.append(
                    Cell(
                        dataset_id=dataset_id,
                        column_id=new_column_id,
                        row=row,
                        value=None,
                        value_infos=json.dumps({"reason": str(e)}),
                        status=CellStatus.ERROR.value,
                    )
                )

    if new_cells:
        Cell.objects.bulk_create(new_cells)
    Column.objects.filter(id=new_column_id).update(status=StatusType.COMPLETED.value)
    # insert_embeddings_task.delay(column_ids=[str(new_column_id)])


class GetEmbeddingsListView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        try:
            # Define available embeddings and their configurations
            embeddings_list = {
                "openai": {
                    "name": "OpenAI Embeddings",
                    "description": "OpenAI's text embedding models",
                    "requires_api_key": True,
                    "config_schema": {
                        "api_key": {
                            "type": "string",
                            "required": True,
                            "description": "Your OpenAI API key",
                        },
                        "model": {
                            "type": "string",
                            "required": True,
                            "description": "Model ID to use for embeddings",
                            "default": "text-embedding-ada-002",
                        },
                    },
                },
                "huggingface": {
                    "name": "Hugging Face Embeddings",
                    "description": "Access to Hugging Face's vast collection of embedding models",
                    "requires_api_key": True,
                    "config_schema": {
                        "api_key": {
                            "type": "string",
                            "required": True,
                            "description": "Your Hugging Face API key",
                        },
                        "model": {
                            "type": "string",
                            "required": True,
                            "description": "Model ID to use for embeddings",
                            "default": "sentence-transformers/all-mpnet-base-v2",
                        },
                    },
                },
                "sentence_transformers": {
                    "name": "Sentence Transformers",
                    "description": "Local sentence transformer models for embedding generation",
                    "requires_api_key": False,
                    "config_schema": {
                        "model": {
                            "type": "string",
                            "required": True,
                            "description": "Model ID to use for embeddings",
                            "default": "all-mpnet-base-v2",
                        }
                    },
                },
            }

            # Get embedding type from query params if specified
            embedding_type = request.query_params.get("type")
            if embedding_type:
                if embedding_type not in embeddings_list:
                    return self._gm.bad_request(
                        f"Invalid embedding type: {embedding_type}"
                    )
                return self._gm.success_response(
                    {"embedding": embeddings_list[embedding_type]}
                )

            return self._gm.success_response({"embeddings": embeddings_list})

        except Exception as e:
            logger.exception(f"Error in fetching embeddings: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_EMBEDDINGS")
            )


class FeedbackViewSet(viewsets.ModelViewSet):
    serializer_class = FeedbackSerializer
    permission_classes = [IsAuthenticated]
    _gm = GeneralMethods()

    def create(self, request, *args, **kwargs):
        try:
            serializer = self.get_serializer(data=request.data)
            serializer.is_valid(raise_exception=True)
            feedback = serializer.save(
                user=request.user,
                organization=getattr(request, "organization", None)
                or request.user.organization,
                workspace=getattr(request, "workspace", None),
            )

            return self._gm.success_response({"id": feedback.id})

        except ValidationError:
            return self._gm.bad_request(get_error_message("FAILED_TO_CREATE_FEEDBACK"))
        except Exception as e:
            logger.exception(f"Error in lsubmitting the feedback: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_CREATE_FEEDBACK")
            )

    @action(detail=False, methods=["GET"])
    def get_template(self, request):
        """
        Get evaluation template details based on user_eval_metric_id.
        """
        user_eval_metric_id = request.query_params.get("user_eval_metric_id")

        if not user_eval_metric_id:
            return self._gm.bad_request(
                get_error_message("USER_EVAL_METRIC_ID_REQUIRED")
            )

        try:
            user_eval_metric = UserEvalMetric.objects.get(id=user_eval_metric_id)
        except (UserEvalMetric.DoesNotExist, ValidationError):
            return self._gm.bad_request(
                get_error_message("MISSING_USER_EVAL_METRIC_ID")
            )

        eval_template = user_eval_metric.template

        if not eval_template:
            return self._gm.not_found(get_error_message("EVAL_TEMP_NOT_FOUND"))

        try:
            template_data = {
                "output_type": eval_template.config.get("output"),
                "eval_description": eval_template.description,
                "eval_name": eval_template.name,
                "user_eval_name": user_eval_metric.name,
            }

            if template_data["output_type"] == EVAL_OUTPUT_TYPES["PASS_FAIL"]:
                template_data["choices"] = ["Passed", "Failed"]

            elif template_data["output_type"] == EVAL_OUTPUT_TYPES["CHOICES"]:
                if (
                    user_eval_metric.config
                    and isinstance(user_eval_metric.config, dict)
                    and "config" in user_eval_metric.config
                    and "choices" in user_eval_metric.config["config"]
                    and user_eval_metric.config["config"]["choices"]
                ):
                    template_data["choices"] = user_eval_metric.config["config"][
                        "choices"
                    ]
                    template_data["multi_choice"] = user_eval_metric.config[
                        "config"
                    ].get("multi_choice", False)

                elif hasattr(eval_template, "choices") and eval_template.choices:
                    template_data["choices"] = eval_template.choices
                    template_data["multi_choice"] = eval_template.config.get(
                        "multi_choice", False
                    )

                else:
                    template_data["choices"] = []
                    template_data["multi_choice"] = False

            return self._gm.success_response(template_data)

        except UserEvalMetric.DoesNotExist:
            return self._gm.bad_request(get_error_message("USER_EVAL_METRIC_NOT_EXIST"))
        except Exception as e:
            logger.exception(f"Error in fetching the user eval metric: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_USER_EVAL_METRIC")
            )

    @action(detail=False, methods=["POST"], url_path="submit-feedback")
    def submit_feedback_action(self, request):
        """
        Submit feedback action for re-tuning or re-calculating metrics
        """
        try:
            action_type = request.data.get("action_type")
            feedback_id = request.data.get("feedback_id")
            user_eval_metric_id = request.data.get("user_eval_metric_id")
            value = request.data.get("value") if request.data.get("value") else None
            explanation = (
                request.data.get("explanation")
                if request.data.get("explanation")
                else None
            )

            if not action_type or not user_eval_metric_id or not feedback_id:
                return self._gm.bad_request(
                    get_error_message("MISSING_METRIC_ID_FEEDBACK_ID_AND_ACTION_TYPE")
                )

            # Validate action type
            valid_actions = [
                "retune",
                "recalculate_row",
                "recalculate_dataset",
                "retune_recalculate",
            ]
            if action_type not in valid_actions:
                return self._gm.bad_request(
                    f"Invalid action_type. Must be one of: {', '.join(valid_actions)}"
                )

            feedback = Feedback.objects.get(
                id=feedback_id,
                organization=getattr(request, "organization", None)
                or request.user.organization,
            )
            feedback.action_type = action_type

            row_id = str(feedback.row_id)

            # Get the user eval metric
            eval_column = Column.objects.get(id=feedback.source_id)
            dataset = Dataset.objects.get(id=eval_column.dataset_id)
            try:
                user_eval_metric = UserEvalMetric.objects.get(id=user_eval_metric_id)
            except UserEvalMetric.DoesNotExist:
                return self._gm.bad_request(
                    get_error_message("MISSING_USER_EVAL_METRIC_ID")
                )

            feedback.eval_template = user_eval_metric.template
            feedback.value = value if value else feedback.value
            feedback.explanation = explanation if explanation else feedback.explanation
            feedback.save()
            # get_fewshots = RAG()
            embedding_manager = EmbeddingManager()
            # Get all cells for this row
            row_cells = Cell.objects.filter(
                row_id=feedback.row_id, dataset_id=dataset.id, deleted=False
            ).select_related("column")

            # Initialize row data
            row_dict = {}

            # Add cell values for this feedback
            for cell in row_cells:
                column_id = str(cell.column.id)
                if column_id != str(eval_column.id):
                    row_dict[column_id] = cell.value

            # Add feedback information
            row_dict["feedback_comment"] = feedback.explanation
            row_dict["feedback_value"] = feedback.value
            futureagi_eval = (
                True
                if user_eval_metric.template.config.get("eval_type_id")
                in FUTUREAGI_EVAL_TYPES
                else False
            )
            runner = EvaluationRunner(
                user_eval_metric.template.config.get("eval_type_id"),
                format_output=True,
                futureagi_eval=futureagi_eval,
            )
            source_config = {
                "reference_id": str(user_eval_metric.id),
                "dataset_id": str(feedback.user_eval_metric.dataset.id),
                "row_id": str(feedback.row_id),
                "feedback_id": str(feedback.id),
                "value": feedback.value,
                "explanation": feedback.explanation,
            }

            required_field, mapping = runner._get_required_fields_and_mappings(
                user_eval_metric=user_eval_metric
            )
            # print(required_field,"required_field12*****")
            # print(mapping,"mapping12*****")
            embedding_manager.parallel_process_metadata(
                eval_id=user_eval_metric.template.id,
                metadatas=row_dict,
                inputs_formater=required_field,
                organization_id=dataset.organization.id,
                workspace_id=dataset.workspace.id if dataset.workspace else None,
            )
            embedding_manager.close()
            # Handle different action types
            if action_type == "retune":
                message = "Metric queued for retuning"

            elif action_type == "recalculate_row":
                if not row_id:
                    return self._gm.bad_request(get_error_message("ROW_ID_MISSING"))

                # Prepare data for batch processing
                evaluation_data = {
                    "metric_ids": [user_eval_metric_id],
                    "row_ids": [row_id],
                    "source": "feedback",
                    "source_id": user_eval_metric.template.id,
                    "source_config": source_config,
                    "column_source_id": eval_column.source_id,
                    "column_source": eval_column.source,
                    "column": eval_column.id,
                    "feedback_id": feedback.id,
                }

                # Run all evaluations in a single async task
                run_evaluation_task.apply_async(args=(evaluation_data,))

                message = "Row queued for recalculation"

            else:  # recalculate_dataset
                if eval_column.source == SourceChoices.OPTIMISATION_EVALUATION.value:
                    UserEvalMetric.objects.filter(
                        id=user_eval_metric_id,
                    ).update(status=StatusType.OPTIMIZATION_EVALUATION.value)
                    column = Column.objects.get(id=feedback.source_id)
                    Cell.objects.filter(
                        column__source_id=eval_column.source_id, deleted=False
                    ).update(status=CellStatus.RUNNING.value)

                else:
                    UserEvalMetric.objects.filter(
                        id=user_eval_metric_id,
                    ).update(status=StatusType.NOT_STARTED.value)
                    column = Column.objects.get(source_id=user_eval_metric_id)
                    Cell.objects.filter(
                        column__source_id__in=[
                            user_eval_metric_id,
                            f"{column.id}-sourceid-{user_eval_metric_id}",
                        ],
                        deleted=False,
                    ).update(status=CellStatus.RUNNING.value)

                message = "Dataset queued for recalculation"

            return self._gm.success_response(
                {
                    "message": message,
                    "action_type": action_type,
                    "user_eval_metric_id": str(user_eval_metric_id),
                }
            )

        except Exception as e:
            logger.exception(f"Error in submitting the feedback: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_CREATE_FEEDBACK")
            )

    @action(detail=False, methods=["GET"], url_path="get-feedback-details")
    def get_feedback_details(self, request):
        """
        Get feedback details based on filters like user_eval_metric_id, row_id, etc.
        """
        try:
            # Get filter parameters
            user_eval_metric_id = request.query_params.get("user_eval_metric_id")
            row_id = request.query_params.get("row_id")

            # Build base queryset
            queryset = Feedback.objects.select_related("user").filter(deleted=False)

            # Apply filters if provided
            if user_eval_metric_id:
                queryset = queryset.filter(user_eval_metric_id=user_eval_metric_id)
            if row_id:
                queryset = queryset.filter(row_id=row_id)

            # Order by most recent first
            queryset = queryset.order_by("-created_at")

            # Serialize the feedback data
            feedback_data = []
            for feedback in queryset:
                feedback_info = {
                    "id": str(feedback.id),
                    # 'user': {
                    #     'id': str(feedback.user.id),
                    #     'email': feedback.user.email,
                    #     'name': f"{feedback.user.first_name} {feedback.user.last_name}".strip()
                    # },
                    "value": feedback.value,
                    "comment": feedback.explanation,
                    "created_at": feedback.created_at.isoformat(),
                    "action_type": feedback.action_type,
                }
                feedback_data.append(feedback_info)

            return self._gm.success_response(
                {"feedback": feedback_data, "total_count": len(feedback_data)}
            )

        except Exception as e:
            logger.exception(f"Error in fetching the feedbacks: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_FEEDBACKS")
            )

    @action(detail=False, methods=["GET"], url_path="get-feedback-summary")
    def get_feedback_summary(self, request):
        """
        Get summary statistics for feedback on a specific metric
        """
        try:
            user_eval_metric_id = request.query_params.get("user_eval_metric_id")

            if not user_eval_metric_id:
                return self._gm.bad_request(
                    get_error_message("USER_EVAL_METRIC_ID_REQUIRED")
                )

            # Get all feedback for this metric
            feedback_qs = Feedback.objects.filter(
                user_eval_metric_id=user_eval_metric_id,
                organization=getattr(request, "organization", None)
                or request.user.organization,
                deleted=False,
            )

            # Calculate summary statistics
            summary = {
                "total_feedback": feedback_qs.count(),
                "unique_users": feedback_qs.values("user").distinct().count(),
                "action_types": {
                    "retune": feedback_qs.filter(action_type="retune").count(),
                    "recalculate_row": feedback_qs.filter(
                        action_type="recalculate_row"
                    ).count(),
                    "recalculate_dataset": feedback_qs.filter(
                        action_type="recalculate_dataset"
                    ).count(),
                },
                "status_breakdown": {
                    "pending": feedback_qs.filter(
                        status=StatusType.PENDING.value
                    ).count(),
                    "completed": feedback_qs.filter(
                        status=StatusType.COMPLETED.value
                    ).count(),
                    "failed": feedback_qs.filter(
                        status=StatusType.FAILED.value
                    ).count(),
                },
                "recent_feedback": [],
            }

            # Add most recent feedback entries
            recent_feedback = feedback_qs.select_related("user").order_by(
                "-created_at"
            )[:5]
            for feedback in recent_feedback:
                summary["recent_feedback"].append(
                    {
                        "id": str(feedback.id),
                        "user": f"{feedback.user.first_name} {feedback.user.last_name}".strip(),
                        "action_type": feedback.action_type,
                        "created_at": feedback.created_at.isoformat(),
                        "status": feedback.status,
                    }
                )

            return self._gm.success_response(summary)

        except Exception as e:
            logger.exception(f"Error in fetching the feedback summary: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_FEEDBACK_SUMMARY")
            )


class SingleRowEvaluationView(APIView):
    permission_classes = [IsAuthenticated]
    _gm = GeneralMethods()

    def post(self, request):
        try:
            # Extract the user_eval_metric_ids and row_ids from the request data
            user_eval_metric_ids = request.data.get("user_eval_metric_ids", [])
            row_ids = request.data.get("row_ids", [])
            selected_all_rows = request.data.get("selected_all_rows", False)
            if not user_eval_metric_ids:
                return self._gm.bad_request(
                    get_error_message("USER_EVAL_METRIC_IDs_REQUIRED")
                )
            if not row_ids and not selected_all_rows:
                return self._gm.bad_request(get_error_message("MISSING_ROW_IDS"))

            if selected_all_rows:
                user_eval_metric = UserEvalMetric.objects.get(
                    id=user_eval_metric_ids[0]
                )
                if row_ids and len(row_ids) > 0:
                    row_ids = list(
                        map(
                            str,
                            Row.objects.filter(
                                dataset=user_eval_metric.dataset, deleted=False
                            )
                            .exclude(id__in=row_ids)
                            .values_list("id", flat=True),
                        )
                    )
                else:
                    row_ids = list(
                        map(
                            str,
                            Row.objects.filter(
                                dataset=user_eval_metric.dataset, deleted=False
                            ).values_list("id", flat=True),
                        )
                    )

            # Prepare data for batch processing
            evaluation_data = {"metric_ids": user_eval_metric_ids, "row_ids": row_ids}

            Cell.objects.filter(
                row_id__in=row_ids,
                column__source_id__in=user_eval_metric_ids,
                deleted=False,
            ).update(
                status=CellStatus.RUNNING.value, value=None, value_infos=json.dumps({})
            )

            # Run all evaluations in a single async task
            run_evaluation_task.apply_async(args=(evaluation_data,))

            return self._gm.success_response(
                {"success": "Evaluations queued for processing."}
            )
        except Exception as e:
            logger.exception(f"Error in evaluation of row: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_EVALUATE_ROW")
            )


@temporal_activity(time_limit=3600, queue="tasks_s")
def run_evaluation_task(evaluation_data):
    try:
        logger.info("-----INSIDE run_evaluation_task -----")
        metric_ids = evaluation_data["metric_ids"]
        row_ids = evaluation_data["row_ids"]

        # Update status for all metrics
        metrics = UserEvalMetric.objects.filter(id__in=metric_ids)
        metric_map = {str(metric.id): metric for metric in list(metrics)}
        metrics.update(status=StatusType.RUNNING.value)

        if (
            evaluation_data.get("column_source", "")
            == SourceChoices.OPTIMISATION_EVALUATION.value
        ):
            Cell.objects.filter(
                row_id__in=row_ids,
                column_id=evaluation_data.get("column"),
                deleted=False,
            ).update(status=CellStatus.RUNNING.value)
            optimize = OptimizationDataset.objects.filter(
                id=evaluation_data.get("column_source_id").split("-sourceid-")[0]
            ).first()
            col = Column.objects.filter(id=evaluation_data.get("column")).first()
            args = {
                "optimize": optimize,
                "source": evaluation_data.get("source", None),
                "source_id": evaluation_data.get("source_id", None),
                "column": col,
            }
        else:
            # Mark existing cells as running
            Cell.objects.filter(
                row_id__in=row_ids, column__source_id__in=metric_ids, deleted=False
            ).update(
                status=CellStatus.RUNNING.value, value=None, value_infos=json.dumps({})
            )
            args = {
                "is_only_eval": True,
                "source": evaluation_data.get("source", None),
                "source_id": evaluation_data.get("source_id", None),
                "source_configs": {
                    "feedback_id": str(evaluation_data.get("feedback_id", None))
                },
            }

        futures = []
        blocked_metric_ids = set()
        with ThreadPoolExecutor(max_workers=5) as executor:
            for metric_id in metric_ids:
                metric = metric_map.get(metric_id)
                if metric:
                    try:
                        from ee.usage.services.metering import check_usage
                    except ImportError:
                        check_usage = None

                    if check_usage is not None:
                        api_call_type = _get_api_call_type(
                            metric.model or ModelChoices.TURING_LARGE.value
                        )
                        usage_check = check_usage(
                            str(metric.organization.id), api_call_type
                        )
                        if not usage_check.allowed:
                            from model_hub.tasks.user_evaluation import (
                                _mark_cells_usage_limit_error,
                            )

                            UserEvalMetric.objects.filter(id=metric_id).update(
                                status=StatusType.FAILED.value
                            )
                            blocked_metric_ids.add(str(metric_id))
                            _mark_cells_usage_limit_error(metric, usage_check)
                            logger.warning(
                                "dataset_eval_rerun_usage_limit_exceeded",
                                eval_id=str(metric_id),
                                reason=usage_check.reason,
                            )
                            continue

                    properties = get_mixpanel_properties(
                        org=metric.organization,
                        eval=metric,
                        source=evaluation_data.get("source"),
                        dataset=metric.dataset,
                        count=len(row_ids),
                    )
                    track_mixpanel_event(
                        MixpanelEvents.EVAL_RUN_STARTED.value, properties
                    )
                try:
                    # Initialize the EvaluationRunner for each metric
                    logger.info(
                        " ----- INSIDE run_evaluation_task | Initializing the EvaluationRunner for each metric -----"
                    )
                    runner_args = dict(args)
                    if not runner_args.get("source"):
                        runner_args["source"] = "dataset_evaluation"
                    if not runner_args.get("source_id"):
                        runner_args["source_id"] = metric.template.id
                    runner_source_configs = dict(
                        runner_args.get("source_configs") or {}
                    )
                    if metric.dataset_id:
                        runner_source_configs.setdefault(
                            "dataset_id", str(metric.dataset_id)
                        )
                    runner_source_configs.setdefault("source", "dataset")
                    runner_args["source_configs"] = runner_source_configs

                    evaluation_runner = EvaluationRunner(
                        user_eval_metric_id=metric_id,
                        **runner_args,
                    )

                    # Wrap function with OTel context propagation for thread safety
                    wrapped_run_evaluation = wrap_for_thread(
                        evaluation_runner.run_evaluation_for_row
                    )

                    # Process each row for this metric
                    for row_id in row_ids:
                        futures.append(executor.submit(wrapped_run_evaluation, row_id))

                except Exception as e:
                    logger.error(
                        f"Error initializing evaluation for metric {metric_id}: {str(e)}"
                    )
                    UserEvalMetric.objects.filter(id=metric_id).update(
                        status=StatusType.FAILED.value
                    )

            # Wait for all futures to complete
            success = 0
            for future in as_completed(futures):
                try:
                    future.result()
                    success += 1
                except Exception as e:
                    logger.error(f"Error in evaluation: {str(e)}")

        for metric_id in metric_ids:
            metric = metric_map.get(metric_id)
            if metric:
                properties = get_mixpanel_properties(
                    eval=metric,
                    org=metric.organization,
                    source=evaluation_data.get("source"),
                    dataset=metric.dataset,
                    count=success,
                    failed=len(row_ids) - success,
                )
                track_mixpanel_event(
                    MixpanelEvents.EVAL_RUN_COMPLETED.value, properties
                )

        # Update status to completed for successful metrics
        completed_metric_ids = [
            metric_id
            for metric_id in metric_ids
            if str(metric_id) not in blocked_metric_ids
        ]
        if completed_metric_ids:
            UserEvalMetric.objects.filter(id__in=completed_metric_ids).update(
                status=StatusType.COMPLETED.value
            )
    except Exception as e:
        # Handle exceptions and log errors
        UserEvalMetric.objects.filter(id__in=metric_ids).update(
            status=StatusType.FAILED.value
        )
        logger.error(f"Error in evaluation task: {str(e)}")


class DuplicateRowsView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            row_ids = request.data.get("row_ids", [])
            selected_all_rows = request.data.get("selected_all_rows", False)
            num_copies = request.data.get("num_copies", 1)

            if not row_ids and not selected_all_rows:
                return self._gm.bad_request(get_error_message("MISSING_ROW_IDS"))

            if not isinstance(num_copies, int) or num_copies < 1:
                return self._gm.bad_request(
                    get_error_message("NUM_COPIES_NOT_POSITIVE")
                )

            from model_hub.services.dataset_validators import MAX_DUPLICATE_COPIES

            if num_copies > MAX_DUPLICATE_COPIES:
                return self._gm.bad_request(
                    f"Number of copies cannot exceed {MAX_DUPLICATE_COPIES}"
                )

            # Get the dataset and verify it exists
            dataset = get_object_or_404(Dataset, id=dataset_id, deleted=False)

            # Get max order to append new rows at the end
            last_row = (
                Row.all_objects.filter(dataset=dataset).order_by("-created_at").first()
            )
            if last_row:
                max_order = last_row.order
            else:
                max_order = -1

            new_rows = []
            new_cells = []

            if selected_all_rows:
                if row_ids and len(row_ids) > 0:
                    source_rows = Row.objects.exclude(id__in=row_ids).filter(
                        dataset=dataset, deleted=False
                    )
                else:
                    source_rows = Row.objects.filter(dataset=dataset, deleted=False)
            else:
                source_rows = Row.objects.filter(
                    id__in=row_ids, dataset=dataset, deleted=False
                )

            if log_and_deduct_cost_for_resource_request is not None:
                call_log_row = log_and_deduct_cost_for_resource_request(
                    organization=getattr(request, "organization", None)
                    or request.user.organization,
                    api_call_type=APICallTypeChoices.ROW_ADD.value,
                    config={"total_rows": source_rows.count() * num_copies},
                    workspace=request.workspace,
                )
                if (
                    call_log_row is None
                    or call_log_row.status == APICallStatusChoices.RESOURCE_LIMIT.value
                ):
                    return self._gm.too_many_requests(ROW_LIMIT_REACHED_MESSAGE)
                call_log_row.status = APICallStatusChoices.SUCCESS.value
                call_log_row.save()
            source_cells = Cell.objects.filter(
                row__in=source_rows, deleted=False
            ).select_related("column")

            # Group cells by row for efficient processing
            cells_by_row: dict[Any, Any] = {}
            for cell in source_cells:
                if cell.row_id not in cells_by_row:
                    cells_by_row[cell.row_id] = []
                cells_by_row[cell.row_id].append(cell)

            # Create multiple copies
            current_order = max_order + 1
            for source_row in source_rows:
                # Create specified number of copies for each source row
                for _copy_num in range(num_copies):
                    # Create new row
                    new_row = Row(id=uuid.uuid4(), dataset=dataset, order=current_order)
                    new_rows.append(new_row)
                    current_order += 1

                    # Create new cells for this row
                    if source_row.id in cells_by_row:
                        for source_cell in cells_by_row[source_row.id]:
                            new_cells.append(
                                Cell(
                                    id=uuid.uuid4(),
                                    dataset=dataset,
                                    column=source_cell.column,
                                    row=new_row,
                                    value=source_cell.value,
                                    value_infos=(
                                        source_cell.value_infos
                                        if source_cell.value_infos
                                        else json.dumps({})
                                    ),
                                    status=source_cell.status,
                                )
                            )

            # Bulk create all new rows and cells
            Row.objects.bulk_create(new_rows)
            Cell.objects.bulk_create(new_cells)

            return self._gm.success_response(
                {
                    "message": "Rows duplicated successfully",
                    "source_rows": len(source_rows),
                    "copies_per_row": num_copies,
                    "total_new_rows": len(new_rows),
                    "new_row_ids": [str(row.id) for row in new_rows],
                }
            )

        except Exception as e:
            logger.exception(f"Error in duplication of row: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_DUPLICATE_ROW")
            )


class DuplicateDatasetView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            new_name = request.data.get("name")
            row_ids = request.data.get("row_ids", [])
            selected_all_rows = request.data.get("selected_all_rows", False)
            if not new_name:
                return self._gm.bad_request(
                    get_error_message("MISSING_NEW_DATASET_NAME")
                )

            if log_and_deduct_cost_for_resource_request is not None:
                call_log_row_entry = log_and_deduct_cost_for_resource_request(
                    organization=getattr(request, "organization", None)
                    or request.user.organization,
                    api_call_type=APICallTypeChoices.DATASET_ADD.value,
                    workspace=request.workspace,
                )
                if (
                    call_log_row_entry is None
                    or call_log_row_entry.status
                    == APICallStatusChoices.RESOURCE_LIMIT.value
                ):
                    return self._gm.too_many_requests(
                        get_error_message("DATASET_CREATE_LIMIT_REACHED")
                    )
                call_log_row_entry.status = APICallStatusChoices.SUCCESS.value
                call_log_row_entry.save()

            # Get source dataset and verify it exists
            source_dataset = get_object_or_404(Dataset, id=dataset_id, deleted=False)

            # Create new dataset with copied attributes
            new_dataset = Dataset.objects.create(
                id=uuid.uuid4(),
                name=new_name,
                organization=source_dataset.organization,
                model_type=source_dataset.model_type,
                column_order=(
                    source_dataset.column_order.copy()
                    if source_dataset.column_order
                    else []
                ),
                column_config=(
                    source_dataset.column_config.copy()
                    if source_dataset.column_config
                    else {}
                ),
                user=request.user,
            )

            # Copy columns
            column_id_mapping = {}  # Map old column IDs to new ones
            new_columns = []

            source_columns = Column.objects.filter(
                dataset=source_dataset, deleted=False
            ).exclude(
                source__in=[
                    SourceChoices.EXPERIMENT.value,
                    SourceChoices.EXPERIMENT_EVALUATION.value,
                    SourceChoices.EXPERIMENT_EVALUATION_TAGS.value,
                ]
            )
            for column in source_columns:
                new_column_id = uuid.uuid4()
                column_id_mapping[str(column.id)] = str(new_column_id)

                new_columns.append(
                    Column(
                        id=new_column_id,
                        name=column.name,
                        data_type=column.data_type,
                        source=SourceChoices.OTHERS.value,
                        dataset=new_dataset,
                        deleted=False,
                    )
                )

            Column.objects.bulk_create(new_columns)

            # Update column_order with new column IDs
            new_column_order = []
            new_column_config = {}

            for old_col_id in new_dataset.column_order:
                if old_col_id in column_id_mapping:
                    new_column_order.append(column_id_mapping[old_col_id])
                    if old_col_id in new_dataset.column_config:
                        new_column_config[column_id_mapping[old_col_id]] = (
                            new_dataset.column_config[old_col_id]
                        )

            new_dataset.column_order = new_column_order
            new_dataset.column_config = new_column_config
            new_dataset.save()

            # Copy rows and cells in batches
            if selected_all_rows:
                if row_ids and len(row_ids) > 0:
                    source_rows = Row.objects.exclude(id__in=row_ids).filter(
                        dataset=source_dataset, deleted=False
                    )
                else:
                    source_rows = Row.objects.filter(
                        dataset=source_dataset, deleted=False
                    )
            else:
                source_rows = Row.objects.filter(
                    id__in=row_ids, dataset=source_dataset, deleted=False
                )
            new_rows = []
            new_cells = []

            if log_and_deduct_cost_for_resource_request is not None:
                call_log_row = log_and_deduct_cost_for_resource_request(
                    organization=getattr(request, "organization", None)
                    or request.user.organization,
                    api_call_type=APICallTypeChoices.ROW_ADD.value,
                    config={"total_rows": source_rows.count()},
                    workspace=request.workspace,
                )
                if (
                    call_log_row is None
                    or call_log_row.status == APICallStatusChoices.RESOURCE_LIMIT.value
                ):
                    return self._gm.too_many_requests(ROW_LIMIT_REACHED_MESSAGE)
                call_log_row.status = APICallStatusChoices.SUCCESS.value
                call_log_row.save()
            # Process in batches of 1000 rows
            batch_size = 1000
            for i in range(0, source_rows.count(), batch_size):
                batch_rows = source_rows[i : i + batch_size]
                row_id_mapping = {}  # Map old row IDs to new ones

                # Create new rows
                for row in batch_rows:
                    new_row_id = uuid.uuid4()
                    row_id_mapping[row.id] = new_row_id
                    new_rows.append(
                        Row(
                            id=new_row_id,
                            dataset=new_dataset,
                            order=row.order,
                            deleted=False,
                        )
                    )

                # Bulk create rows
                Row.objects.bulk_create(new_rows)

                # Get cells for current batch of rows
                batch_cells = Cell.objects.filter(
                    row__in=batch_rows, deleted=False
                ).select_related("column")

                # Create new cells
                for cell in batch_cells:
                    if str(cell.column.id) in column_id_mapping:
                        new_cells.append(
                            Cell(
                                id=uuid.uuid4(),
                                dataset=new_dataset,
                                column_id=column_id_mapping[str(cell.column.id)],
                                row_id=row_id_mapping[cell.row_id],
                                value=cell.value,
                                value_infos=(
                                    cell.value_infos
                                    if cell.value_infos
                                    else json.dumps({})
                                ),
                                status=cell.status,
                                deleted=False,
                            )
                        )

                # Bulk create cells
                Cell.objects.bulk_create(new_cells)

                # Clear lists for next batch
                new_rows = []
                new_cells = []

            return self._gm.success_response(
                {
                    "message": "Dataset duplicated successfully",
                    "new_dataset_id": str(new_dataset.id),
                    "new_dataset_name": new_dataset.name,
                    "columns_copied": len(column_id_mapping),
                    "rows_copied": source_rows.count(),
                }
            )

        except Exception as e:
            logger.exception(str(e))
            # Cleanup on error
            try:
                if "new_dataset" in locals():
                    new_dataset.delete()
            except Exception:
                pass
            logger.exception(f"Error in duplication of dataset: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_DUPLICATE_DATASET")
            )


class MergeDatasetView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            target_dataset_id = request.data.get("target_dataset_id")
            row_ids = request.data.get("row_ids") or []
            selected_all_rows = request.data.get("selected_all_rows", False)
            if not target_dataset_id:
                return self._gm.bad_request(
                    get_error_message("MISSING_SOURCE_DATASET_ID")
                )

            # Get both datasets and verify they exist
            target_dataset = get_object_or_404(
                Dataset, id=target_dataset_id, deleted=False
            )
            source_dataset = get_object_or_404(Dataset, id=dataset_id, deleted=False)

            # Get max order of target dataset to append rows at the end
            last_row = (
                Row.all_objects.filter(dataset=source_dataset)
                .order_by("-created_at")
                .first()
            )
            if last_row:
                max_order = last_row.order
            else:
                max_order = -1

            # Get columns in the order specified by source_dataset's column_order
            source_column_order = source_dataset.column_order or []
            source_columns = []
            for col_id in source_column_order:
                col = Column.objects.filter(
                    id=col_id, dataset=source_dataset, deleted=False
                ).first()
                if col:
                    source_columns.append(col)

            # Add any columns that might not be in column_order
            remaining_cols = (
                Column.objects.filter(dataset=source_dataset, deleted=False)
                .exclude(id__in=source_column_order)
                .exclude(
                    source__in=[
                        SourceChoices.EXPERIMENT.value,
                        SourceChoices.EXPERIMENT_EVALUATION.value,
                        SourceChoices.EXPERIMENT_EVALUATION_TAGS.value,
                    ]
                )
            )
            source_columns.extend(remaining_cols)

            target_columns = Column.objects.filter(
                dataset=target_dataset, deleted=False
            ).exclude(
                source__in=[
                    SourceChoices.EXPERIMENT.value,
                    SourceChoices.EXPERIMENT_EVALUATION.value,
                    SourceChoices.EXPERIMENT_EVALUATION_TAGS.value,
                ]
            )

            column_mapping = {}  # Maps source column IDs to target column IDs
            new_columns = []

            # First, map existing columns while maintaining order
            for source_col in source_columns:
                matching_col = next(
                    (
                        col
                        for col in target_columns
                        if col.name == source_col.name
                        and col.data_type == source_col.data_type
                    ),
                    None,
                )

                if matching_col:
                    column_mapping[str(source_col.id)] = str(matching_col.id)
                else:
                    # Create new column for unmatched source columns
                    new_col_id = uuid.uuid4()
                    new_columns.append(
                        Column(
                            id=new_col_id,
                            name=source_col.name,
                            data_type=DataTypeChoices.TEXT.value,
                            source=SourceChoices.OTHERS.value,
                            dataset=target_dataset,
                        )
                    )
                    column_mapping[str(source_col.id)] = str(new_col_id)

            # Bulk create new columns
            if new_columns:
                Column.objects.bulk_create(new_columns)

                # Update target dataset's column order and config
                column_order = target_dataset.column_order or []
                column_config = target_dataset.column_config or {}

                # Add new columns in the same order as source dataset
                for source_col in source_columns:
                    if str(source_col.id) in column_mapping:
                        new_col_id = column_mapping[str(source_col.id)]
                        # Only add to column_order if it's a new column
                        if new_col_id not in column_order:
                            column_order.append(new_col_id)
                            column_config[new_col_id] = {
                                "is_visible": True,
                                "is_frozen": None,
                            }

                target_dataset.column_order = column_order
                target_dataset.column_config = column_config
                target_dataset.save()

            # Copy rows and cells in batches
            if selected_all_rows:
                if row_ids and len(row_ids) > 0:
                    source_rows = Row.objects.exclude(id__in=row_ids).filter(
                        dataset=source_dataset, deleted=False
                    )
                else:
                    source_rows = Row.objects.filter(
                        dataset=source_dataset, deleted=False
                    )
            else:
                source_rows = Row.objects.filter(
                    id__in=row_ids, dataset=source_dataset, deleted=False
                )
            batch_size = 1000
            current_order = max_order + 1

            for i in range(0, source_rows.count(), batch_size):
                batch_rows = source_rows[i : i + batch_size]
                new_rows = []
                new_cells = []
                row_id_mapping = {}

                # Create new rows
                for row in batch_rows:
                    new_row_id = uuid.uuid4()
                    row_id_mapping[row.id] = new_row_id
                    new_rows.append(
                        Row(
                            id=new_row_id,
                            dataset=target_dataset,
                            order=current_order,
                        )
                    )
                    current_order += 1

                # Bulk create rows
                Row.objects.bulk_create(new_rows)

                # Get cells for current batch of rows
                batch_cells = Cell.objects.filter(
                    row__in=batch_rows, deleted=False
                ).select_related("column")

                # Create new cells
                for cell in batch_cells:
                    if str(cell.column.id) in column_mapping:
                        new_cells.append(
                            Cell(
                                id=uuid.uuid4(),
                                dataset=target_dataset,
                                column_id=column_mapping[str(cell.column.id)],
                                row_id=row_id_mapping[cell.row_id],
                                value=cell.value,
                                value_infos=(
                                    cell.value_infos
                                    if cell.value_infos
                                    else json.dumps({})
                                ),
                                status=cell.status,
                            )
                        )

                # Bulk create cells
                if new_cells:
                    Cell.objects.bulk_create(new_cells)

            return self._gm.success_response(
                {
                    "message": "Datasets merged successfully",
                    "rows_added": source_rows.count(),
                    "new_columns_created": len(new_columns),
                    "columns_mapped": len(column_mapping) - len(new_columns),
                }
            )

        except Exception as e:
            logger.exception(f"Error in merging the datasets: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_MERGE_DATASETS")
            )


class GetDerivedDatasets(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    # parser_classes = (MultiPartParser, FormParser, JSONParser)

    def get(self, request, dataset_id, *args, **kwargs):
        try:
            dataset = get_object_or_404(Dataset, id=dataset_id)

            # Filter datasets and exclude those with null experiments
            derived_datasets = ExperimentDatasetTable.objects.filter(
                experiment__dataset=dataset, deleted=False
            )

            serializer = DerivedDatasetSerializer(derived_datasets, many=True)
            filtered_data = [
                data for data in serializer.data if data["experiment"] is not None
            ]
            return self._gm.success_response(filtered_data)

        except Exception as e:
            logger.exception(f"Error in fetching the derived dataset: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_DERIVED_DATASET")
            )


class GetBaseColumnsView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            dataset_ids = request.query_params.getlist(
                "dataset_ids", []
            ) or request.query_params.getlist("datasetIds", [])

            # Verify all datasets exist and aren't deleted in one query
            dataset_count = Dataset.objects.filter(
                id__in=dataset_ids, deleted=False
            ).count()
            if dataset_count != len(dataset_ids):
                return self._gm.bad_request(get_error_message("INVALID_DATASET_IDS"))

            # Get all columns per dataset in a single query
            excluded_sources = [
                SourceChoices.EXPERIMENT.value,
                SourceChoices.EXPERIMENT_EVALUATION_TAGS.value,
                SourceChoices.EVALUATION_TAGS.value,
                SourceChoices.OPTIMISATION_EVALUATION_TAGS.value,
                SourceChoices.EVALUATION.value,
                SourceChoices.EXPERIMENT_EVALUATION.value,
                SourceChoices.OPTIMISATION_EVALUATION.value,
                SourceChoices.EVALUATION_REASON.value,
            ]

            columns = (
                Column.objects.filter(dataset_id__in=dataset_ids, deleted=False)
                .exclude(source__in=excluded_sources)
                .values("dataset_id", "name")
            )

            # Group column names by dataset_id
            columns_by_dataset = {dataset_id: set() for dataset_id in dataset_ids}
            for col in columns:
                dataset_id = str(col["dataset_id"])
                columns_by_dataset[dataset_id].add(col["name"])

            # Find intersection of column names across all datasets
            common_columns = (
                set.intersection(*columns_by_dataset.values())
                if columns_by_dataset
                else set()
            )

            return self._gm.success_response({"base_columns": list(common_columns)})

        except Exception as e:
            logger.exception(f"Error in fetching base columns: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_BASE_COLUMNS")
            )


class GetCompareDatasetRow(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def get(self, request, compare_id, row_id, *args, **kwargs):
        try:
            metadata = {"status": "processing", "file_row_ids": {}}

            loop_start = time.time()
            while (not metadata.get("status") == "completed") and (
                str(row_id) not in metadata.get("file_row_ids")
            ):
                logger.info(
                    f"Waiting for metadata to be ready for compare_id: {compare_id} and row_id: {row_id}"
                )
                with open(f"compare/{compare_id}/metadata.json") as f:
                    metadata = json.load(f)
                if time.time() - loop_start > 300:
                    return self._gm.bad_request(
                        "Timeout while waiting for metadata to be ready"
                    )
                time.sleep(0.1)

            response = defaultdict()
            row_found = False

            # Find the page for the current row_id
            page_name = metadata.get("file_row_ids").get(str(row_id))
            if not page_name:
                return self._gm.bad_request("Row ID not found")

            # Extract page index from the page_name (e.g., "page_1" -> 1)
            page_index = int(page_name.split("_")[1])

            # Collect rows for the current page
            rows_in_page = [
                row_id
                for row_id, page in metadata["file_row_ids"].items()
                if page == page_name
            ]

            # Find the position of the current row in the page
            try:
                row_pos = rows_in_page.index(str(row_id))
            except ValueError:
                return self._gm.bad_request("Row not found in current page")

            # Determine previous and next row IDs
            prev_row_id = None
            next_row_id = None

            # Handle previous row navigation
            if row_pos > 0:
                prev_row_id = rows_in_page[row_pos - 1]
            else:
                # If it's the first row in the page, get the last row from the previous page (if any)
                if page_index > 1:
                    prev_page_name = f"page_{page_index - 1}"
                    prev_page_rows = [
                        row_id
                        for row_id, page in metadata["file_row_ids"].items()
                        if page == prev_page_name
                    ]
                    prev_row_id = prev_page_rows[-1] if prev_page_rows else None

            # Handle next row navigation
            if row_pos < len(rows_in_page) - 1:
                next_row_id = rows_in_page[row_pos + 1]
            else:
                # If it's the last row in the page, get the first row from the next page (if any)
                if page_index < int(metadata.get("total_pages")):
                    next_page_name = f"page_{page_index + 1}"
                    next_page_rows = [
                        row_id
                        for row_id, page in metadata["file_row_ids"].items()
                        if page == next_page_name
                    ]
                    next_row_id = next_page_rows[0] if next_page_rows else None

            response["prev_row_id"] = prev_row_id
            response["next_row_id"] = next_row_id

            # Fetch the table data for the current page (page_index)
            data = download_json_from_s3(
                object_key=f"compare/{compare_id}/page_{page_index}.json"
            )

            # Find the specific row data on this page
            row_found = False
            for table_data in data["table"]:
                if str(table_data["row_id"]) == str(row_id):
                    row_found = True
                    response["table"] = [table_data]
                    break

            if not row_found:
                return self._gm.bad_request(get_error_message("ROW_NOT_FOUND"))

            comparison_datasets = metadata.get("comparison_datasets")
            dynamic_sources = [
                SourceChoices.RUN_PROMPT.value,
                SourceChoices.EXPERIMENT_EVALUATION_TAGS.value,
                SourceChoices.EVALUATION_TAGS.value,
                SourceChoices.OPTIMISATION_EVALUATION_TAGS.value,
                SourceChoices.EVALUATION.value,
                SourceChoices.EXPERIMENT_EVALUATION.value,
                SourceChoices.OPTIMISATION_EVALUATION.value,
                SourceChoices.EVALUATION_REASON.value,
            ]
            base_column_name = metadata.get("base_column_name")
            common_columns = metadata.get("common_column_names")
            dataset_info = metadata.get("dataset_info")
            dynamic_columns = list(
                Column.objects.filter(
                    name__in=common_columns,
                    dataset__id__in=comparison_datasets,
                    source__in=dynamic_sources,
                ).select_related("dataset")
            )

            dataset_priority = {ds: i for i, ds in enumerate(comparison_datasets)}

            dynamic_columns = sorted(
                dynamic_columns,
                key=lambda col: (
                    dataset_priority.get(str(col.dataset_id), len(dataset_priority)),
                    col.name,
                ),
            )

            if dynamic_columns:
                eval_metrics_needed = []
                for common_name in common_columns:
                    for ds in comparison_datasets:
                        comp_col = Column.objects.get(
                            dataset=ds, name=common_name, deleted=False
                        )
                        # comp_col = columns_lookup.get((ds.id, common_name))
                        if comp_col and comp_col.source in dynamic_sources[1:]:
                            eval_id = (
                                comp_col.source_id.split("-sourceid-")[1]
                                if "-sourceid-" in comp_col.source_id
                                else comp_col.source_id
                            )
                            if eval_id:
                                eval_metrics_needed.append(eval_id)

                eval_metrics = {}
                if eval_metrics_needed:
                    for metric in UserEvalMetric.objects.filter(
                        id__in=eval_metrics_needed, deleted=False
                    ):
                        eval_metrics[str(metric.id)] = metric.status

                cells = Cell.objects.filter(
                    column__in=dynamic_columns,
                    dataset__in=comparison_datasets,
                    deleted=False,
                    row__deleted=False,
                ).all()
                cell_lookup = {
                    (str(cell.dataset_id), cell.column.name, str(cell.row_id)): cell
                    for cell in cells
                }

                columns_qs = Column.objects.filter(
                    name__in=common_columns, dataset__id__in=comparison_datasets
                )
                base_columns = {}
                for col in columns_qs:
                    if col.name == base_column_name:
                        base_columns[str(col.dataset.id)] = col

                for row_data in response["table"]:
                    base_cell_value = row_data[
                        str(base_columns[metadata.get("base_dataset_id")].id)
                    ]["cell_value"]
                    row_info = dataset_info[base_cell_value]
                    base_values = {}
                    for dyn_col in dynamic_columns:
                        current_cell = cell_lookup.get(
                            (
                                str(dyn_col.dataset.id),
                                dyn_col.name,
                                row_info[str(dyn_col.dataset.id)],
                            )
                        )
                        # current_cell = cells.get(dataset=ds, column__name=dyn_col.name, row_id=row_info[str(ds.id)])
                        if current_cell:
                            value_infos = (
                                json.loads(current_cell.value_infos)
                                if current_cell.value_infos
                                else {}
                            )
                            current_cell_metadata = value_infos.get("metadata") or {}
                            if isinstance(metadata, str):
                                current_cell_metadata = json.loads(
                                    current_cell_metadata
                                )

                            if str(dyn_col.dataset.id) == metadata.get(
                                "base_dataset_id"
                            ):
                                base_values[dyn_col.name] = current_cell.value

                            cell_diff = (
                                get_diff(base_values[dyn_col.name], current_cell.value)
                                if str(dyn_col.dataset.id)
                                != metadata.get("base_dataset_id")
                                and dyn_col.source == dynamic_sources[0]
                                else None
                            )
                            row_data[str(dyn_col.id)] = {
                                "cell_value": current_cell.value,
                                "cell_diff_value": cell_diff,
                                "status": current_cell.status,
                                "value_infos": value_infos,
                                "metadata": current_cell_metadata,
                                "cell_row_id": str(current_cell.row_id),
                            }

            return self._gm.success_response(response)

        except Exception as e:
            logger.exception(f"Error in fetching compare dataset row: {str(e)}")
            return self._gm.internal_server_error_response(str(e))

    def thread_delete(self, compare_id):
        if os.path.isdir(f"compare/{compare_id}"):
            shutil.rmtree(f"compare/{compare_id}")
        delete_compare_folder(compare_id)

    def delete(self, request, compare_id, *args, **kwargs):
        try:
            # Import the activity to register it
            import tfc.temporal.background_tasks.activities  # noqa: F401
            from tfc.temporal.drop_in import start_activity

            start_activity(
                "delete_compare_folder_activity",
                args=(str(compare_id),),
                queue="default",
            )
            return self._gm.success_response(
                {"message": "File(s) deleted successfully"}
            )
        except Exception as e:
            logger.exception(f"Error in deleting compare dataset file: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_DELETE_COMPARE_DATASET_FILE")
            )


class CompareDatasetsView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def read_metadata_safely(self, file_path, max_retries=5, retry_delay=0.1):
        """
        Read a JSON file with retry logic to handle concurrent write operations.
        """
        retries = 0
        while retries < max_retries:
            try:
                with open(file_path) as f:
                    return json.load(f)
            except (json.JSONDecodeError, ValueError):
                # If reading fails, wait briefly and try again
                retries += 1
                if retries >= max_retries:
                    # Return a safe default if all retries fail
                    return {
                        "status": "processing",
                        "total_processed": 0,
                        "total_pages": 0,
                    }
                time.sleep(retry_delay)

    def generate_dataset_info(self, comparison_datasets, base_columns):
        try:
            close_old_connections()
            column_map = {ds.id: base_columns[ds.id].id for ds in comparison_datasets}
            column_ids = list(column_map.values())

            # Single query to fetch all relevant cells
            all_cells = Cell.objects.filter(
                column__in=column_ids, deleted=False, row__deleted=False
            ).values_list("column_id", "value", "row_id")

            # Create a mapping of dataset_id -> {value -> row_id}
            dataset_value_map = defaultdict(dict)
            # Also track which datasets have each value
            value_to_datasets = defaultdict(set)

            column_to_dataset = {col_id: ds_id for ds_id, col_id in column_map.items()}

            for column_id, value, row_id in all_cells:
                ds_id = column_to_dataset[column_id]
                dataset_value_map[ds_id][value] = row_id
                value_to_datasets[value].add(ds_id)

            # Only keep values present in all datasets
            num_datasets = len(comparison_datasets)
            common_values = {
                val
                for val, datasets in value_to_datasets.items()
                if len(datasets) == num_datasets
            }

            # Build final dataset_info
            dataset_info = {}
            for value in common_values:
                dataset_info[value] = {}
                for ds_id in dataset_value_map:
                    if value in dataset_value_map[ds_id]:
                        dataset_info[value][str(ds_id)] = str(
                            dataset_value_map[ds_id][value]
                        )

            return dataset_info
        finally:
            close_old_connections()

    def process_other_datasets(
        self,
        base_val,
        col_name,
        og_cell,
        columns_lookup,
        data_by_dataset,
        ds,
        dynamic_sources,
        i,
    ):
        try:
            close_old_connections()
            ds_id = str(ds.id)
            if ds_id in data_by_dataset and base_val in data_by_dataset[ds_id]:
                cells_dict = data_by_dataset[ds_id][base_val]
                cell = cells_dict.get(col_name)
                if cell:
                    col_obj = columns_lookup.get(
                        (ds.id, col_name)
                    ) or Column.objects.get(dataset=ds, name=col_name)
                    if col_obj.source in dynamic_sources:
                        return None
                    value_infos = (
                        json.loads(cell.value_infos) if cell.value_infos else {}
                    )
                    metadata = {}
                    if isinstance(value_infos, dict):
                        metadata = value_infos.get("metadata") or {}
                        if isinstance(metadata, str):
                            metadata = json.loads(metadata)
                    diff_value = None
                    cell_value = cell.value
                    included_sources = [
                        SourceChoices.EXTRACTED_JSON.value,
                        SourceChoices.CLASSIFICATION.value,
                        SourceChoices.EXTRACTED_ENTITIES.value,
                        SourceChoices.API_CALL.value,
                        SourceChoices.PYTHON_CODE.value,
                        SourceChoices.VECTOR_DB.value,
                        SourceChoices.CONDITIONAL.value,
                        SourceChoices.OTHERS.value,
                        SourceChoices.RUN_PROMPT.value,
                    ]
                    if (
                        not (i == 0 or "-reason" in col_name)
                        and (col_obj.source in included_sources)
                        and (
                            col_obj.data_type
                            not in [
                                DataTypeChoices.AUDIO.value,
                                DataTypeChoices.IMAGE.value,
                            ]
                        )
                    ):
                        diff_value = get_diff(og_cell.value, cell.value)
                    return {
                        "col_id": str(col_obj.id),
                        "cell_value": cell_value,
                        "cell_diff_value": diff_value,
                        "status": cell.status,
                        "value_infos": value_infos,
                        "metadata": metadata,
                        "cell_row_id": str(cell.row_id),
                    }
            return None
        finally:
            close_old_connections()

    def process_base_values(
        self, ds, common_base_values, dataset_info, base_column_name, columns_qs
    ):
        try:
            close_old_connections()
            tstart = time.time()
            averages = defaultdict(dict)
            ds_id = str(ds.id)
            # Collect all row IDs for this dataset from dataset_info
            row_ids = [
                dataset_info[base_val][ds_id]
                for base_val in common_base_values
                if ds_id in dataset_info[base_val]
            ]
            if not row_ids:
                return {}, {}

            # col_names = [base_column_name] + list(common_columns)
            columns = columns_qs.filter(dataset=ds).prefetch_related("cell_set")
            cells = Cell.objects.filter(
                row__id__in=row_ids, column__in=columns, deleted=False
            ).select_related("column", "row")

            logger.info(
                f"Time Taken to fetch columns, cells in process_base_values: {time.time() - tstart}"
            )

            # Get column names for each column_id for faster lookup
            column_id_to_name = {col.id: col.name for col in columns}
            base_column_id = None
            for col in columns:
                if col.name == base_column_name:
                    base_column_id = col.id
                    break

            cells_list = list(cells)

            def process_cell_batch(batch, column_id_to_name, base_column_id):
                local_row_to_cells = defaultdict(dict)
                local_row_to_base_val = {}

                for cell in batch:
                    row_id = cell.row_id
                    column_id = cell.column_id
                    value = cell.value
                    column_name = column_id_to_name.get(column_id)

                    local_row_to_cells[row_id][column_name] = cell

                    # Check if this is the base column
                    if column_id == base_column_id:
                        local_row_to_base_val[row_id] = value

                return local_row_to_cells, local_row_to_base_val

            batch_size = max(1, len(cells_list) // 5)

            batches = [
                cells_list[i : i + batch_size]
                for i in range(0, len(cells_list), batch_size)
            ]
            # Group cells by row_id and column
            row_to_cells = defaultdict(dict)
            row_to_base_val = {}

            with ThreadPoolExecutor(max_workers=5) as executor:
                process_func = partial(
                    process_cell_batch,
                    column_id_to_name=column_id_to_name,
                    base_column_id=base_column_id,
                )

                # Submit all batches for processing
                future_results = [
                    executor.submit(process_func, batch) for batch in batches
                ]

                # Collect results as they complete
                for future in as_completed(future_results):
                    try:
                        local_row_to_cells, local_row_to_base_val = future.result()

                        # Merge results from this batch into main dictionaries
                        for row_id, col_names in local_row_to_cells.items():
                            row_to_cells[row_id].update(col_names)

                        row_to_base_val.update(local_row_to_base_val)

                    except Exception as e:
                        logger.error(f"Error processing batch: {str(e)}")

            logger.info(
                f"Time Taken to process cells in process_base_values: {time.time() - tstart}"
            )

            for column in columns:
                avg_result = calculate_column_average(column, row_ids=row_ids)
                averages[(ds_id, column.name)] = avg_result.get("average")

            logger.info(
                f"Time Taken to fetch process columns in process_base_values: {time.time() - tstart}"
            )

            # Map base values to row data
            base_val_to_row = {}
            for _row_id, cells_dict in row_to_cells.items():
                if base_column_name in cells_dict:
                    base_val = cells_dict[base_column_name].value
                    if base_val in common_base_values:
                        base_val_to_row[base_val] = cells_dict

            logger.info(
                f"Time Taken to create base_val_to_row in process_base_values: {time.time() - tstart}"
            )

            # data_by_dataset[ds_id] = base_val_to_row
            return {ds_id: base_val_to_row}, averages
        finally:
            connection.close()
            close_old_connections()

    def prepare_compare_dataset(
        self,
        dataset_id,
        common_base_values,
        base_column_name,
        data_by_dataset,
        comparison_datasets,
        columns_lookup,
        main_base_column,
        common_columns,
        compare_id,
        column_config,
        dataset_info,
        dynamic_sources,
    ):
        try:
            close_old_connections()
            # Build table rows using pre-fetched data
            table = []
            main_ds_id = str(dataset_id)
            with open(f"compare/{compare_id}/metadata.json", "w") as f:
                json.dump(
                    {
                        "status": "processing",
                        "total_rows": len(common_base_values),
                        "total_pages": (len(common_base_values) + 9) // 10,
                        "total_processed": 0,
                        "dataset_info": dataset_info,
                        "base_column_name": base_column_name,
                        "base_dataset_id": main_ds_id,
                        "comparison_datasets": [
                            str(comparison_dataset.id)
                            for comparison_dataset in comparison_datasets
                        ],
                        "common_column_names": list(
                            common_columns.union({base_column_name})
                        ),
                    },
                    f,
                    indent=4,
                )
                f.truncate()

            last_index_processed = None
            rowid_in_file = {}
            for index, base_val in enumerate(common_base_values):
                row_data = {"row_id": str(uuid.uuid4())}
                if (
                    main_ds_id in data_by_dataset
                    and base_val in data_by_dataset[main_ds_id]
                ):
                    main_cells = data_by_dataset[main_ds_id][base_val]
                    main_cell = main_cells.get(base_column_name)
                    if main_cell:
                        value_infos = (
                            json.loads(main_cell.value_infos)
                            if main_cell.value_infos
                            else {}
                        )
                        metadata = value_infos.get("metadata") or {}
                        if isinstance(metadata, str):
                            metadata = json.loads(metadata)
                        row_data[str(main_base_column.id)] = {
                            "cell_value": main_cell.value,
                            "status": main_cell.status,
                            "value_infos": value_infos,
                            "metadata": metadata,
                            "cell_row_id": str(main_cell.row_id),
                        }

                for col_name in common_columns:
                    main_cells = data_by_dataset[main_ds_id][base_val]
                    og_cell = main_cells.get(col_name)

                    with ThreadPoolExecutor(max_workers=10) as executor:
                        futures = []
                        for i, ds in enumerate(comparison_datasets):
                            future = executor.submit(
                                self.process_other_datasets,
                                base_val,
                                col_name,
                                og_cell,
                                columns_lookup,
                                data_by_dataset,
                                ds,
                                dynamic_sources,
                                i,
                            )
                            futures.append(future)

                        for future in as_completed(futures):
                            try:
                                result = future.result()
                                if result is not None:
                                    row_data[result.pop("col_id")] = result
                            except Exception as e:
                                logger.exception(
                                    f"Error in processing other datasets: {str(e)}"
                                )

                table.append(row_data)
                if index % 10 == 0 and index != 0:
                    compare_json = {
                        "column_config": column_config,
                        "table": table[index - 10 : index],
                    }

                    logger.info(f"Writing page {index // 10} to file")
                    rowid_in_file.update(
                        {
                            str(table["row_id"]): f"page_{index // 10}"
                            for table in compare_json["table"]
                        }
                    )
                    upload_compare_json_to_s3(
                        compare_id=compare_id,
                        compare_json=compare_json,
                        page_name=f"page_{index // 10}.json",
                    )

                    with open(f"compare/{compare_id}/metadata.json", "r+") as f:
                        metadata = json.load(f)
                        metadata["total_processed"] = index // 10
                        metadata["file_row_ids"] = rowid_in_file
                        f.seek(0)
                        json.dump(metadata, f, indent=4)
                        f.truncate()  # Important to prevent potential file corruption if new data is shorter

                    last_index_processed = index

            leftover_compare_json = None
            if last_index_processed and last_index_processed != index:
                leftover_compare_json = {
                    "column_config": column_config,
                    "table": table[last_index_processed:],
                }
                rowid_in_file.update(
                    {
                        str(table["row_id"]): f"page_{index // 10 + 1}"
                        for table in leftover_compare_json["table"]
                    }
                )
            else:
                leftover_compare_json = {"column_config": column_config, "table": table}
                rowid_in_file.update(
                    {
                        str(table["row_id"]): f"page_{index // 10 + 1}"
                        for table in leftover_compare_json["table"]
                    }
                )

            if leftover_compare_json:
                upload_compare_json_to_s3(
                    compare_id=compare_id,
                    compare_json=leftover_compare_json,
                    page_name=f"page_{(index // 10) + 1}.json",
                )

            with open(f"compare/{compare_id}/metadata.json", "r+") as f:
                metadata = json.load(f)
                metadata["status"] = "completed"
                metadata["total_processed"] = len(common_base_values) // 10
                metadata["file_row_ids"] = rowid_in_file
                f.seek(0)
                json.dump(metadata, f, indent=4)
                f.truncate()  # Important to prevent potential file corruption if new data is shorter

        except Exception as e:
            logger.exception(f"Error in preparing compare dataset: {str(e)}")
        finally:
            close_old_connections()

    def get_paginated_compare_json(
        self,
        compare_id,
        start,
        end,
        start_page,
        end_page,
        common_columns,
        comparison_datasets,
        columns_lookup,
        dataset_id,
        columns_qs,
        common_base_values,
        dataset_info,
        base_column_name,
        dynamic_sources,
        result,
    ):
        time_start = time.time()
        result.setdefault("table", [])
        page_range = range(start_page, end_page)

        # Dictionary to store future results keyed by page number
        results_by_page: dict[Any, Any] = {}
        args_list = [
            (compare_id, page, start, end, start_page, end_page) for page in page_range
        ]
        max_workers = min(5, len(page_range))

        def fetch_page(args):
            compare_id, page, start, end, start_page, end_page = args
            object_key = f"compare/{compare_id}/page_{page}.json"
            json_data = download_json_from_s3(object_key)

            start_idx = start % 10 if page == start_page else 0
            end_idx = (
                (end % 10 if end % 10 != 0 else 10) if page == end_page - 1 else 10
            )
            sliced_table = json_data.get("table", [])[start_idx:end_idx]

            return page, sliced_table, json_data.get("column_config")

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {executor.submit(fetch_page, args): args[1] for args in args_list}

            for future in as_completed(futures):
                page, table_slice, column_config = future.result()
                results_by_page[page] = (table_slice, column_config)

        # Re-assemble results in page order
        for page in sorted(results_by_page.keys()):
            table_slice, column_config = results_by_page[page]
            result["table"].extend(table_slice)
            if "column_config" not in result and column_config:
                result["column_config"] = column_config

        logger.info(
            f"Time taken for extracting paginated data: {time.time() - time_start}"
        )

        dynamic_columns = list(
            Column.objects.filter(
                name__in=common_columns,
                dataset__in=comparison_datasets,
                source__in=dynamic_sources,
            ).select_related("dataset")
        )

        dataset_priority = {ds.id: i for i, ds in enumerate(comparison_datasets)}

        dynamic_columns = sorted(
            dynamic_columns,
            key=lambda col: (
                dataset_priority.get(col.dataset_id, len(dataset_priority)),
                col.name,
            ),
        )
        if dynamic_columns:
            base_ds = comparison_datasets[0]
            base_columns = {}
            for col in columns_qs:
                if col.name == base_column_name:
                    base_columns[str(col.dataset_id)] = col
            col_conf = result["column_config"]

            eval_metrics_needed = []
            for common_name in common_columns:
                for ds in comparison_datasets:
                    comp_col = columns_lookup.get((ds.id, common_name))
                    if comp_col and comp_col.source in dynamic_sources[1:]:
                        eval_id = (
                            comp_col.source_id.split("-sourceid-")[1]
                            if "-sourceid-" in comp_col.source_id
                            else comp_col.source_id
                        )
                        if eval_id:
                            eval_metrics_needed.append(eval_id)

            eval_metrics = {}
            if eval_metrics_needed:
                for metric in UserEvalMetric.objects.filter(
                    id__in=eval_metrics_needed, deleted=False
                ):
                    eval_metrics[str(metric.id)] = metric.status

            averages = {}
            max_workers = min(5, len(comparison_datasets))
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures = []

                for ds in comparison_datasets:
                    future = executor.submit(
                        self.process_base_values,
                        ds,
                        common_base_values,
                        dataset_info,
                        base_column_name,
                        columns_qs,
                    )
                    futures.append(future)

                for future in as_completed(futures):
                    check, avg_result = future.result()
                    if check is not None:
                        averages.update(avg_result)

            cells = Cell.objects.filter(
                column__in=dynamic_columns,
                dataset__in=comparison_datasets,
                deleted=False,
                row__deleted=False,
            ).all()
            cell_lookup = {
                (cell.dataset_id, cell.column.name, str(cell.row_id)): cell
                for cell in cells
            }
            for dyn in dynamic_columns:
                status = dyn.status
                if (
                    dyn.source == dynamic_sources[0]
                    and RunPrompter.objects.filter(
                        id=dyn.source_id, deleted=False
                    ).exists()
                ):
                    status = RunPrompter.objects.get(id=dyn.source_id).status
                else:
                    eval_id = (
                        dyn.source_id.split("-sourceid-")[1]
                        if dyn.source_id and "-sourceid-" in dyn.source_id
                        else dyn.source_id
                    )
                    if eval_id and eval_id in eval_metrics:
                        status = eval_metrics[eval_id]
                main_common_col = columns_lookup.get((dataset_id, dyn.name))
                if not main_common_col:
                    main_common_col = columns_qs.get(
                        dataset_id=dataset_id, name=dyn.name
                    )
                avg_score = averages.get((str(dyn.dataset.id), dyn.name), None)
                col_conf.append(
                    {
                        "id": str(dyn.id),
                        "name": dyn.dataset.name,
                        "data_type": dyn.data_type,
                        "origin_type": dyn.source,
                        "dataset_id": str(dyn.dataset.id),
                        "status": status,
                        "source_id": dyn.source_id,
                        "group": {
                            "id": str(main_common_col.id),
                            "name": main_common_col.name,
                            "data_type": main_common_col.data_type,
                            "origin": (
                                "Evaluation"
                                if main_common_col.source in dynamic_sources[1:]
                                else "Dataset"
                            ),
                        },
                        "average_score": avg_score,
                    }
                )
                base_values = {}
                for row_data in result["table"]:
                    base_cell_value = row_data[
                        str(base_columns[str(comparison_datasets[0].id)].id)
                    ]["cell_value"]
                    row_info = dataset_info[base_cell_value]
                    og_cell = cell_lookup.get(
                        (base_ds.id, dyn.name, row_info.get(str(base_ds.id)))
                    )
                    if og_cell:
                        base_values[row_info.get(str(base_ds.id))] = og_cell.value

                    current_cell = cell_lookup.get(
                        (dyn.dataset.id, dyn.name, row_info.get(str(dyn.dataset.id)))
                    )

                    if current_cell:
                        # Process the cell value and metadata
                        value_infos = {}
                        if current_cell.value_infos:
                            try:
                                value_infos = json.loads(current_cell.value_infos)
                            except (json.JSONDecodeError, TypeError):
                                value_infos = {}

                        current_cell_metadata = value_infos.get("metadata", {})
                        if isinstance(current_cell_metadata, str):
                            try:
                                current_cell_metadata = json.loads(
                                    current_cell_metadata
                                )
                            except (json.JSONDecodeError, TypeError):
                                current_cell_metadata = {}

                        # Calculate cell_diff only for non-base datasets and only for certain columns
                        cell_diff = None
                        if (
                            dyn.dataset != base_ds
                            and dyn.source == dynamic_sources[0]
                            and row_info.get(str(base_ds.id)) in base_values
                        ):
                            cell_diff = get_diff(
                                base_values[row_info.get(str(base_ds.id))],
                                current_cell.value,
                            )

                        # Set the cell data
                        row_data[str(dyn.id)] = {
                            "cell_value": current_cell.value,
                            "cell_diff_value": cell_diff,
                            "status": current_cell.status,
                            "value_infos": value_infos,
                            "metadata": current_cell_metadata,
                            "cell_row_id": str(current_cell.row_id),
                        }

            result["column_config"] = col_conf

        return result

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            start_time = time.time()
            # Pagination parameters
            serializer = CompareDatasetSerializer(data=request.data)
            if not serializer.is_valid():
                return self._gm.bad_request(serializer.errors)

            validated_data = serializer.validated_data
            page_size = validated_data["page_size"]
            current_page = validated_data["current_page_index"]
            base_column_name = validated_data["base_column_name"]
            compare_id = validated_data["compare_id"]

            if not compare_id:
                compare_id = str(uuid.uuid4())
            start = current_page * page_size
            end = start + page_size

            os.makedirs(f"compare/{compare_id}/", exist_ok=True)

            dynamic_sources = [
                SourceChoices.RUN_PROMPT.value,
                SourceChoices.EXPERIMENT_EVALUATION_TAGS.value,
                SourceChoices.EVALUATION_TAGS.value,
                SourceChoices.OPTIMISATION_EVALUATION_TAGS.value,
                SourceChoices.EVALUATION.value,
                SourceChoices.EXPERIMENT_EVALUATION.value,
                SourceChoices.OPTIMISATION_EVALUATION.value,
                SourceChoices.EVALUATION_REASON.value,
            ]

            num_pages_to_fetch = (page_size + 9) // 10
            start_page = math.ceil((start + 1) / 10)
            end_page = start_page + num_pages_to_fetch

            if os.path.exists(f"compare/{compare_id}/metadata.json"):
                metadata = self.read_metadata_safely(
                    f"compare/{compare_id}/metadata.json"
                )
                if metadata.get("total_pages") < num_pages_to_fetch:
                    return self._gm.bad_request(
                        get_error_message("FAILED_TO_COMPARE_DATASETS")
                    )

                loop_start = time.time()
                while (
                    metadata.get("status") == "processing"
                    and metadata.get("total_processed") >= start_page
                ):
                    logger.info("waiting for files to be created")
                    metadata = self.read_metadata_safely(
                        f"compare/{compare_id}/metadata.json"
                    )
                    if time.time() - loop_start > 300:
                        return self._gm.bad_request(
                            get_error_message("FAILED_TO_COMPARE_DATASETS")
                        )
                    time.sleep(0.1)
                result = {}
                result["metadata"] = {
                    "compare_id": compare_id,
                    "total_rows": int(metadata.get("total_rows")),
                    "total_pages": (int(metadata.get("total_rows")) + page_size - 1)
                    // page_size,
                }

                dataset_ids = validated_data["dataset_ids"]
                if not dataset_ids:
                    return self._gm.bad_request("No dataset IDs provided.")
                ordered_ids = [dataset_id] + [str(did) for did in dataset_ids]

                # Create the Case/When expressions for ordering
                preserved_order = Case(
                    *[When(id=id, then=pos) for pos, id in enumerate(ordered_ids)],
                    output_field=IntegerField(),
                )

                datasets = Dataset.objects.filter(
                    id__in=ordered_ids, deleted=False
                ).order_by(preserved_order)
                if len(datasets) != len(dataset_ids) + 1:
                    return self._gm.bad_request(
                        get_error_message("INVALID_DATASET_IDS")
                    )

                comparison_datasets = list(datasets)
                columns_qs = (
                    Column.objects.filter(
                        dataset__in=comparison_datasets, deleted=False
                    )
                    .select_related("dataset")
                    .prefetch_related("cell_set")
                    .all()
                )

                averages = {}
                max_workers = min(5, len(comparison_datasets))
                dataset_info = metadata.get("dataset_info")
                common_base_values = list(dataset_info.keys())
                columns_qs = (
                    Column.objects.filter(
                        dataset__in=comparison_datasets, deleted=False
                    )
                    .select_related("dataset")
                    .prefetch_related("cell_set")
                    .all()
                )
                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    futures = []

                    for ds in comparison_datasets:
                        future = executor.submit(
                            self.process_base_values,
                            ds,
                            common_base_values,
                            dataset_info,
                            base_column_name,
                            columns_qs,
                        )
                        futures.append(future)

                    for future in as_completed(futures):
                        check, avg_result = future.result()
                        if check is not None:
                            averages.update(avg_result)

                columns_by_dataset = {
                    dataset.id: set() for dataset in comparison_datasets
                }
                columns_lookup = {}
                base_columns = {}
                for col in columns_qs:
                    if col.name == base_column_name:
                        base_columns[str(col.dataset_id)] = col
                    else:
                        columns_by_dataset[col.dataset_id].add(col.name)
                        columns_lookup[(col.dataset_id, col.name)] = col

                common_columns = (
                    set.intersection(*columns_by_dataset.values())
                    if columns_by_dataset
                    else set()
                )

                result = self.get_paginated_compare_json(
                    compare_id,
                    start,
                    end,
                    start_page,
                    end_page,
                    common_columns,
                    comparison_datasets,
                    columns_lookup,
                    dataset_id,
                    columns_qs,
                    common_base_values,
                    dataset_info,
                    base_column_name,
                    dynamic_sources,
                    result,
                )

                return self._gm.success_response(result)

            # Check if dataset_info and common_column_names are provided
            dataset_info = validated_data["dataset_info"]
            common_column_names = validated_data["common_column_names"]

            if dataset_info and common_column_names:
                # Use provided data
                dataset_ids = list(
                    {
                        ds_id
                        for base_val_dict in dataset_info.values()
                        for ds_id in base_val_dict.keys()
                    }
                )
                if str(dataset_id) not in dataset_ids:
                    dataset_ids.insert(0, str(dataset_id))

                common_columns = set(common_column_names)
            else:
                dataset_ids = validated_data["dataset_ids"]
                if not dataset_ids:
                    return self._gm.bad_request("No dataset IDs provided.")

            # list dataset ids in order main_datset id then other dataset ids
            ordered_ids = [dataset_id] + [str(did) for did in dataset_ids]

            # Create the Case/When expressions for ordering
            preserved_order = Case(
                *[When(id=id, then=pos) for pos, id in enumerate(ordered_ids)],
                output_field=IntegerField(),
            )

            # Apply the filter and ordering
            datasets = Dataset.objects.filter(
                id__in=ordered_ids, deleted=False
            ).order_by(preserved_order)
            for dst in datasets:
                logger.info(f"Datasets: {dst.id}  {dst.name}")
            if len(datasets) != len(dataset_ids) + 1:
                return self._gm.bad_request(get_error_message("INVALID_DATASET_IDS"))

            comparison_datasets = list(datasets)
            columns_qs = (
                Column.objects.filter(dataset__in=comparison_datasets, deleted=False)
                .select_related("dataset")
                .prefetch_related("cell_set")
                .all()
            )

            columns_by_dataset = {dataset.id: set() for dataset in comparison_datasets}
            columns_lookup = {}
            base_columns = {}
            for col in columns_qs:
                if col.name == base_column_name:
                    base_columns[col.dataset_id] = col
                else:
                    columns_by_dataset[col.dataset_id].add(col.name)
                    columns_lookup[(col.dataset_id, col.name)] = col

            common_columns = (
                set.intersection(*columns_by_dataset.values())
                if columns_by_dataset
                else set()
            )

            if len(base_columns) != len(comparison_datasets):
                return self._gm.bad_request(
                    f"Base column {base_column_name} not present in provided datasets"
                )

            logger.info(
                f"Time taken to fetch base columns: {time.time() - start_time} seconds"
            )
            # If dataset_info not provided, find common base values and corresponding row IDs
            if not dataset_info:
                dataset_info = self.generate_dataset_info(
                    comparison_datasets, base_columns
                )

            logger.info(
                f"Time taken to fetch dataset info: {time.time() - start_time} seconds"
            )

            common_base_values = set(dataset_info.keys())

            # Pre-fetch all relevant cells in bulk for each dataset
            data_by_dataset: dict[Any, Any] = {}
            averages = {}
            max_workers = min(5, len(comparison_datasets))
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures = []

                for ds in comparison_datasets:
                    future = executor.submit(
                        self.process_base_values,
                        ds,
                        common_base_values,
                        dataset_info,
                        base_column_name,
                        columns_qs,
                    )
                    futures.append(future)

                for future in as_completed(futures):
                    result, avg_result = future.result()
                    if result is not None:
                        data_by_dataset.update(result)
                        averages.update(avg_result)

            logger.info(
                f"Time taken to get averages: {time.time() - start_time} seconds"
            )

            # Build column configuration using precomputed averages
            column_config = []
            main_base_column = base_columns[dataset_id]
            base_avg_score = averages.get((str(dataset_id), base_column_name), None)
            column_config.append(
                {
                    "id": str(main_base_column.id),
                    "name": main_base_column.name,
                    "data_type": main_base_column.data_type,
                    "origin_type": main_base_column.source,
                    "dataset_id": str(dataset_id),
                    "status": main_base_column.status,
                    "source_id": main_base_column.source_id,
                    "group": {
                        "id": str(main_base_column.id),
                        "name": main_base_column.name,
                        "data_type": main_base_column.data_type,
                        "origin": "Dataset",
                    },
                    "average_score": base_avg_score,
                }
            )

            # Build common column configurations
            # logger.info(f"eval_metrics {eval_metrics}")
            for common_name in common_columns:
                # Get the main column for this common name just once
                main_common_col = columns_lookup.get((dataset_id, common_name))
                if not main_common_col:
                    main_common_col = columns_qs.get(
                        dataset_id=dataset_id, name=common_name
                    )

                for ds in comparison_datasets:
                    ds_id = str(ds.id)
                    comp_col = columns_lookup.get((ds.id, common_name))
                    if not comp_col:
                        comp_col = columns_qs.get(dataset=ds, name=common_name)

                    avg_score = averages.get((ds_id, common_name), None)

                    if comp_col.source in dynamic_sources:
                        continue

                    column_config.append(
                        {
                            "id": str(comp_col.id),
                            "name": ds.name,
                            "data_type": comp_col.data_type,
                            "origin_type": comp_col.source,
                            "dataset_id": str(ds.id),
                            "status": comp_col.status,
                            "source_id": comp_col.source_id,
                            "group": {
                                "id": str(main_common_col.id),
                                "name": main_common_col.name,
                                "data_type": main_common_col.data_type,
                                "origin": "Dataset",
                            },
                            "average_score": avg_score,
                        }
                    )

            logger.info(
                f"Time taken to prepare column config: {time.time() - start_time} seconds"
            )
            compare_json = {}

            if len(common_base_values) > 0:
                # Use ThreadPoolExecutor directly for complex objects that can't be serialized
                import concurrent.futures

                with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                    executor.submit(
                        _prepare_compare_dataset_impl,
                        dataset_id,
                        common_base_values,
                        base_column_name,
                        data_by_dataset,
                        comparison_datasets,
                        columns_lookup,
                        main_base_column,
                        common_columns,
                        compare_id,
                        column_config,
                        dataset_info,
                        dynamic_sources,
                    )

                loop_start = time.time()
                while not os.path.exists(f"compare/{compare_id}/metadata.json"):
                    logger.info("waiting for metadata file to be created")
                    if time.time() - loop_start > 300:
                        return self._gm.bad_request(
                            get_error_message("FAILED_TO_COMPARE_DATASETS")
                        )
                    time.sleep(0.1)
                metadata = self.read_metadata_safely(
                    f"compare/{compare_id}/metadata.json"
                )

                loop_start = time.time()
                while (not metadata.get("status") == "completed") and (
                    not metadata.get("total_processed") >= end_page - 1
                ):
                    if time.time() - loop_start > 300:
                        return self._gm.bad_request(
                            get_error_message("FAILED_TO_COMPARE_DATASETS")
                        )
                    metadata = self.read_metadata_safely(
                        f"compare/{compare_id}/metadata.json"
                    )
                    logger.info("waiting for metadata file to be updated")
                    time.sleep(0.1)

                logger.info(
                    f"Time taken to get actual data: {time.time() - start_time} seconds"
                )

                compare_json = self.get_paginated_compare_json(
                    compare_id,
                    start,
                    end,
                    start_page,
                    end_page,
                    common_columns,
                    comparison_datasets,
                    columns_lookup,
                    dataset_id,
                    columns_qs,
                    common_base_values,
                    dataset_info,
                    base_column_name,
                    dynamic_sources,
                    compare_json,
                )

            compare_json.update(
                {
                    "metadata": {
                        "compare_id": compare_id,
                        "total_rows": (
                            int(metadata.get("total_rows", 0))
                            if "metadata" in locals()
                            else 0
                        ),
                        "total_pages": (
                            (int(metadata.get("total_rows", 0)) + page_size - 1)
                            // page_size
                            if "metadata" in locals()
                            else 0
                        ),
                    }
                }
            )

            return self._gm.success_response(compare_json)

        except Exception as e:
            logger.exception(f"Error in comparing the datasets: {str(e)}")
            # return self._gm.bad_request(get_error_message('FAILED_TO_COMPARE_DATASETS'))
            return self._gm.bad_request(str(e))


class DownloadComparisonDatasetView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            # Get main dataset and validate request data
            get_object_or_404(Dataset, id=dataset_id, deleted=False)
            serializer = CompareDatasetSerializer(data=request.data)
            compare_id = request.data.get("compare_id", None)

            if not serializer.is_valid():
                return self._gm.bad_request(serializer.errors)

            validated_data = serializer.validated_data
            base_column_name = validated_data["base_column_name"]
            dataset_ids = validated_data["dataset_ids"]
            if compare_id:
                if os.path.exists(f"compare/{compare_id}"):
                    metadata = {"status ": "processing"}
                    loop_start = time.time()
                    while not metadata.get("status") == "completed":
                        with open(f"compare/{compare_id}/metadata.json") as f:
                            metadata = json.load(f)
                        if time.time() - loop_start > 300:
                            return self._gm.bad_request(
                                get_error_message("FAILED_TO_DOWNLOAD_DATASET")
                            )
                        time.sleep(0.2)

                    dataset_info = metadata.get("dataset_info")
                    common_column_names = metadata.get("common_column_names")

            if not base_column_name:
                return self._gm.bad_request(
                    get_error_message("FAILED_TO_GET_BASE_COLUMNS")
                )

            # Ensure main dataset is included in comparison
            if str(dataset_id) not in dataset_ids:
                dataset_ids.insert(0, str(dataset_id))
            comparison_datasets = [
                get_object_or_404(Dataset, id=ds_id, deleted=False)
                for ds_id in dataset_ids
            ]

            if not dataset_info or not common_column_names:
                return self._gm.bad_request(
                    get_error_message("FAILED_TO_DOWNLOAD_DATASET")
                )

            common_base_values = []
            # if not dataset_info:
            #     return self._gm.bad_request(get_error_message("NO_DATASET_INFO_PROVIDED"))

            common_base_values = list(dataset_info.keys())

            # if not common_base_values:
            #     return self._gm.bad_request("No common base values found.")

            # Prepare the data for CSV
            data = defaultdict(list)

            # For each common base value
            for base_val in common_base_values:
                row_data = {base_column_name: base_val}

                # For each dataset, fetch the cells for this base value
                for ds in comparison_datasets:
                    ds_id = str(ds.id)

                    # Get row_id from dataset_info if available
                    row_id = None
                    if base_val in dataset_info and ds_id in dataset_info[base_val]:
                        row_id = dataset_info[base_val][ds_id]

                    # If we have a row_id, fetch cells directly
                    if row_id:
                        # Fetch all common columns for this row in one query
                        cells = Cell.objects.filter(
                            row_id=row_id,
                            column__name__in=common_column_names,
                            deleted=False,
                        ).select_related("column")

                        for cell in cells:
                            col_name = cell.column.name + "-" + ds.name
                            row_data[col_name] = cell.value

                # Add this row to our data
                for col in row_data:
                    data[col].append(row_data.get(col, ""))

            # Convert to pandas DataFrame
            df = pd.DataFrame(data)

            # Create CSV in-memory
            buffer = io.BytesIO()
            df.to_csv(buffer, index=False, encoding="utf-8")
            buffer.seek(0)

            # Create the response with the file
            filename = f"dataset_comparison_{dataset_id}.csv"
            response = FileResponse(
                buffer, as_attachment=True, filename=filename, content_type="text/csv"
            )
            return response

        except Exception as e:
            logger.exception(f"Error in downloading the dataset: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_DOWNLOAD_DATASET")
            )


class CompareDatasetsStatsView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            # Extract parameters from the request

            base_column_name = request.data.get("base_column_name")
            dataset_ids = request.data.get("dataset_ids", [])
            stat_type = request.data.get("stat_type", "evaluation")

            if str(dataset_id) not in dataset_ids:
                dataset_ids.insert(0, str(dataset_id))

            base_columns = Column.objects.filter(
                dataset__in=dataset_ids, name=base_column_name, deleted=False
            ).values_list("id", flat=True)

            if not base_columns or len(base_columns) == 0:
                return self._gm.bad_request(
                    f"Base column {base_column_name} not found in any of the datasets"
                )

            if len(base_columns) != len(dataset_ids):
                return self._gm.bad_request(
                    f"Base column {base_column_name} not found in all datasets"
                )

            cells_qs = Cell.objects.filter(
                dataset_id__in=dataset_ids,
                column_id__in=base_columns,
                status=CellStatus.PASS.value,
                deleted=False,
            ).select_related("dataset", "column")

            # Group cells by dataset_id
            cells_by_value: dict[Any, Any] = {}
            row_ids_by_value: dict[Any, Any] = {}
            row_ids = []
            for cell in cells_qs:
                if str(cell.value) not in cells_by_value:
                    cells_by_value[str(cell.value)] = []
                    row_ids_by_value[str(cell.value)] = set()
                cells_by_value[str(cell.value)].append(cell)
                row_ids_by_value[str(cell.value)].add(str(cell.row_id))

            for _value, cells in cells_by_value.items():
                length_of_datasets = len(dataset_ids)
                length_of_cells = len(cells)

            if length_of_cells % length_of_datasets == 0:
                row_ids_set = row_ids_by_value[str(cell.value)]
                row_ids.extend(list(row_ids_set))

            match stat_type:
                case "evaluation":
                    template_ids = UserEvalMetric.objects.filter(
                        dataset_id__in=dataset_ids,
                        deleted=False,
                        template__deleted=False,
                    ).values_list("template_id", flat=True)
                    templates = EvalTemplate.no_workspace_objects.filter(
                        id__in=template_ids, deleted=False
                    )

                    response = {}
                    for id in dataset_ids:
                        final_data = []

                        with ThreadPoolExecutor(max_workers=10) as executor:
                            results = list(
                                executor.map(
                                    lambda template: get_eval_stats(
                                        template,
                                        id,
                                        None,
                                        row_ids,  # noqa: B023
                                    ),
                                    templates,
                                )
                            )
                            final_data.extend(results)

                        response[id] = final_data

                    return self._gm.success_response(response)

                case "run_prompt":
                    run_prompters = RunPrompter.objects.filter(
                        dataset_id__in=dataset_ids, deleted=False
                    )

                    response = {}
                    for id in dataset_ids:
                        final_data = get_prompt_stats(run_prompters, id, row_ids)
                        response[id] = final_data

                    return self._gm.success_response(response)

                case _:
                    return self._gm.bad_request(get_error_message("INVALID_STAT_TYPE"))

        except Exception as e:
            logger.exception(f"Error in comparing the datasets stats: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_COMPARE_DATASETS_STATS")
            )


class AddCompareExperimentEvalView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, dataset_id, *args, **kwargs):
        organization = (
            getattr(request, "organization", None) or request.user.organization
        )

        serializer = UserEvalSerializer(data=request.data)
        run = request.data.get("run", False)
        save_as_template = request.data.get("save_as_template", False)
        if serializer.is_valid():
            validated_data = serializer.validated_data
            dataset_ids = request.data.get("dataset_ids", [])

            if str(dataset_id) not in dataset_ids:
                dataset_ids.append(str(dataset_id))

            template_id = validated_data.get("template_id")
            (validated_data.get("config", {}).get("mapping", {}).get("input"))
            # for dataset_id in dataset_ids:
            #     for key in validated_data.get("config", {}).get("mapping", {}).keys():
            #         validated_data["config"]["mapping"][key] = str(Column.objects.filter(dataset_id=dataset_id, name=validated_data.get("config", {}).get("mapping", {}).get(key), deleted=False).id)

            if UserEvalMetric.objects.filter(
                name=validated_data.get("name"),
                organization=organization,
                dataset_id__in=dataset_ids,
                deleted=False,
            ).exists():
                return self._gm.bad_request(get_error_message("EVAL_NAME_EXISTS"))

            if save_as_template:
                if (
                    EvalTemplate.objects.filter(
                        name=validated_data.get("name"),
                        organization=organization,
                        deleted=False,
                    ).exists()
                    or EvalTemplate.no_workspace_objects.filter(
                        name=validated_data.get("name"),
                        owner=OwnerChoices.SYSTEM.value,
                        deleted=False,
                    ).exists()
                ):
                    return self._gm.bad_request(get_error_message("EVAL_NAME_EXISTS"))

                template = EvalTemplate.no_workspace_objects.get(
                    id=validated_data.get("template_id")
                )
                new_template = EvalTemplate(
                    name=validated_data.get("name"),
                    description=template.description,
                    config=template.config,
                    eval_tags=template.eval_tags,
                    organization=getattr(request, "organization", None)
                    or request.user.organization,
                    owner=OwnerChoices.USER.value,
                    criteria=template.criteria,
                    choices=template.choices,
                    multi_choice=template.multi_choice,
                )
                new_config = template.config
                try:
                    runtime_config = normalize_eval_runtime_config(
                        template.config, validated_data.get("config", {})
                    )
                except ValueError as e:
                    return self._gm.bad_request(str(e))
                input_config = runtime_config.get("config", {})
                input_params = runtime_config.get("params", {})
                for key in input_config:
                    if key in new_config.get("config", {}):
                        new_config["config"][key]["default"] = input_config[key]
                if has_function_params_schema(new_config):
                    for key, value in input_params.items():
                        if key in new_config.get("function_params_schema", {}):
                            new_config["function_params_schema"][key]["default"] = value

                new_template.config = new_config
                new_template.save()
                template_id = new_template.id

            # Create a deep copy of the config to avoid modifying the original across iterations
            selected_template = EvalTemplate.no_workspace_objects.get(id=template_id)
            try:
                normalized_config = normalize_eval_runtime_config(
                    selected_template.config, validated_data.get("config", {})
                )
            except ValueError as e:
                return self._gm.bad_request(str(e))
            original_config = copy.deepcopy(normalized_config)

            # Helper to extract base column name and JSON path from column_name.path format
            def extract_column_name_and_path(column_name):
                if not column_name or "." not in column_name:
                    return column_name, None
                dot_index = column_name.find(".")
                return column_name[:dot_index], column_name[dot_index + 1 :]

            for dataset_id in set(dataset_ids):
                config = copy.deepcopy(original_config)

                for key in config.get("mapping", {}).keys():
                    column_name = config["mapping"][key]
                    # Handle JSON paths (e.g., "Input.prompt" -> "column_id.prompt")
                    base_column_name, json_path = extract_column_name_and_path(
                        column_name
                    )
                    column_obj = Column.objects.filter(
                        dataset_id=dataset_id,
                        name=base_column_name,
                        deleted=False,
                    ).first()
                    if column_obj:
                        # Append JSON path to column ID if present
                        config["mapping"][key] = (
                            f"{column_obj.id}.{json_path}"
                            if json_path
                            else str(column_obj.id)
                        )
                    else:
                        return self._gm.bad_request(
                            get_error_message("COLUMN_NOT_FOUND") + column_name
                        )

                user_eval_metric = UserEvalMetric.objects.create(
                    name=validated_data.get("name"),
                    organization=organization,
                    dataset_id=dataset_id,
                    template_id=template_id,
                    config=config,
                    status=(
                        StatusType.NOT_STARTED.value
                        if run
                        else StatusType.INACTIVE.value
                    ),
                    user=request.user,
                    model=validated_data.get("model", ModelChoices.TURING_LARGE.value),
                    composite_weight_overrides=validated_data.get(
                        "composite_weight_overrides"
                    ),
                )

                if run:
                    with transaction.atomic():
                        # Use no_workspace_objects manager to avoid the outer join issue with select_for_update
                        dataset = Dataset.no_workspace_objects.select_for_update().get(
                            id=dataset_id
                        )
                        data_type = infer_eval_result_column_data_type(
                            user_eval_metric.template
                        )

                        column = Column.objects.create(
                            name=validated_data.get("name"),
                            data_type=data_type,
                            source=SourceChoices.EVALUATION.value,
                            dataset=dataset,
                            source_id=user_eval_metric.id,
                        )
                        column_order = dataset.column_order
                        column_order.append(str(column.id))
                        dataset.column_order = column_order
                        dataset.save()

                        if validated_data.get("config").get("reason_column"):
                            reason_column, created = Column.objects.get_or_create(
                                name=f"{user_eval_metric.name}-reason",
                                data_type=DataTypeChoices.TEXT.value,
                                source=SourceChoices.EVALUATION_REASON.value,
                                dataset=dataset,
                                source_id=f"{column.id}-sourceid-{user_eval_metric.id}",
                            )
                            column_order.append(str(reason_column.id))
                            dataset.column_order = column_order
                            dataset.save()

            return self._gm.success_response("success")
        return self._gm.bad_request(parse_serialized_errors(serializer))


class CompareDatasetsStartEvalsProcess(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, dataset_id, *args, **kwargs):
        try:
            user_eval_names = request.data.get("user_eval_names", [])
            dataset_ids = request.data.get("dataset_ids", [])

            if not user_eval_names:
                return self._gm.bad_request(get_error_message("MISSSING_EVAL_IDS"))

            if str(dataset_id) not in dataset_ids:
                dataset_ids.append(str(dataset_id))

            # Fetch all UserEvalMetrics in one query
            user_eval_metrics = UserEvalMetric.objects.filter(
                name__in=user_eval_names,
                dataset_id__in=dataset_ids,
                organization=getattr(request, "organization", None)
                or request.user.organization,
                deleted=False,
                template__deleted=False,
            ).select_related("dataset")

            # Validate metrics existence
            if user_eval_metrics.count() % len(dataset_ids) != 0:
                return self._gm.bad_request(get_error_message("EVALS_NOT_FOUND"))

            # Check for deleted columns
            for metric in user_eval_metrics:
                if metric.column_deleted:
                    return self._gm.bad_request(
                        f"{get_error_message('COLUMN_DELETED')} {metric.name}"
                    )

            # Extract user eval IDs and unique dataset IDs
            user_eval_ids = [metric.id for metric in user_eval_metrics]
            {metric.dataset_id for metric in user_eval_metrics}

            # Bulk update UserEvalMetric status
            updated_count = UserEvalMetric.objects.filter(id__in=user_eval_ids).update(
                status=StatusType.NOT_STARTED.value
            )

            if updated_count != len(user_eval_ids):
                return self._gm.bad_request(get_error_message("EVALS_NOT_FOUND"))

            # Bulk update Cell status for evaluation columns
            # Cell.objects.filter(
            #     dataset_id__in=unique_dataset_ids,
            #     column__source_id__in=user_eval_ids,
            #     deleted=False
            # ).update(status=CellStatus.RUNNING.value)

            # Process each metric for column creation
            for metric in user_eval_metrics:
                data_type = infer_eval_result_column_data_type(metric.template)

                # Get or create evaluation column
                column, created = Column.objects.get_or_create(
                    source_id=metric.id,
                    defaults={
                        "name": f"{metric.name}",
                        "data_type": data_type,
                        "source": SourceChoices.EVALUATION.value,
                        "dataset": metric.dataset,
                    },
                )

                if created:
                    # Update dataset column order
                    metric.dataset.column_order.append(str(column.id))
                    metric.dataset.save()

                    # Create reason column
                    reason_column, _ = Column.objects.create(
                        name=f"{metric.name}-reason",
                        data_type=DataTypeChoices.TEXT.value,
                        source=SourceChoices.EVALUATION_REASON.value,
                        dataset=metric.dataset,
                        source_id=f"{column.id}-sourceid-{metric.id}",
                    )
                    metric.dataset.column_order.append(str(reason_column.id))
                    metric.dataset.save()

                # Update Cell status for reason column
                Cell.objects.filter(
                    column__source_id=f"{column.id}-sourceid-{metric.id}", deleted=False
                ).update(status=CellStatus.RUNNING.value)

            return self._gm.success_response(
                f"Successfully updated {updated_count} eval(s) status"
            )

        except Exception as e:
            logger.exception(f"Error in starting evaluation process: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_START_EVAL_PROCESS")
            )


class GetCompareEvalsListView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        try:
            search_text = request.data.get("search_text", "").strip()
            eval_type = request.data.get("eval_type")
            dataset_ids = request.data.get("dataset_ids", [])
            if not eval_type or eval_type != "user":
                return self._gm.bad_request(
                    get_error_message("INVALID_OR_MISSING_EVAL_TYPE")
                )

            if eval_type == "user":
                if not dataset_ids:
                    return self._gm.bad_request(get_error_message("MISSING_DATASET_ID"))
                return self._get_user_evals(
                    request, dataset_ids, search_text=search_text
                )

        except Exception as e:
            logger.exception(f"Error in fetching the eval lists: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_EVAL_LISTS")
            )

    def _get_user_evals(self, request, dataset_ids, search_text):
        # Fetch all relevant UserEvalMetrics in one query with preloaded templates
        user_evals = UserEvalMetric.objects.filter(
            dataset_id__in=dataset_ids,
            show_in_sidebar=True,
            organization=getattr(request, "organization", None)
            or request.user.organization,
            deleted=False,
            template__deleted=False,
        ).select_related("template")

        if search_text:
            user_evals = user_evals.filter(name__icontains=search_text)

        # Count occurrences of eval names across datasets
        eval_name_count = defaultdict(int)
        for eval in user_evals:
            eval_name_count[eval.name] += 1

        # Identify eval names present in all datasets
        common_eval_names = {
            name for name, count in eval_name_count.items() if count == len(dataset_ids)
        }

        # Filter evaluations to only those with common names (in memory)
        user_evals_list = [
            eval for eval in user_evals if eval.name in common_eval_names
        ]

        # Fetch all relevant columns in one query
        column_ids = Column.objects.filter(
            source_id__in=[str(eval.id) for eval in user_evals_list], deleted=False
        ).values_list("source_id", "id")
        column_map = dict(column_ids)

        # Build response with unique eval names
        run_evals = []
        seen_names = set()
        for eval in user_evals_list:
            if eval.name not in seen_names:
                seen_names.add(eval.name)
                template = eval.template
                run_evals.append(
                    {
                        "id": eval.id,
                        "name": eval.name,
                        "template_name": template.name,
                        "eval_template_name": template.name,
                        "eval_required_keys": template.config.get("required_keys", []),
                        "eval_template_tags": template.eval_tags,
                        "description": template.description,
                        "model": eval.config.get("config", {}).get("model", ""),
                        "column_id": column_map.get(str(eval.id), None),
                        "last_updated": eval.updated_at,
                        "mapping": eval.config.get("mapping", {}),
                        "params": eval.config.get("params", {}),
                    }
                )

        return self._gm.success_response({"evals": run_evals})


class ComparePreviewRunEvalView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        try:
            config = request.data.get("config")
            model = request.data.get("model", ModelChoices.TURING_LARGE.value)
            template_id = request.data.get("template_id")
            dataset_ids = request.data.get("dataset_ids", [])
            dataset_info = request.data.get("dataset_info", {})

            protect = False
            is_only_eval = True

            # Get eval template and setup
            eval_template = EvalTemplate.no_workspace_objects.get(id=template_id)
            eval_class = globals().get(eval_template.config.get("eval_type_id"))
            responses = []

            # Create a deep copy of the config
            original_config = copy.deepcopy(config)

            # Evaluation runner setup
            self.futureagi_eval = (
                True
                if eval_template.config.get("eval_type_id") in FUTUREAGI_EVAL_TYPES
                else False
            )
            eval_id = eval_template.config.get("eval_type_id")
            runner = EvaluationRunner(
                eval_id,
                is_only_eval=is_only_eval,
                format_output=True,
                futureagi_eval=self.futureagi_eval,
                source=request.data.get("source", "dataset_evaluation"),
                source_id=template_id,
                protect=protect,
            )
            data_config = config.get("config")
            run_prompt_column = eval_template.config.get("run_prompt_column", False)

            # Convert dataset_ids to a set for faster lookups
            dataset_ids_set = set(dataset_ids)

            # Determine the rows to process
            if not (dataset_info and isinstance(dataset_info, dict)):
                return self._gm.bad_request(
                    get_error_message("NO_DATASET_INFO_PROVIDED")
                )

            row_ids = []
            if dataset_info:
                first_base_value = next(iter(dataset_info.keys()), None)
                if first_base_value:
                    for ds_id in dataset_ids_set:
                        if ds_id in dataset_info[first_base_value]:
                            row_ids.append(dataset_info[first_base_value][ds_id])

            rows = Row.objects.filter(id__in=row_ids, deleted=False)

            # Fetch source from the first dataset if available
            source = (
                Dataset.objects.get(id=dataset_ids[0]).source
                if dataset_ids and rows
                else None
            )

            # Fetch all relevant columns in one query
            all_columns = Column.objects.filter(
                dataset_id__in=dataset_ids, deleted=False
            )
            column_map = {
                (str(col.dataset_id), col.name): str(col.id) for col in all_columns
            }

            # Group rows by dataset
            rows_by_dataset = defaultdict(list)
            for row in rows:
                rows_by_dataset[str(row.dataset.id)].append(row)

            # Process each dataset's rows
            futures = []

            # Helper to extract base column name and JSON path from column_name.path format
            def extract_column_name_and_path(column_name):
                if not column_name or "." not in column_name:
                    return column_name, None
                # Split on first dot to get base column name and JSON path
                dot_index = column_name.find(".")
                return column_name[:dot_index], column_name[dot_index + 1 :]

            # First validate all mappings before starting any threads
            for ds_id, _ds_rows in rows_by_dataset.items():
                dataset_config = copy.deepcopy(original_config)
                mapping = dataset_config.get("mapping", {})
                for key in mapping.keys():
                    column_name = mapping[key]
                    # Extract base column name for JSON paths (e.g., "Input.prompt" -> "Input")
                    base_column_name, _ = extract_column_name_and_path(column_name)
                    if (ds_id, base_column_name) not in column_map:
                        return self._gm.bad_request(
                            get_error_message("COLUMN_NOT_FOUND") + column_name
                        )

            # Wrap function with OTel context propagation for thread safety
            wrapped_process_eval = wrap_for_thread(process_eval_for_single_row)

            with ThreadPoolExecutor(max_workers=10) as executor:
                for ds_id, ds_rows in rows_by_dataset.items():
                    dataset_config = copy.deepcopy(original_config)
                    mapping = dataset_config.get("mapping", {})
                    for key in mapping.keys():
                        column_name = mapping[key]
                        # Handle JSON paths (e.g., "Input.prompt" -> "column_id.prompt")
                        base_column_name, json_path = extract_column_name_and_path(
                            column_name
                        )
                        column_id = column_map[(ds_id, base_column_name)]
                        # Append JSON path to column ID if present
                        mapping[key] = (
                            f"{column_id}.{json_path}" if json_path else column_id
                        )

                    # Process each row in the dataset in parallel
                    for row in ds_rows:
                        future = executor.submit(
                            wrapped_process_eval,
                            runner,
                            request.user,
                            row,
                            mapping,
                            data_config,
                            run_prompt_column,
                            eval_class,
                            eval_template,
                            self.futureagi_eval,
                            source,
                            ds_id,
                            model,
                        )
                        futures.append(future)

                # Collect results from all threads
                for future in futures:
                    response = future.result()
                    responses.append(response)

            return self._gm.success_response({"responses": responses})

        except Exception as e:
            logger.exception(f"Error in preview the evaluation: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_PREVIEW_EVAL")
            )


class CreateKnowledgeBaseView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, JSONParser]

    # Check total size of the KB
    def calculate_total_size(self, files, kb_id=None):
        kb = None
        if kb_id:
            kb = KnowledgeBaseFile.objects.filter(id=kb_id).first()
        size = kb.size if kb is not None else 0
        return size + sum(f.size for f in files)  # 1 GB

    def validate_all_files(self, files):
        """
        Validate ALL files using ThreadPoolExecutor (same pattern as original dev).

        Returns only after all validations complete. Does NOT create any DB records.
        KB should only be created if this returns valid=True.

        Args:
            files: List of uploaded file objects

        Returns:
            dict with valid (bool) and files_with_issues (list)
        """
        from concurrent.futures import ThreadPoolExecutor, as_completed

        files_with_issues = []
        max_workers = max(1, min(len(files), 10))
        futures = {}

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            for file_obj in files:
                futures[executor.submit(self.is_file_readable, file_obj)] = file_obj

            for future in as_completed(futures):
                file_obj = futures[future]
                file_name = file_obj.name

                try:
                    check = future.result()
                except Exception as e:
                    files_with_issues.append(
                        {"name": file_name, "error": f"Unexpected error: {str(e)}"}
                    )
                    continue

                if not check.get("status"):
                    files_with_issues.append(
                        {
                            "name": str(file_name),
                            "error": check.get(
                                "error", f"Unable to upload {file_name}"
                            ),
                        }
                    )

        return {
            "valid": len(files_with_issues) == 0,
            "files_with_issues": files_with_issues,
        }

    # Bounded thread pool for S3 uploads - prevents resource exhaustion
    _upload_executor = None

    @classmethod
    def _get_upload_executor(cls):
        """Get or create the bounded ThreadPoolExecutor for S3 uploads."""
        if cls._upload_executor is None:
            from concurrent.futures import ThreadPoolExecutor

            cls._upload_executor = ThreadPoolExecutor(
                max_workers=10, thread_name_prefix="kb-s3-upload"
            )
        return cls._upload_executor

    def create_files_and_upload(self, files, user, kb_id, org_id=None):
        created_files = []
        file_metadata = {}
        executor = self._get_upload_executor()

        for file_obj in files:
            file_name = file_obj.name
            file_size = file_obj.size
            extension = file_name.split(".")[-1].lower() if "." in file_name else ""

            # Create Files record
            file_instance = Files.objects.create(
                name=file_name,
                status=StatusType.PROCESSING.value,
                metadata=json.dumps({"size": file_size}),
                updated_by=user,
            )

            # Read file bytes for S3 upload
            file_obj.seek(0)
            file_bytes = file_obj.read()

            # Submit to bounded thread pool (max 10 concurrent uploads)
            executor.submit(
                self._upload_file_to_s3_background,
                file_bytes,
                file_name,
                str(kb_id),
                str(file_instance.id),
                org_id,
            )

            file_metadata[str(file_instance.id)] = {
                "name": file_name,
                "extension": extension,
            }
            created_files.append(str(file_instance.id))

        return {
            "files": created_files,
            "file_metadata": file_metadata,
        }

    def _upload_file_to_s3_background(
        self, file_bytes, file_name, kb_id, file_id, org_id=None
    ):
        from django.db import close_old_connections, connection

        from tfc.utils.storage import upload_file_to_s3

        try:
            close_old_connections()
            connection.ensure_connection()

            upload_file_to_s3(
                file_bytes=file_bytes,
                file_name=file_name,
                kb_id=kb_id,
                file_id=file_id,
                org_id=org_id,
            )
            logger.info(f"Background S3 upload completed for file {file_id}")

        except Exception as e:
            logger.error(f"Background S3 upload failed for file {file_id}: {e}")
            try:
                Files.objects.filter(id=file_id).update(
                    status=StatusType.FAILED.value,
                    metadata=json.dumps({"error": str(e)}),
                )
            except Exception as db_error:
                logger.error(f"Failed to update file status: {db_error}")
        finally:
            close_old_connections()

    # Check if file is valid
    def is_file_readable(self, file_obj):
        try:
            file_name = file_obj.name
            extension = file_name.split(".")[-1].lower()

            try:
                file_obj.seek(0)
                file_bytes = file_obj.read()
                file_stream = io.BytesIO(file_bytes)
                file_stream.seek(0)
            except Exception as e:
                logger.error(f"Error reading file: {e}")
                return {
                    "status": False,
                    "error": "File is either password-protected or corrupted",
                }

            if extension == "pdf":
                try:
                    file_stream = io.BytesIO(file_bytes)
                    reader = PdfReader(file_stream)
                    if reader.is_encrypted:
                        try:
                            reader.decrypt("")
                        except Exception:
                            return {
                                "status": False,
                                "error": "File is password-protected",
                            }

                    # Process only the first page
                    text = reader.pages[0].extract_text() if reader.pages else ""
                    text = text.strip() if text else ""
                    if not text:
                        return {"status": False, "error": "File is empty or corrupted"}

                    return {"status": True}

                except PdfReadError:
                    return {"status": False, "error": "Invalid or corrupted PDF file"}
                except Exception as e:
                    return {
                        "status": False,
                        "error": f"Unexpected error: {str(e)}",
                    }

            elif extension == "docx":
                try:
                    doc = Document(io.BytesIO(file_bytes))
                    has_text = any(p.text.strip() for p in doc.paragraphs[:5])
                    has_tables = len(doc.tables) > 0
                    if not has_text and not has_tables:
                        return {
                            "status": False,
                            "error": "DOCX file is empty or corrupted",
                        }
                except Exception as e:
                    error_message = str(e).lower()
                    if "encrypted" in error_message or "password" in error_message:
                        return {"status": False, "error": "File is password-protected"}
                    return {"status": False, "error": "Invalid DOCX file"}

            elif extension == "txt":
                try:
                    content = file_bytes.decode("utf-8")
                    if not content.strip():
                        return {
                            "status": False,
                            "error": "Text file is empty or corrupted",
                        }
                except Exception:
                    return {"status": False, "error": "Invalid text file"}

            elif extension == "rtf":
                try:
                    from striprtf.striprtf import rtf_to_text

                    snippet = rtf_to_text(file_bytes.decode("utf-8", errors="ignore"))
                    if not snippet.strip():
                        return {
                            "status": False,
                            "error": "RTF file is empty or corrupted",
                        }
                except Exception:
                    return {"status": False, "error": "Invalid RTF file"}

            return {"status": True}

        except Exception:
            return {"status": False, "error": "Unknown error while checking the file"}

    def _generate_unique_name(self, org):
        """Generate the next auto-increment name "Knowledge Base - N" for an organization."""
        existing = KnowledgeBaseFile.objects.filter(
            name__startswith="Knowledge Base -", organization=org
        ).values_list("name", flat=True)

        used_numbers = set()
        for name in existing:
            try:
                num = int(name.split("-")[1])
                used_numbers.add(num)
            except (IndexError, ValueError):
                continue

        next_number = 1
        while next_number in used_numbers:
            next_number += 1

        return f"Knowledge Base - {next_number}"

    def upload_created_files(self, file_paths, kb_id, org):
        if not file_paths:
            return None

        max_workers = max(1, min(len(file_paths), 10))
        futures = {}
        uploaded_file_paths = {}
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            for file_id, file_path in file_paths.items():
                futures[
                    executor.submit(upload_file_to_s3, file_path, kb_id, file_id)
                ] = file_id

            for future in as_completed(futures):
                file_id = futures[future]
                try:
                    result = future.result()
                    uploaded_file_paths[file_id] = result
                except Exception as e:
                    logger.info(f"Error in uploading file '{file_id}' to s3: {str(e)}")

        if uploaded_file_paths:
            ingest_files_to_s3.delay(uploaded_file_paths, str(kb_id), str(org))

    # Api for fetching sdk code
    def get(self, request, *args, **kwargs):
        try:
            org = getattr(request, "organization", None) or request.user.organization
            request.query_params.get("kb_id", "YOUR_KB_ID")
            name = request.query_params.get("name", "")
            type = request.query_params.get("type", "create")

            apiKeys = OrgApiKey.objects.filter(
                organization=org, type="user", enabled=True, user=request.user
            )
            if len(apiKeys) == 0:
                org_api_key = OrgApiKey.objects.create(
                    organization=org, type="user", user=request.user
                )
                serialized_keys = OrgApiKeySerializer(
                    org_api_key,
                ).data

            else:
                apiKeys = OrgApiKey.objects.filter(
                    organization=org, type="user", enabled=True, user=request.user
                )
                serialized_keys = OrgApiKeySerializer(
                    apiKeys[0],
                ).data

            if type == "create":
                code = CREATE_KB_SDK_CODE.format(
                    serialized_keys["api_key"],
                    serialized_keys["secret_key"],
                    serialized_keys["api_key"],
                    serialized_keys["secret_key"],
                    name,
                )
            else:
                code = UPDATE_KB_SDK_CODE.format(
                    serialized_keys["api_key"],
                    serialized_keys["secret_key"],
                    name,
                    serialized_keys["api_key"],
                    serialized_keys["secret_key"],
                    "UPDATED_KB_NAME",
                )
            response = {
                "code": code,
            }
            return self._gm.success_response(response)

        except Exception as e:
            logger.exception(f"Error in getting the kb sdk code: {e}")
            return self._gm.internal_server_error_response(
                "Error in getting the kb sdk code"
            )

    def post(self, request, *args, **kwargs):
        try:
            start_time = time.time()
            logger.info(f"START TIME: {str(start_time)}")
            org = getattr(request, "organization", None) or request.user.organization
            kb_name = None

            data = request.data
            created_by = User.objects.get(id=request.user.id).name
            uploaded_files = request.FILES.getlist("file")
            file_names = {file.name for file in uploaded_files}
            if len(file_names) != len(uploaded_files):
                return self._gm.bad_request(get_error_message("DUPLICATE_FILES"))

            updated_size = self.calculate_total_size(uploaded_files)
            if updated_size > MAX_KB_SIZE:
                return self._gm.bad_request(get_error_message("MAX_KB_SIZE_EXCEEDED"))

            entitlements_checked = False
            try:
                from ee.usage.services.entitlements import Entitlements

                kb_count = KnowledgeBaseFile.objects.filter(
                    organization=org, deleted=False
                ).count()
                if Entitlements is not None:
                    ent_check = Entitlements.can_create(
                    str(org.id), "knowledge_bases", kb_count
                )
                if not ent_check.allowed:
                    return self._gm.forbidden_response(ent_check.reason)

                if Entitlements is not None:
                    feat_check = Entitlements.check_feature(
                        str(org.id), "has_knowledge_base"
                    )
                    if not feat_check.allowed:
                        return self._gm.forbidden_response(feat_check.reason)
                entitlements_checked = True
            except ImportError:
                pass

            if not entitlements_checked:
                if log_and_deduct_cost_for_resource_request is not None:
                    call_log_row = log_and_deduct_cost_for_resource_request(
                        organization=org,
                        api_call_type=APICallTypeChoices.KNOWLEDGE_BASE.value,
                        workspace=request.workspace,
                    )
                    if (
                        call_log_row is None
                        or call_log_row.status == APICallStatusChoices.RESOURCE_LIMIT.value
                    ):
                        return self._gm.too_many_requests(
                            get_error_message("KB_CREATION_LIMIT_REACHED")
                        )
                    call_log_row.status = APICallStatusChoices.SUCCESS.value
                    call_log_row.save()

            # Validate ALL files FIRST (before creating KB)
            # Uses is_file_readable for full validation (password check, content parsing)
            validation_result = self.validate_all_files(uploaded_files)
            if not validation_result["valid"]:
                return self._gm.bad_request(
                    {
                        "message": "Some files could not be processed",
                        "not_uploaded": True,
                        "files": validation_result["files_with_issues"],
                    }
                )

            with transaction.atomic():
                # Generate KB name first
                if not data.get("name"):
                    kb_name = self._generate_unique_name(org)

                final_kb_name = data.get("name") if not kb_name else kb_name

                if KnowledgeBaseFile.objects.filter(
                    name=final_kb_name.strip() if final_kb_name else final_kb_name,
                    organization=org,
                ).exists():
                    return self._gm.bad_request(
                        get_error_message("KNOWLEDGE_BASE_ALREADY_EXISTS")
                    )

                # All files are valid - create KB
                kb = KnowledgeBaseFile.objects.create(
                    organization=org,
                    created_by=created_by,
                    name=final_kb_name,
                    size=updated_size,
                )

                # Create file records and start S3 upload (fire-and-forget)
                created_files = self.create_files_and_upload(
                    uploaded_files, created_by, kb.id, org_id=str(org.id)
                )

                kb.files.set(Files.objects.filter(id__in=created_files["files"]))
                kb.save()

                # Schedule ingestion after transaction commits
                schedule_kb_ingestion_on_commit(
                    created_files.get("file_metadata", {}),
                    kb.id,
                    org.id,
                )

            end_time = time.time()
            logger.info(f"END TIME: {str(end_time)}")
            logger.info(f"RESPONSE TIME {str(end_time - start_time)}")

            if request.headers.get("X-Api-Key") is not None:
                properties = get_mixpanel_properties(
                    user=request.user,
                    knowledge_base=kb.id,
                    file_ids=created_files["files"],
                )
                track_mixpanel_event(MixpanelEvents.SDK_KB_CREATE.value, properties)

            response_data = {
                "detail": "Creating Knowledge Base",
                "kb_id": kb.id,
                "kb_name": kb.name,
                "file_ids": created_files["files"],
            }

            return self._gm.success_response(response_data)

        except Exception as e:
            logger.exception(f"Error in creating the knowledge base: {str(e)}")
            return self._gm.bad_request(
                get_error_message("FAILED_TO_CREATE_KNOWLEDGE_BASE")
            )

    # Update knowledge base name and/or Add files
    def patch(self, request, *args, **kwargs):
        try:
            org = getattr(request, "organization", None) or request.user.organization
            kb_id = request.data.get("kb_id")
            files = request.FILES.getlist("file")
            kb_name = None
            user = request.user.name

            if not kb_id or not org:
                return self._gm.bad_request(
                    get_error_message("MISSING_KNOWLEDGE_BASE_ID_OR_ORGANIZATION")
                )

            kb_instance = KnowledgeBaseFile.objects.filter(
                id=kb_id, organization=org, deleted=False
            ).first()
            if not kb_instance:
                return self._gm.bad_request(
                    get_error_message("KNOWLEDGE_BASE_NOT_FOUND")
                )

            try:
                try:
                    from ee.usage.services.entitlements import Entitlements
                except ImportError:
                    Entitlements = None

                if Entitlements is not None:
                    feat_check = Entitlements.check_feature(
                        str(org.id), "has_knowledge_base"
                    )
                    if not feat_check.allowed:
                        return self._gm.forbidden_response(feat_check.reason)
            except ImportError:
                pass

            file_names = {file.name for file in files}
            if len(file_names) != len(files):
                return self._gm.bad_request(get_error_message("DUPLICATE_FILES"))

            updated_size = self.calculate_total_size(files, kb_id)
            if updated_size > MAX_KB_SIZE:
                return self._gm.bad_request(get_error_message("MAX_KB_SIZE_EXCEEDED"))

            if Files.objects.filter(
                knowledge_base_files__id=kb_id,
                name__in=list(file_names),
                knowledge_base_files__organization=org,
            ).exists():
                return self._gm.bad_request(get_error_message("FILE_ALREADY_EXISTS"))

            if request.data.get("name") and not (
                request.data.get("name").strip() == kb_instance.name.strip()
            ):
                kb_name = request.data.get("name").strip()

                if KnowledgeBaseFile.objects.filter(
                    name=kb_name, organization=org, deleted=False
                ).exists():
                    return self._gm.bad_request(
                        get_error_message("KNOWLEDGE_BASE_ALREADY_EXISTS")
                    )

            if not request.data.get("name"):
                kb_name = self._generate_unique_name(org)

            # Validate ALL files FIRST (same as POST)
            if files:
                validation_result = self.validate_all_files(files)
                if not validation_result["valid"]:
                    return self._gm.bad_request(
                        {
                            "message": "Some files could not be processed",
                            "not_uploaded": True,
                            "files": validation_result["files_with_issues"],
                        }
                    )

            with transaction.atomic():
                kb_instance.name = kb_name if kb_name else kb_instance.name
                if files:
                    # All files valid - create records and start S3 upload
                    created_files = self.create_files_and_upload(
                        files, user, kb_instance.id, org_id=str(org.id)
                    )

                    kb_instance.files.add(
                        *Files.objects.filter(id__in=created_files["files"])
                    )
                    kb_instance.size = updated_size

                    # Schedule ingestion after transaction commits
                    schedule_kb_ingestion_on_commit(
                        created_files.get("file_metadata", {}),
                        kb_instance.id,
                        org.id,
                    )

                kb_instance.save()

            response_data = KnowledgeBaseFileSerializer(kb_instance).data
            return self._gm.success_response(response_data)

        except Exception as e:
            logger.exception(f"Error in updating the knowledge base: {str(e)}")
            return self._gm.bad_request(
                get_error_message("FAILED_TO_UPDATE_KNOWLEDGE_BASE")
            )

    # Delete knowledge base
    def delete(self, request, *args, **kwargs):
        try:
            kb_ids = request.data.get("kb_ids", [])
            org = getattr(request, "organization", None) or request.user.organization

            if not kb_ids or not org:
                return self._gm.bad_request(
                    get_error_message("MISSING_KNOWLEDGE_BASE_ID_OR_ORGANIZATION")
                )

            # Cancel ingestion workflows for KBs in PROCESSING state
            processing_kbs = KnowledgeBaseFile.objects.filter(
                id__in=kb_ids,
                organization=org,
                status=StatusType.PROCESSING.value,
            ).values_list("id", flat=True)

            for kb_id in processing_kbs:
                cancel_kb_ingestion_workflow(kb_id)

            remove_kb_files.delay(None, str(org.id), kb_ids)
            KnowledgeBaseFile.objects.filter(id__in=kb_ids, organization=org).update(
                deleted=True
            )
            return self._gm.success_response("Successfully deleted the Knowledge Base")

        except Exception as e:
            logger.exception(f"Error in deleting the knowledge base: {str(e)}")
            return self._gm.bad_request(
                get_error_message("FAILED_TO_DELETE_KNOWLEDGE_BASE")
            )


class GetKnowledgeBaseDetailsView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    # List knowledge base table data
    def get(self, request, *args, **kwargs):
        try:
            org = getattr(request, "organization", None) or request.user.organization
            search = request.query_params.get("search", None)
            sort_config = request.query_params.get("sort", [])
            page_number = int(request.query_params.get("page_number", 0)) or int(
                request.query_params.get("pageNumber", 0)
            )
            page_size = int(request.query_params.get("page_size", 10)) or int(
                request.query_params.get("pageSize", 10)
            )

            if search:
                search = json.loads(search)

            start = page_size * page_number
            end = start + page_size

            if not org:
                return self._gm.bad_request(
                    get_error_message("ORGANIZATION_ID_MISSING")
                )

            else:
                required_columns = [
                    "id",
                    "name",
                    "files_uploaded",
                    "status",
                    "updated_at",
                    "created_by",
                ]
                column_config = [
                    {"id": f"column_{i}", "name": col}
                    for i, col in enumerate(required_columns)
                ]
                kbs = KnowledgeBaseFile.objects.filter(
                    organization_id=org.id, deleted=False
                ).all()
                if not kbs:
                    if KnowledgeBaseFile.all_objects.filter(
                        organization=org, deleted=True
                    ).exists():
                        return self._gm.success_response(
                            {"column_config": column_config}
                        )
                    else:
                        return self._gm.success_response({})

                if search:
                    kbs = kbs.filter(name__icontains=search)

                serializer = KnowledgeBaseFileSerializer(kbs, many=True)

                response = {}
                data = serializer.data
                col_type = {
                    "name": DataTypeChoices.TEXT.value,
                    "files_uploaded": DataTypeChoices.INTEGER.value,
                    "status": DataTypeChoices.TEXT.value,
                    "updated_at": DataTypeChoices.DATETIME.value,
                    "created_by": DataTypeChoices.TEXT.value,
                }

                if sort_config:
                    for sort_item in sort_config:
                        try:
                            column_id = sort_item.get("column_id")
                            sort_type = sort_item.get("type")
                            reverse = sort_type == "descending"

                            if not column_id or not sort_type:
                                continue

                            if column_id not in col_type.keys():
                                return self._gm.bad_request(
                                    "Invalid Column ID for sorting."
                                )

                            column_type = col_type.get(column_id)
                            if column_type == DataTypeChoices.TEXT.value:
                                sort_data = sorted(
                                    data,
                                    key=lambda x: x[column_id] or "",
                                    reverse=reverse,
                                )

                            elif column_type == DataTypeChoices.INTEGER.value:
                                sort_data = sorted(
                                    data, key=lambda x: len(x["files"]), reverse=reverse
                                )

                            elif column_type == DataTypeChoices.DATETIME.value:
                                sort_data = sorted(
                                    data,
                                    key=lambda x: (
                                        datetime.strptime(
                                            x[column_id], "%Y-%m-%dT%H:%M:%S.%fZ"
                                        )
                                        if x.get(column_id)
                                        else datetime.min
                                    ),
                                    reverse=reverse,
                                )

                        except Exception as e:
                            logger.exception(f"Error in sorting: {e}")

                    table_data = sort_data[start:end]
                else:
                    table_data = data[start:end]

                table_rows = []
                for entry in table_data:
                    table_rows.append(
                        {
                            "id": entry["id"],
                            "name": entry["name"],
                            "files_uploaded": len(entry["files"]),
                            "status": entry["status"],
                            "error": (
                                entry.get("last_error", None)
                                if entry["status"]
                                not in [
                                    StatusType.COMPLETED.value,
                                    StatusType.PROCESSING.value,
                                ]
                                else None
                            ),
                            "updated_at": entry["updated_at"],
                            "created_by": entry["created_by"],
                        }
                    )

                response = {"table_data": table_rows, "total_rows": len(data)}
            return self._gm.success_response(response)
        except Exception as e:
            logger.exception(f"Error in fetching knowledge bases: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_KNOWLEDGE_BASE")
            )


class ListKnowledgeBaseDetailsView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    # List knowledge base in org for dropdowns
    def get(self, request, *args, **kwargs):
        try:
            org = getattr(request, "organization", None) or request.user.organization
            search = request.GET.get("search", None)

            status = request.GET.get("status", False)
            if search:
                kbs = KnowledgeBaseFile.objects.filter(
                    name__icontains=search, organization_id=org.id, deleted=False
                )
            else:
                kbs = KnowledgeBaseFile.objects.filter(
                    organization_id=org.id, deleted=False
                )

            if status:
                kbs = list(
                    kbs.filter(status=StatusType.COMPLETED.value).values("id", "name")
                )
            else:
                kbs = list(kbs.values("id", "name"))

            table_rows = [
                {
                    "id": entry["id"],
                    "name": entry["name"],
                }
                for entry in kbs
            ]

            response = {
                "table_data": table_rows,
            }
            return self._gm.success_response(response)
        except Exception as e:
            logger.exception(f"Error in fetching knowledge bases: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_KNOWLEDGE_BASE")
            )


class ExistingKnowledgeBaseView(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    # List files present in the KB
    def post(self, request, *args, **kwargs):
        try:
            org = getattr(request, "organization", None) or request.user.organization
            kb_id = request.data.get("kb_id", None)
            search = request.data.get("search", None)
            sort_config = request.data.get("sort", [])
            page_number = int(request.data.get("page_number", 0))
            page_size = int(request.data.get("page_size", 10))

            if not kb_id or not org:
                return self._gm.bad_request(
                    get_error_message("MISSING_KNOWLEDGE_BASE_ID_OR_ORGANIZATION")
                )
            kb = KnowledgeBaseFile.objects.prefetch_related("files").filter(
                id=kb_id, organization=org
            )

            if not kb:
                return self._gm.bad_request(
                    get_error_message("KNOWLEDGE_BASE_NOT_FOUND")
                )

            files = (
                kb.first().files.filter(
                    Q(name__icontains=search) & ~Q(status=StatusType.DELETING.value)
                )
                if search
                else kb.first().files.exclude(status=StatusType.DELETING.value)
            )
            serializer = FileSerializer(files, many=True)

            data = serializer.data
            start = page_size * page_number
            end = start + page_size if start + page_size < len(data) else len(data)

            col_type = {
                "name": DataTypeChoices.TEXT.value,
                "file_size": DataTypeChoices.INTEGER.value,
                "status": DataTypeChoices.TEXT.value,
                "updated_at": DataTypeChoices.DATETIME.value,
                "updated_by": DataTypeChoices.TEXT.value,
            }

            if sort_config:
                for sort_item in sort_config:
                    try:
                        column_id = sort_item.get("column_id")
                        sort_type = sort_item.get("type")
                        reverse = sort_type == "descending"

                        if not column_id or not sort_type:
                            continue

                        if column_id not in col_type.keys():
                            return self._gm.bad_request(
                                "Invalid Column ID for sorting."
                            )

                        column_type = col_type.get(column_id)
                        if column_type == DataTypeChoices.TEXT.value:
                            sort_data = sorted(
                                data, key=lambda x: x[column_id] or "", reverse=reverse
                            )

                        elif column_type == DataTypeChoices.INTEGER.value:
                            sort_data = sorted(
                                data,
                                key=lambda x: json.loads(x["metadata"]).get("size", 0),
                                reverse=reverse,
                            )

                        elif column_type == DataTypeChoices.DATETIME.value:
                            sort_data = sorted(
                                data,
                                key=lambda x: (
                                    datetime.strptime(
                                        x[column_id], "%Y-%m-%dT%H:%M:%S.%fZ"
                                    )
                                    if x.get(column_id)
                                    else datetime.min
                                ),
                                reverse=reverse,
                            )

                    except Exception as e:
                        logger.exception(f"Error in sorting: {e}")

                table_data = sort_data[start:end]
            else:
                table_data = data[start:end]

            table_rows = []
            for entry in table_data:
                if entry.get("status") == StatusType.DELETING.value:
                    continue
                metadata = json.loads(entry["metadata"])
                size = metadata.get("size", 0)
                error = metadata.get("error", None)
                table_rows.append(
                    {
                        "id": entry["id"],
                        "name": entry["name"],
                        "file_size": size,
                        "status": entry["status"],
                        "updated": entry["updated_at"],
                        "updated_by": entry["updated_by"],
                    }
                )
                if error is not None:
                    table_rows[-1]["error"] = error

            count = 0
            kb_status = StatusType.COMPLETED.value
            processing_files = (
                kb.first()
                .files.filter(status=StatusType.PROCESSING.value, deleted=False)
                .all()
                .count()
            )
            deleting_files = (
                kb.first()
                .files.filter(status=StatusType.DELETING.value, deleted=False)
                .all()
                .count()
            )
            if processing_files > 0:
                count = processing_files
                kb_status = StatusType.PROCESSING.value
            if deleting_files > 0:
                count = deleting_files
                kb_status = StatusType.DELETING.value

            response = {
                "table_data": table_rows,
                "last_updated": kb.first().updated_at,
                "status": kb_status,
                "status_count": count,
                "total_rows": len(data),
            }

            return self._gm.success_response(response)
        except Exception as e:
            logger.exception(f"Error in fetching the knowledge base: {str(e)}")
            return self._gm.internal_server_error_response(
                get_error_message("FAILED_TO_GET_KNOWLEDGE_BASE")
            )

    # Delete files from kb
    def delete(self, request, *args, **kwargs):
        try:
            org = getattr(request, "organization", None) or request.user.organization.id
            kb_id = request.data.get("kb_id", None)
            delete_all = request.data.get("delete_all", False)
            file_ids = request.data.get("file_ids", [])
            excluded_file_ids = request.data.get("excluded_file_ids", [])
            file_names = request.data.get("file_names", [])

            if not kb_id or not org:
                return self._gm.bad_request(
                    get_error_message("MISSING_KNOWLEDGE_BASE_ID_OR_ORGANIZATION")
                )

            # Case 1: Delete all files except excluded ones
            if delete_all:
                deleted_files = Files.objects.filter(
                    knowledge_base_files__id=kb_id,
                    knowledge_base_files__organization=org,
                )
                if excluded_file_ids:
                    deleted_files = deleted_files.exclude(id__in=excluded_file_ids)
            # Case 2: Delete specific files by IDs or names
            else:
                if file_ids:
                    deleted_files = Files.objects.filter(id__in=file_ids)
                elif file_names:
                    deleted_files = Files.objects.filter(
                        knowledge_base_files__id=kb_id,
                        name__in=file_names,
                        knowledge_base_files__organization=org,
                    )
                else:
                    return self._gm.bad_request(
                        get_error_message("MISSING_FILE_IDS_OR_NAMES")
                    )

            deleted_files.update(status=StatusType.DELETING.value)
            remove_kb_files.delay(
                list(deleted_files.values_list("id", flat=True)), str(org), kb_id
            )

            return self._gm.success_response("Deleting selected files")
        except Exception as e:
            logger.exception(f"Error in deleting the knowledge base files: {str(e)}")
            return self._gm.bad_request(
                get_error_message("FAILED_TO_DELETE_KNOWLEDGE_BASE_FILES")
            )


class GetDatasetExplanationSummary(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def get(self, request, dataset_id):
        try:
            organization = (
                getattr(request, "organization", None) or request.user.organization
            )
            dataset = get_object_or_404(
                Dataset, id=dataset_id, organization=organization
            )
            eval_reasons = dataset.eval_reasons
            eval_reason_last_updated = dataset.eval_reason_last_updated
            status = dataset.eval_reason_status

            row_count = Row.objects.filter(dataset_id=dataset_id, deleted=False).count()

            if row_count < MIN_ROWS_FOR_CRITICAL_ISSUES:
                status = EvalExplanationSummaryStatus.INSUFFICIENT_DATA
                dataset.eval_reason_status = status
                dataset.save(update_fields=["eval_reason_status"])
            elif eval_reason_last_updated is None:
                dataset.eval_reason_status = EvalExplanationSummaryStatus.PENDING
                dataset.save(update_fields=["eval_reason_status"])
                get_explanation_summary.delay(str(dataset_id))

            return self._gm.success_response(
                {
                    "response": eval_reasons,
                    "last_updated": eval_reason_last_updated,
                    "status": status,
                    "row_count": row_count,
                    "min_rows_required": MIN_ROWS_FOR_CRITICAL_ISSUES,
                }
            )

        except Exception as e:
            logger.exception(
                f"Error getting explanation summary: {e}", dataset_id=dataset_id
            )
            return self._gm.bad_request(
                get_error_message("UNABLE_TO_FETCH_EVAL_REASON_SUMMARY")
            )


class RefreshDatasetExplanationSummary(APIView):
    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def post(self, request, dataset_id):
        try:
            organization = (
                getattr(request, "organization", None) or request.user.organization
            )
            dataset = get_object_or_404(
                Dataset, id=dataset_id, organization=organization
            )

            row_count = Row.objects.filter(dataset_id=dataset_id, deleted=False).count()

            if row_count < MIN_ROWS_FOR_CRITICAL_ISSUES:
                dataset.eval_reason_status = (
                    EvalExplanationSummaryStatus.INSUFFICIENT_DATA
                )
                dataset.save(update_fields=["eval_reason_status"])
                return self._gm.success_response(
                    {
                        "response": dataset.eval_reasons,
                        "last_updated": dataset.eval_reason_last_updated,
                        "status": dataset.eval_reason_status,
                        "row_count": row_count,
                        "min_rows_required": MIN_ROWS_FOR_CRITICAL_ISSUES,
                    }
                )

            dataset.eval_reason_status = EvalExplanationSummaryStatus.PENDING
            dataset.save(update_fields=["eval_reason_status"])

            get_explanation_summary.delay(str(dataset_id))

            return self._gm.success_response(
                {
                    "response": dataset.eval_reasons,
                    "last_updated": dataset.eval_reason_last_updated,
                    "status": dataset.eval_reason_status,
                    "row_count": row_count,
                    "min_rows_required": MIN_ROWS_FOR_CRITICAL_ISSUES,
                }
            )

        except Exception as e:
            logger.exception(
                f"Error refreshing explanation summary: {e}", dataset_id=dataset_id
            )
            return self._gm.bad_request(
                get_error_message("UNABLE_TO_FETCH_EVAL_REASON_SUMMARY")
            )


def get_json_column_schemas(dataset):
    """
    Get JSON schemas and images metadata for columns in a dataset.
    Used for autocomplete suggestions when accessing JSON properties
    and for indexed access to images columns.

    Args:
        dataset: Dataset model instance

    Returns:
        dict: JSON schemas keyed by column ID
    """
    from model_hub.utils.json_path_resolver import (
        extract_json_schema_for_column,
        parse_json_safely,
    )

    # Get JSON-type columns + text columns that may contain JSON values
    json_columns = Column.objects.filter(
        dataset=dataset,
        data_type__in=["json", "text"],
        deleted=False,
    )

    result = {}

    for column in json_columns:
        # Check if schema already exists in metadata
        metadata = column.metadata or {}
        json_schema = metadata.get("json_schema")

        if json_schema and (json_schema.get("keys") or json_schema.get("max_array_count")):
            # Use cached schema
            entry = {
                "name": column.name,
                "keys": json_schema.get("keys", []),
                "sample": json_schema.get("sample"),
            }
            if json_schema.get("max_array_count"):
                entry["max_array_count"] = json_schema["max_array_count"]
            result[str(column.id)] = entry
        else:
            # Extract schema from cells. For text columns, check a small
            # sample first to avoid wasting time on non-JSON content.
            base_qs = (
                Cell.objects.filter(column=column, deleted=False)
                .exclude(value__isnull=True)
                .exclude(value="")
            )

            if column.data_type == "text":
                # Quick check: peek at first 3 non-empty cells
                peek = list(base_qs.values_list("value", flat=True)[:3])
                has_json = any(
                    parse_json_safely(v)[1] for v in peek
                )
                if not has_json:
                    continue
                sample_cells = list(base_qs.values_list("value", flat=True)[:500])
            else:
                sample_cells = list(base_qs.values_list("value", flat=True)[:500])

            if sample_cells:
                schema = extract_json_schema_for_column(list(sample_cells))

                if schema.get("keys") or schema.get("max_array_count"):
                    # Store schema in column metadata
                    metadata["json_schema"] = schema
                    column.metadata = metadata
                    column.save(update_fields=["metadata"])

                    entry = {
                        "name": column.name,
                        "keys": schema.get("keys", []),
                        "sample": schema.get("sample"),
                    }
                    if schema.get("max_array_count"):
                        entry["max_array_count"] = schema["max_array_count"]
                    result[str(column.id)] = entry

    # Get images-type columns and calculate max_images_count
    images_columns = Column.objects.filter(
        dataset=dataset,
        data_type=DataTypeChoices.IMAGES.value,
        deleted=False,
    )

    for column in images_columns:
        try:
            # Get all cell values for this column and find max array length
            images_cells = (
                Cell.objects.filter(
                    column=column,
                    deleted=False,
                    status=CellStatus.PASS.value,
                )
                .exclude(value__isnull=True)
                .exclude(value="")
            )

            max_count = 0
            for cell in images_cells[:500]:  # Limit to first 500 cells for performance
                try:
                    images_list = (
                        json.loads(cell.value)
                        if isinstance(cell.value, str)
                        else cell.value
                    )
                    if isinstance(images_list, list):
                        max_count = max(max_count, len(images_list))
                except (json.JSONDecodeError, TypeError):
                    pass

            if max_count > 0:
                result[str(column.id)] = {
                    "name": column.name,
                    "max_images_count": max_count,
                }
        except Exception as e:
            logger.warning(
                f"Error calculating max_images_count for column {column.id}: {e}"
            )

    return result


class GetJsonColumnSchemaView(APIView):
    """
    API endpoint to get JSON schemas and images metadata for columns in a dataset.
    Used by frontend for autocomplete suggestions when accessing JSON properties
    and for indexed access to images columns.

    Returns:
    - For JSON-type columns: schema from Column.metadata["json_schema"] with keys for autocomplete
    - For images-type columns: max_images_count for indexed access (e.g., images[0], images[1])

    If schema not yet extracted, computes it from sample cells and stores it.
    """

    _gm = GeneralMethods()
    permission_classes = [IsAuthenticated]

    def get(self, request, dataset_id, *args, **kwargs):
        try:
            dataset = get_object_or_404(
                Dataset,
                id=dataset_id,
                organization=getattr(request, "organization", None)
                or request.user.organization,
                deleted=False,
            )

            result = get_json_column_schemas(dataset)
            return self._gm.success_response(result)

        except Exception as e:
            logger.exception(f"Error getting JSON schema: {e}", dataset_id=dataset_id)
            return self._gm.bad_request(str(e))
